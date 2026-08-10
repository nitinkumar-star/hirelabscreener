from flask import Flask, request, jsonify, send_file, Response
from flask_cors import CORS
import sqlite3, json, os, datetime, requests, shutil, io, re, smtplib, time
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from pathlib import Path
try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
try:
    import invoice_engine
    HAS_INVOICE = True
except Exception:
    HAS_INVOICE = False
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

app = Flask(__name__, static_folder='.', static_url_path='')
# NOTE: flask-cors' CORS(app) allows ALL origins, which is unsafe with
# credentialed requests. CORS is handled by add_cors_headers() below with a
# strict allow-list (app origin + Naukri extension). Do NOT re-enable CORS(app).
# CORS(app)

# ── Safety net: never let a raw BLOB (e.g. embedding_vec bytes) crash jsonify.
# Any bytes that reach serialization become null instead of a 500. Defense in
# depth — list endpoints also strip heavy embedding columns before returning.
try:
    from flask.json.provider import DefaultJSONProvider

    class _SafeJSONProvider(DefaultJSONProvider):
        def default(self, o):
            if isinstance(o, (bytes, bytearray, memoryview)):
                return None
            return super().default(o)

    app.json = _SafeJSONProvider(app)
except Exception as _json_err:
    print(f'[json-provider] safe provider not installed: {_json_err}')


# Heavy per-candidate columns the UI never needs; stripped from list/detail
# responses to keep payloads small and memory low (fixes OOM on large mandates).
_HEAVY_CAND_COLS = ('embedding', 'embedding_text', 'embedding_vec')

def _cand_public(row):
    """Row -> client-safe dict without the heavy embedding columns."""
    d = dict(row)
    for k in _HEAVY_CAND_COLS:
        d.pop(k, None)
    return d
app.config['MAX_CONTENT_LENGTH'] = 64 * 1024 * 1024  # 64MB max upload

def _cors_origin_allowed(origin):
    """Only trusted origins may make credentialed cross-site calls.
    Defaults cover the app itself + the Naukri Chrome extension. Extra origins
    can be added via the CORS_ALLOWED_ORIGINS env var (comma-separated)."""
    if not origin:
        return False
    # Chrome extension always trusted (that's the whole point of CORS here).
    if origin.startswith('chrome-extension://'):
        return True
    if origin in ('https://www.naukri.com', 'https://naukri.com'):
        return True
    try:
        if origin == request.host_url.rstrip('/'):
            return True
    except Exception:
        pass
    extra = os.environ.get('CORS_ALLOWED_ORIGINS', '') or ''
    allow = [o.strip() for o in extra.split(',') if o.strip()]
    return origin in allow


@app.after_request
def add_cors_headers(resp):
    # Echo the Origin ONLY for trusted callers (see _cors_origin_allowed).
    # Arbitrary websites get no ACAO header, so a logged-in user's browser
    # will not expose credentialed API responses to them.
    origin = request.headers.get('Origin')
    if origin and _cors_origin_allowed(origin):
        resp.headers['Access-Control-Allow-Origin'] = origin
        resp.headers['Access-Control-Allow-Credentials'] = 'true'
        resp.headers['Vary'] = 'Origin'
    resp.headers['Access-Control-Allow-Methods'] = 'GET, POST, PUT, DELETE, OPTIONS'
    resp.headers['Access-Control-Allow-Headers'] = 'Content-Type'
    return resp


# ── Login Protection ──────────────────────────────────────────────────────────
from functools import wraps
from flask import session, redirect as flask_redirect
import hashlib, secrets, hmac

def _get_secret_key():
    # 1) Prefer env var (set SECRET_KEY in Render for best security)
    env_secret = os.environ.get('SECRET_KEY', '').strip()
    if env_secret:
        return env_secret
    # 2) Else use/create a random secret stored on the persistent disk so login
    #    sessions survive restarts. This is far safer than a hardcoded fallback.
    try:
        secret_path = os.path.join(DATA_DIR, '.secret_key')
        if os.path.exists(secret_path):
            with open(secret_path) as f:
                v = f.read().strip()
                if v:
                    return v
        import secrets as _secrets
        v = _secrets.token_hex(32)
        with open(secret_path, 'w') as f:
            f.write(v)
        return v
    except Exception:
        # 3) Last resort (ephemeral) — sessions reset on restart, but never hardcoded
        import secrets as _secrets
        return _secrets.token_hex(32)

app.secret_key = _get_secret_key()

# ── Session cookie hardening ────────────────────────────────────────────────
# SameSite must stay 'None' because the Naukri Chrome extension makes
# cross-site credentialed calls (its session cookie must be sent). 'None'
# REQUIRES Secure=True. On Render everything is HTTPS; set COOKIE_SECURE=0 only
# for local HTTP testing.
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SECURE=(os.environ.get('COOKIE_SECURE', '1') != '0'),
    SESSION_COOKIE_SAMESITE='None',
)

# ── Login brute-force throttle (in-memory, per-IP) ──────────────────────────
_LOGIN_FAILS = {}          # ip -> [failure_timestamps]
_LOGIN_MAX_FAILS = 8       # allowed failures within the window
_LOGIN_WINDOW = 900        # 15 minutes

def _login_ip():
    fwd = request.headers.get('X-Forwarded-For', '')
    if fwd:
        return fwd.split(',')[0].strip()
    return request.remote_addr or 'unknown'

def _login_is_blocked(ip):
    now = time.time()
    arr = [t for t in _LOGIN_FAILS.get(ip, []) if now - t < _LOGIN_WINDOW]
    _LOGIN_FAILS[ip] = arr
    return len(arr) >= _LOGIN_MAX_FAILS

def _login_note_fail(ip):
    _LOGIN_FAILS.setdefault(ip, []).append(time.time())

def _login_clear(ip):
    _LOGIN_FAILS.pop(ip, None)

# ── Generic per-IP rate limiter (for public/unauthenticated endpoints) ──────
_RL_BUCKETS = {}   # (bucket, ip) -> [timestamps]

def _rate_ok(bucket, max_n, window):
    """Return True if this IP is still under the limit for `bucket`, else False.
    Sliding window of `window` seconds, `max_n` requests allowed."""
    ip = _login_ip()
    now = time.time()
    key = (bucket, ip)
    arr = [t for t in _RL_BUCKETS.get(key, []) if now - t < window]
    arr.append(now)
    _RL_BUCKETS[key] = arr
    return len(arr) <= max_n

# ─────────────────────────────────────────────────────────────────────────
#  AUTH HELPERS (multi-user)
# ─────────────────────────────────────────────────────────────────────────
def hash_password(pw, salt=None):
    # PBKDF2-HMAC-SHA256 — strong, slow, and stdlib (no external dependency).
    # Format:  pbkdf2$<iterations>$<salt_hex>$<hash_hex>
    iterations = 210000
    if salt is None:
        salt = secrets.token_hex(16)
    dk = hashlib.pbkdf2_hmac('sha256', pw.encode(), salt.encode(), iterations)
    return 'pbkdf2$' + str(iterations) + '$' + salt + '$' + dk.hex()


def verify_password(pw, stored):
    try:
        if (stored or '').startswith('pbkdf2$'):
            _, iters, salt, h = stored.split('$', 3)
            dk = hashlib.pbkdf2_hmac('sha256', pw.encode(), salt.encode(), int(iters))
            return secrets.compare_digest(dk.hex(), h)
        # Legacy format: <salt>$<sha256hex>. Still verifies so existing users can
        # log in; login transparently re-hashes them to PBKDF2 (see auth_login).
        salt, h = stored.split('$', 1)
        return secrets.compare_digest(hashlib.sha256((salt + pw).encode()).hexdigest(), h)
    except Exception:
        return False


def password_needs_rehash(stored):
    """True if the stored hash is the old (weaker) SHA-256 format."""
    return not (stored or '').startswith('pbkdf2$')

def any_users_exist():
    conn = get_db()
    n = conn.execute('SELECT COUNT(*) n FROM users').fetchone()['n']
    conn.close()
    return n > 0

def current_user():
    """The logged-in user. If admin is 'viewing as' another tenant, the
    EFFECTIVE workspace is that tenant, but real identity stays admin."""
    uid = session.get('user_id')
    if not uid:
        return None
    conn = get_db()
    u = conn.execute('SELECT * FROM users WHERE id=?', (uid,)).fetchone()
    conn.close()
    return dict(u) if u else None

def current_company_id():
    """The company (tenant) of the logged-in user, ignoring any view-as."""
    u = current_user()
    return u.get('company_id') if u else None

def effective_user_id():
    """TENANT-ID this request operates on. Historically named for the
    single-user era; it now returns the effective COMPANY (tenant) id, which
    is what every data row's owner_id column stores. A super-admin can
    impersonate another tenant via 'view_as_company'. Kept under the old name
    so all existing `WHERE owner_id = effective_user_id()` filters keep working
    and enforce tenant isolation automatically."""
    va = session.get('view_as_company')
    if va:
        return va
    return current_company_id()

# Clearer alias for new code.
def effective_company_id():
    return effective_user_id()


# ── Canonical "successful placement" definition ────────────────────────────
# A candidate counts as placed if they're in the Placed/Joined stage, OR they
# carry a real offer/joining record (offered CTC, placement fee, or a joining
# date) — unless they sit in an explicitly-negative stage (offer fell through).
# This catches candidates who were placed via the Billing module but whose
# pipeline stage was never moved. Reused by Analytics, Command Center, the admin
# summary and the placed-list endpoint so the KPI number and the click-through
# list ALWAYS agree.
_PLACED_REJECT_STAGES = ('not interested', 'not suitable',
                         'client rejected on paper', 'client rejected after interview')
PLACED_SQL = (
    "( TRIM(LOWER(COALESCE(stage,''))) IN ('placed','joined') "
    "  OR ( (COALESCE(offered_ctc,0)>0 OR COALESCE(placement_fee,0)>0 "
    "        OR TRIM(COALESCE(joining_date,''))!='') "
    "       AND TRIM(LOWER(COALESCE(stage,''))) NOT IN "
    "       ('not interested','not suitable','client rejected on paper','client rejected after interview') ) )"
)

def is_admin():
    u = current_user()
    return u and u.get('role') == 'admin'

def is_company_admin():
    """True if the logged-in user can manage their whole company (see all
    company mandates, assign them, delete jobs). This is the platform owner OR
    a user flagged as their company's admin. When a super-admin is viewing-as a
    company, they act as that company's admin."""
    if session.get('view_as_company'):
        return True  # super-admin impersonating a tenant acts as its admin
    u = current_user()
    if not u:
        return False
    return u.get('role') == 'admin' or u.get('is_company_admin') == 1

def real_user_id():
    """The actual logged-in user id (NOT the company id that effective_user_id
    returns). Used for per-recruiter mandate assignment."""
    return session.get('user_id')


def _tenant_owns_candidate(conn, cid):
    """True if candidate `cid` belongs to the current tenant (company).
    owner_id stores the company id, so this enforces cross-agency isolation."""
    r = conn.execute('SELECT owner_id FROM candidates WHERE id=?', (cid,)).fetchone()
    return bool(r) and r['owner_id'] == effective_company_id()


def _tenant_owns_mandate(conn, mid):
    """True if mandate `mid` belongs to the current tenant (company)."""
    r = conn.execute('SELECT owner_id FROM mandates WHERE id=?', (mid,)).fetchone()
    return bool(r) and r['owner_id'] == effective_company_id()

def log_activity(action, detail='', entity_type='', entity_id=0, meta=None,
                 actor_type='user', actor_name=''):
    """Universal activity timeline. Backward-compatible: existing callers that
    pass only (action, detail) keep working. New callers can attach an entity
    (candidate/client/invoice), an actor (user/client/system) and a JSON meta
    payload that future workflow-automation can consume without schema changes."""
    try:
        u = current_user()
        uid = u['id'] if u else 0
        uname = actor_name or (u['username'] if u else 'system')
        try:
            company_id = effective_company_id()
        except Exception:
            company_id = 0
        meta_json = ''
        if meta is not None:
            try:
                meta_json = json.dumps(meta)
            except Exception:
                meta_json = ''
        conn = get_db()
        conn.execute(
            'INSERT INTO activity_log (user_id,username,action,detail,created_at,'
            'company_id,entity_type,entity_id,actor_type,actor_name,meta) '
            'VALUES (?,?,?,?,?,?,?,?,?,?,?)',
            (uid, uname, action, detail, ts(), company_id, entity_type, entity_id,
             actor_type, uname, meta_json))
        conn.commit(); conn.close()
    except Exception:
        pass


def log_audit(entity_type, entity_id, field, old_value, new_value,
              actor_type='user', actor_name=''):
    """Record a single field change (old → new) for audit history."""
    try:
        u = current_user()
        uid = u['id'] if u else 0
        uname = actor_name or (u['username'] if u else 'system')
        try:
            company_id = effective_company_id()
        except Exception:
            company_id = 0
        conn = get_db()
        conn.execute(
            'INSERT INTO audit_log (company_id,entity_type,entity_id,field,old_value,'
            'new_value,actor_type,actor_id,actor_name,created_at) VALUES (?,?,?,?,?,?,?,?,?,?)',
            (company_id, entity_type, entity_id, field, str(old_value or ''),
             str(new_value or ''), actor_type, uid, uname, ts()))
        conn.commit(); conn.close()
    except Exception:
        pass


def record_changes(entity_type, entity_id, before: dict, after: dict, fields,
                   actor_type='user', actor_name=''):
    """Diff two dicts across `fields` and write one audit row per changed field.
    Returns a human-readable summary list of the changes (for activity detail)."""
    changes = []
    for f in fields:
        old_v = before.get(f) if before else None
        new_v = after.get(f) if after else None
        if str(old_v or '') != str(new_v or ''):
            log_audit(entity_type, entity_id, f, old_v, new_v, actor_type, actor_name)
            changes.append(f'{f}: {old_v or "\u2014"} \u2192 {new_v or "\u2014"}')
    return changes


def _ist_now():
    """Current time in IST (India Standard Time, UTC+5:30).
    Render servers run in UTC, so we add the offset to get correct local time."""
    return datetime.datetime.utcnow() + datetime.timedelta(hours=5, minutes=30)


def utcnow_iso():
    # Despite the name, we return IST so timestamps match the user's local time.
    return _ist_now().isoformat()


def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        # API calls return 401 JSON; page loads redirect to /login
        if not session.get('user_id'):
            if request.path.startswith('/api/'):
                return jsonify({'error': 'auth_required'}), 401
            return flask_redirect('/login')
        return f(*args, **kwargs)
    return decorated

def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('user_id'):
            return jsonify({'error': 'auth_required'}), 401
        if not is_admin():
            return jsonify({'error': 'Admin access required'}), 403
        return f(*args, **kwargs)
    return decorated


@app.route('/api/diag')
@admin_required
def diag():
    """Diagnostic: shows whether the persistent disk and data are intact.
    Helps debug 'data disappeared / logged out / asked to create new account'."""
    info = {'data_dir': DATA_DIR, 'db_path': DB_PATH}
    try:
        info['db_exists'] = os.path.exists(DB_PATH)
        info['data_dir_writable'] = os.access(DATA_DIR, os.W_OK)
        info['secret_key_file_exists'] = os.path.exists(os.path.join(DATA_DIR, '.secret_key'))
        info['reset_marker'] = os.path.exists(os.path.join(DATA_DIR, '.last_reset'))
        info['reset_data_env'] = bool(os.environ.get('RESET_DATA'))
        info['secret_key_env'] = bool(os.environ.get('SECRET_KEY'))
        info['data_dir_env_set'] = bool(os.environ.get('DATA_DIR'))

        # Storage persistence (the key signal for the 'data keeps disappearing' bug)
        info['storage_persistent'] = _PERSISTENCE.get('persistent')
        info['restarts_survived'] = _PERSISTENCE.get('boots_seen', 0)

        # Backups present on disk
        try:
            baks = sorted(Path(BAK_DIR).glob('hirelab_*.db'), reverse=True)
            info['backup_count'] = len(baks)
            info['latest_backup'] = baks[0].name if baks else None
            info['latest_backup_users'] = _db_user_count(str(baks[0])) if baks else 0
        except Exception as e:
            info['backup_error'] = str(e)

        conn = get_db(); c = conn.cursor()
        for t in ['users', 'companies', 'mandates', 'candidates']:
            try: info[t + '_count'] = c.execute(f'SELECT COUNT(*) FROM {t}').fetchone()[0]
            except Exception as e: info[t + '_count'] = f'err: {e}'
        conn.close()

        # Plain-language diagnosis
        looks_ephemeral = (not info['data_dir_env_set']) and ('HireLab' in DATA_DIR)
        if looks_ephemeral:
            info['diagnosis'] = ('DATA_DIR is not set to the mounted disk — data is on TEMPORARY '
                                 'storage and WILL be lost on restart. On Render: attach a disk at '
                                 '/data and set env var DATA_DIR=/data.')
        elif info.get('storage_persistent') is True:
            info['diagnosis'] = 'OK — storage is persistent and has survived restarts.'
        elif info.get('users_count') == 0 and info.get('latest_backup_users', 0) > 0:
            info['diagnosis'] = 'DB is empty but a backup with users exists — auto-restore should recover on next start.'
        else:
            info['diagnosis'] = ('Persistence not yet confirmed. Restart the service once and re-check; '
                                 'restarts_survived should increase if the disk is persistent.')
    except Exception as e:
        info['error'] = str(e)
    return jsonify(info)



@app.route('/api/auth/status')
def auth_status():
    """Tells the frontend whether to show: first-run admin setup, login, or app."""
    if not any_users_exist():
        return jsonify({'state': 'setup'})
    u = current_user()
    if not u:
        return jsonify({'state': 'login'})
    va = session.get('view_as_company')
    viewing = None
    conn = get_db()
    if va:
        vc = conn.execute('SELECT id,name FROM companies WHERE id=?', (va,)).fetchone()
        viewing = {'id': vc['id'], 'name': vc['name']} if vc else None
    # The user's own company name (for the top bar)
    own_company = None
    if u.get('company_id'):
        oc = conn.execute('SELECT id,name FROM companies WHERE id=?', (u['company_id'],)).fetchone()
        own_company = {'id': oc['id'], 'name': oc['name']} if oc else None
    pending_count = 0
    if u.get('role') == 'admin':
        pending_count = conn.execute("SELECT COUNT(*) n FROM users WHERE status='pending'").fetchone()['n']
    conn.close()
    return jsonify({'state': 'app', 'user': {
        'id': u['id'], 'username': u['username'], 'display_name': u['display_name'],
        'role': u['role'], 'company': own_company,
        'is_company_admin': (u.get('role') == 'admin' or u.get('is_company_admin') == 1),
        'workflow_mode': (get_setting('workflow_mode', 'agency') or 'agency')
    }, 'viewing_as': viewing, 'pending_count': pending_count})


@app.route('/api/auth/setup', methods=['POST'])
def auth_setup():
    """First-run: create the first admin. Only works if no users exist."""
    if any_users_exist():
        return jsonify({'error': 'Setup already complete'}), 400
    d = request.json or {}
    username = (d.get('username') or '').strip()
    password = d.get('password') or ''
    display = (d.get('display_name') or username).strip()
    if not username or len(password) < 8:
        return jsonify({'error': 'Username required and password min 8 chars'}), 400
    conn = get_db()
    display_company = (d.get('company_name') or 'HireLab').strip() or 'HireLab'
    conn.execute("INSERT INTO companies (name,status,plan,billing_status,created_at) VALUES (?,?,?,?,?)",
                 (display_company, 'active', 'owner', 'owner', ts()))
    company_id = conn.execute('SELECT id FROM companies ORDER BY id DESC LIMIT 1').fetchone()['id']
    conn.execute('INSERT INTO users (username,password_hash,display_name,role,created_at,status,company_name,company_id,is_company_admin) VALUES (?,?,?,?,?,?,?,?,1)',
                 (username, hash_password(password), display, 'admin', ts(), 'approved', display_company, company_id))
    conn.commit()
    uid = conn.execute('SELECT id FROM users WHERE username=?', (username,)).fetchone()['id']
    # This is the FIRST admin (no users existed before this). Any data already in
    # the DB therefore belongs to a previous, now-deleted user. Claim ALL of it
    # for this admin's COMPANY (tenant), so imported data is visible. Safe because
    # this branch only runs when no users existed.
    conn.execute('UPDATE mandates SET owner_id=?', (company_id,))
    conn.execute('UPDATE candidates SET owner_id=?', (company_id,))
    conn.execute('UPDATE reminders SET owner_id=?', (company_id,))
    conn.commit(); conn.close()
    session['user_id'] = uid
    return jsonify({'ok': True})


@app.route('/api/auth/signup', methods=['POST'])
def auth_signup():
    """Public self-signup. Creates the account as 'pending' — it cannot log
    in until a super-admin approves it. (If no users exist yet at all, this
    is the very first account, so it is created as an approved admin instead
    — see auth_setup for that bootstrap path.)"""
    if not any_users_exist():
        return jsonify({'error': 'No admin account exists yet. Use the initial setup screen instead.'}), 400
    d = request.json or {}
    username = (d.get('username') or '').strip().lower()
    password = d.get('password') or ''
    display = (d.get('display_name') or username).strip()
    company = (d.get('company_name') or '').strip()
    email = (d.get('email') or '').strip().lower()
    if not username or len(password) < 8:
        return jsonify({'error': 'Username required and password min 8 chars'}), 400
    if not re.match(r'^[a-z0-9._-]{3,40}$', username):
        return jsonify({'error': 'Username can only contain letters, numbers, dots, dashes and underscores'}), 400
    if not email or '@' not in email:
        return jsonify({'error': 'A valid email is required (used for password reset)'}), 400
    if not d.get('accept_terms'):
        return jsonify({'error': 'Please accept the Terms of Service and Privacy Policy to continue.'}), 400
    log_activity('signup_consent', username + ' accepted terms @ ' + ts())
    conn = get_db()
    exists = conn.execute('SELECT id FROM users WHERE username=?', (username,)).fetchone()
    if exists:
        conn.close()
        return jsonify({'error': 'Username already taken'}), 400
    # Create the agency's company (tenant) up front, in 'pending' status. It is
    # activated when the super-admin approves the user. New signups are regular
    # 'user' accounts within their own company — they get full access to their
    # own workspace but NOT the platform-level super-admin panel (which is
    # reserved for the platform owner).
    company_label = company or (display + "'s agency")
    conn.execute("INSERT INTO companies (name,status,plan,created_at) VALUES (?,?,?,?)",
                 (company_label, 'pending', 'standard', ts()))
    new_company_id = conn.execute('SELECT id FROM companies ORDER BY id DESC LIMIT 1').fetchone()['id']
    conn.execute('''INSERT INTO users (username,password_hash,display_name,role,created_at,status,company_name,requested_at,company_id,is_company_admin,email)
                     VALUES (?,?,?,?,?,?,?,?,?,1,?)''',
                 (username, hash_password(password), display, 'user', ts(), 'pending', company, ts(), new_company_id, email))
    conn.commit(); conn.close()
    log_activity('signup_requested', username + (' (' + company + ')' if company else ''))
    return jsonify({'ok': True, 'pending': True})


@app.route('/api/admin/pending-users', methods=['GET'])
@admin_required
def list_pending_users():
    conn = get_db()
    rows = conn.execute('''SELECT id, username, display_name, company_name, requested_at
                            FROM users WHERE status='pending' ORDER BY id''').fetchall()
    conn.close()
    return jsonify({'ok': True, 'pending': [dict(r) for r in rows]})


def company_user_limit(cid):
    if not cid:
        return 0
    conn = get_db()
    r = conn.execute("SELECT user_limit FROM companies WHERE id=?", (cid,)).fetchone()
    conn.close()
    try:
        return int((r['user_limit'] if r else 0) or 0)
    except Exception:
        return 0


def company_approved_users(cid):
    if not cid:
        return 0
    conn = get_db()
    r = conn.execute("SELECT COUNT(*) n FROM users WHERE company_id=? AND status='approved'", (cid,)).fetchone()
    conn.close()
    return int((r['n'] if r else 0) or 0)


@app.route('/api/admin/pending-users/<int:uid>/approve', methods=['POST'])
@admin_required
def approve_pending_user(uid):
    conn = get_db()
    u = conn.execute('SELECT username, status, company_id FROM users WHERE id=?', (uid,)).fetchone()
    if not u:
        conn.close()
        return jsonify({'error': 'User not found'}), 404
    _lim = company_user_limit(u['company_id'])
    if _lim and company_approved_users(u['company_id']) >= _lim:
        conn.close()
        return jsonify({'error': f'This agency has reached its user limit ({_lim}). Increase the limit before approving more users.'}), 403
    conn.execute("UPDATE users SET status='approved' WHERE id=?", (uid,))
    if u['company_id']:
        # Activate the company and start its free trial.
        try:
            trial_days = int(get_setting('billing_trial_days', '14') or 14)
        except Exception:
            trial_days = 14
        trial_end = (datetime.datetime.now() + datetime.timedelta(days=trial_days)).isoformat()
        conn.execute("UPDATE companies SET status='active', billing_status='trial', trial_ends_at=? WHERE id=? AND (trial_ends_at IS NULL OR trial_ends_at='')",
                     (trial_end, u['company_id']))
        # If trial was already set (re-approval), just activate.
        conn.execute("UPDATE companies SET status='active' WHERE id=?", (u['company_id'],))
    conn.commit(); conn.close()
    log_activity('approve_user', u['username'])
    return jsonify({'ok': True})


@app.route('/api/admin/pending-users/<int:uid>/reject', methods=['POST'])
@admin_required
def reject_pending_user(uid):
    conn = get_db()
    u = conn.execute('SELECT username, status FROM users WHERE id=?', (uid,)).fetchone()
    if not u:
        conn.close()
        return jsonify({'error': 'User not found'}), 404
    conn.execute("UPDATE users SET status='rejected' WHERE id=?", (uid,))
    conn.commit(); conn.close()
    log_activity('reject_user', u['username'])
    return jsonify({'ok': True})


@app.route('/api/auth/login', methods=['POST'])
def auth_login():
    d = request.json or {}
    username = (d.get('username') or '').strip()
    password = d.get('password') or ''
    ip = _login_ip()
    if _login_is_blocked(ip):
        return jsonify({'error': 'Too many failed attempts. Please wait about 15 minutes and try again.'}), 429
    conn = get_db()
    u = conn.execute('SELECT * FROM users WHERE username=?', (username,)).fetchone()
    if not u or not verify_password(password, u['password_hash']):
        conn.close()
        _login_note_fail(ip)
        return jsonify({'error': 'Invalid username or password'}), 401
    # Transparent upgrade: if this account still has the old SHA-256 hash,
    # re-hash to PBKDF2 now that we have the plaintext password.
    if password_needs_rehash(u['password_hash']):
        try:
            conn.execute('UPDATE users SET password_hash=? WHERE id=?', (hash_password(password), u['id']))
            conn.commit()
        except Exception:
            pass
    if u['status'] == 'pending':
        conn.close()
        return jsonify({'error': 'Your account is awaiting admin approval. You will be able to sign in once approved.'}), 403
    if u['status'] == 'rejected':
        conn.close()
        return jsonify({'error': 'This account request was declined. Contact your admin for access.'}), 403
    # Block login if the tenant company is suspended (super-admin can suspend
    # an agency e.g. for non-payment). The platform owner is never blocked.
    if u['role'] != 'admin' and u['company_id']:
        comp = conn.execute('SELECT status, billing_status, trial_ends_at, plan FROM companies WHERE id=?', (u['company_id'],)).fetchone()
        if comp and comp['status'] == 'suspended':
            conn.close()
            return jsonify({'error': 'Your agency account is currently suspended. Please contact support.'}), 403
        # Trial expiry: if on trial and the trial period has passed without
        # converting to a paid subscription, block until they subscribe.
        if comp and comp['plan'] != 'owner' and comp['billing_status'] == 'trial' and comp['trial_ends_at']:
            try:
                te = datetime.datetime.fromisoformat(comp['trial_ends_at'])
                if datetime.datetime.now() > te:
                    conn.execute("UPDATE companies SET billing_status='past_due' WHERE id=?", (u['company_id'],))
                    conn.commit(); conn.close()
                    return jsonify({'error': 'Your free trial has ended. Please subscribe to continue using HireLab.'}), 402
            except Exception:
                pass
    conn.execute('UPDATE users SET last_login=? WHERE id=?', (ts(), u['id']))
    conn.commit(); conn.close()
    session['user_id'] = u['id']
    session.pop('view_as_company', None)
    session.permanent = True
    _login_clear(ip)
    log_activity('login', username)
    return jsonify({'ok': True})


@app.route('/api/auth/logout', methods=['POST'])
def auth_logout():
    log_activity('logout')
    session.clear()
    return jsonify({'ok': True})


@app.route('/api/auth/forgot', methods=['POST'])
def auth_forgot():
    """Email a password-reset link. Always returns a generic success so we never
    reveal whether an account exists. Reset link points at /?reset=<token>."""
    d = request.json or {}
    ident = (d.get('identifier') or d.get('username') or d.get('email') or '').strip()
    generic = {'ok': True, 'message': 'If an account matches, a reset link has been sent to its email.'}
    if not ident:
        return jsonify(generic)
    conn = get_db()
    u = conn.execute('SELECT id, username, COALESCE(email, "") AS email FROM users WHERE lower(username)=lower(?) OR lower(COALESCE(email,""))=lower(?)',
                     (ident, ident)).fetchone()
    if not u:
        conn.close(); return jsonify(generic)
    target_email = (u['email'] or '').strip()
    if not target_email and '@' in u['username']:
        target_email = u['username'].strip()
    if not target_email:
        conn.close()
        print('[forgot] user "%s" has NO email on file — cannot send reset. (Old accounts created before the email field have no email.)' % ident)
        return jsonify(generic)
    token = secrets.token_urlsafe(32)
    expires = (datetime.datetime.now() + datetime.timedelta(hours=1)).strftime('%Y-%m-%d %H:%M:%S')
    conn.execute('INSERT INTO password_resets (token,user_id,expires_at,used,created_at) VALUES (?,?,?,0,?)',
                 (token, u['id'], expires, ts()))
    conn.commit(); conn.close()
    base = request.host_url.rstrip('/')
    link = f'{base}/?reset={token}'
    subject = 'Reset your HireLab Screener password'
    plain = (f'Hi {u["username"]},\n\nWe received a request to reset your password. '
             f'Click the link below (valid for 1 hour):\n\n{link}\n\n'
             f'If you did not request this, you can safely ignore this email.')
    html = (f'<div style="font-family:sans-serif;font-size:14px;color:#333">'
            f'<p>Hi {u["username"]},</p>'
            f'<p>We received a request to reset your password. This link is valid for <b>1 hour</b>:</p>'
            f'<p><a href="{link}" style="display:inline-block;background:#13A37E;color:#fff;'
            f'padding:10px 18px;border-radius:8px;text-decoration:none;font-weight:600">Reset password</a></p>'
            f'<p style="font-size:12px;color:#888">Or paste this URL:<br>{link}</p>'
            f'<p style="font-size:12px;color:#888">If you did not request this, you can ignore this email.</p></div>')
    ok_send, err_send = _platform_smtp_send(target_email, subject, plain, html)
    if ok_send:
        print('[forgot] reset email SENT to %s' % target_email)
    else:
        print('[forgot] FAILED to send reset to %s: %s' % (target_email, err_send))
    return jsonify(generic)


@app.route('/api/auth/reset', methods=['POST'])
def auth_reset():
    """Complete a password reset using a valid token."""
    d = request.json or {}
    token = (d.get('token') or '').strip()
    password = d.get('password') or ''
    if not token or len(password) < 8:
        return jsonify({'error': 'Invalid link or password too short (min 8 chars)'}), 400
    conn = get_db()
    row = conn.execute('SELECT token,user_id,expires_at,used FROM password_resets WHERE token=?', (token,)).fetchone()
    if not row or row['used']:
        conn.close(); return jsonify({'error': 'This reset link is invalid or already used.'}), 400
    try:
        exp = datetime.datetime.strptime(row['expires_at'], '%Y-%m-%d %H:%M:%S')
    except Exception:
        exp = datetime.datetime.now() - datetime.timedelta(seconds=1)
    if datetime.datetime.now() > exp:
        conn.close(); return jsonify({'error': 'This reset link has expired. Please request a new one.'}), 400
    conn.execute('UPDATE users SET password_hash=? WHERE id=?', (hash_password(password), row['user_id']))
    conn.execute('UPDATE password_resets SET used=1 WHERE token=?', (token,))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


# ═══════════════════════════════════════════════════════════════════════════
#  CLICK-TO-CALL: device registration + push-to-dial
# ═══════════════════════════════════════════════════════════════════════════
@app.route('/api/devices/register', methods=['POST'])
@login_required
def register_device():
    """Android app sends its FCM token after login. We store it so we can push
    call requests to the user's phone later."""
    d = request.json or {}
    fcm_token = (d.get('fcm_token') or '').strip()
    device_name = (d.get('device_name') or 'Unknown device').strip()
    if not fcm_token:
        return jsonify({'error': 'fcm_token required'}), 400
    uid = real_user_id()
    conn = get_db()
    # Upsert: if this token already exists for this user, update; otherwise insert
    existing = conn.execute('SELECT id FROM devices WHERE user_id=? AND fcm_token=?',
                            (uid, fcm_token)).fetchone()
    if existing:
        conn.execute('UPDATE devices SET is_active=1, device_name=?, updated_at=? WHERE id=?',
                     (device_name, ts(), existing['id']))
    else:
        conn.execute('INSERT INTO devices (user_id,fcm_token,device_name,is_active,created_at,updated_at) '
                     'VALUES (?,?,?,1,?,?)', (uid, fcm_token, device_name, ts(), ts()))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/push-call', methods=['POST'])
@login_required
def push_call():
    """Desktop webapp calls this when the user clicks "Call via Phone".
    Sends an FCM push to the user's phone(s).

    NOTE: we send a VISIBLE notification (not data-only). Android 10+ forbids an
    app from launching the dialer from the background, so a silent data push
    appears to do nothing. With a notification the phone always shows it, and
    tapping it opens the app with the number ready to dial."""
    d = request.json or {}
    phone = (d.get('phone') or '').strip()
    name = (d.get('name') or 'Candidate').strip()
    if not phone:
        return jsonify({'error': 'Phone number required'}), 400

    uid = real_user_id()
    conn = get_db()
    devs = conn.execute('SELECT fcm_token, device_name FROM devices WHERE user_id=? AND is_active=1',
                        (uid,)).fetchall()
    conn.close()
    if not devs:
        return jsonify({'error': 'No phone connected. Open the HireLab Dialer app on your phone and login first.'}), 400

    payload = {
        'action': 'call',
        'phone': phone,
        'name': name,
        'kind': 'call',
        'title': 'Call ' + name,
    }
    # IMPORTANT: send DATA-ONLY (no 'notification' block). The app's CallReceiver
    # builds the notification itself with a tap-to-dial (ACTION_DIAL tel:) intent.
    # If we attach a 'notification' block, Android draws it and the app's dial
    # intent never runs, so tapping only opens the app instead of the dialer.
    sent = _send_fcm_to_user(uid, payload)
    if sent > 0:
        return jsonify({'ok': True, 'sent': sent, 'devices': len(devs),
                        'message': 'Call push sent to %d device(s) - check your phone.' % sent})
    return jsonify({'error': 'Push failed - try re-opening the Dialer app on your phone.'}), 500


@app.route('/api/devices/list')
@login_required
def devices_list():
    """Which phones are registered for pushes (helps diagnose stale devices)."""
    conn = get_db()
    rows = conn.execute('SELECT id, device_name, is_active, created_at, updated_at '
                        'FROM devices WHERE user_id=? ORDER BY id DESC', (real_user_id(),)).fetchall()
    conn.close()
    return jsonify({'ok': True, 'devices': [dict(r) for r in rows]})


# ── FCM V1 helper: get OAuth2 access token from service account JSON ─────
_fcm_token_cache = {'token': '', 'expires': 0, 'project_id': ''}


def _send_fcm_to_user(uid, data_payload, notification=None):
    """Send an FCM push to all of a user's active devices.
    data_payload: dict of string->string.
    notification: optional {'title':..., 'body':...}. When provided, Android
    shows a system-tray notification itself even if the app has no handler for
    this payload (needed because apps can't launch the dialer from the
    background on Android 10+). Returns number of devices reached."""
    conn = get_db()
    tokens = conn.execute('SELECT fcm_token FROM devices WHERE user_id=? AND is_active=1',
                          (uid,)).fetchall()
    conn.close()
    if not tokens:
        return 0
    access_token, project_id, err = _get_fcm_access_token()
    if err:
        print(f'[fcm] token error: {err}')
        return 0
    # FCM data values must all be strings
    data = {k: (str(v) if v is not None else '') for k, v in data_payload.items()}
    sent = 0
    for row in tokens:
        try:
            _msg = {
                'token': row['fcm_token'],
                'data': data,
                'android': {'priority': 'high'},
            }
            if notification:
                _msg['notification'] = {
                    'title': str(notification.get('title', '')),
                    'body': str(notification.get('body', '')),
                }
                # Tapping the notification opens the app and hands it the data.
                _msg['android']['notification'] = {'click_action': 'FLUTTER_NOTIFICATION_CLICK'}
            resp = requests.post(
                f'https://fcm.googleapis.com/v1/projects/{project_id}/messages:send',
                json={'message': _msg},
                headers={'Authorization': f'Bearer {access_token}',
                         'Content-Type': 'application/json'},
                timeout=10)
            if resp.status_code == 200:
                sent += 1
            else:
                print(f'[fcm] error {resp.status_code}: {resp.text[:150]}')
        except Exception as e:
            print(f'[fcm] send failed: {e}')
    return sent


def _reminder_candidate_details(conn, r):
    """Build a rich detail payload for a reminder's candidate (for the push)."""
    cid = r['candidate_id']
    cand = conn.execute(
        'SELECT name, phone, company, designation, stage FROM candidates WHERE id=?',
        (cid,)).fetchone()
    # last note = most recent candidate event of type 'note' or the reminder note
    last_note = r['note'] or ''
    try:
        ev = conn.execute(
            "SELECT detail FROM candidate_events WHERE candidate_id=? AND event_type='note' "
            "ORDER BY created_at DESC LIMIT 1", (cid,)).fetchone()
        if ev and ev['detail']:
            last_note = ev['detail']
    except Exception:
        pass
    name = (cand['name'] if cand else '') or r['candidate_name'] or 'Candidate'
    _tok = _reminder_token(r['id'])
    try:
        _base = request.host_url.rstrip('/')
    except Exception:
        _base = (os.environ.get('PUBLIC_BASE_URL', '') or 'https://hirelabscreener.onrender.com').rstrip('/')
    return {
        'action': 'reminder',
        'reminder_id': str(r['id']),
        'reminder_token': _tok,
        'done_url': (_base + '/api/reminders/%d/done?token=%s' % (r['id'], _tok)) if _base else '',
        'snooze_url': (_base + '/api/reminders/%d/snooze?token=%s' % (r['id'], _tok)) if _base else '',
        'candidate_id': str(cid),
        'name': name,
        'phone': (cand['phone'] if cand else '') or '',
        'company': (cand['company'] if cand else '') or '',
        'designation': (cand['designation'] if cand else '') or '',
        'stage': (cand['stage'] if cand else '') or '',
        'mandate_label': r['mandate_label'] or '',
        'note': (last_note or '')[:400],
        'due_at': r['due_at'] or '',
    }


def _reminder_scheduler_loop():
    """Background loop: every 60s, check for due reminders and send push
    notifications. Repeats every 5 min until the reminder is done or snoozed.
    Also sends an early warning 5-10 min before the due time."""
    import time as _time
    # Small delay so the app is fully up before first check
    _time.sleep(15)
    while True:
        try:
            now = _ist_now()
            now_iso = now.isoformat(timespec='seconds')
            conn = get_db()
            # All active reminders not yet done
            rows = conn.execute(
                'SELECT * FROM reminders WHERE done=0 AND due_at IS NOT NULL AND due_at!=""'
            ).fetchall()
            for r in rows:
                try:
                    due_str = r['due_at']
                    # normalise: reminders may be 'YYYY-MM-DDTHH:MM' or with seconds
                    due = None
                    for fmt in ('%Y-%m-%dT%H:%M:%S', '%Y-%m-%dT%H:%M'):
                        try:
                            due = datetime.datetime.strptime(due_str[:19], fmt)
                            break
                        except Exception:
                            continue
                    if not due:
                        continue

                    owner = r['owner_id'] if ('owner_id' in r.keys()) else 0
                    if not owner:
                        continue

                    snoozed_until = r['snoozed_until'] if ('snoozed_until' in r.keys()) else ''
                    if snoozed_until:
                        try:
                            su = datetime.datetime.strptime(snoozed_until[:19], '%Y-%m-%dT%H:%M:%S')
                            if now < su:
                                continue  # still snoozed
                        except Exception:
                            pass

                    early_warned = r['early_warned'] if ('early_warned' in r.keys()) else 0
                    notified_at = r['notified_at'] if ('notified_at' in r.keys()) else ''

                    mins_to_due = (due - now).total_seconds() / 60.0

                    # EARLY WARNING: 5-10 min before due
                    if not early_warned and 0 < mins_to_due <= 10:
                        payload = _reminder_candidate_details(conn, r)
                        payload['kind'] = 'early'
                        payload['title'] = 'Upcoming: ' + payload['name']
                        _send_fcm_to_user(owner, payload)
                        conn.execute('UPDATE reminders SET early_warned=1 WHERE id=?', (r['id'],))
                        conn.commit()
                        continue

                    # DUE NOW (or overdue): notify, then repeat every 5 min —
                    # but STOP after 3 pushes so a reminder never spams forever
                    # (e.g. if the phone app can't reach the server to mark done).
                    if mins_to_due <= 0:
                        ncount = r['notify_count'] if ('notify_count' in r.keys() and r['notify_count'] is not None) else 0
                        if ncount >= 3:
                            continue  # already nudged enough; it stays in the ATS list
                        send = False
                        if not notified_at:
                            send = True
                        else:
                            try:
                                last = datetime.datetime.strptime(notified_at[:19], '%Y-%m-%dT%H:%M:%S')
                                if (now - last).total_seconds() >= 300:  # 5 min
                                    send = True
                            except Exception:
                                send = True
                        if send:
                            payload = _reminder_candidate_details(conn, r)
                            payload['kind'] = 'due'
                            payload['title'] = 'Reminder: ' + payload['name']
                            _send_fcm_to_user(owner, payload)
                            conn.execute('UPDATE reminders SET notified_at=?, notify_count=? WHERE id=?',
                                         (now_iso, ncount + 1, r['id']))
                            conn.commit()
                except Exception as _re:
                    print(f'[reminder-scheduler] row error: {_re}')
            conn.close()
        except Exception as e:
            print(f'[reminder-scheduler] loop error: {e}')
        _time.sleep(60)


_reminder_thread_started = False
def _start_reminder_scheduler():
    global _reminder_thread_started
    if _reminder_thread_started:
        return
    _reminder_thread_started = True
    import threading
    t = threading.Thread(target=_reminder_scheduler_loop, daemon=True)
    t.start()
    print('[reminder-scheduler] background thread started')


def _get_fcm_access_token():
    """Get a short-lived OAuth2 access token for FCM V1 API.
    Reads the service account JSON from either:
      1. FCM_SERVICE_ACCOUNT_JSON env var (the entire JSON string), or
      2. A file at DATA_DIR/firebase-service-account.json
    Caches the token until it expires."""
    import time
    now = time.time()
    if _fcm_token_cache['token'] and _fcm_token_cache['expires'] > now + 60:
        return _fcm_token_cache['token'], _fcm_token_cache['project_id'], None

    # Load service account credentials
    sa_json = os.environ.get('FCM_SERVICE_ACCOUNT_JSON', '').strip()
    sa_path = os.path.join(DATA_DIR, 'firebase-service-account.json')

    try:
        if sa_json:
            import io
            sa_info = json.loads(sa_json)
        elif os.path.exists(sa_path):
            with open(sa_path) as f:
                sa_info = json.load(f)
        else:
            return None, None, ('FCM not configured. Either:\n'
                                '1. Upload firebase-service-account.json to your data folder, or\n'
                                '2. Set FCM_SERVICE_ACCOUNT_JSON env var on Render with the full JSON content.')
    except Exception as e:
        return None, None, f'Failed to read service account: {e}'

    project_id = sa_info.get('project_id', '')
    if not project_id:
        return None, None, 'Service account JSON missing project_id.'

    # Build a JWT and exchange for an access token (no external library needed)
    try:
        import jwt as _jwt_lib
        _has_pyjwt = True
    except ImportError:
        _has_pyjwt = False

    if _has_pyjwt:
        token, exp = _fcm_token_via_pyjwt(sa_info, now)
    else:
        token, exp = _fcm_token_via_manual_jwt(sa_info, now)

    if token:
        _fcm_token_cache['token'] = token
        _fcm_token_cache['expires'] = exp
        _fcm_token_cache['project_id'] = project_id
        return token, project_id, None
    return None, None, 'Failed to generate FCM access token. Check service account JSON.'


def _fcm_token_via_pyjwt(sa_info, now):
    """Use PyJWT library if available."""
    import jwt, time
    payload = {
        'iss': sa_info['client_email'],
        'scope': 'https://www.googleapis.com/auth/firebase.messaging',
        'aud': 'https://oauth2.googleapis.com/token',
        'iat': int(now),
        'exp': int(now) + 3600,
    }
    signed = jwt.encode(payload, sa_info['private_key'], algorithm='RS256')
    resp = requests.post('https://oauth2.googleapis.com/token', data={
        'grant_type': 'urn:ietf:params:oauth:grant-type:jwt-bearer',
        'assertion': signed,
    }, timeout=15)
    data = resp.json()
    return data.get('access_token'), int(now) + data.get('expires_in', 3500)


def _fcm_token_via_manual_jwt(sa_info, now):
    """Build JWT manually without any external library (pure Python + stdlib)."""
    import base64, hashlib, hmac, struct, time
    from cryptography.hazmat.primitives import hashes, serialization
    from cryptography.hazmat.primitives.asymmetric import padding

    header = base64.urlsafe_b64encode(json.dumps(
        {'alg': 'RS256', 'typ': 'JWT'}).encode()).rstrip(b'=')
    claims = base64.urlsafe_b64encode(json.dumps({
        'iss': sa_info['client_email'],
        'scope': 'https://www.googleapis.com/auth/firebase.messaging',
        'aud': 'https://oauth2.googleapis.com/token',
        'iat': int(now), 'exp': int(now) + 3600,
    }).encode()).rstrip(b'=')
    signing_input = header + b'.' + claims

    private_key = serialization.load_pem_private_key(
        sa_info['private_key'].encode(), password=None)
    signature = private_key.sign(signing_input, padding.PKCS1v15(), hashes.SHA256())
    sig_b64 = base64.urlsafe_b64encode(signature).rstrip(b'=')

    jwt_token = (signing_input + b'.' + sig_b64).decode()
    resp = requests.post('https://oauth2.googleapis.com/token', data={
        'grant_type': 'urn:ietf:params:oauth:grant-type:jwt-bearer',
        'assertion': jwt_token,
    }, timeout=15)
    data = resp.json()
    return data.get('access_token'), int(now) + data.get('expires_in', 3500)


@app.route('/api/users', methods=['GET'])
@admin_required
def list_users():
    conn = get_db()
    rows = conn.execute('''SELECT u.id, u.username, u.display_name, u.role, u.created_at,
                                  u.last_login, u.status, u.company_name, u.company_id,
                                  co.name AS company_label, co.status AS company_status
                           FROM users u LEFT JOIN companies co ON co.id = u.company_id
                           ORDER BY u.id''').fetchall()
    out = []
    for u in rows:
        d = dict(u)
        # Counts are per-tenant (company), since owner_id stores the company id.
        cid = u['company_id']
        if cid:
            d['mandate_count'] = conn.execute('SELECT COUNT(*) n FROM mandates WHERE owner_id=?', (cid,)).fetchone()['n']
            d['candidate_count'] = conn.execute('SELECT COUNT(*) n FROM candidates WHERE owner_id=?', (cid,)).fetchone()['n']
        else:
            d['mandate_count'] = 0
            d['candidate_count'] = 0
        out.append(d)
    conn.close()
    return jsonify({'ok': True, 'users': out})

@app.route('/api/users', methods=['POST'])
@admin_required
def create_user():
    """Add a RECRUITER to a company. This never creates a super-admin and, by
    default, adds the user to the caller's OWN company (agency). To onboard a
    NEW agency (separate tenant) use signup+approve or the owner create-agency
    flow — NOT this endpoint."""
    d = request.json or {}
    username = (d.get('username') or '').strip()
    password = d.get('password') or ''
    display = (d.get('display_name') or username).strip()
    # SECURITY: this endpoint can never mint a super-admin. Agency members are
    # recruiters (role='user'); optionally flag as their company's admin.
    role = 'user'
    is_company_admin = 1 if d.get('is_company_admin') else 0
    # Only the platform owner may place a user in a DIFFERENT company; everyone
    # else (incl. agency admins) can only add to their own company.
    if is_admin() and d.get('company_id'):
        company_id = d.get('company_id')
    else:
        company_id = current_company_id()
    if not username or len(password) < 8:
        return jsonify({'error': 'Username required and password min 8 chars'}), 400
    # Enforce the agency's seat (user) limit set by the platform owner.
    _lim = company_user_limit(company_id)
    if _lim and company_approved_users(company_id) >= _lim:
        return jsonify({'error': f'User limit reached ({_lim} recruiters). Ask the platform owner to increase your limit.'}), 403
    conn = get_db()
    exists = conn.execute('SELECT id FROM users WHERE username=?', (username,)).fetchone()
    if exists:
        conn.close()
        return jsonify({'error': 'Username already taken'}), 400
    conn.execute('INSERT INTO users (username,password_hash,display_name,role,created_at,status,company_id,is_company_admin) VALUES (?,?,?,?,?,?,?,?)',
                 (username, hash_password(password), display, role, ts(), 'approved', company_id, is_company_admin))
    conn.commit(); conn.close()
    log_activity('create_user', username + ' (recruiter)')
    return jsonify({'ok': True})

@app.route('/api/users/<int:uid>/password', methods=['POST'])
@admin_required
def reset_user_password(uid):
    d = request.json or {}
    password = d.get('password') or ''
    if len(password) < 8:
        return jsonify({'error': 'Password min 8 chars'}), 400
    conn = get_db()
    u = conn.execute('SELECT username FROM users WHERE id=?', (uid,)).fetchone()
    if not u:
        conn.close()
        return jsonify({'error': 'User not found'}), 404
    conn.execute('UPDATE users SET password_hash=? WHERE id=?', (hash_password(password), uid))
    conn.commit(); conn.close()
    log_activity('reset_password', u['username'])
    return jsonify({'ok': True})

@app.route('/api/users/<int:uid>', methods=['DELETE'])
@admin_required
def delete_user(uid):
    me = current_user()
    if me and me['id'] == uid:
        return jsonify({'error': "You can't delete your own account"}), 400
    conn = get_db()
    u = conn.execute('SELECT username, role FROM users WHERE id=?', (uid,)).fetchone()
    if not u:
        conn.close()
        return jsonify({'error': 'User not found'}), 404
    # Safety: don't allow deleting the last admin
    if u['role'] == 'admin':
        admins = conn.execute("SELECT COUNT(*) n FROM users WHERE role='admin'").fetchone()['n']
        if admins <= 1:
            conn.close()
            return jsonify({'error': 'Cannot delete the only admin'}), 400
    # Data belongs to the user's COMPANY (tenant), not to the individual user,
    # so deleting a user does NOT touch any mandates/candidates — the company
    # keeps all its data for its remaining (or future) members.
    conn.execute('DELETE FROM users WHERE id=?', (uid,))
    conn.commit(); conn.close()
    log_activity('delete_user', u['username'])
    return jsonify({'ok': True})


@app.route('/api/admin/claim-orphans', methods=['POST'])
@admin_required
def claim_orphans():
    """Assign any data with owner_id=0/NULL to the admin's COMPANY (tenant) so
    it shows up in their workspace. owner_id stores the tenant/company id."""
    tenant = current_company_id()
    conn = get_db(); c = conn.cursor()
    n_m = c.execute('UPDATE mandates SET owner_id=? WHERE owner_id IS NULL OR owner_id=0', (tenant,)).rowcount
    n_c = c.execute('UPDATE candidates SET owner_id=? WHERE owner_id IS NULL OR owner_id=0', (tenant,)).rowcount
    n_r = c.execute('UPDATE reminders SET owner_id=? WHERE owner_id IS NULL OR owner_id=0', (tenant,)).rowcount
    conn.commit(); conn.close()
    return jsonify({'ok': True, 'mandates': n_m, 'candidates': n_c, 'reminders': n_r})



# ── Super-Admin: view-as (impersonate a tenant's workspace) ───────────────
@app.route('/api/admin/view-as', methods=['POST'])
@admin_required
def admin_view_as():
    """Super-admin impersonates a COMPANY (tenant). Accepts a company_id (or a
    legacy user_id, resolved to that user's company)."""
    d = request.json or {}
    company_id = d.get('company_id')
    # Backward-compat: if a user_id is sent, resolve it to that user's company.
    if not company_id and d.get('user_id'):
        conn = get_db()
        ur = conn.execute('SELECT company_id FROM users WHERE id=?', (d.get('user_id'),)).fetchone()
        conn.close()
        company_id = ur['company_id'] if ur else None
    conn = get_db()
    if company_id:
        comp = conn.execute('SELECT name FROM companies WHERE id=?', (company_id,)).fetchone()
        conn.close()
        if not comp:
            return jsonify({'error': 'Company not found'}), 404
        session['view_as_company'] = company_id
        log_activity('view_as', comp['name'])
    else:
        conn.close()
        session.pop('view_as_company', None)
        log_activity('view_as', 'exited')
    return jsonify({'ok': True})


@app.route('/api/admin/api-usage', methods=['GET'])
@admin_required
def admin_api_usage():
    """Per-agency API usage & estimated cost. Optional ?days=N (default 30)."""
    try:
        days = int(request.args.get('days', 30))
    except Exception:
        days = 30
    since = (_ist_now() - datetime.timedelta(days=days)).isoformat()
    conn = get_db()
    companies = conn.execute('SELECT id, name FROM companies ORDER BY id').fetchall()
    name_map = {c['id']: c['name'] for c in companies}

    rows = conn.execute('''
        SELECT company_id, provider,
               SUM(input_tokens) in_tok, SUM(output_tokens) out_tok,
               SUM(audio_seconds) audio_sec, SUM(cost_usd) cost, COUNT(*) calls
        FROM api_usage WHERE created_at >= ?
        GROUP BY company_id, provider''', (since,)).fetchall()
    conn.close()

    agg = {}
    grand_total = 0.0
    for r in rows:
        cid = r['company_id']
        if cid not in agg:
            agg[cid] = {'company_id': cid,
                        'company': name_map.get(cid, '(unknown / deleted)'),
                        'total_cost': 0.0, 'total_calls': 0, 'providers': {}}
        agg[cid]['providers'][r['provider']] = {
            'calls': r['calls'], 'input_tokens': r['in_tok'] or 0,
            'output_tokens': r['out_tok'] or 0,
            'audio_minutes': round((r['audio_sec'] or 0) / 60.0, 1),
            'cost_usd': round(r['cost'] or 0, 4)}
        agg[cid]['total_cost'] += (r['cost'] or 0)
        agg[cid]['total_calls'] += r['calls']
        grand_total += (r['cost'] or 0)

    out = sorted(agg.values(), key=lambda x: x['total_cost'], reverse=True)
    for a in out:
        a['total_cost'] = round(a['total_cost'], 4)
    return jsonify({'ok': True, 'days': days, 'grand_total_usd': round(grand_total, 4),
                    'usage': out})


def compute_company_bill(company_id, days=30):
    """Compute one company's bill: (recruiters x price) + token charges (API
    cost passed through, USD→INR x markup), then GST. The platform owner's own
    company is never billed."""
    conn = get_db()
    comp = conn.execute('SELECT * FROM companies WHERE id=?', (company_id,)).fetchone()
    if not comp:
        conn.close(); return None
    recruiters = conn.execute("SELECT COUNT(*) n FROM users WHERE company_id=? AND status='approved'",
                              (company_id,)).fetchone()['n']
    since = (_ist_now() - datetime.timedelta(days=days)).isoformat()
    token_usd = conn.execute('SELECT COALESCE(SUM(cost_usd),0) c FROM api_usage WHERE company_id=? AND created_at>=?',
                             (company_id, since)).fetchone()['c'] or 0
    conn.close()

    price = int(get_setting('billing_price_per_recruiter', '700') or 700)
    usd_inr = float(get_setting('billing_usd_inr', '88') or 88)
    markup = float(get_setting('billing_token_markup', '1.0') or 1.0)
    gst_rate = float(get_setting('billing_gst_rate', '18') or 18)

    base = recruiters * price
    token_inr = round(token_usd * usd_inr * markup, 2)
    subtotal = round(base + token_inr, 2)
    gst = round(subtotal * gst_rate / 100.0, 2)
    total = round(subtotal + gst, 2)

    trial_left = None
    if comp['trial_ends_at']:
        try:
            te = datetime.datetime.fromisoformat(comp['trial_ends_at'])
            trial_left = max(0, (te - datetime.datetime.now()).days)
        except Exception:
            trial_left = None

    return {
        'company_id': company_id, 'company': comp['name'],
        'billing_status': comp['billing_status'], 'status': comp['status'],
        'recruiters': recruiters, 'price_per_recruiter': price,
        'base_inr': base, 'token_usd': round(token_usd, 4), 'token_inr': token_inr,
        'subtotal_inr': subtotal, 'gst_rate': gst_rate, 'gst_inr': gst, 'total_inr': total,
        'trial_ends_at': comp['trial_ends_at'], 'trial_days_left': trial_left,
        'is_owner': (comp['plan'] == 'owner' or comp['billing_status'] == 'owner'),
        'token_cap': company_token_cap(company_id),
        'tokens_used_month': tokens_used_this_month(company_id),
        'user_limit': company_user_limit(company_id),
    }


@app.route('/api/admin/companies/<int:cid>/limits', methods=['POST'])
@admin_required
def set_company_limits(cid):
    """Platform owner sets a company's monthly AI token cap and user (seat) limit.
    Both 0 = unlimited."""
    d = request.json or {}
    conn = get_db()
    comp = conn.execute('SELECT id FROM companies WHERE id=?', (cid,)).fetchone()
    if not comp:
        conn.close(); return jsonify({'error': 'Company not found'}), 404
    fields, vals = [], []
    if 'token_cap' in d:
        try:
            fields.append('token_cap=?'); vals.append(max(0, int(d.get('token_cap') or 0)))
        except Exception:
            pass
    if 'user_limit' in d:
        try:
            fields.append('user_limit=?'); vals.append(max(0, int(d.get('user_limit') or 0)))
        except Exception:
            pass
    if not fields:
        conn.close(); return jsonify({'error': 'Nothing to update'}), 400
    vals.append(cid)
    conn.execute('UPDATE companies SET ' + ','.join(fields) + ' WHERE id=?', vals)
    conn.commit(); conn.close()
    log_activity('set_limits', f'company {cid}: ' + ', '.join(f'{f.split("=")[0]}={v}' for f, v in zip(fields, vals[:-1])))
    return jsonify({'ok': True})


@app.route('/api/admin/agencies', methods=['POST'])
@admin_required
def create_agency():
    """Platform owner creates a new agency (tenant): an isolated company plus its
    admin login, active immediately, with optional token cap and user limit."""
    d = request.json or {}
    company_name = (d.get('company_name') or '').strip()
    username = (d.get('username') or '').strip().lower()
    password = d.get('password') or ''
    email = (d.get('email') or '').strip().lower()
    display = (d.get('display_name') or company_name or username).strip()
    try:
        token_cap = max(0, int(d.get('token_cap') or 0))
    except Exception:
        token_cap = 0
    try:
        user_limit = max(0, int(d.get('user_limit') or 0))
    except Exception:
        user_limit = 0
    if not company_name:
        return jsonify({'error': 'Agency / company name is required'}), 400
    if not username or len(password) < 8:
        return jsonify({'error': 'Admin username and password (min 4 chars) are required'}), 400
    if not re.match(r'^[a-z0-9._-]{3,40}$', username):
        return jsonify({'error': 'Username can only contain letters, numbers, dots, dashes and underscores'}), 400
    if email and '@' not in email:
        return jsonify({'error': 'Enter a valid admin email (for password reset)'}), 400
    conn = get_db()
    if conn.execute('SELECT id FROM users WHERE username=?', (username,)).fetchone():
        conn.close(); return jsonify({'error': 'Username already taken'}), 400
    try:
        trial_days = int(get_setting('billing_trial_days', '14') or 14)
    except Exception:
        trial_days = 14
    trial_end = (datetime.datetime.now() + datetime.timedelta(days=trial_days)).isoformat()
    conn.execute("INSERT INTO companies (name,status,plan,billing_status,trial_ends_at,created_at,token_cap,user_limit) VALUES (?,?,?,?,?,?,?,?)",
                 (company_name, 'active', 'standard', 'trial', trial_end, ts(), token_cap, user_limit))
    new_cid = conn.execute('SELECT id FROM companies ORDER BY id DESC LIMIT 1').fetchone()['id']
    conn.execute('INSERT INTO users (username,password_hash,display_name,role,created_at,status,company_name,company_id,is_company_admin,email) VALUES (?,?,?,?,?,?,?,?,1,?)',
                 (username, hash_password(password), display, 'user', ts(), 'approved', company_name, new_cid, email))
    conn.commit(); conn.close()
    log_activity('create_agency', f'{company_name} (admin: {username})')
    return jsonify({'ok': True, 'company_id': new_cid})


@app.route('/api/admin/billing', methods=['GET'])
@admin_required
def admin_billing():
    """Super-admin billing dashboard: every agency's monthly bill + status."""
    try:
        days = int(request.args.get('days', 30))
    except Exception:
        days = 30
    conn = get_db()
    ids = [r['id'] for r in conn.execute('SELECT id FROM companies ORDER BY id').fetchall()]
    conn.close()
    bills = []
    revenue = 0.0
    for cid in ids:
        b = compute_company_bill(cid, days)
        if not b:
            continue
        bills.append(b)
        if not b['is_owner'] and b['billing_status'] in ('active', 'past_due'):
            revenue += b['total_inr']
    return jsonify({'ok': True, 'days': days, 'monthly_revenue_inr': round(revenue, 2),
                    'gstin': get_setting('billing_gstin', ''), 'bills': bills})


@app.route('/api/billing/me', methods=['GET'])
@login_required
def my_billing():
    """An agency admin sees their own current bill + trial status."""
    if not is_company_admin():
        return jsonify({'error': 'Not allowed'}), 403
    b = compute_company_bill(effective_company_id(), 30)
    return jsonify({'ok': True, 'bill': b}) if b else (jsonify({'error': 'No company'}), 404)


@app.route('/api/admin/billing/<int:cid>/status', methods=['POST'])
@admin_required
def set_billing_status(cid):
    """Super-admin sets an agency's billing status (active/suspended/trial/past_due)."""
    d = request.json or {}
    status = d.get('status', '')
    if status not in ('active', 'suspended', 'trial', 'past_due'):
        return jsonify({'error': 'Invalid status'}), 400
    conn = get_db()
    comp = conn.execute('SELECT name FROM companies WHERE id=?', (cid,)).fetchone()
    if not comp:
        conn.close(); return jsonify({'error': 'Company not found'}), 404
    # billing_status drives the badge; company.status controls actual login block.
    company_status = 'suspended' if status == 'suspended' else 'active'
    conn.execute('UPDATE companies SET billing_status=?, status=? WHERE id=?',
                 (status, company_status, cid))
    conn.commit(); conn.close()
    log_activity('billing_status', f"{comp['name']} → {status}")
    return jsonify({'ok': True})


@app.route('/api/admin/billing/<int:cid>/pay', methods=['POST'])
@admin_required
def record_payment(cid):
    """Record a manual payment (UPI/bank transfer) and activate the agency."""
    d = request.json or {}
    conn = get_db()
    comp = conn.execute('SELECT name FROM companies WHERE id=?', (cid,)).fetchone()
    if not comp:
        conn.close(); return jsonify({'error': 'Company not found'}), 404
    bill = compute_company_bill(cid, 30)
    amount = float(d.get('amount') or (bill['total_inr'] if bill else 0))
    note = d.get('note', '')
    invoice_no = _next_invoice_no(conn)
    period = datetime.date.today().strftime('%b %Y')
    conn.execute('''INSERT INTO payments (company_id,invoice_no,amount_inr,period,method,note,created_at)
                    VALUES (?,?,?,?,?,?,?)''',
                 (cid, invoice_no, amount, period, d.get('method', 'manual'), note, ts()))
    # Mark active and clear trial so they're a paying customer now.
    conn.execute("UPDATE companies SET billing_status='active', status='active' WHERE id=?", (cid,))
    conn.commit(); conn.close()
    log_activity('payment', f"{comp['name']} ₹{amount} ({invoice_no})")
    return jsonify({'ok': True, 'invoice_no': invoice_no, 'amount': amount})


def _next_invoice_no(conn):
    """Sequential invoice number like HL-2026-0001."""
    yr = datetime.date.today().year
    n = conn.execute("SELECT COUNT(*) c FROM payments").fetchone()['c'] + 1
    return f"HL-{yr}-{n:04d}"


@app.route('/api/admin/billing/<int:cid>/payments', methods=['GET'])
@admin_required
def list_payments(cid):
    conn = get_db()
    rows = conn.execute('SELECT * FROM payments WHERE company_id=? ORDER BY created_at DESC', (cid,)).fetchall()
    conn.close()
    return jsonify({'ok': True, 'payments': [dict(r) for r in rows]})


@app.route('/api/billing/invoice/<int:cid>', methods=['GET'])
@login_required
def billing_invoice(cid):
    """Invoice data for a company's current bill. Super-admin for any company;
    a company admin only for their own."""
    if not is_admin() and not (is_company_admin() and effective_company_id() == cid):
        return jsonify({'error': 'Not allowed'}), 403
    bill = compute_company_bill(cid, 30)
    if not bill:
        return jsonify({'error': 'No company'}), 404
    conn = get_db()
    last = conn.execute('SELECT invoice_no FROM payments WHERE company_id=? ORDER BY id DESC LIMIT 1', (cid,)).fetchone()
    conn.close()
    line_items = [{'desc': f"Subscription — {bill['recruiters']} recruiter(s) × ₹{bill['price_per_recruiter']}",
                   'amount': bill['base_inr']}]
    if bill['token_inr'] > 0:
        line_items.append({'desc': 'AI / API usage charges (this period)', 'amount': bill['token_inr']})
    return jsonify({'ok': True, 'invoice': {
        'invoice_no': (last['invoice_no'] if last else 'DRAFT'),
        'date': datetime.date.today().isoformat(),
        'seller_name': get_setting('billing_legal_name', 'HireLab Talent Resource'),
        'seller_address': get_setting('billing_address', ''),
        'seller_gstin': get_setting('billing_gstin', ''),
        'buyer': bill['company'],
        'line_items': line_items,
        'subtotal': bill['subtotal_inr'], 'gst_rate': bill['gst_rate'],
        'gst': bill['gst_inr'], 'total': bill['total_inr'],
        'billing_status': bill['billing_status'],
    }})



@app.route('/api/admin/summary', methods=['GET'])
@admin_required
def admin_summary():
    conn = get_db()
    companies = conn.execute("SELECT id, name, status, plan, created_at, expires_at FROM companies ORDER BY id").fetchall()
    summary = []
    for comp in companies:
        cid = comp['id']
        mand = conn.execute('SELECT COUNT(*) n FROM mandates WHERE owner_id=?', (cid,)).fetchone()['n']
        active_mand = conn.execute("SELECT COUNT(*) n FROM mandates WHERE owner_id=? AND status='active'", (cid,)).fetchone()['n']
        cands = conn.execute('SELECT COUNT(*) n FROM candidates WHERE owner_id=?', (cid,)).fetchone()['n']
        placed = conn.execute("SELECT COUNT(*) n FROM candidates WHERE owner_id=? AND " + PLACED_SQL, (cid,)).fetchone()['n']
        members = conn.execute("SELECT COUNT(*) n FROM users WHERE company_id=? AND status='approved'", (cid,)).fetchone()['n']
        last_login = conn.execute("SELECT MAX(last_login) m FROM users WHERE company_id=?", (cid,)).fetchone()['m']
        summary.append({
            'id': cid, 'name': comp['name'], 'status': comp['status'], 'plan': comp['plan'],
            'created_at': comp['created_at'], 'expires_at': comp['expires_at'],
            'members': members, 'last_login': last_login,
            'mandates': mand, 'active_mandates': active_mand,
            'candidates': cands, 'placed': placed,
        })
    recent = conn.execute('SELECT username, action, detail, created_at FROM activity_log ORDER BY id DESC LIMIT 50').fetchall()
    conn.close()
    return jsonify({'ok': True, 'summary': summary, 'recent_activity': [dict(r) for r in recent]})


@app.route('/api/admin/companies/<int:cid>/suspend', methods=['POST'])
@admin_required
def suspend_company(cid):
    """Suspend an agency (e.g. non-payment). Its users can't log in until
    reactivated. The platform owner's own company can't be suspended."""
    me = current_user()
    if me and me.get('company_id') == cid:
        return jsonify({'error': "You can't suspend your own company"}), 400
    conn = get_db()
    comp = conn.execute('SELECT name FROM companies WHERE id=?', (cid,)).fetchone()
    if not comp:
        conn.close()
        return jsonify({'error': 'Company not found'}), 404
    conn.execute("UPDATE companies SET status='suspended' WHERE id=?", (cid,))
    conn.commit(); conn.close()
    log_activity('suspend_company', comp['name'])
    return jsonify({'ok': True})


@app.route('/api/admin/companies/<int:cid>/activate', methods=['POST'])
@admin_required
def activate_company(cid):
    conn = get_db()
    comp = conn.execute('SELECT name FROM companies WHERE id=?', (cid,)).fetchone()
    if not comp:
        conn.close()
        return jsonify({'error': 'Company not found'}), 404
    conn.execute("UPDATE companies SET status='active' WHERE id=?", (cid,))
    conn.commit(); conn.close()
    log_activity('activate_company', comp['name'])
    return jsonify({'ok': True})


@app.route('/login')
def login_page():
    # Single-page app handles login UI; just serve the app which checks auth_status
    return flask_redirect('/')


# Data lives in user home — survives app updates
# Railway: set DATA_DIR=/data in env vars (persistent volume)
# Local: defaults to ~/HireLab
DATA_DIR = os.environ.get('DATA_DIR',
    '/data' if os.environ.get('RAILWAY_ENVIRONMENT') else
    os.path.join(os.path.expanduser('~'), 'HireLab'))
DB_PATH  = os.path.join(DATA_DIR, 'hirelab.db')
CV_DIR   = os.path.join(DATA_DIR, 'cvs')
BAK_DIR  = os.path.join(DATA_DIR, 'backups')
CRM_FILES_DIR = os.path.join(DATA_DIR, 'crm_files')

CLAUDE_URL   = 'https://api.anthropic.com/v1/messages'
CLAUDE_MODEL = 'claude-sonnet-4-20250514'

for d in [DATA_DIR, CV_DIR, BAK_DIR, CRM_FILES_DIR]:
    os.makedirs(d, exist_ok=True)

def get_db():
    conn = sqlite3.connect(DB_PATH, timeout=60, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    # IMPORTANT: DELETE journal mode + FULL sync (not WAL+NORMAL).
    # WAL keeps recent writes in a separate -wal file that can fail to flush
    # into the main DB file on a cloud-host restart/redeploy — this was the
    # root cause of a critical bug where the users table (and other recent
    # writes) appeared empty after a restart, forcing a fresh "Create Admin"
    # setup and orphaning the previous data. DELETE+FULL writes every change
    # straight into the main database file, so there is no separate WAL file
    # that can be lost.
    conn.execute('PRAGMA journal_mode=DELETE')
    conn.execute('PRAGMA busy_timeout=60000')
    conn.execute('PRAGMA synchronous=FULL')
    return conn

def esc_html(s):
    """Escape a string for safe inclusion in HTML email bodies."""
    return (str(s or '').replace('&', '&amp;').replace('<', '&lt;')
            .replace('>', '&gt;').replace('"', '&quot;'))

def ts():
    return _ist_now().isoformat(timespec='seconds')

def html_to_text(html):
    """Convert JD rich-text HTML to clean plain text for AI prompts / exports."""
    if not html:
        return ''
    import re as _re
    txt = html
    # Lists: prefix items with bullet/number markers before stripping tags
    txt = _re.sub(r'<li[^>]*>', '\n- ', txt, flags=_re.I)
    # Block-level tags -> newlines
    txt = _re.sub(r'</(p|div|h[1-6]|li|ul|ol)>', '\n', txt, flags=_re.I)
    txt = _re.sub(r'<br\s*/?>', '\n', txt, flags=_re.I)
    # Strip remaining tags
    txt = _re.sub(r'<[^>]+>', '', txt)
    # Decode common HTML entities
    txt = (txt.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
              .replace('&quot;', '"').replace('&#039;', "'").replace('&nbsp;', ' '))
    # Collapse excess blank lines/spaces
    txt = _re.sub(r'\n[ \t]+', '\n', txt)
    txt = _re.sub(r'\n{3,}', '\n\n', txt)
    return txt.strip()

# ═══════════════════════════════════════════════════════════════════════════
#  SKILL GRAPH — Intelligence Layer · Phase 1
#  Domain skill graph for electrical / automation / solar / renewable / BMS.
#  Terms are lowercase & space-separated (no hyphens/slashes) so they match the
#  query n-grams and candidate tags, which are normalised the same way.
# ═══════════════════════════════════════════════════════════════════════════
_SKILL_GRAPH_SEED_VERSION = '1'
_SKILL_GRAPH_SEED = [
    # ── Building Management / BMS ──────────────────────────────────────────
    {'c': 'bms', 'd': 'BMS', 'cat': 'Building Automation',
     'a': ['building management system', 'building management', 'bems'], 'p': ['building automation'],
     'r': ['bacnet', 'niagara', 'ddc', 'hvac', 'modbus', 'lonworks', 'honeywell', 'johnson controls',
           'desigo', 'ecostruxure', 'commissioning', 'scada']},
    {'c': 'building automation', 'd': 'Building Automation', 'cat': 'Building Automation',
     'a': ['building automation system', 'bas'], 'p': [], 'r': ['bms', 'ddc', 'hvac', 'bacnet']},
    {'c': 'bacnet', 'd': 'BACnet', 'cat': 'Building Automation',
     'a': ['bacnet ip', 'bacnet mstp'], 'p': ['building automation'], 'r': ['bms', 'modbus', 'lonworks', 'niagara']},
    {'c': 'niagara', 'd': 'Niagara (Tridium)', 'cat': 'Building Automation',
     'a': ['tridium', 'tridium niagara', 'niagara framework', 'niagara 4', 'jace'],
     'p': ['building automation'], 'r': ['bms', 'bacnet', 'ddc', 'honeywell']},
    {'c': 'ddc', 'd': 'DDC', 'cat': 'Building Automation',
     'a': ['direct digital control'], 'p': ['building automation'], 'r': ['bms', 'plc', 'hvac']},
    {'c': 'hvac', 'd': 'HVAC', 'cat': 'Building Automation',
     'a': ['heating ventilation air conditioning'], 'p': [], 'r': ['bms', 'chillers', 'ahu', 'vav', 'ddc']},
    {'c': 'lonworks', 'd': 'LonWorks', 'cat': 'Building Automation',
     'a': [], 'p': ['building automation'], 'r': ['bacnet', 'bms']},
    {'c': 'desigo', 'd': 'Siemens Desigo', 'cat': 'Building Automation',
     'a': ['siemens desigo'], 'p': ['building automation'], 'r': ['bms', 'siemens']},
    # ── PLC / SCADA / Industrial Automation ────────────────────────────────
    {'c': 'plc', 'd': 'PLC', 'cat': 'Industrial Automation',
     'a': ['programmable logic controller'], 'p': ['industrial automation'],
     'r': ['scada', 'hmi', 'dcs', 'ladder logic', 'siemens', 'allen bradley', 'modbus', 'vfd']},
    {'c': 'scada', 'd': 'SCADA', 'cat': 'Industrial Automation',
     'a': ['supervisory control and data acquisition'], 'p': ['industrial automation'],
     'r': ['plc', 'hmi', 'dcs', 'wincc', 'wonderware', 'ignition', 'modbus']},
    {'c': 'hmi', 'd': 'HMI', 'cat': 'Industrial Automation',
     'a': ['human machine interface'], 'p': ['industrial automation'], 'r': ['plc', 'scada', 'wincc']},
    {'c': 'dcs', 'd': 'DCS', 'cat': 'Industrial Automation',
     'a': ['distributed control system'], 'p': ['industrial automation'],
     'r': ['plc', 'scada', '800xa', 'experion', 'yokogawa']},
    {'c': 'industrial automation', 'd': 'Industrial Automation', 'cat': 'Industrial Automation',
     'a': ['factory automation', 'process automation'], 'p': [],
     'r': ['plc', 'scada', 'dcs', 'instrumentation', 'robotics']},
    {'c': 'siemens', 'd': 'Siemens', 'cat': 'OEM / Ecosystem',
     'a': ['siemens automation'], 'p': [], 'r': ['tia portal', 's7', 'simatic', 'wincc', 'profinet', 'plc', 'desigo']},
    {'c': 'tia portal', 'd': 'TIA Portal', 'cat': 'Industrial Automation',
     'a': ['tia'], 'p': ['siemens'], 'r': ['s7', 'simatic', 'wincc', 'plc']},
    {'c': 's7', 'd': 'Siemens S7', 'cat': 'Industrial Automation',
     'a': ['s7 1200', 's7 1500', 's7 300', 's7 400', 'simatic s7', 'simatic'], 'p': ['siemens'],
     'r': ['tia portal', 'plc']},
    {'c': 'wincc', 'd': 'WinCC', 'cat': 'Industrial Automation',
     'a': [], 'p': ['siemens'], 'r': ['scada', 'hmi', 'tia portal']},
    {'c': 'profinet', 'd': 'PROFINET', 'cat': 'Industrial Automation',
     'a': [], 'p': ['siemens'], 'r': ['profibus', 'industrial ethernet', 'plc']},
    {'c': 'profibus', 'd': 'PROFIBUS', 'cat': 'Industrial Automation',
     'a': [], 'p': [], 'r': ['profinet', 'modbus', 'fieldbus']},
    {'c': 'allen bradley', 'd': 'Allen-Bradley (Rockwell)', 'cat': 'OEM / Ecosystem',
     'a': ['rockwell', 'rockwell automation'], 'p': [], 'r': ['rslogix', 'studio 5000', 'plc', 'factorytalk']},
    {'c': 'rslogix', 'd': 'RSLogix / Studio 5000', 'cat': 'Industrial Automation',
     'a': ['studio 5000', 'logix'], 'p': ['allen bradley'], 'r': ['plc']},
    {'c': '800xa', 'd': 'ABB 800xA', 'cat': 'Industrial Automation',
     'a': ['abb 800xa', 'system 800xa'], 'p': ['abb'], 'r': ['dcs']},
    {'c': 'modbus', 'd': 'Modbus', 'cat': 'Industrial Automation',
     'a': ['modbus tcp', 'modbus rtu'], 'p': [], 'r': ['plc', 'scada', 'bacnet', 'profibus']},
    {'c': 'ladder logic', 'd': 'Ladder Logic', 'cat': 'Industrial Automation',
     'a': ['ladder programming'], 'p': [], 'r': ['plc', 'function block']},
    {'c': 'vfd', 'd': 'VFD / Drives', 'cat': 'Industrial Automation',
     'a': ['variable frequency drive', 'variable speed drive', 'vsd', 'ac drive', 'drives'],
     'p': [], 'r': ['motor control', 'plc', 'servo', 'altivar']},
    {'c': 'servo', 'd': 'Servo / Motion', 'cat': 'Industrial Automation',
     'a': ['servo drive', 'servo motor', 'motion control'], 'p': [], 'r': ['plc', 'vfd']},
    {'c': 'instrumentation', 'd': 'Instrumentation', 'cat': 'Industrial Automation',
     'a': ['field instrumentation'], 'p': [], 'r': ['transmitters', 'control valves', 'plc', 'scada', 'calibration', 'dcs']},
    # ── Electrical / Power ─────────────────────────────────────────────────
    {'c': 'switchgear', 'd': 'Switchgear', 'cat': 'Electrical',
     'a': [], 'p': [], 'r': ['lv switchgear', 'mv switchgear', 'panels', 'circuit breaker']},
    {'c': 'lv switchgear', 'd': 'LV Switchgear', 'cat': 'Electrical',
     'a': ['low voltage switchgear', 'lv panels', 'lt switchgear', 'lt panels'], 'p': ['switchgear'],
     'r': ['mcc', 'panels', 'acb', 'mccb', 'schneider', 'abb', 'siemens']},
    {'c': 'mv switchgear', 'd': 'MV Switchgear', 'cat': 'Electrical',
     'a': ['medium voltage switchgear', 'ht switchgear', 'ht panels'], 'p': ['switchgear'],
     'r': ['vcb', 'protection relays', 'ring main unit', 'gis']},
    {'c': 'protection relays', 'd': 'Protection Relays', 'cat': 'Electrical',
     'a': ['protection relay', 'numerical relay', 'relay coordination'], 'p': [],
     'r': ['switchgear', 'sld', 'substation', 'siprotec']},
    {'c': 'sld', 'd': 'Single Line Diagram', 'cat': 'Electrical',
     'a': ['single line diagram'], 'p': [], 'r': ['electrical design', 'power distribution', 'etap']},
    {'c': 'etap', 'd': 'ETAP', 'cat': 'Electrical',
     'a': [], 'p': [], 'r': ['power system study', 'load flow', 'sld', 'electrical design']},
    {'c': 'mcc', 'd': 'MCC', 'cat': 'Electrical',
     'a': ['motor control center', 'motor control centre'], 'p': [], 'r': ['switchgear', 'vfd', 'panels', 'plc']},
    {'c': 'power distribution', 'd': 'Power Distribution', 'cat': 'Electrical',
     'a': ['electrical distribution'], 'p': [], 'r': ['switchgear', 'transformers', 'sld', 'substation']},
    {'c': 'transformers', 'd': 'Transformers', 'cat': 'Electrical',
     'a': ['transformer'], 'p': [], 'r': ['power distribution', 'substation']},
    {'c': 'substation', 'd': 'Substation', 'cat': 'Electrical',
     'a': [], 'p': [], 'r': ['transformers', 'switchgear', 'protection relays', 'power distribution']},
    {'c': 'panels', 'd': 'Panels / Panel Design', 'cat': 'Electrical',
     'a': ['electrical panels', 'panel design', 'control panels'], 'p': [], 'r': ['switchgear', 'mcc', 'wiring']},
    {'c': 'electrical design', 'd': 'Electrical Design', 'cat': 'Electrical',
     'a': [], 'p': [], 'r': ['sld', 'etap', 'autocad electrical', 'power distribution']},
    # ── Solar / Renewable ──────────────────────────────────────────────────
    {'c': 'solar pv', 'd': 'Solar PV', 'cat': 'Renewable',
     'a': ['solar', 'solar power', 'photovoltaic', 'pv', 'solar energy'], 'p': ['renewable energy'],
     'r': ['inverters', 'mppt', 'string design', 'pvsyst', 'epc', 'net metering', 'bess', 'on grid', 'off grid']},
    {'c': 'renewable energy', 'd': 'Renewable Energy', 'cat': 'Renewable',
     'a': ['renewables', 'clean energy', 'green energy'], 'p': [], 'r': ['solar pv', 'wind', 'bess']},
    {'c': 'inverters', 'd': 'Inverters', 'cat': 'Renewable',
     'a': ['solar inverter', 'string inverter', 'central inverter'], 'p': ['solar pv'],
     'r': ['mppt', 'sungrow', 'sma']},
    {'c': 'mppt', 'd': 'MPPT', 'cat': 'Renewable',
     'a': ['maximum power point tracking'], 'p': ['solar pv'], 'r': ['inverters']},
    {'c': 'pvsyst', 'd': 'PVsyst', 'cat': 'Renewable',
     'a': [], 'p': ['solar pv'], 'r': ['string design', 'energy yield']},
    {'c': 'string design', 'd': 'String Design', 'cat': 'Renewable',
     'a': ['string sizing'], 'p': ['solar pv'], 'r': ['pvsyst', 'inverters']},
    {'c': 'bess', 'd': 'BESS / Energy Storage', 'cat': 'Renewable',
     'a': ['battery energy storage', 'battery storage', 'energy storage'], 'p': [],
     'r': ['solar pv', 'renewable energy', 'inverters', 'lithium ion']},
    {'c': 'wind', 'd': 'Wind Energy', 'cat': 'Renewable',
     'a': ['wind energy', 'wind power', 'wind turbine'], 'p': ['renewable energy'], 'r': ['scada']},
    {'c': 'epc', 'd': 'EPC', 'cat': 'Renewable',
     'a': ['engineering procurement construction'], 'p': [], 'r': ['solar pv', 'substation', 'project management']},
    {'c': 'net metering', 'd': 'Net Metering', 'cat': 'Renewable',
     'a': [], 'p': [], 'r': ['solar pv', 'discom', 'on grid']},
    {'c': 'on grid', 'd': 'On-Grid', 'cat': 'Renewable',
     'a': ['grid tied', 'grid connected'], 'p': [], 'r': ['solar pv', 'net metering', 'inverters']},
    {'c': 'off grid', 'd': 'Off-Grid', 'cat': 'Renewable',
     'a': ['standalone solar'], 'p': [], 'r': ['solar pv', 'bess', 'inverters']},
    # ── OEM ecosystems ─────────────────────────────────────────────────────
    {'c': 'abb', 'd': 'ABB', 'cat': 'OEM / Ecosystem',
     'a': [], 'p': [], 'r': ['800xa', 'dcs', 'robotics', 'switchgear', 'drives']},
    {'c': 'schneider', 'd': 'Schneider Electric', 'cat': 'OEM / Ecosystem',
     'a': ['schneider electric'], 'p': [], 'r': ['ecostruxure', 'modicon', 'lv switchgear', 'altivar', 'bms']},
    {'c': 'ecostruxure', 'd': 'EcoStruxure', 'cat': 'OEM / Ecosystem',
     'a': [], 'p': ['schneider'], 'r': ['bms', 'scada']},
    {'c': 'modicon', 'd': 'Modicon', 'cat': 'Industrial Automation',
     'a': ['unity pro'], 'p': ['schneider'], 'r': ['plc']},
    {'c': 'honeywell', 'd': 'Honeywell', 'cat': 'OEM / Ecosystem',
     'a': [], 'p': [], 'r': ['experion', 'niagara', 'bms', 'dcs']},
    {'c': 'johnson controls', 'd': 'Johnson Controls (JCI)', 'cat': 'OEM / Ecosystem',
     'a': ['jci', 'metasys'], 'p': [], 'r': ['bms', 'hvac', 'building automation']},
    # ── Cross-cutting ──────────────────────────────────────────────────────
    {'c': 'commissioning', 'd': 'Commissioning', 'cat': 'Functional',
     'a': ['testing and commissioning'], 'p': [], 'r': ['bms', 'plc', 'scada', 'hvac', 'installation']},
]


def _seed_skill_graph(c):
    """Idempotent, version-guarded seed of the shared (owner_id=0) skill graph."""
    try:
        row = c.execute("SELECT value FROM settings WHERE key='skill_graph_seed_version'").fetchone()
        cur = row['value'] if row else None
    except Exception:
        cur = None
    if cur == _SKILL_GRAPH_SEED_VERSION:
        return
    now = ts()
    for n in _SKILL_GRAPH_SEED:
        c.execute(
            "INSERT INTO skill_graph (owner_id,canonical,display,category,aliases,parents,related,created_at,updated_at) "
            "VALUES (0,?,?,?,?,?,?,?,?) "
            "ON CONFLICT(owner_id,canonical) DO UPDATE SET "
            "display=excluded.display, category=excluded.category, aliases=excluded.aliases, "
            "parents=excluded.parents, related=excluded.related, updated_at=excluded.updated_at",
            (n['c'], n.get('d', ''), n.get('cat', ''),
             json.dumps(n.get('a', [])), json.dumps(n.get('p', [])), json.dumps(n.get('r', [])), now, now))
    c.execute("INSERT OR REPLACE INTO settings (key,value) VALUES ('skill_graph_seed_version',?)",
              (_SKILL_GRAPH_SEED_VERSION,))
    print(f'[skill-graph] seeded {len(_SKILL_GRAPH_SEED)} nodes (v{_SKILL_GRAPH_SEED_VERSION})')


_SKILL_GRAPH_CACHE = {'ts': 0.0, 'oid': None, 'nodes': None, 'lookup': None, 'children': None}


def _load_skill_graph(conn, oid):
    """Load shared (owner_id=0) + this tenant's skill-graph nodes. Cached briefly.
    Returns dict with: nodes{canonical->{display,aliases,parents,related}},
    lookup{term->canonical}, children{canonical->set(child canonicals)}."""
    import time as _t
    now = _t.time()
    ce = _SKILL_GRAPH_CACHE
    if ce['nodes'] is not None and ce['oid'] == oid and now - ce['ts'] < 300:
        return ce
    nodes, lookup, children = {}, {}, {}
    # Self-heal: make sure the table exists and is seeded even on an older DB
    # that was migrated before Phase 1 shipped.
    try:
        conn.execute('''CREATE TABLE IF NOT EXISTS skill_graph (
            id INTEGER PRIMARY KEY AUTOINCREMENT, owner_id INTEGER DEFAULT 0,
            canonical TEXT NOT NULL, display TEXT DEFAULT '', category TEXT DEFAULT '',
            aliases TEXT DEFAULT '[]', parents TEXT DEFAULT '[]', related TEXT DEFAULT '[]',
            created_at TEXT DEFAULT '', updated_at TEXT DEFAULT '')''')
        conn.execute("CREATE UNIQUE INDEX IF NOT EXISTS uq_skill_graph ON skill_graph(owner_id, canonical)")
    except Exception:
        pass
    try:
        rows = conn.execute(
            "SELECT canonical,display,aliases,parents,related FROM skill_graph WHERE owner_id=0 OR owner_id=?",
            (oid,)).fetchall()
    except sqlite3.OperationalError:
        rows = []
    if not rows:
        try:
            _seed_skill_graph(conn); conn.commit()
            rows = conn.execute(
                "SELECT canonical,display,aliases,parents,related FROM skill_graph WHERE owner_id=0 OR owner_id=?",
                (oid,)).fetchall()
        except Exception as _e:
            print('[skill-graph] lazy seed failed:', _e); rows = []
    def _jl(v):
        try:
            return [str(x).strip().lower() for x in json.loads(v or '[]') if str(x).strip()]
        except Exception:
            return []
    for r in rows:
        can = (r['canonical'] or '').strip().lower()
        if not can:
            continue
        nodes[can] = {'display': r['display'] or can, 'aliases': _jl(r['aliases']),
                      'parents': _jl(r['parents']), 'related': _jl(r['related'])}
        lookup[can] = can
        for a in nodes[can]['aliases']:
            lookup.setdefault(a, can)
    for can, n in nodes.items():
        for p in n['parents']:
            children.setdefault(p, set()).add(can)
    if nodes:
        ce.update({'ts': now, 'oid': oid, 'nodes': nodes, 'lookup': lookup, 'children': children})
        return ce
    # empty (not seeded yet) — return without caching so the next call retries
    return {'ts': 0.0, 'oid': oid, 'nodes': nodes, 'lookup': lookup, 'children': children}


def _skill_expand(query, conn, oid):
    """Use the domain skill graph to expand recognised query skills into aliases
    (same skill), related, child (more specific) and parent (broader) skills.
    Returns {'expand': {term: weight}, 'recognized': [display], 'related': [display]}.
    Stays empty when the query has no known skills, so search is unaffected."""
    g = _load_skill_graph(conn, oid)
    nodes, lookup, children = g['nodes'], g['lookup'], g['children']
    if not nodes:
        return {'expand': {}, 'recognized': [], 'related': []}
    grams, _toks = _query_grams(query)
    hit = set(lookup[gr] for gr in grams if gr in lookup)
    expand, recognized = {}, []
    def _bump(term, w):
        term = (term or '').strip().lower()
        if term and term not in hit and expand.get(term, 0) < w:
            expand[term] = w
    for can in hit:
        n = nodes.get(can) or {}
        recognized.append(n.get('display', can))
        for a in n.get('aliases', []):
            _bump(a, 1.0)
        for rterm in n.get('related', []):
            _bump(rterm, 0.5)
        for ch in children.get(can, ()):
            _bump(ch, 0.5)
            for a in (nodes.get(ch) or {}).get('aliases', []):
                _bump(a, 0.5)
        for p in n.get('parents', []):
            _bump(p, 0.35)
    rel_display, seen = [], set()
    for term, w in sorted(expand.items(), key=lambda kv: -kv[1]):
        if w >= 1.0:
            continue
        d = (nodes.get(lookup.get(term, term)) or {}).get('display', term)
        if d not in seen:
            seen.add(d); rel_display.append(d)
    return {'expand': expand, 'recognized': sorted(set(recognized)), 'related': rel_display[:12]}


def init_db():
    conn = get_db(); c = conn.cursor()
    c.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            display_name TEXT DEFAULT '',
            role TEXT DEFAULT 'user',
            created_at TEXT,
            last_login TEXT,
            status TEXT DEFAULT 'approved',
            company_name TEXT DEFAULT '',
            requested_at TEXT,
            company_id INTEGER DEFAULT 0
        );
        CREATE TABLE IF NOT EXISTS companies (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL DEFAULT '',
            status TEXT DEFAULT 'active',
            plan TEXT DEFAULT 'standard',
            created_at TEXT,
            expires_at TEXT DEFAULT '',
            notes TEXT DEFAULT '',
            billing_status TEXT DEFAULT 'trial',
            trial_ends_at TEXT DEFAULT '',
            price_per_recruiter INTEGER DEFAULT 700,
            cf_subscription_id TEXT DEFAULT '',
            gstin TEXT DEFAULT ''
        );
        CREATE TABLE IF NOT EXISTS activity_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            username TEXT DEFAULT '',
            action TEXT DEFAULT '',
            detail TEXT DEFAULT '',
            created_at TEXT
        );
        CREATE TABLE IF NOT EXISTS mandates (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            client TEXT NOT NULL DEFAULT '',
            role TEXT NOT NULL DEFAULT '',
            location TEXT DEFAULT '',
            division TEXT DEFAULT '',
            ctc_min REAL DEFAULT 0,
            ctc_max REAL DEFAULT 0,
            jd TEXT DEFAULT '',
            sop_text TEXT DEFAULT '',
            sop_version INTEGER DEFAULT 1,
            sop_changelog TEXT DEFAULT '[]',
            status TEXT DEFAULT 'active',
            email_templates TEXT DEFAULT '[]',
            created_at TEXT DEFAULT ''
        );
        CREATE TABLE IF NOT EXISTS candidates (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            mandate_id INTEGER NOT NULL,
            name TEXT DEFAULT '',
            company TEXT DEFAULT '',
            designation TEXT DEFAULT '',
            experience REAL DEFAULT 0,
            ctc_current REAL DEFAULT 0,
            ctc_expected REAL DEFAULT 0,
            notice_period INTEGER DEFAULT 0,
            location TEXT DEFAULT '',
            preferred_location TEXT DEFAULT '',
            phone TEXT DEFAULT '',
            email TEXT DEFAULT '',
            qualification TEXT DEFAULT '',
            key_skills TEXT DEFAULT '[]',
            secondary_skills TEXT DEFAULT '[]',
            career_summary TEXT DEFAULT '',
            industry_background TEXT DEFAULT '',
            is_mnc INTEGER DEFAULT 0,
            screening_decision TEXT DEFAULT '',
            ai_score REAL DEFAULT 0,
            ai_reasoning TEXT DEFAULT '',
            stage TEXT DEFAULT 'Screening',
            recruiter_feedback TEXT DEFAULT '',
            client_feedback TEXT DEFAULT '',
            general_comments TEXT DEFAULT '',
            cv_path TEXT DEFAULT '',
            cv_original_name TEXT DEFAULT '',
            msg1_sent_at TEXT DEFAULT '',
            fu1_sent_at TEXT DEFAULT '',
            fu2_sent_at TEXT DEFAULT '',
            wa_response TEXT DEFAULT '',
            wa_response_note TEXT DEFAULT '',
            wa_response_at TEXT DEFAULT '',
            key_skill_tags TEXT DEFAULT '[]',
            domain_tags TEXT DEFAULT '[]',
            created_at TEXT DEFAULT '',
            updated_at TEXT DEFAULT '',
            FOREIGN KEY (mandate_id) REFERENCES mandates(id)
        );
        CREATE TABLE IF NOT EXISTS reminders (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            candidate_id INTEGER NOT NULL,
            mandate_id INTEGER,
            candidate_name TEXT DEFAULT '',
            mandate_label TEXT DEFAULT '',
            note TEXT DEFAULT '',
            due_at TEXT NOT NULL,
            done INTEGER DEFAULT 0,
            created_at TEXT
        );

        CREATE TABLE IF NOT EXISTS work_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            candidate_id INTEGER NOT NULL,
            company TEXT DEFAULT '',
            designation TEXT DEFAULT '',
            start_date TEXT DEFAULT '',
            end_date TEXT DEFAULT '',
            is_current INTEGER DEFAULT 0,
            description TEXT DEFAULT '',
            sort_order INTEGER DEFAULT 0
        );
        CREATE TABLE IF NOT EXISTS submissions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT DEFAULT '',
            phone TEXT DEFAULT '',
            email TEXT DEFAULT '',
            company TEXT DEFAULT '',
            designation TEXT DEFAULT '',
            experience REAL DEFAULT 0,
            ctc_current REAL DEFAULT 0,
            ctc_expected REAL DEFAULT 0,
            notice_period INTEGER DEFAULT 0,
            location TEXT DEFAULT '',
            key_skills TEXT DEFAULT '[]',
            domain_tags TEXT DEFAULT '[]',
            custom_fields TEXT DEFAULT '{}',
            cv_path TEXT DEFAULT '',
            cv_original_name TEXT DEFAULT '',
            resume_parsed INTEGER DEFAULT 0,
            status TEXT DEFAULT 'new',
            notes TEXT DEFAULT '',
            created_at TEXT
        );
        CREATE TABLE IF NOT EXISTS stage_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            candidate_id INTEGER NOT NULL,
            from_stage TEXT DEFAULT '',
            to_stage TEXT DEFAULT '',
            note TEXT DEFAULT '',
            created_at TEXT DEFAULT '',
            FOREIGN KEY (candidate_id) REFERENCES candidates(id)
        );
        CREATE TABLE IF NOT EXISTS candidate_events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            candidate_id INTEGER NOT NULL,
            event_type TEXT DEFAULT '',
            detail TEXT DEFAULT '',
            created_at TEXT DEFAULT ''
        );
        CREATE TABLE IF NOT EXISTS interviews (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            candidate_id INTEGER NOT NULL,
            mandate_id INTEGER,
            owner_id INTEGER,
            round_name TEXT DEFAULT '',
            mode TEXT DEFAULT '',
            location TEXT DEFAULT '',
            interviewer TEXT DEFAULT '',
            scheduled_at TEXT DEFAULT '',
            status TEXT DEFAULT 'scheduled',
            result TEXT DEFAULT '',
            task_snoozed_until TEXT DEFAULT '',
            created_at TEXT DEFAULT ''
        );
        CREATE TABLE IF NOT EXISTS devices (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            fcm_token TEXT NOT NULL,
            device_name TEXT DEFAULT '',
            is_active INTEGER DEFAULT 1,
            created_at TEXT DEFAULT '',
            updated_at TEXT DEFAULT ''
        );
        CREATE TABLE IF NOT EXISTS settings (
            key TEXT PRIMARY KEY,
            value TEXT DEFAULT ''
        );
        CREATE TABLE IF NOT EXISTS tenant_settings (
            company_id INTEGER NOT NULL,
            key TEXT NOT NULL,
            value TEXT DEFAULT '',
            PRIMARY KEY (company_id, key)
        );
        CREATE TABLE IF NOT EXISTS api_usage (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            company_id INTEGER DEFAULT 0,
            provider TEXT DEFAULT '',
            model TEXT DEFAULT '',
            endpoint TEXT DEFAULT '',
            input_tokens INTEGER DEFAULT 0,
            output_tokens INTEGER DEFAULT 0,
            audio_seconds REAL DEFAULT 0,
            cost_usd REAL DEFAULT 0,
            created_at TEXT
        );
        CREATE TABLE IF NOT EXISTS payments (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            company_id INTEGER NOT NULL,
            invoice_no TEXT DEFAULT '',
            amount_inr REAL DEFAULT 0,
            period TEXT DEFAULT '',
            method TEXT DEFAULT 'manual',
            note TEXT DEFAULT '',
            created_at TEXT
        );
    """)
    defaults = [
        ('recruiter_name', 'Nitin Kumar'),
        ('company_name', 'HireLab'),
        ('submission_cc_emails', ''),   # internal team auto-CC on client submissions
        ('company_website', ''),        # shown in the submission email signature (W:)
        ('submission_signature', ''),   # optional: exact signature block (plain text); auto-built if empty
        ('seller_gstin', '09ECWPP1647A1Z9'),
        ('seller_address', 'Office no: GF064, B-128, First Floor, Sector-2, Gautam Buddha Nagar, Noida, Uttar Pradesh 201301.'),
        ('seller_udyam', 'UDYAM : UDYAM-UP-29-0178859 (Micro/Services)'),
        ('seller_state', 'Uttar Pradesh'),
        ('seller_state_code', '09'),
        ('seller_reg_office', 'Building No.36, Tronica City, Ghaziabad, Uttar Pradesh-201102'),
        ('invoice_signatory', 'Pavitra'),
        ('invoice_hsn', '998512'),
        ('invoice_fy', '2026-27'),
        ('imap_host', 'imap.gmail.com'),
        ('cc_bank_cash', ''), ('cc_monthly_fixed', ''), ('cc_team_size', '1'),
        ('cc_year_target', ''), ('cc_funding_available', ''),
        ('cc_target_total', '1000000000'), ('cc_target_years', '3'),
        ('cc_notes', ''), ('cc_last_plan', ''), ('cc_task_prefs', ''), ('cc_last_review', ''),
        ('embedding_api_key', ''), ('embedding_base_url', 'https://api.jina.ai/v1'),
        ('embedding_model', 'jina-embeddings-v3'), ('rag_enabled', '1'),
        ('seller_name', 'HireLab Talent Resource'),
        ('seller_gstin', '09ECWPP1647A1Z9'),
        ('seller_address', 'Office no: GF064, B-128, First Floor, Sector-2, Gautam Buddha Nagar, Noida, Uttar Pradesh 201301.'),
        ('seller_udyam', 'UDYAM : UDYAM-UP-29-0178859 (Micro/Services)'),
        ('seller_state', 'Uttar Pradesh'),
        ('seller_state_code', '09'),
        ('seller_reg_office', 'Building No.36, Tronica City, Ghaziabad, Uttar Pradesh-201102'),
        ('invoice_signatory', 'Pavitra'),
        ('invoice_hsn', '998512'),
        ('invoice_prefix', ''),   # optional prefix like "2026-27/"; auto-computed from FY if blank
        ('claude_api_key', os.environ.get('CLAUDE_API_KEY', '')),
        ('deepseek_api_key', os.environ.get('DEEPSEEK_API_KEY', '')),
        ('groq_api_key', os.environ.get('GROQ_API_KEY', '')),
        ('fu1_hours', '8'),
        ('fu2_hours', '24'),
        ('stale_days', '7'),
        ('promise_hours', '24'),
        ('analytics_stale_days', '7'),
        ('bd_stale_days', '21'),
        ('interview_template',
         'Dear {name},\n\n'
         'We are pleased to inform you that your interview for the position of {role} '
         'has been scheduled.\n\n'
         'Round: {round}\n'
         'Date & Time: {datetime}\n'
         'Mode: {mode}\n'
         '{location_line}\n\n'
         'Please be available on time. Kindly confirm your availability.\n\n'
         'Best regards,\n{recruiter}'),
        ('template_msg1', 'Hi {Name}, this is {RecruiterName} from HireLab. I wanted to speak about a {Position} opportunity at {Location}.\n\nIf you are interested, please suggest the best time to connect.'),
        ('template_fu1', 'Hi {Name}, I had messaged you earlier about a {Position} role at {Location}.\n\nJust following up — would love to connect for a quick 10-minute call.\n\nLooking forward to hearing from you!'),
        ('template_fu2', 'Hi {Name}, this is my last follow up regarding the {Position} opportunity at {Location}.\n\nIf the timing is not right, no worries. But do let me know if you would like to explore this.\n\nHave a great day!'),
        # ── Billing config (super-admin editable) ──
        ('billing_price_per_recruiter', '700'),   # INR per recruiter / month
        ('billing_trial_days', '14'),
        ('billing_usd_inr', '88'),                # rate to convert API cost USD→INR
        ('billing_token_markup', '1.0'),          # multiplier on pass-through token cost
        ('billing_gst_rate', '18'),               # GST % on the invoice
        ('billing_gstin', ''),                    # your GST number (for invoices)
        ('billing_legal_name', 'HireLab Talent Resource'),
        ('billing_address', 'Ghaziabad / NCR, India'),
    ]
    for k, v in defaults:
        c.execute('INSERT OR IGNORE INTO settings (key,value) VALUES (?,?)', (k, v))
    # Migrate: add owner_id to mandates, candidates, reminders (multi-user)
    for tbl in ['mandates', 'candidates', 'reminders']:
        try:
            c.execute(f'ALTER TABLE {tbl} ADD COLUMN owner_id INTEGER DEFAULT 0')
        except sqlite3.OperationalError:
            pass
    # Add created_by to candidates/mandates for tracking
    for tbl in ['mandates', 'candidates']:
        try:
            c.execute(f'ALTER TABLE {tbl} ADD COLUMN created_by TEXT DEFAULT ""')
        except sqlite3.OperationalError:
            pass
    # Per-recruiter mandate assignment (within a company) + company-admin flag
    try:
        c.execute('ALTER TABLE mandates ADD COLUMN assigned_user_id INTEGER DEFAULT 0')
    except sqlite3.OperationalError:
        pass
    try:
        c.execute("ALTER TABLE mandates ADD COLUMN email_templates TEXT DEFAULT '[]'")
    except sqlite3.OperationalError:
        pass
    # Link a mandate to a CRM client (Option B: proper foreign key)
    try:
        c.execute("ALTER TABLE mandates ADD COLUMN crm_client_id INTEGER DEFAULT 0")
    except sqlite3.OperationalError:
        pass
    # Timestamped client notes per mandate (hidden from candidates, used in AI rating)
    c.execute('''CREATE TABLE IF NOT EXISTS mandate_client_notes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        mandate_id INTEGER NOT NULL,
        owner_id INTEGER NOT NULL,
        note TEXT NOT NULL DEFAULT '',
        created_by INTEGER DEFAULT 0,
        created_at TEXT DEFAULT '',
        is_active INTEGER DEFAULT 1
    )''')
    try:
        c.execute('CREATE INDEX IF NOT EXISTS idx_mcn_mandate ON mandate_client_notes(mandate_id, is_active)')
    except sqlite3.OperationalError:
        pass
    # 2-way email: stores both sent and received messages, threaded by Message-ID
    c.execute('''CREATE TABLE IF NOT EXISTS email_messages (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        company_id INTEGER NOT NULL,
        candidate_id INTEGER DEFAULT 0,
        direction TEXT DEFAULT 'sent',
        from_addr TEXT DEFAULT '',
        to_addr TEXT DEFAULT '',
        subject TEXT DEFAULT '',
        body TEXT DEFAULT '',
        message_id TEXT DEFAULT '',
        in_reply_to TEXT DEFAULT '',
        sent_at TEXT DEFAULT '',
        created_at TEXT DEFAULT ''
    )''')
    for sql in [
        'CREATE INDEX IF NOT EXISTS idx_em_candidate ON email_messages(candidate_id, sent_at)',
        'CREATE INDEX IF NOT EXISTS idx_em_company ON email_messages(company_id)',
        'CREATE UNIQUE INDEX IF NOT EXISTS idx_em_msgid ON email_messages(company_id, message_id)',
    ]:
        try:
            c.execute(sql)
        except sqlite3.OperationalError:
            pass
    try:
        c.execute("ALTER TABLE submissions ADD COLUMN task_snoozed_until TEXT DEFAULT ''")
    except sqlite3.OperationalError:
        pass
    # ── Audit & Activity Foundation (PRD-0) ─────────────────────────────────
    # Extend the existing activity_log (non-destructively) so any module can log
    # structured, entity-scoped events for the universal timeline.
    for col, defn in [
        ('company_id', 'INTEGER DEFAULT 0'),   # tenant scope
        ('entity_type', "TEXT DEFAULT ''"),    # e.g. 'candidate','client','invoice'
        ('entity_id', 'INTEGER DEFAULT 0'),    # id of that entity
        ('actor_type', "TEXT DEFAULT 'user'"), # 'user' | 'client' | 'system'
        ('actor_name', "TEXT DEFAULT ''"),     # display name (client contacts have no user row)
        ('meta', "TEXT DEFAULT ''"),           # optional JSON payload for automation
    ]:
        try:
            c.execute(f'ALTER TABLE activity_log ADD COLUMN {col} {defn}')
        except sqlite3.OperationalError:
            pass
    # Field-level audit table: every important change records old → new value.
    c.execute('''CREATE TABLE IF NOT EXISTS audit_log (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        company_id INTEGER DEFAULT 0,
        entity_type TEXT DEFAULT '',
        entity_id INTEGER DEFAULT 0,
        field TEXT DEFAULT '',
        old_value TEXT DEFAULT '',
        new_value TEXT DEFAULT '',
        actor_type TEXT DEFAULT 'user',
        actor_id INTEGER DEFAULT 0,
        actor_name TEXT DEFAULT '',
        created_at TEXT DEFAULT ''
    )''')
    # CRM attachments (Sprint 2) — files attached to a company or a contact.
    c.execute('''CREATE TABLE IF NOT EXISTS crm_attachments (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        company_id INTEGER NOT NULL,              -- tenant
        client_id INTEGER DEFAULT 0,              -- always resolved (contact -> its client)
        entity_type TEXT DEFAULT 'client',        -- client | contact
        entity_id INTEGER DEFAULT 0,
        category TEXT DEFAULT 'Other',            -- NDA | Agreement | PO | ...
        original_name TEXT DEFAULT '',
        stored_name TEXT DEFAULT '',              -- filename on disk (CRM_FILES_DIR)
        size_bytes INTEGER DEFAULT 0,
        mime TEXT DEFAULT '',
        uploaded_by INTEGER DEFAULT 0,
        is_active INTEGER DEFAULT 1,              -- soft delete
        created_at TEXT DEFAULT ''
    )''')
    for idx_sql in [
        'CREATE INDEX IF NOT EXISTS idx_activity_entity ON activity_log(entity_type, entity_id)',
        'CREATE INDEX IF NOT EXISTS idx_activity_company ON activity_log(company_id, created_at)',
        'CREATE INDEX IF NOT EXISTS idx_audit_entity ON audit_log(entity_type, entity_id)',
        'CREATE INDEX IF NOT EXISTS idx_audit_company ON audit_log(company_id, created_at)',
        'CREATE INDEX IF NOT EXISTS idx_crm_att_entity ON crm_attachments(company_id, entity_type, entity_id, is_active)',
        'CREATE INDEX IF NOT EXISTS idx_crm_att_client ON crm_attachments(company_id, client_id, is_active)',
    ]:
        try:
            c.execute(idx_sql)
        except sqlite3.OperationalError:
            pass
    try:
        c.execute('ALTER TABLE users ADD COLUMN is_company_admin INTEGER DEFAULT 0')
    except sqlite3.OperationalError:
        pass
    try:
        c.execute("ALTER TABLE users ADD COLUMN email TEXT DEFAULT ''")
    except sqlite3.OperationalError:
        pass
    c.execute('''CREATE TABLE IF NOT EXISTS password_resets (
        token TEXT PRIMARY KEY,
        user_id INTEGER,
        expires_at TEXT,
        used INTEGER DEFAULT 0,
        created_at TEXT
    )''')
    # Billing columns on companies (for existing DBs)
    for col, defn in [
        ('billing_status', "TEXT DEFAULT 'trial'"),
        ('trial_ends_at', "TEXT DEFAULT ''"),
        ('price_per_recruiter', 'INTEGER DEFAULT 700'),
        ('cf_subscription_id', "TEXT DEFAULT ''"),
        ('gstin', "TEXT DEFAULT ''"),
        ('token_cap', 'INTEGER DEFAULT 0'),   # monthly AI token limit; 0 = unlimited
        ('user_limit', 'INTEGER DEFAULT 0'),  # max approved users (recruiters); 0 = unlimited
    ]:
        try:
            c.execute(f'ALTER TABLE companies ADD COLUMN {col} {defn}')
        except sqlite3.OperationalError:
            pass
    # The platform owner's own company (the first one) is not on trial — it's
    # the owner, mark it active so the owner never bills themselves.
    try:
        c.execute("UPDATE companies SET billing_status='owner' WHERE plan='owner'")
    except sqlite3.OperationalError:
        pass
    # Backfill company-admin: platform super-admins, and the first (lowest-id)
    # user of each company (the agency's own admin from signup/setup).
    try:
        c.execute("UPDATE users SET is_company_admin=1 WHERE role='admin'")
        c.execute('''UPDATE users SET is_company_admin=1 WHERE id IN (
            SELECT MIN(id) FROM users WHERE company_id>0 GROUP BY company_id)''')
    except sqlite3.OperationalError:
        pass

    # Trigger: a candidate inherits its mandate's owner_id automatically, so
    # every insert path (manual, extension, bulk, central-db, import) is covered
    # without touching each one. Only fills when owner_id is 0/NULL.
    try:
        c.execute('''CREATE TRIGGER IF NOT EXISTS candidate_inherit_owner
                     AFTER INSERT ON candidates
                     FOR EACH ROW WHEN (NEW.owner_id IS NULL OR NEW.owner_id=0)
                     BEGIN
                       UPDATE candidates SET owner_id =
                         (SELECT owner_id FROM mandates WHERE id = NEW.mandate_id)
                       WHERE id = NEW.id;
                     END''')
    except sqlite3.OperationalError:
        pass

    # Migrate: billing details on CRM clients (so invoices auto-fill, fill-once).
    for col in ['gstin', 'bill_address', 'bill_state', 'bill_state_code']:
        try:
            c.execute(f'ALTER TABLE crm_clients ADD COLUMN {col} TEXT DEFAULT ""')
        except sqlite3.OperationalError:
            pass

    # Migrate: placement & billing lifecycle fields (close the loop after Joined).
    for col, typ in [('placement_fee', 'REAL DEFAULT 0'), ('joining_date', 'TEXT DEFAULT ""'),
                     ('guarantee_days', 'INTEGER DEFAULT 90'), ('replacement_flag', 'INTEGER DEFAULT 0'),
                     ('billing_notes', 'TEXT DEFAULT ""'),
                     ('offered_ctc', 'REAL DEFAULT 0'), ('fee_percent', 'REAL DEFAULT 8.33')]:
        try:
            c.execute(f'ALTER TABLE candidates ADD COLUMN {col} {typ}')
        except sqlite3.OperationalError:
            pass

    # Migrate: candidate specialization (branch/stream, e.g. "Electronics") —
    # separate from the degree stored in qualification.
    try:
        c.execute('ALTER TABLE candidates ADD COLUMN specialization TEXT DEFAULT ""')
    except sqlite3.OperationalError:
        pass  # already exists

    # Migrate: add deep AI-analysis cache column (single additive column).
    # Stores a JSON blob {"md": "<markdown report>", "at": "<ts>", "model": "..."}
    # so the "AI Analysis" profile tab opens instantly; "Re-run" refreshes it.
    try:
        c.execute('ALTER TABLE candidates ADD COLUMN deep_analysis TEXT DEFAULT ""')
    except sqlite3.OperationalError:
        pass  # already exists

    # Invoicing: GST tax invoices (perm) + business expenses.
    c.execute("""
        CREATE TABLE IF NOT EXISTS invoices (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            owner_id INTEGER DEFAULT 0,
            invoice_no TEXT DEFAULT '', invoice_date TEXT DEFAULT '',
            fy TEXT DEFAULT '', seq INTEGER DEFAULT 0,
            buyer_name TEXT DEFAULT '', buyer_address TEXT DEFAULT '',
            buyer_gstin TEXT DEFAULT '', buyer_state TEXT DEFAULT '', buyer_state_code TEXT DEFAULT '',
            consignee_same INTEGER DEFAULT 1,
            con_name TEXT DEFAULT '', con_address TEXT DEFAULT '',
            con_gstin TEXT DEFAULT '', con_state TEXT DEFAULT '', con_state_code TEXT DEFAULT '',
            candidate_name TEXT DEFAULT '', role TEXT DEFAULT '',
            description TEXT DEFAULT '', extra_lines TEXT DEFAULT '[]',
            hsn TEXT DEFAULT '998512', quantity TEXT DEFAULT '', rate TEXT DEFAULT '',
            per TEXT DEFAULT 'CTC', total_qty TEXT DEFAULT '',
            amount REAL DEFAULT 0, gst_rate REAL DEFAULT 18,
            place_of_supply TEXT DEFAULT '', ref_no TEXT DEFAULT '', other_ref TEXT DEFAULT '',
            order_no TEXT DEFAULT '', order_date TEXT DEFAULT '',
            status TEXT DEFAULT 'sent', due_date TEXT DEFAULT '',
            received_date TEXT DEFAULT '', received_amount REAL DEFAULT 0,
            client_id INTEGER DEFAULT 0, candidate_id INTEGER DEFAULT 0, mandate_id INTEGER DEFAULT 0,
            created_at TEXT, updated_at TEXT
        );
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS expenses (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            owner_id INTEGER DEFAULT 0,
            date TEXT DEFAULT '', category TEXT DEFAULT '',
            payee TEXT DEFAULT '', amount REAL DEFAULT 0, note TEXT DEFAULT '',
            invoice_id INTEGER DEFAULT 0, created_at TEXT
        );
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS vec_chunks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            owner_id INTEGER DEFAULT 0,
            source_type TEXT DEFAULT '', source_id INTEGER DEFAULT 0,
            chash TEXT DEFAULT '', text TEXT DEFAULT '', embedding TEXT DEFAULT '',
            updated_at TEXT
        );
    """)
    try:
        c.execute("CREATE INDEX IF NOT EXISTS idx_vec_owner ON vec_chunks(owner_id, source_type, source_id)")
    except Exception:
        pass

    c.execute("""
        CREATE TABLE IF NOT EXISTS command_tasks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            owner_id INTEGER DEFAULT 0,
            text TEXT DEFAULT '', category TEXT DEFAULT '', priority TEXT DEFAULT 'medium',
            done INTEGER DEFAULT 0, source TEXT DEFAULT 'ai',
            reason TEXT DEFAULT '', ref TEXT DEFAULT '', snooze_until TEXT DEFAULT '',
            task_date TEXT DEFAULT '', created_at TEXT
        );
    """)
    for col in ['reason', 'ref', 'snooze_until']:
        try:
            c.execute(f'ALTER TABLE command_tasks ADD COLUMN {col} TEXT DEFAULT ""')
        except sqlite3.OperationalError:
            pass

    c.execute("""
        CREATE TABLE IF NOT EXISTS command_chat (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            owner_id INTEGER DEFAULT 0,
            role TEXT DEFAULT 'user', content TEXT DEFAULT '', created_at TEXT
        );
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS emails (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            owner_id INTEGER DEFAULT 0,
            msg_id TEXT DEFAULT '', folder TEXT DEFAULT '',
            from_addr TEXT DEFAULT '', from_name TEXT DEFAULT '', to_addr TEXT DEFAULT '',
            subject TEXT DEFAULT '', date_str TEXT DEFAULT '', date_ts REAL DEFAULT 0,
            snippet TEXT DEFAULT '', body TEXT DEFAULT '', body_html TEXT DEFAULT '', is_read INTEGER DEFAULT 0,
            created_at TEXT
        );
    """)

    # Client-submission drafts: composed 'Share to Client' emails saved for later.
    c.execute("""
        CREATE TABLE IF NOT EXISTS submission_drafts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            owner_id INTEGER DEFAULT 0,
            mandate_id INTEGER,
            to_emails TEXT DEFAULT '',
            cc_emails TEXT DEFAULT '',
            subject TEXT DEFAULT '',
            greeting TEXT DEFAULT '',
            intro TEXT DEFAULT '',
            candidate_ids TEXT DEFAULT '[]',
            body_html TEXT DEFAULT '',
            created_at TEXT,
            updated_at TEXT
        );
    """)
    try:
        c.execute("ALTER TABLE submission_drafts ADD COLUMN body_html TEXT DEFAULT ''")
    except sqlite3.OperationalError:
        pass

    # Migrate: SPOC (single point of contact) on a mandate → a CRM contact id.
    try:
        c.execute('ALTER TABLE mandates ADD COLUMN spoc_contact_id INTEGER DEFAULT 0')
    except sqlite3.OperationalError:
        pass  # already exists
    # Migrate: mandate required experience + per-mandate submission CC lists.
    for col, typ in [('experience', 'TEXT DEFAULT ""'),
                     ('cc_external_ids', "TEXT DEFAULT '[]'"),
                     ('cc_internal_emails', "TEXT DEFAULT '[]'")]:
        try:
            c.execute(f'ALTER TABLE mandates ADD COLUMN {col} {typ}')
        except sqlite3.OperationalError:
            pass  # already exists

    # Migrate: add recruiter-pitch cache column on mandates (single additive col).
    # Stores JSON {"md": "<markdown pitch>", "at": "<ts>", "model": "..."} so the
    # "Recruiter Pitch" mandate sub-tab opens instantly; "Re-run" refreshes it.
    try:
        c.execute('ALTER TABLE mandates ADD COLUMN recruiter_pitch TEXT DEFAULT ""')
    except sqlite3.OperationalError:
        pass  # already exists

    # Migrate: add embedding columns to candidates for semantic search
    for col, typ in [('embedding', 'TEXT'), ('embedding_text', 'TEXT'), ('embedded_at', 'TEXT')]:
        try:
            c.execute(f'ALTER TABLE candidates ADD COLUMN {col} {typ} DEFAULT ""')
        except sqlite3.OperationalError:
            pass  # already exists

    # Migrate: embedding METADATA columns (Feature 2 — versioning). Lets us
    # migrate embedding models later without breaking existing vectors.
    # embedded_at already added above. All idempotent.
    for col, typ in [
        ('embedding_model',        'TEXT DEFAULT ""'),
        ('embedding_version',      'TEXT DEFAULT ""'),
        ('embedding_dimension',    'INTEGER DEFAULT 0'),
        ('embedding_status',       'TEXT DEFAULT ""'),
        ('embedding_error',        'TEXT DEFAULT ""'),
        ('embedding_duration_ms',  'INTEGER DEFAULT 0'),
        ('embedding_text_version', 'TEXT DEFAULT ""'),
        ('embedding_vec',          'BLOB'),  # Sprint 4: float32 fast-read cache
    ]:
        try:
            c.execute(f'ALTER TABLE candidates ADD COLUMN {col} {typ}')
        except sqlite3.OperationalError:
            pass  # already exists

    # Backfill metadata for pre-existing vectors (one-time, idempotent: only
    # rows with no status yet are touched, so this is a no-op on later boots).
    # Existing vectors were built by the OLD text builder -> template v1.
    try:
        unstamped = c.execute(
            "SELECT id, embedding FROM candidates "
            "WHERE embedding_status IS NULL OR embedding_status=''").fetchall()
        bf_completed = bf_missing = bf_pending = 0
        for row in unstamped:
            emb = row['embedding']
            if emb is None or emb == '':
                # never embedded
                c.execute("UPDATE candidates SET embedding_status='pending' WHERE id=?", (row['id'],))
                bf_pending += 1
                continue
            dim = 0
            try:
                v = json.loads(emb)
                dim = len(v) if isinstance(v, list) else 0
            except Exception:
                dim = 0
            if dim > 0:
                c.execute(
                    "UPDATE candidates SET embedding_status='completed', embedding_model=?, "
                    "embedding_version=?, embedding_dimension=?, embedding_text_version=? WHERE id=?",
                    ('gemini-embedding-001', 'v1', dim, 'candidate-template-v1', row['id']))
                bf_completed += 1
            else:
                # embedding was '[]' (blank text) or unparseable -> missing
                c.execute(
                    "UPDATE candidates SET embedding_status='missing', embedding_model=?, "
                    "embedding_version=?, embedding_dimension=0, embedding_text_version=? WHERE id=?",
                    ('gemini-embedding-001', 'v1', 'candidate-template-v1', row['id']))
                bf_missing += 1
        if unstamped:
            conn.commit()
            print(f'[embed-migrate] backfilled metadata: completed={bf_completed} '
                  f'missing={bf_missing} pending={bf_pending}')
    except sqlite3.OperationalError:
        pass  # candidates table not ready yet on a fresh DB; nothing to backfill

    # Migrate: embedding job QUEUE (Sprint 3). Persistent so queued/failed jobs
    # survive a cloud restart and resume automatically. One row per embedding
    # attempt lifecycle for a candidate.
    c.execute('''CREATE TABLE IF NOT EXISTS embedding_jobs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        candidate_id INTEGER NOT NULL,
        status TEXT DEFAULT 'pending',      -- pending/processing/completed/failed/retrying/cancelled
        retry_count INTEGER DEFAULT 0,
        max_retries INTEGER DEFAULT 5,
        last_error TEXT DEFAULT '',
        duration_ms INTEGER DEFAULT 0,
        created_at TEXT DEFAULT '',
        started_at TEXT DEFAULT '',
        completed_at TEXT DEFAULT '',
        next_attempt_at TEXT DEFAULT ''
    )''')
    # Poll index (find due jobs fast) + candidate index (idempotency checks).
    for idx_sql in (
        "CREATE INDEX IF NOT EXISTS idx_embjobs_due ON embedding_jobs(status, next_attempt_at)",
        "CREATE INDEX IF NOT EXISTS idx_embjobs_cand ON embedding_jobs(candidate_id)",
    ):
        try:
            c.execute(idx_sql)
        except sqlite3.OperationalError:
            pass

    # Migrate: MULTI-VECTOR embeddings (Phase 2 / Sprint 8). One row per
    # (candidate, facet) so a query can match the RIGHT part of a profile —
    # skills vs experience vs projects — instead of one diluted whole-profile
    # vector. The whole-profile 'full' vector stays on candidates.embedding_vec
    # (search is unchanged); these are additive, blob-only (no JSON) to save disk.
    c.execute('''CREATE TABLE IF NOT EXISTS candidate_vectors (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        candidate_id  INTEGER NOT NULL,
        facet         TEXT NOT NULL,        -- skills / experience / projects
        embedding_vec BLOB,                 -- float32 bytes (same format as candidates.embedding_vec)
        embedding_text TEXT DEFAULT '',
        embedding_model TEXT DEFAULT '',
        embedding_version TEXT DEFAULT '',
        embedding_dimension INTEGER DEFAULT 0,
        embedding_text_version TEXT DEFAULT '',
        status        TEXT DEFAULT 'pending',
        embedded_at   TEXT DEFAULT ''
    )''')
    for idx_sql in (
        "CREATE UNIQUE INDEX IF NOT EXISTS uq_candvec ON candidate_vectors(candidate_id, facet)",
        "CREATE INDEX IF NOT EXISTS idx_candvec_cand ON candidate_vectors(candidate_id)",
        "CREATE INDEX IF NOT EXISTS idx_candvec_facet ON candidate_vectors(facet)",
    ):
        try:
            c.execute(idx_sql)
        except sqlite3.OperationalError:
            pass

    # Migrate: persistent JOB-DESCRIPTION embeddings (Phase 2 / Sprint 9).
    # One reusable vector per mandate so candidate<->JD matching never re-embeds
    # the JD. Kept in a SEPARATE table (not a BLOB on mandates) so the mandates
    # list stays light and JSON-safe.
    c.execute('''CREATE TABLE IF NOT EXISTS mandate_vectors (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        mandate_id    INTEGER NOT NULL,
        embedding_vec BLOB,
        embedding_text TEXT DEFAULT '',
        embedding_model TEXT DEFAULT '',
        embedding_version TEXT DEFAULT '',
        embedding_dimension INTEGER DEFAULT 0,
        embedding_text_version TEXT DEFAULT '',
        status        TEXT DEFAULT 'pending',
        embedded_at   TEXT DEFAULT ''
    )''')
    for idx_sql in (
        "CREATE UNIQUE INDEX IF NOT EXISTS uq_mandvec ON mandate_vectors(mandate_id)",
        "CREATE INDEX IF NOT EXISTS idx_mandvec_status ON mandate_vectors(status)",
    ):
        try:
            c.execute(idx_sql)
        except sqlite3.OperationalError:
            pass

    # ══════════════════════════════════════════════════════════════════
    #  RECRUITMENT MEMORY ENGINE (Sprint 5) — architecture foundation.
    #  Three generic, cross-entity tables that let ANY business object
    #  (candidate, job, client, company, recruiter, skill, industry, ...)
    #  carry long-term memories, an event timeline, and relationships.
    #  Namespaced rme_* so they COMPLEMENT (not replace) the existing
    #  candidate_events / activity_log / stage_history tables. Nothing is
    #  auto-generated yet — this sprint builds the shape only.
    #  Extensibility: JSON metadata columns + schema_version mean future
    #  fields need no migration.
    # ══════════════════════════════════════════════════════════════════
    c.execute('''CREATE TABLE IF NOT EXISTS rme_memories (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        entity_type TEXT NOT NULL,          -- candidate/job/client/company/recruiter/skill/...
        entity_id   TEXT NOT NULL,          -- id or slug of the entity (stored as text = flexible)
        memory_type TEXT DEFAULT 'fact',    -- fact/event/relationship/observation/preference/ai_insight/recruiter_note/system_note/interaction
        title       TEXT DEFAULT '',
        content     TEXT DEFAULT '',
        source      TEXT DEFAULT 'system',  -- recruiter/system/ai/import
        created_by  TEXT DEFAULT '',        -- username / user id / 'system'
        created_at  TEXT DEFAULT '',
        updated_at  TEXT DEFAULT '',
        visibility  TEXT DEFAULT 'internal',-- internal/team/private/candidate_visible
        confidence  REAL DEFAULT 1.0,       -- 0..1 (mainly for ai_insight)
        importance  INTEGER DEFAULT 0,      -- 0..5 priority
        tags        TEXT DEFAULT '[]',      -- JSON array
        metadata    TEXT DEFAULT '{}',      -- JSON, free-form future fields
        status      TEXT DEFAULT 'active',  -- active/archived/deleted
        schema_version INTEGER DEFAULT 1
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS rme_events (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        event_type  TEXT NOT NULL,          -- candidate_created/resume_uploaded/interview_scheduled/...
        entity_type TEXT NOT NULL,
        entity_id   TEXT NOT NULL,
        actor       TEXT DEFAULT 'system',  -- who/what caused it
        summary     TEXT DEFAULT '',
        data        TEXT DEFAULT '{}',      -- JSON payload
        related_entity_type TEXT DEFAULT '',-- optional cross-link (e.g. candidate event about a job)
        related_entity_id   TEXT DEFAULT '',
        created_at  TEXT DEFAULT '',
        schema_version INTEGER DEFAULT 1
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS rme_relationships (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        from_type TEXT NOT NULL, from_id TEXT NOT NULL,
        to_type   TEXT NOT NULL, to_id   TEXT NOT NULL,
        rel_type  TEXT NOT NULL,            -- works_at/has_skill/uses_technology/in_industry/managed_by/prefers/...
        weight     REAL DEFAULT 1.0,        -- strength
        confidence REAL DEFAULT 1.0,
        source     TEXT DEFAULT 'system',
        metadata   TEXT DEFAULT '{}',
        created_at TEXT DEFAULT '',
        updated_at TEXT DEFAULT '',
        status     TEXT DEFAULT 'active',
        schema_version INTEGER DEFAULT 1
    )''')
    for idx_sql in (
        "CREATE INDEX IF NOT EXISTS idx_rme_mem_entity ON rme_memories(entity_type, entity_id)",
        "CREATE INDEX IF NOT EXISTS idx_rme_mem_type   ON rme_memories(memory_type)",
        "CREATE INDEX IF NOT EXISTS idx_rme_mem_status ON rme_memories(status)",
        "CREATE INDEX IF NOT EXISTS idx_rme_evt_entity ON rme_events(entity_type, entity_id)",
        "CREATE INDEX IF NOT EXISTS idx_rme_evt_type   ON rme_events(event_type)",
        "CREATE INDEX IF NOT EXISTS idx_rme_evt_time   ON rme_events(created_at)",
        "CREATE INDEX IF NOT EXISTS idx_rme_rel_from   ON rme_relationships(from_type, from_id)",
        "CREATE INDEX IF NOT EXISTS idx_rme_rel_to     ON rme_relationships(to_type, to_id)",
        "CREATE INDEX IF NOT EXISTS idx_rme_rel_type   ON rme_relationships(rel_type)",
        # one row per (from, to, rel_type) — enables upsert, prevents dup edges
        "CREATE UNIQUE INDEX IF NOT EXISTS uq_rme_rel ON rme_relationships(from_type, from_id, to_type, to_id, rel_type)",
    ):
        try:
            c.execute(idx_sql)
        except sqlite3.OperationalError:
            pass

    # ══════════════════════════════════════════════════════════════════
    #  RECRUITMENT KNOWLEDGE GRAPH (Sprint 7) — architecture foundation.
    #  A canonical CONCEPT registry (rkg_entities) + typed edges (rkg_edges).
    #  Distinct from rme_relationships (which links business objects by id):
    #  the RKG deduplicates concepts via normalization + aliases, so "MCC",
    #  "Motor Control Centre" and "Motor Control Center" collapse to ONE node.
    #  All access goes through rkg_* functions (a repository layer) so a real
    #  graph engine (Neo4j, etc.) can replace the storage later without any
    #  business-logic change. No extraction/reasoning yet — shape only.
    # ══════════════════════════════════════════════════════════════════
    c.execute('''CREATE TABLE IF NOT EXISTS rkg_entities (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        entity_type     TEXT NOT NULL,       -- skill/technology/company/certification/industry/product/location/...
        display_name    TEXT NOT NULL,       -- human form, e.g. "Motor Control Centre"
        normalized_name TEXT NOT NULL,       -- canonical dedup key, e.g. "motorcontrolcentre"
        aliases         TEXT DEFAULT '[]',   -- JSON array of surface forms
        description     TEXT DEFAULT '',
        status          TEXT DEFAULT 'active',
        metadata        TEXT DEFAULT '{}',
        created_at      TEXT DEFAULT '',
        updated_at      TEXT DEFAULT ''
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS rkg_edges (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        source_id  INTEGER NOT NULL,         -- -> rkg_entities.id
        target_id  INTEGER NOT NULL,         -- -> rkg_entities.id
        rel_type   TEXT NOT NULL,            -- WORKED_AT/HAS_SKILL/MANUFACTURES/REQUIRES_SKILL/...
        confidence REAL DEFAULT 1.0,
        source     TEXT DEFAULT 'system',
        status     TEXT DEFAULT 'active',
        metadata   TEXT DEFAULT '{}',
        created_at TEXT DEFAULT '',
        updated_at TEXT DEFAULT ''
    )''')
    for idx_sql in (
        "CREATE UNIQUE INDEX IF NOT EXISTS uq_rkg_entity ON rkg_entities(entity_type, normalized_name)",
        "CREATE INDEX IF NOT EXISTS idx_rkg_entity_type ON rkg_entities(entity_type)",
        "CREATE INDEX IF NOT EXISTS idx_rkg_entity_norm ON rkg_entities(normalized_name)",
        "CREATE INDEX IF NOT EXISTS idx_rkg_edge_src  ON rkg_edges(source_id)",
        "CREATE INDEX IF NOT EXISTS idx_rkg_edge_tgt  ON rkg_edges(target_id)",
        "CREATE INDEX IF NOT EXISTS idx_rkg_edge_type ON rkg_edges(rel_type)",
        "CREATE UNIQUE INDEX IF NOT EXISTS uq_rkg_edge ON rkg_edges(source_id, target_id, rel_type)",
    ):
        try:
            c.execute(idx_sql)
        except sqlite3.OperationalError:
            pass

    # Alias resolution table — makes normalization DATA-driven (not a hardcoded
    # dictionary): every surface form maps to its canonical entity so future
    # lookups of a learned alias resolve to the same node.
    c.execute('''CREATE TABLE IF NOT EXISTS rkg_aliases (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        entity_id        INTEGER NOT NULL,
        entity_type      TEXT NOT NULL,
        normalized_alias TEXT NOT NULL,
        created_at       TEXT DEFAULT ''
    )''')
    for idx_sql in (
        "CREATE UNIQUE INDEX IF NOT EXISTS uq_rkg_alias ON rkg_aliases(entity_type, normalized_alias)",
        "CREATE INDEX IF NOT EXISTS idx_rkg_alias_entity ON rkg_aliases(entity_id)",
    ):
        try:
            c.execute(idx_sql)
        except sqlite3.OperationalError:
            pass

    # ── Skill Graph (Intelligence Layer · Phase 1) ─────────────────────────
    # A domain skill graph for electrical / automation / solar recruiting.
    # owner_id=0 rows are the shared system seed; owner_id=<company> rows are a
    # tenant's own additions. Powers query-time skill expansion in search
    # (e.g. "BMS" also matches BACnet / Niagara / DDC candidates).
    c.execute('''CREATE TABLE IF NOT EXISTS skill_graph (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        owner_id  INTEGER DEFAULT 0,
        canonical TEXT NOT NULL,
        display   TEXT DEFAULT '',
        category  TEXT DEFAULT '',
        aliases   TEXT DEFAULT '[]',
        parents   TEXT DEFAULT '[]',
        related   TEXT DEFAULT '[]',
        created_at TEXT DEFAULT '',
        updated_at TEXT DEFAULT ''
    )''')
    try:
        c.execute("CREATE UNIQUE INDEX IF NOT EXISTS uq_skill_graph ON skill_graph(owner_id, canonical)")
    except sqlite3.OperationalError:
        pass
    try:
        _seed_skill_graph(c)
    except Exception as _e:
        print('[skill-graph] seed skipped:', _e)

    # Migrate: add reminders table if not exists
    c.execute('''CREATE TABLE IF NOT EXISTS reminders (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        candidate_id INTEGER NOT NULL,
        mandate_id INTEGER,
        candidate_name TEXT DEFAULT '',
        mandate_label TEXT DEFAULT '',
        note TEXT DEFAULT '',
        due_at TEXT NOT NULL,
        done INTEGER DEFAULT 0,
        created_at TEXT
    )''')

    # Migrate: add new tag columns to existing DBs
    for col, defn in [
        ('product_handles', 'TEXT DEFAULT "[]"'),
        ('function_tags', 'TEXT DEFAULT "[]"'),
        ('status_tags', 'TEXT DEFAULT "[]"'),
        ('preferred_location', "TEXT DEFAULT ''"),
        ('task_snoozed_until', "TEXT DEFAULT ''"),
        ('update_token', "TEXT DEFAULT ''"),
        ('update_requested_at', "TEXT DEFAULT ''"),
        ('update_submitted_at', "TEXT DEFAULT ''"),
        ('linkedin_url', "TEXT DEFAULT ''"),
        ('ai_insight_cv', "TEXT DEFAULT ''"),
        ('sourced_by', "INTEGER DEFAULT 0"),
        ('sourced_at', "TEXT DEFAULT ''"),
    ]:
        try:
            c.execute(f'ALTER TABLE candidates ADD COLUMN {col} {defn}')
        except Exception:
            pass

    # Migrate: reminder scheduler columns (push notifications + snooze)
    for col, defn in [
        ('notified_at', "TEXT DEFAULT ''"),       # last time a push was sent
        ('snoozed_until', "TEXT DEFAULT ''"),      # if snoozed, don't notify until this time
        ('early_warned', "INTEGER DEFAULT 0"),     # sent the 5-10 min advance warning?
        ('owner_id', "INTEGER DEFAULT 0"),
        ('notify_count', "INTEGER DEFAULT 0"),     # how many due-notifications sent (cap the spam)
        ('stage', "TEXT DEFAULT 'todo'"),          # Kanban column: todo / doing / done
    ]:
        try:
            c.execute(f'ALTER TABLE reminders ADD COLUMN {col} {defn}')
        except Exception:
            pass

    # Migrate: add signup-approval + company columns to existing 'users' table
    for col, defn in [
        ('status', "TEXT DEFAULT 'approved'"),
        ('company_name', "TEXT DEFAULT ''"),
        ('requested_at', 'TEXT'),
        ('company_id', 'INTEGER DEFAULT 0'),
        ('profile_phone', "TEXT DEFAULT ''"),
        ('profile_designation', "TEXT DEFAULT ''"),
        ('profile_email', "TEXT DEFAULT ''"),
    ]:
        try:
            c.execute(f'ALTER TABLE users ADD COLUMN {col} {defn}')
        except Exception:
            pass
    # Backfill: any pre-existing user (created before this column existed)
    # must default to 'approved' so nobody already in the system gets locked
    # out by the new approval gate.
    try:
        c.execute("UPDATE users SET status='approved' WHERE status IS NULL OR status=''")
    except Exception:
        pass

    conn.commit()

    # ── ONE-TIME MULTI-TENANT MIGRATION ────────────────────────────────────
    # Phase 2: introduce companies (tenants). Data isolation is by company.
    # We REUSE the existing owner_id column on mandates/candidates/reminders as
    # the TENANT (company) id — every data-filtering query already filters on
    # owner_id, so isolation is enforced everywhere automatically and there is
    # no risk of a missed filter leaking data across tenants.
    #
    # This block runs only once: if there are users but no companies yet, we:
    #   1. Create one company per the user's stated company_name (or a default
    #      "HireLab" company for the admin), so existing data stays together as
    #      a single agency exactly as it is today.
    #   2. Point every user at their company.
    #   3. Remap all existing data rows: owner_id (currently a USER id) becomes
    #      the COMPANY id of whoever owned it.
    try:
        have_users = c.execute('SELECT COUNT(*) n FROM users').fetchone()['n']
        have_companies = c.execute('SELECT COUNT(*) n FROM companies').fetchone()['n']
    except Exception:
        have_users, have_companies = 0, 0

    if have_users > 0 and have_companies == 0:
        print('*** Phase-2 tenant migration: creating companies for existing users ***')
        # All existing users belong to ONE agency: "HireLab" (the original
        # single-company system). This matches the owner's mental model that
        # the existing 290 candidates / 10 mandates are HireLab's own data.
        # New signups (post-migration) each get their OWN company.
        admin_row = c.execute("SELECT id, company_name FROM users WHERE role='admin' ORDER BY id LIMIT 1").fetchone()
        default_name = ''
        if admin_row and (admin_row['company_name'] or '').strip():
            default_name = admin_row['company_name'].strip()
        if not default_name:
            default_name = 'HireLab'
        c.execute("INSERT INTO companies (name,status,plan,created_at) VALUES (?,?,?,?)",
                  (default_name, 'active', 'owner', ts()))
        hirelab_company_id = c.lastrowid
        # Point every existing user at this company.
        c.execute('UPDATE users SET company_id=?', (hirelab_company_id,))
        # Remap ALL existing data to this single company (owner_id was a user id
        # before; now it is the company id). Existing data becomes HireLab's.
        for tbl in ['mandates', 'candidates', 'reminders']:
            try:
                c.execute(f'UPDATE {tbl} SET owner_id=?', (hirelab_company_id,))
            except Exception:
                pass
        conn.commit()
        print(f'*** Tenant migration complete: company "{default_name}" (id={hirelab_company_id}) now owns all existing data ***')

    # ── RecruitOS modules: import them first (registers their migrations),
    #    then build their tables alongside core schema ──────────────────────
    try:
        import modules
        modules.import_all_modules()
        modules.run_migrations(conn)
    except Exception as e:
        print(f'[modules] migration hook skipped: {e}')

    conn.commit(); conn.close()

    # One-time safety migration: if this DB file still has a pending -wal file
    # on disk from before the journal-mode fix, force a full checkpoint so
    # those writes land in the main DB file before we proceed. Harmless no-op
    # if the DB was already in DELETE mode (no -wal file exists).
    try:
        wal_path = DB_PATH + '-wal'
        if os.path.exists(wal_path):
            _c2 = sqlite3.connect(DB_PATH, timeout=60)
            _c2.execute('PRAGMA wal_checkpoint(TRUNCATE)')
            _c2.execute('PRAGMA journal_mode=DELETE')
            _c2.close()
            print('*** One-time WAL checkpoint completed: pending writes flushed to main DB file ***')
    except Exception as _wal_err:
        print(f'WAL checkpoint warning (non-fatal): {_wal_err}')


def migrate_old():
    if os.path.exists(DB_PATH):
        return
    script_dir = os.path.dirname(os.path.abspath(__file__))
    for old_path in [os.path.join(script_dir, 'hirelab.db'), os.path.join(os.getcwd(), 'hirelab.db')]:
        if os.path.exists(old_path):
            shutil.copy2(old_path, DB_PATH)
            print(f'[MIGRATE] {old_path} -> {DB_PATH}')
            old_cvs = os.path.join(os.path.dirname(old_path), 'cvs')
            if os.path.exists(old_cvs):
                for f in os.listdir(old_cvs):
                    src = os.path.join(old_cvs, f)
                    dst = os.path.join(CV_DIR, f)
                    if os.path.isfile(src) and not os.path.exists(dst):
                        shutil.copy2(src, dst)
            break

def log_candidate_event(cid, event_type, detail=''):
    """Record a journey event (tag added, call analysed, etc.) for a candidate."""
    try:
        conn = get_db()
        conn.execute('INSERT INTO candidate_events (candidate_id,event_type,detail,created_at) VALUES (?,?,?,?)',
                     (cid, event_type, detail, ts()))
        conn.commit(); conn.close()
    except Exception as e:
        print(f'log_candidate_event warning: {e}')


def daily_backup():
    """Snapshot the DB once per day. NEVER snapshot an empty DB over an existing
    good backup, and refresh today's backup if the live DB now has more users
    than the stored snapshot (so a startup-time empty snapshot can't 'stick')."""
    if not os.path.exists(DB_PATH):
        return
    live_users = _db_user_count(DB_PATH)
    if live_users == 0:
        return  # Never back up an empty DB — it could clobber a good backup.
    bak = os.path.join(BAK_DIR, f'hirelab_{datetime.date.today()}.db')
    existing_users = _db_user_count(bak) if os.path.exists(bak) else -1
    if (not os.path.exists(bak)) or live_users >= existing_users:
        shutil.copy2(DB_PATH, bak)
        print(f'[BACKUP] {bak} ({live_users} users)')
    for old in sorted(Path(BAK_DIR).glob('hirelab_*.db'))[:-7]:
        old.unlink()


def _db_user_count(path):
    """How many users a given SQLite file holds (0 if unreadable/missing)."""
    try:
        if not os.path.exists(path):
            return 0
        c = sqlite3.connect(path, timeout=10)
        n = c.execute('SELECT COUNT(*) FROM users').fetchone()[0]
        c.close()
        return n
    except Exception:
        return 0


def auto_restore_if_empty():
    """SAFETY NET against the 'logged out → asked to create a new account →
    old data gone' problem. If the live DB has zero users (e.g. a restart lost
    recent writes, or the file was recreated empty), but a backup on disk DOES
    contain users, restore the newest such backup automatically. This runs at
    every startup, including under gunicorn on the cloud."""
    try:
        live_users = _db_user_count(DB_PATH)
        if live_users > 0:
            return  # DB is healthy, nothing to do.
        # Pick the backup with the MOST users (most complete), newest as tiebreak.
        backups = sorted(Path(BAK_DIR).glob('hirelab_*.db'),
                         key=lambda p: (_db_user_count(str(p)), p.stat().st_mtime),
                         reverse=True)
        for bak in backups:
            if _db_user_count(str(bak)) > 0:
                # Keep a copy of the (empty) current file just in case.
                try:
                    if os.path.exists(DB_PATH):
                        shutil.copy2(DB_PATH, DB_PATH + '.empty-before-restore')
                except Exception:
                    pass
                shutil.copy2(str(bak), DB_PATH)
                print(f'*** AUTO-RESTORE: live DB had 0 users — restored from backup {bak.name} '
                      f'({_db_user_count(DB_PATH)} users recovered). ***')
                return
        if not backups:
            print('*** AUTO-RESTORE: DB empty and NO backups found on disk. '
                  'If this is a restart, your storage may NOT be persistent — see /api/diag. ***')
    except Exception as e:
        print(f'Auto-restore warning (non-fatal): {e}')


def check_storage_persistence():
    """Definitively detect whether DATA_DIR survives restarts. We write a marker
    file containing a boot counter. If the marker is MISSING on a later boot,
    the storage is ephemeral (data WILL be lost on every restart) — the true
    root cause of the 'asks me to create a new account again' problem. Returns
    a dict used by /api/diag and the startup banner."""
    marker = os.path.join(DATA_DIR, '.persistence_check')
    result = {'marker_path': marker, 'persistent': None, 'boots_seen': 0}
    try:
        prev = ''
        if os.path.exists(marker):
            with open(marker) as f:
                prev = f.read().strip()
        if prev:
            # Marker survived a previous boot → storage IS persistent.
            try:
                boots = int(prev.split('|')[0]) + 1
            except Exception:
                boots = 1
            result['persistent'] = True
            result['boots_seen'] = boots
        else:
            # First boot (or marker was wiped). Can't conclude persistence yet.
            boots = 1
            result['persistent'] = None  # unknown until we see it survive once
            result['boots_seen'] = 1
        with open(marker, 'w') as f:
            f.write(f'{boots}|{datetime.datetime.now().isoformat()}')
    except Exception as e:
        result['error'] = str(e)
    return result


# Cache the persistence result computed at startup.
_PERSISTENCE = {'persistent': None, 'boots_seen': 0}

def check_timers():
    conn = get_db(); c = conn.cursor()
    r1 = c.execute("SELECT value FROM settings WHERE key='fu1_hours'").fetchone()
    r2 = c.execute("SELECT value FROM settings WHERE key='fu2_hours'").fetchone()
    fu1_h = float(r1['value']) if r1 else 8.0
    fu2_h = float(r2['value']) if r2 else 24.0
    n = datetime.datetime.utcnow()
    for cand in c.execute("SELECT id,msg1_sent_at FROM candidates WHERE msg1_sent_at!='' AND stage='Screening'").fetchall():
        try:
            if (n - datetime.datetime.fromisoformat(cand['msg1_sent_at'])).total_seconds() >= fu1_h * 3600:
                c.execute("UPDATE candidates SET stage='Follow Up 1',updated_at=? WHERE id=?", (ts(), cand['id']))
                c.execute("INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) VALUES (?,?,?,?,?)",
                          (cand['id'], 'Screening', 'Follow Up 1', f'Auto-moved after {fu1_h}h', ts()))
        except Exception:
            pass
    for cand in c.execute("SELECT id,fu1_sent_at FROM candidates WHERE fu1_sent_at!='' AND stage='Follow Up 1'").fetchall():
        try:
            if (n - datetime.datetime.fromisoformat(cand['fu1_sent_at'])).total_seconds() >= fu2_h * 3600:
                c.execute("UPDATE candidates SET stage='Follow Up 2',updated_at=? WHERE id=?", (ts(), cand['id']))
                c.execute("INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) VALUES (?,?,?,?,?)",
                          (cand['id'], 'Follow Up 1', 'Follow Up 2', f'Auto-moved after {fu2_h}h', ts()))
        except Exception:
            pass
    conn.commit(); conn.close()

# ── Per-tenant API usage & cost tracking ──────────────────────────────────
# Rates are EDITABLE ESTIMATES in USD. Token COUNTS logged below are exact
# (read from each API response); only these per-unit rates are approximate and
# can be adjusted any time without touching the logged history.
#   *_in  / *_out = USD per 1,000,000 tokens
#   whisper       = USD per second of audio
API_PRICING = {
    'claude':   {'in': 3.00,  'out': 15.00},   # Claude Sonnet (per 1M tokens)
    'deepseek': {'in': 0.27,  'out': 1.10},    # DeepSeek chat   (per 1M tokens)
    'gemini':   {'in': 0.075, 'out': 0.30},    # Gemini embed/flash (per 1M tokens)
    'groq':     {'audio_per_sec': 0.111/3600}, # Groq Whisper ~ $0.111 / audio-hour
}

def log_api_usage(provider, model='', input_tokens=0, output_tokens=0, audio_seconds=0, endpoint=''):
    """Record one AI API call against the current tenant, with an estimated
    cost. Fully defensive — never raises into the calling request."""
    try:
        p = API_PRICING.get(provider, {})
        cost = 0.0
        if audio_seconds:
            cost += float(audio_seconds) * p.get('audio_per_sec', 0)
        if input_tokens:
            cost += (float(input_tokens) / 1_000_000.0) * p.get('in', 0)
        if output_tokens:
            cost += (float(output_tokens) / 1_000_000.0) * p.get('out', 0)
        try:
            company_id = effective_company_id() or 0
        except Exception:
            company_id = 0
        conn = get_db()
        conn.execute('''INSERT INTO api_usage
            (company_id,provider,model,endpoint,input_tokens,output_tokens,audio_seconds,cost_usd,created_at)
            VALUES (?,?,?,?,?,?,?,?,?)''',
            (company_id, provider, model, endpoint,
             int(input_tokens or 0), int(output_tokens or 0),
             float(audio_seconds or 0), round(cost, 6), ts()))
        conn.commit(); conn.close()
    except Exception as e:
        print(f'log_api_usage warning (non-fatal): {e}')


class TokenCapError(Exception):
    """Raised when a company has hit its monthly AI token cap."""
    pass


@app.errorhandler(TokenCapError)
def _handle_token_cap(e):
    return jsonify({'error': 'Monthly AI usage limit reached for your agency. Please contact your admin to increase the cap.', 'code': 'token_cap'}), 429


def tokens_used_this_month(company_id):
    if not company_id:
        return 0
    month_start = datetime.datetime.now().strftime('%Y-%m-01 00:00:00')
    conn = get_db()
    r = conn.execute("SELECT COALESCE(SUM(input_tokens+output_tokens),0) t FROM api_usage WHERE company_id=? AND created_at >= ?",
                     (company_id, month_start)).fetchone()
    conn.close()
    return int((r['t'] if r else 0) or 0)


def company_token_cap(company_id):
    if not company_id:
        return 0
    conn = get_db()
    r = conn.execute("SELECT token_cap FROM companies WHERE id=?", (company_id,)).fetchone()
    conn.close()
    try:
        return int((r['token_cap'] if r else 0) or 0)
    except Exception:
        return 0


def over_token_cap(company_id=None):
    """True if the company is at/over its monthly AI token cap (0 = unlimited).
    Safe outside a request context (returns False)."""
    try:
        if company_id is None:
            company_id = effective_company_id()
    except Exception:
        return False
    cap = company_token_cap(company_id)
    if not cap:
        return False
    return tokens_used_this_month(company_id) >= cap


def call_claude(api_key, system_msg, messages, max_tokens=8000, endpoint='claude'):
    if over_token_cap():
        raise TokenCapError()
    resp = requests.post(CLAUDE_URL,
        headers={'x-api-key': api_key, 'anthropic-version': '2023-06-01', 'content-type': 'application/json'},
        json={'model': CLAUDE_MODEL, 'max_tokens': max_tokens, 'system': system_msg, 'messages': messages},
        timeout=120)
    # Log token usage (per tenant) without breaking the caller.
    try:
        u = resp.json().get('usage', {})
        if u:
            log_api_usage('claude', CLAUDE_MODEL,
                          input_tokens=u.get('input_tokens', 0),
                          output_tokens=u.get('output_tokens', 0),
                          endpoint=endpoint)
    except Exception:
        pass
    return resp

def call_deepseek(api_key, payload, timeout=60, endpoint='deepseek'):
    """POST to DeepSeek (OpenAI-compatible) and log token usage per tenant.
    Returns the raw requests response, so existing callers work unchanged."""
    if over_token_cap():
        raise TokenCapError()
    resp = requests.post('https://api.deepseek.com/chat/completions',
        headers={'Authorization': 'Bearer ' + api_key, 'Content-Type': 'application/json'},
        json=payload, timeout=timeout)
    try:
        u = resp.json().get('usage', {})
        if u:
            log_api_usage('deepseek', payload.get('model', 'deepseek-chat'),
                          input_tokens=u.get('prompt_tokens', 0),
                          output_tokens=u.get('completion_tokens', 0),
                          endpoint=endpoint)
    except Exception:
        pass
    return resp

def parse_json(text):
    text = text.strip()
    if '```' in text:
        for part in text.split('```'):
            p = part.strip()
            if p.startswith('json'): p = p[4:].strip()
            if p.startswith(('{', '[')):
                text = p; break
    for bracket in [('[', ']'), ('{', '}')]:
        s = text.find(bracket[0])
        if s >= 0:
            e = text.rfind(bracket[1]) + 1
            if e > s:
                try: return json.loads(text[s:e])
                except Exception: pass
    return None

# Map sensitive setting keys to environment variable names. If the env var is
# set (e.g. on Render), it OVERRIDES whatever is stored in the DB. This lets API
# keys live safely in the host environment instead of in code or the database.
_ENV_KEY_MAP = {
    'groq_api_key': 'GROQ_API_KEY',
    'claude_api_key': 'CLAUDE_API_KEY',
    'deepseek_api_key': 'DEEPSEEK_API_KEY',
    'gemini_api_key': 'GEMINI_API_KEY',
    'embedding_api_key': 'JINA_API_KEY',
}

def get_setting(key, default=''):
    # Env var takes priority for sensitive keys
    env_name = _ENV_KEY_MAP.get(key)
    if env_name:
        env_val = os.environ.get(env_name, '').strip()
        if env_val:
            return env_val
    conn = get_db()
    r = conn.execute('SELECT value FROM settings WHERE key=?', (key,)).fetchone()
    conn.close()
    return r['value'] if r else default

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  ROUTES
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.route('/')
def index():
    return send_file(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'index.html'))
@app.route('/v2')
def recruitos_v2():
    return send_file(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'recruitos-live.html'))

# ── Legal pages (Privacy / Terms / Refund) ─────────────────────────────────
# Served from here so deployment stays single-file. Company details are
# pre-filled from HireLab's records — confirm the support email + jurisdiction
# before going live, and have a lawyer review for your final terms.
_LEGAL_COMPANY = 'HireLab Talent Resource'
_LEGAL_EMAIL   = 'support@hirelabtalent.com'
_LEGAL_SITE    = 'hirelabtalent.com'
_LEGAL_JURIS   = 'Ghaziabad, Uttar Pradesh, India'

def _legal_page(title, body_html):
    updated = datetime.date.today().strftime('%d %B %Y')
    return (
        '<!doctype html><html lang="en"><head><meta charset="utf-8">'
        '<meta name="viewport" content="width=device-width,initial-scale=1">'
        '<title>' + title + ' — HireLab Screener</title>'
        '<style>body{font-family:-apple-system,Segoe UI,Roboto,sans-serif;max-width:760px;'
        'margin:0 auto;padding:40px 22px;color:#1f2a37;line-height:1.6;font-size:15px}'
        'h1{color:#0E2A47;font-size:26px;margin-bottom:4px}h2{color:#0E2A47;font-size:18px;'
        'margin-top:28px}.upd{color:#8A92A0;font-size:12.5px;margin-bottom:24px}'
        'a{color:#13A37E}.back{display:inline-block;margin-top:32px;font-size:13px}'
        'ul{padding-left:20px}li{margin:5px 0}</style></head><body>'
        '<h1>' + title + '</h1><div class="upd">Last updated: ' + updated + '</div>'
        + body_html +
        '<a class="back" href="/">&larr; Back to HireLab Screener</a></body></html>'
    )

@app.route('/privacy')
def privacy_page():
    body = (
        '<p>' + _LEGAL_COMPANY + ' ("we", "us") operates HireLab Screener, a recruitment '
        'software platform. This policy explains what data we handle and how we protect it.</p>'
        '<h2>1. Information we process</h2><ul>'
        '<li><b>Account data</b> of subscribing recruiters/agencies: name, username, email, agency name.</li>'
        '<li><b>Candidate data</b> that you (the recruiter) upload or import: names, contact details, '
        'CVs, work history, compensation and related recruitment information.</li>'
        '<li><b>Usage/technical data</b>: log entries, IP address, and diagnostic information.</li></ul>'
        '<h2>2. How we use it</h2><p>To provide and operate the platform, authenticate users, '
        'process recruitment workflows you initiate, keep the service secure, and provide support. '
        'We do not sell your data or your candidates\' data.</p>'
        '<h2>3. Data isolation</h2><p>Each subscribing agency\'s data is logically isolated. '
        'One agency cannot access another agency\'s candidates, mandates, or files.</p>'
        '<h2>4. AI processing</h2><p>Some features (parsing, ranking, drafting) send the relevant '
        'text to third-party AI providers to perform that task. Only the data needed for the '
        'requested feature is sent.</p>'
        '<h2>5. Candidate data &amp; your responsibility</h2><p>Recruiters using the platform are '
        'the data controllers for the candidate information they upload and are responsible for '
        'having a lawful basis to process it. We act as a processor on your behalf.</p>'
        '<h2>6. Retention &amp; deletion</h2><p>Data is retained while your subscription is active. '
        'On cancellation, data is retained during any grace period and then deleted. You may request '
        'export or deletion of your data by contacting us.</p>'
        '<h2>7. Security</h2><p>We use authenticated access, per-agency isolation, hashed passwords, '
        'and encrypted transport (HTTPS). No system is perfectly secure, but we take reasonable steps '
        'to protect your data.</p>'
        '<h2>8. Contact</h2><p>Questions or data requests: <a href="mailto:' + _LEGAL_EMAIL + '">'
        + _LEGAL_EMAIL + '</a> &middot; ' + _LEGAL_SITE + '</p>'
    )
    return _legal_page('Privacy Policy', body)

@app.route('/terms')
def terms_page():
    body = (
        '<p>These Terms govern your use of HireLab Screener, provided by ' + _LEGAL_COMPANY + '. '
        'By creating an account or using the service you agree to these Terms.</p>'
        '<h2>1. The service</h2><p>HireLab Screener is a subscription-based recruitment software '
        'platform. We may add, change, or remove features over time.</p>'
        '<h2>2. Accounts</h2><p>You are responsible for your login credentials and all activity under '
        'your account. Accounts are approved by an administrator. Keep your password confidential.</p>'
        '<h2>3. Subscription &amp; billing</h2><p>Access is provided on a per-recruiter subscription '
        'basis. Fees, billing cycle, and payment method are as agreed at sign-up. Access may be '
        'suspended for non-payment or on expiry of a trial.</p>'
        '<h2>4. Acceptable use</h2><ul>'
        '<li>Only upload candidate data you are lawfully entitled to process.</li>'
        '<li>Do not attempt to access other agencies\' data, probe, or disrupt the service.</li>'
        '<li>Do not use the service for unlawful, abusive, or spam activity.</li></ul>'
        '<h2>5. Your data</h2><p>You retain ownership of the data you upload. You grant us a limited '
        'licence to process it solely to provide the service. See our '
        '<a href="/privacy">Privacy Policy</a>.</p>'
        '<h2>6. Availability</h2><p>We aim for high availability but do not guarantee uninterrupted '
        'service. Maintenance and occasional downtime may occur.</p>'
        '<h2>7. Limitation of liability</h2><p>The service is provided "as is". To the extent permitted '
        'by law, we are not liable for indirect or consequential losses, or for loss of data beyond our '
        'reasonable control. Our total liability is limited to the fees paid in the preceding 3 months.</p>'
        '<h2>8. Termination</h2><p>Either party may terminate as per the subscription arrangement. On '
        'termination, access ends and data is handled per the Privacy Policy.</p>'
        '<h2>9. Governing law</h2><p>These Terms are governed by the laws of India, with jurisdiction '
        'in the courts of ' + _LEGAL_JURIS + '.</p>'
        '<h2>10. Contact</h2><p><a href="mailto:' + _LEGAL_EMAIL + '">' + _LEGAL_EMAIL + '</a> &middot; '
        + _LEGAL_SITE + '</p>'
    )
    return _legal_page('Terms of Service', body)

@app.route('/refund')
def refund_page():
    body = (
        '<p>This policy explains billing, cancellation, and refunds for HireLab Screener '
        'subscriptions provided by ' + _LEGAL_COMPANY + '.</p>'
        '<h2>1. Subscriptions</h2><p>Subscriptions are billed per recruiter for the agreed period. '
        'Access continues until the end of the paid period.</p>'
        '<h2>2. Cancellation</h2><p>You may cancel at any time by contacting us. Cancellation stops '
        'future renewals; your access continues until the end of the current paid period.</p>'
        '<h2>3. Refunds</h2><p>Fees already paid for the current period are generally non-refundable, '
        'except where required by law. If you believe you were billed in error, contact us within 7 days '
        'and we will review it in good faith.</p>'
        '<h2>4. Trials</h2><p>Free-trial periods are not charged. If you do not subscribe before the '
        'trial ends, access is paused until you subscribe.</p>'
        '<h2>5. Data after cancellation</h2><p>After cancellation, your data is retained during a short '
        'grace period and then deleted. Request an export before cancelling if you need your data.</p>'
        '<h2>6. Contact</h2><p><a href="mailto:' + _LEGAL_EMAIL + '">' + _LEGAL_EMAIL + '</a> &middot; '
        + _LEGAL_SITE + '</p>'
    )
    return _legal_page('Refund &amp; Cancellation Policy', body)
# ── PWA: installable mobile app (manifest + service worker + icons) ────────
@app.route('/manifest.webmanifest')
def pwa_manifest():
    return jsonify({
        'name': 'HireLab Screener',
        'short_name': 'HireLab',
        'description': 'Recruitment ATS with mobile call assistant',
        'start_url': '/',
        'display': 'standalone',
        'background_color': '#0E2A47',
        'theme_color': '#0E2A47',
        'orientation': 'portrait-primary',
        'icons': [
            {'src': '/icon-192.png', 'sizes': '192x192', 'type': 'image/png', 'purpose': 'any maskable'},
            {'src': '/icon-512.png', 'sizes': '512x512', 'type': 'image/png', 'purpose': 'any maskable'},
        ]
    })

@app.route('/sw.js')
def pwa_sw():
    # Minimal network-first service worker. Its presence (with a fetch handler)
    # is what makes the app installable to the home screen. We deliberately do
    # NOT cache API responses so tenant data is always fresh from the server.
    sw = (
        "self.addEventListener('install', e => self.skipWaiting());\n"
        "self.addEventListener('activate', e => self.clients.claim());\n"
        "self.addEventListener('fetch', function(e){\n"
        "  // Pass through to network; no offline caching of data.\n"
        "  e.respondWith(fetch(e.request).catch(function(){\n"
        "    return new Response('Offline', {status: 503});\n"
        "  }));\n"
        "});\n"
    )
    return app.response_class(sw, mimetype='application/javascript')

@app.route('/icon-192.png')
def pwa_icon_192():
    p = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'icon-192.png')
    return send_file(p) if os.path.exists(p) else ('', 404)

@app.route('/icon-512.png')
def pwa_icon_512():
    p = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'icon-512.png')
    return send_file(p) if os.path.exists(p) else ('', 404)




@app.route('/api/db-status')
@admin_required
def db_status():
    """Check DB health — useful for debugging Railway/Render issues."""
    try:
        conn = get_db()
        tables = conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()
        table_names = [t['name'] for t in tables]
        counts = {}
        for t in table_names:
            try:
                counts[t] = conn.execute(f'SELECT COUNT(*) as c FROM {t}').fetchone()['c']
            except Exception:
                counts[t] = -1
        conn.close()
        return jsonify({
            'ok': True,
            'db_path': DB_PATH,
            'data_dir': DATA_DIR,
            'tables': table_names,
            'counts': counts
        })
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e), 'db_path': DB_PATH, 'data_dir': DATA_DIR}), 500


@app.route('/api/shorten-jd', methods=['POST'])
@login_required
def shorten_jd():
    """Use DeepSeek to shorten a long JD/role-description text so it fits
    within Gmail's compose-URL length limit, while keeping it useful for
    a candidate outreach email (role purpose + 4-6 key responsibilities/
    highlights, in plain text, no markdown)."""
    d = request.json or {}
    text = (d.get('text') or '').strip()
    max_chars = int(d.get('max_chars') or 1200)
    if not text:
        return jsonify({'error': 'No text provided'}), 400

    ds_key = get_setting('deepseek_api_key')
    if not ds_key:
        return jsonify({'error': 'DeepSeek API key not set. Go to Settings.'}), 400

    system_msg = (
        "You are condensing a long job description into a well-formatted, "
        f"recruiter-friendly summary, under roughly {max_chars} characters total. "
        "Be thoughtful — preserve the most important and specific information "
        "(role purpose, key responsibilities, must-have requirements) and cut "
        "only repetitive, generic, or low-value detail. Do not cut sentences "
        "mid-way; every section and bullet must be complete and make sense.\n\n"
        "Output format (plain text, using these exact markers so it can be "
        "converted to formatted HTML):\n"
        "- A line starting with '## ' is a section heading (use short ones like "
        "'## Job Purpose', '## Key Responsibilities', '## Requirements').\n"
        "- A line starting with '- ' is a bullet point.\n"
        "- Any other non-empty line is a normal paragraph.\n"
        "- Separate sections/paragraphs/bullet-groups with a single blank line.\n\n"
        "Structure: 1 short 'Job Purpose' paragraph, then a 'Key Responsibilities' "
        "section with 4-7 bullets, then (if relevant) a short 'Requirements' "
        "section with 2-4 bullets. Do NOT use markdown bold/italic (**, *, _). "
        "Output ONLY the formatted text, nothing else."
    )

    try:
        resp = call_deepseek(ds_key,
            {'model': 'deepseek-chat', 'temperature': 0.3, 'max_tokens': 900,
                  'messages': [{'role': 'system', 'content': system_msg},
                                {'role': 'user', 'content': text[:12000]}]},
            timeout=60, endpoint='screening')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

    if resp.status_code != 200:
        try: err = resp.json().get('error', {}).get('message', 'DeepSeek API error')
        except Exception: err = resp.text[:200]
        return jsonify({'error': err}), 500

    shortened = resp.json()['choices'][0]['message']['content'].strip()
    return jsonify({'ok': True, 'shortened': shortened})



# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  REMINDERS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.route('/api/reminders', methods=['GET'])
@login_required
def get_reminders():
    conn = get_db()
    rows = conn.execute(
        'SELECT r.*, c.phone AS cand_phone, c.company AS cand_company, '
        'c.designation AS cand_designation, c.stage AS cand_stage '
        'FROM reminders r LEFT JOIN candidates c ON c.id = r.candidate_id '
        'WHERE r.done=0 AND r.owner_id=? ORDER BY r.due_at ASC',
        (effective_user_id(),)
    ).fetchall()
    conn.close()
    return jsonify({'ok': True, 'reminders': [dict(r) for r in rows]})

@app.route('/api/reminders', methods=['POST'])
@login_required
def add_reminder():
    d = request.json or {}
    cid   = d.get('candidate_id')
    note  = (d.get('note') or '').strip()
    due   = (d.get('due_at') or '').strip()
    stage = (d.get('stage') or 'todo').strip().lower()
    if stage not in ('todo', 'doing', 'done'):
        stage = 'todo'
    # A task needs either its own text (standalone) or a linked candidate.
    if not note and not cid:
        return jsonify({'error': 'note or candidate_id required'}), 400

    conn = get_db()
    cand_id = 0
    cand_name = ''
    mandate_id = None
    mandate_label = ''
    if cid:
        cand = conn.execute('SELECT * FROM candidates WHERE id=?', (cid,)).fetchone()
        if not cand:
            conn.close()
            return jsonify({'error': 'Candidate not found'}), 404
        cand_id = cid
        cand_name = cand['name'] or ''
        mandate_id = cand['mandate_id']
        mandate = conn.execute('SELECT * FROM mandates WHERE id=?', (cand['mandate_id'],)).fetchone()
        mandate_label = (mandate['role'] + ' — ' + mandate['client']) if mandate else ''

    done_flag = 1 if stage == 'done' else 0
    conn.execute(
        'INSERT INTO reminders (candidate_id,mandate_id,candidate_name,mandate_label,note,due_at,done,stage,created_at,owner_id) '
        'VALUES (?,?,?,?,?,?,?,?,?,?)',
        (cand_id, mandate_id, cand_name, mandate_label, note, due, done_flag, stage, ts(), effective_user_id())
    )
    conn.commit(); conn.close()
    return jsonify({'ok': True})

@app.route('/api/reminders/<int:rid>/stage', methods=['POST'])
@login_required
def set_reminder_stage(rid):
    """Move a task between Kanban columns (todo / doing / done)."""
    d = request.json or {}
    stage = (d.get('stage') or '').strip().lower()
    if stage not in ('todo', 'doing', 'done'):
        return jsonify({'error': 'invalid stage'}), 400
    conn = get_db()
    r = conn.execute('SELECT owner_id FROM reminders WHERE id=?', (rid,)).fetchone()
    if not r or (r['owner_id'] or 0) != effective_user_id():
        conn.close()
        return jsonify({'error': 'not found'}), 404
    done_flag = 1 if stage == 'done' else 0
    conn.execute('UPDATE reminders SET stage=?, done=? WHERE id=?', (stage, done_flag, rid))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/tasks/board', methods=['GET'])
@login_required
def tasks_board():
    """All manual tasks (reminders) grouped into Kanban columns. Includes done
    tasks so the Done column is populated. Standalone tasks have candidate_id=0."""
    conn = get_db()
    uid = effective_user_id()
    rows = conn.execute(
        "SELECT r.*, c.phone AS cand_phone, c.name AS cand_name2 "
        "FROM reminders r LEFT JOIN candidates c ON c.id = r.candidate_id "
        "WHERE r.owner_id=? ORDER BY r.id DESC", (uid,)
    ).fetchall()
    conn.close()
    cols = {'todo': [], 'doing': [], 'done': []}
    for r in rows:
        keys = r.keys()
        # Legacy rows: done=1 always lands in the Done column.
        st = 'done' if r['done'] else ((r['stage'] if 'stage' in keys else None) or 'todo')
        if st not in cols:
            st = 'todo'
        cols[st].append({
            'id': r['id'],
            'note': r['note'] or '',
            'candidate_id': r['candidate_id'] or 0,
            'candidate_name': (r['candidate_name'] or (r['cand_name2'] if 'cand_name2' in keys else '') or ''),
            'mandate_id': r['mandate_id'],
            'mandate_label': r['mandate_label'] or '',
            'phone': (r['cand_phone'] if 'cand_phone' in keys else '') or '',
            'due_at': r['due_at'] or '',
            'stage': st,
        })
    # Keep the Done column tidy — most-recent 60 only.
    cols['done'] = cols['done'][:60]
    return jsonify({'ok': True, 'columns': cols})


def _reminder_token(rid):
    """Unguessable per-reminder token so the mobile app can mark done/snooze
    WITHOUT a login session (which is what was failing → 'could not connect')."""
    key = app.secret_key or 'fallback'
    if isinstance(key, str):
        key = key.encode()
    return hmac.new(key, ('reminder:%d' % rid).encode(), hashlib.sha256).hexdigest()[:24]


def _reminder_action_ok(rid):
    """Authorized if the user has a session OR a valid reminder token."""
    try:
        if session.get('user_id'):
            return True
    except Exception:
        pass
    tok = request.values.get('token', '')
    return bool(tok) and tok == _reminder_token(rid)


@app.route('/api/reminders/<int:rid>/done', methods=['POST', 'GET'])
def mark_reminder_done(rid):
    # Works from webapp (session) AND mobile app (token) — marking done here
    # stops mobile notifications and clears it from the webapp list.
    if not _reminder_action_ok(rid):
        return jsonify({'error': 'unauthorized'}), 401
    conn = get_db()
    conn.execute('UPDATE reminders SET done=1 WHERE id=?', (rid,))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/reminders/<int:rid>/snooze', methods=['POST', 'GET'])
def snooze_reminder(rid):
    """Snooze a reminder. Body/query: minutes=10|30|60 or until=tomorrow.
    Accepts a session OR a valid reminder token (for the mobile app)."""
    if not _reminder_action_ok(rid):
        return jsonify({'error': 'unauthorized'}), 401
    d = request.get_json(silent=True) or request.values or {}
    now = _ist_now()
    until = None
    if d.get('until') == 'tomorrow' or d.get('minutes') == 'tomorrow':
        tomorrow = (now + datetime.timedelta(days=1)).replace(hour=9, minute=0, second=0, microsecond=0)
        until = tomorrow
    else:
        try:
            mins = int(d.get('minutes', 10))
        except Exception:
            mins = 10
        if mins not in (10, 30, 60):
            mins = 10
        until = now + datetime.timedelta(minutes=mins)
    until_iso = until.isoformat(timespec='seconds')
    conn = get_db()
    # Reset notification state so it fires fresh after the snooze window
    conn.execute(
        'UPDATE reminders SET snoozed_until=?, notified_at="", early_warned=0, notify_count=0 WHERE id=?',
        (until_iso, rid))
    conn.commit(); conn.close()
    return jsonify({'ok': True, 'snoozed_until': until_iso})


@app.route('/api/reminders/<int:rid>/edit', methods=['POST'])
@login_required
def edit_reminder(rid):
    """Edit a reminder's due time and/or note. Used by webapp and mobile."""
    d = request.json or {}
    conn = get_db()
    r = conn.execute('SELECT owner_id FROM reminders WHERE id=?', (rid,)).fetchone()
    if not r or r['owner_id'] != effective_user_id():
        conn.close()
        return jsonify({'error': 'Reminder not found'}), 404
    fields = []
    params = []
    if 'due_at' in d and (d.get('due_at') or '').strip():
        fields.append('due_at=?')
        params.append((d.get('due_at') or '').strip())
        # editing the time resets notification state so it fires fresh
        fields.append('notified_at=""')
        fields.append('early_warned=0')
        fields.append('snoozed_until=""')
    if 'note' in d:
        fields.append('note=?')
        params.append((d.get('note') or '').strip())
    if not fields:
        conn.close()
        return jsonify({'error': 'Nothing to update'}), 400
    params.append(rid)
    conn.execute(f'UPDATE reminders SET {", ".join(fields)} WHERE id=?', params)
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


@app.route('/api/reminders/<int:rid>', methods=['DELETE'])
@login_required
def delete_reminder(rid):
    conn = get_db()
    conn.execute('DELETE FROM reminders WHERE id=?', (rid,))
    conn.commit(); conn.close()
    return jsonify({'ok': True})

@app.route('/api/dashboard-tasks')
@login_required
def dashboard_tasks_alias():
    """Legacy alias — the dashboard/badge now reuse the unified task engine."""
    return get_tasks()

PROMISE_TAGS = ['Will Callback', 'Call Later', 'Asked to Send JD', 'Interested']

@app.route('/api/analytics')
@login_required
def analytics():
    """Dashboard analytics. Admin sees the whole company; a recruiter sees only
    their own assigned mandates + candidates. Pass ?scope=me to force self-view."""
    conn = get_db()
    cid = effective_company_id()
    scope_me = request.args.get('scope') == 'me'
    admin = is_company_admin() and not scope_me

    stale_days = float(get_setting('analytics_stale_days', '7') or 7)
    now = _ist_now()

    # ── Date range for time-based metrics (placements, added, time-to-fill, sources) ──
    range_from, range_to = None, None
    rng = request.args.get('range', '7')
    q_from = request.args.get('from', '')
    q_to = request.args.get('to', '')
    if q_from or q_to:
        try:
            if q_from: range_from = datetime.datetime.fromisoformat(q_from + 'T00:00:00')
        except Exception: range_from = None
        try:
            if q_to: range_to = datetime.datetime.fromisoformat(q_to + 'T23:59:59')
        except Exception: range_to = None
        range_label = 'Custom range'
    elif rng == 'all':
        range_label = 'All time'
    else:
        try:
            days = int(rng)
        except Exception:
            days = 7
        range_from = now - datetime.timedelta(days=days)
        range_label = f'Last {days} days'

    def _in_range(iso_str):
        if not iso_str:
            return False
        try:
            dt = datetime.datetime.fromisoformat(iso_str)
        except Exception:
            return False
        if range_from and dt < range_from:
            return False
        if range_to and dt > range_to:
            return False
        return True

    # Determine which mandate ids are in scope
    if admin:
        mrows = conn.execute('SELECT * FROM mandates WHERE owner_id=?', (cid,)).fetchall()
    else:
        mrows = conn.execute('SELECT * FROM mandates WHERE owner_id=? AND assigned_user_id=?',
                             (cid, real_user_id())).fetchall()
    mandates = [dict(m) for m in mrows]
    mandate_ids = [m['id'] for m in mandates]

    def _in(ids):
        return '(' + ','.join('?' for _ in ids) + ')' if ids else '(NULL)'

    # KPI: mandate status counts
    active = sum(1 for m in mandates if (m.get('status') or 'active') == 'active')
    hold = sum(1 for m in mandates if m.get('status') == 'hold')
    closed = sum(1 for m in mandates if m.get('status') == 'closed')

    # Candidates in scope
    if mandate_ids:
        cand_rows = conn.execute(
            f'SELECT id, stage, mandate_id, updated_at, ai_reasoning, created_at FROM candidates WHERE mandate_id IN {_in(mandate_ids)}',
            mandate_ids).fetchall()
    else:
        cand_rows = []
    cands = [dict(c) for c in cand_rows]
    total_pipeline = len([c for c in cands if c['stage'] not in ('Placed', 'Not Interested', 'Not Suitable', 'Client Rejected on Paper', 'Client Rejected After Interview')])

    # Total successful placements = Placed + Joined, ALL-TIME (deliberately NOT
    # date-filtered) so the headline matches the actual Placed list and the
    # Command Center. Admin view counts every candidate for the company (even
    # those not tied to a mandate); a recruiter sees only their in-scope ones.
    # TRIM/LOWER guards against stray whitespace/casing in imported data.
    if admin:
        placed_this_month = conn.execute(
            "SELECT COUNT(*) n FROM candidates WHERE owner_id=? AND " + PLACED_SQL, (cid,)).fetchone()['n']
    elif mandate_ids:
        placed_this_month = conn.execute(
            f"SELECT COUNT(*) n FROM candidates WHERE mandate_id IN {_in(mandate_ids)} AND " + PLACED_SQL,
            mandate_ids).fetchone()['n']
    else:
        placed_this_month = 0

    # Avg time-to-fill (days from candidate created → Placed) within range.
    # Joined candidates count too — they were placed, just moved one step further.
    fill_days = []
    for c in cands:
        if (c['stage'] or '').strip().lower() in ('placed', 'joined'):
            sh = conn.execute("SELECT created_at FROM stage_history WHERE candidate_id=? AND to_stage='Placed' ORDER BY created_at DESC LIMIT 1", (c['id'],)).fetchone()
            if sh and c['created_at']:
                try:
                    d0 = datetime.datetime.fromisoformat(c['created_at'])
                    d1 = datetime.datetime.fromisoformat(sh['created_at'])
                    if _in_range(sh['created_at']):
                        fill_days.append((d1 - d0).days)
                except Exception:
                    pass
    avg_ttf = round(sum(fill_days) / len(fill_days)) if fill_days else None

    # Pipeline funnel (grouped stages)
    funnel_map = [
        ('Screening', ['Screening']),
        ('Follow Up', ['Follow Up 1', 'Follow Up 2', 'Not Contacted', 'Called']),
        ('Interested', ['Interested', 'Updated CV awaited']),
        ('Shared with Client', ['Shared with Client']),
        ('Interview', ['Interview Inprocess']),
        ('Placed', ['Placed']),
    ]
    funnel = []
    for label, stages in funnel_map:
        funnel.append({'label': label, 'count': len([c for c in cands if c['stage'] in stages]), 'stages': stages})

    # Source effectiveness (of placed candidates, by source)
    def _source(reason):
        r = (reason or '').lower()
        if 'naukri' in r: return 'Naukri extension'
        if 'bulk' in r: return 'Bulk paste'
        if 'manual' in r: return 'Manual add'
        return 'Other'
    src_counts = {}
    for c in cands:
        if (c['stage'] or '').strip().lower() in ('placed', 'joined'):
            sh = conn.execute("SELECT created_at FROM stage_history WHERE candidate_id=? AND to_stage='Placed' ORDER BY created_at DESC LIMIT 1", (c['id'],)).fetchone()
            when = sh['created_at'] if sh else c['updated_at']
            if not _in_range(when):
                continue
            s = _source(c['ai_reasoning'])
            src_counts[s] = src_counts.get(s, 0) + 1
    total_placed = sum(src_counts.values())
    sources = [{'source': k, 'count': v, 'pct': round(v * 100 / total_placed) if total_placed else 0}
               for k, v in sorted(src_counts.items(), key=lambda x: -x[1])]

    # Recruiter leaderboard (admin only)
    leaderboard = []
    if admin:
        team = conn.execute("SELECT id, display_name, username FROM users WHERE company_id=? AND status='approved'", (cid,)).fetchall()
        for u in team:
            uid = u['id']
            u_mandates = conn.execute('SELECT id FROM mandates WHERE owner_id=? AND assigned_user_id=?', (cid, uid)).fetchall()
            u_mids = [r['id'] for r in u_mandates]
            added = placed = interviews = 0
            if u_mids:
                u_cands = conn.execute(f'SELECT id, stage, created_at, updated_at FROM candidates WHERE mandate_id IN {_in(u_mids)}', u_mids).fetchall()
                for uc in u_cands:
                    if _in_range(uc['created_at']):
                        added += 1
                    if (uc['stage'] or '').strip().lower() in ('placed', 'joined'):
                        sh = conn.execute("SELECT created_at FROM stage_history WHERE candidate_id=? AND to_stage='Placed' ORDER BY created_at DESC LIMIT 1", (uc['id'],)).fetchone()
                        if _in_range(sh['created_at'] if sh else uc['updated_at']):
                            placed += 1
                iv_rows = conn.execute(f'SELECT created_at FROM interviews WHERE mandate_id IN {_in(u_mids)}', u_mids).fetchall()
                interviews = sum(1 for r in iv_rows if _in_range(r['created_at']))
            leaderboard.append({'name': u['display_name'] or u['username'] or 'User',
                                'added': added, 'interviews': interviews, 'placed': placed})
        leaderboard.sort(key=lambda x: (-x['placed'], -x['added']))

    # Stale mandates (no candidate activity in stale_days) + stale candidates within
    stale_mandates = []
    for m in mandates:
        if (m.get('status') or 'active') != 'active':
            continue
        m_cands = [c for c in cands if c['mandate_id'] == m['id']]
        if not m_cands:
            continue
        latest = None
        for c in m_cands:
            for t in (c['updated_at'],):
                if t:
                    try:
                        dt = datetime.datetime.fromisoformat(t)
                        if not latest or dt > latest: latest = dt
                    except Exception:
                        pass
        if latest and (now - latest).days >= stale_days:
            stale_mandates.append({'id': m['id'], 'role': m['role'], 'client': m['client'],
                                   'days': (now - latest).days})
    stale_mandates.sort(key=lambda x: -x['days'])

    conn.close()
    return jsonify({'ok': True, 'is_admin_view': admin,
                    'kpi': {'open_mandates': len(mandates), 'active': active, 'hold': hold, 'closed': closed,
                            'placed_this_month': placed_this_month, 'avg_time_to_fill': avg_ttf,
                            'pipeline_candidates': total_pipeline},
                    'funnel': funnel, 'sources': sources, 'leaderboard': leaderboard,
                    'stale_mandates': stale_mandates, 'stale_days': int(stale_days),
                    'range_label': range_label})


@app.route('/api/analytics/stage-candidates')
@login_required
def analytics_stage_candidates():
    """All candidates in a given funnel-stage-group, across every in-scope mandate.
    Used when the user clicks a funnel bar (opens in a new tab via hash route)."""
    stages_param = request.args.get('stages', '')
    stages = [s for s in stages_param.split('||') if s]
    if not stages:
        return jsonify({'ok': True, 'candidates': []})
    conn = get_db()
    cid = effective_company_id()
    if is_company_admin():
        mrows = conn.execute('SELECT id, role, client FROM mandates WHERE owner_id=?', (cid,)).fetchall()
    else:
        mrows = conn.execute('SELECT id, role, client FROM mandates WHERE owner_id=? AND assigned_user_id=?',
                             (cid, real_user_id())).fetchall()
    mmap = {m['id']: dict(m) for m in mrows}
    mandate_ids = list(mmap.keys())
    if not mandate_ids:
        conn.close(); return jsonify({'ok': True, 'candidates': []})
    stale_days = float(get_setting('analytics_stale_days', '7') or 7)
    now = _ist_now()
    ph = '(' + ','.join('?' for _ in mandate_ids) + ')'
    sph = '(' + ','.join('?' for _ in stages) + ')'
    rows = conn.execute(
        f'SELECT id, name, company, designation, phone, email, stage, mandate_id, updated_at, cv_path '
        f'FROM candidates WHERE mandate_id IN {ph} AND stage IN {sph} ORDER BY name',
        mandate_ids + stages).fetchall()
    conn.close()
    out = []
    for r in rows:
        d = dict(r)
        m = mmap.get(d['mandate_id'], {})
        d['mandate_role'] = m.get('role', '')
        d['mandate_client'] = m.get('client', '')
        # stale flag
        d['is_stale'] = False
        if d['updated_at']:
            try:
                if (now - datetime.datetime.fromisoformat(d['updated_at'])).days >= stale_days:
                    d['is_stale'] = True
            except Exception:
                pass
        out.append(d)
    return jsonify({'ok': True, 'candidates': out, 'stages': stages, 'stale_days': int(stale_days)})

@app.route('/api/analytics/placed-list')
@login_required
def analytics_placed_list():
    """Every successful placement (Placed/Joined stage OR a real offer/joining
    record), across all positions. Powers the one-click drill-down from the
    Placements KPI, and uses the SAME predicate as the count so the list length
    always equals the headline number."""
    conn = get_db()
    cid = effective_company_id()
    sel = ("SELECT id,name,company,designation,phone,email,stage,mandate_id,updated_at,cv_path,"
           "offered_ctc,placement_fee,joining_date FROM candidates WHERE ")
    order = (" ORDER BY CASE TRIM(LOWER(COALESCE(stage,''))) "
             "WHEN 'joined' THEN 0 WHEN 'placed' THEN 1 ELSE 2 END, name")
    if is_company_admin():
        rows = conn.execute(sel + "owner_id=? AND " + PLACED_SQL + order, (cid,)).fetchall()
    else:
        mrows = conn.execute('SELECT id FROM mandates WHERE owner_id=? AND assigned_user_id=?',
                             (cid, real_user_id())).fetchall()
        mids = [m['id'] for m in mrows]
        if not mids:
            conn.close(); return jsonify({'ok': True, 'candidates': [], 'count': 0})
        ph = '(' + ','.join('?' for _ in mids) + ')'
        rows = conn.execute(sel + f"mandate_id IN {ph} AND " + PLACED_SQL + order, mids).fetchall()
    mmap = {}
    for m in conn.execute('SELECT id, role, client FROM mandates WHERE owner_id=?', (cid,)).fetchall():
        mmap[m['id']] = dict(m)
    conn.close()
    out = []
    for r in rows:
        d = dict(r)
        m = mmap.get(d.get('mandate_id'), {})
        d['mandate_role'] = m.get('role', '')
        d['mandate_client'] = m.get('client', '')
        stg = (d.get('stage') or '').strip().lower()
        d['placed_reason'] = 'joined' if stg == 'joined' else ('placed' if stg == 'placed' else 'offer')
        out.append(d)
    return jsonify({'ok': True, 'candidates': out, 'count': len(out)})

@app.route('/api/tasks')
@login_required
def get_tasks():
    """Unified follow-up task list for the dedicated Tasks tab.
    Sources: manual reminders, stale candidates (no activity in N days),
    promised follow-ups (a promise-tag with no activity after it for N hours),
    and new submissions. Each task carries a 'section' for Overdue/Today/
    Tomorrow/Upcoming grouping on the frontend."""
    conn = get_db()
    now = _ist_now()
    today = now.date()
    uid = effective_user_id()
    stale_days = float(get_setting('stale_days', '7') or 7)
    promise_hours = float(get_setting('promise_hours', '24') or 24)

    tasks = []

    def section_for(due_dt):
        if due_dt < now: return 'overdue'
        if due_dt.date() == today: return 'today'
        if due_dt.date() == today + datetime.timedelta(days=1): return 'tomorrow'
        return 'upcoming'

    # ── 1. Manual reminders ──────────────────────────────────────────────
    rem_rows = conn.execute(
        "SELECT r.*, c.phone AS cand_phone FROM reminders r "
        "LEFT JOIN mandates m ON m.id = r.mandate_id "
        "LEFT JOIN candidates c ON c.id = r.candidate_id "
        "WHERE r.done=0 AND r.owner_id=? AND (m.id IS NULL OR m.status NOT IN ('hold','closed')) "
        "ORDER BY r.due_at ASC", (uid,)
    ).fetchall()
    for r in rem_rows:
        try:
            due = datetime.datetime.fromisoformat(r['due_at'])
        except Exception:
            due = now
        _standalone = not r['candidate_id']
        if _standalone:
            _title = r['note'] or 'Task'
            _subtitle = 'Task'
        else:
            _title = r['candidate_name'] or 'Candidate'
            _subtitle = (r['note'] or 'Reminder') + (' \u00b7 ' + r['mandate_label'] if r['mandate_label'] else '')
        tasks.append({
            'id': 'reminder-' + str(r['id']), 'type': 'reminder', 'ref_id': r['id'],
            'candidate_id': r['candidate_id'], 'mandate_id': r['mandate_id'],
            'title': _title,
            'subtitle': _subtitle,
            'phone': (r['cand_phone'] if ('cand_phone' in r.keys()) else '') or '',
            'due_at': r['due_at'], 'section': section_for(due),
        })

    # ── Candidate base data for stale + promise detection ────────────────
    # Only surface candidates attached to a REAL, active position. This means:
    #   • the job's status must be 'active' (hold / closed jobs are skipped), and
    #   • the job must not be the "Central Database" pool — candidates from
    #     deleted jobs get parked there with an 'active' status, and we don't
    #     want that pool flooding the stale-follow-up list.
    central_mid = get_or_create_central_mandate()
    cand_rows = conn.execute(
        "SELECT c.id, c.name, c.phone, c.mandate_id, c.updated_at, c.task_snoozed_until, c.stage, "
        "m.role, m.client FROM candidates c LEFT JOIN mandates m ON m.id=c.mandate_id "
        "WHERE c.owner_id=? AND m.status='active' AND m.id != ?", (uid, central_mid)
    ).fetchall()

    # Stale follow-up only applies to candidates still in an active, meaningful
    # stage — dead/terminal-negative stages are not worth chasing.
    STALE_FOLLOWUP_STAGES = {
        'Screening', 'Follow Up 1', 'Follow Up 2', 'Not Contacted', 'Called',
        'Interested', 'Updated CV awaited', 'Shared with Client',
        'Interview Inprocess', 'Placed', 'Joined',
    }
    for c in cand_rows:
        snoozed = c['task_snoozed_until']
        if snoozed:
            try:
                if datetime.datetime.fromisoformat(snoozed) > now:
                    continue  # suppressed until this candidate's snooze passes
            except Exception:
                pass

        # Most recent event (any kind) for this candidate
        ev = conn.execute(
            "SELECT event_type, detail, created_at FROM candidate_events "
            "WHERE candidate_id=? ORDER BY created_at DESC LIMIT 1", (c['id'],)
        ).fetchone()
        stg = conn.execute(
            "SELECT created_at FROM stage_history WHERE candidate_id=? ORDER BY created_at DESC LIMIT 1", (c['id'],)
        ).fetchone()

        candidates_ts = [c['updated_at'] or '']
        if ev: candidates_ts.append(ev['created_at'] or '')
        if stg: candidates_ts.append(stg['created_at'] or '')
        candidates_ts = [t for t in candidates_ts if t]
        try:
            last_activity = max(datetime.datetime.fromisoformat(t) for t in candidates_ts) if candidates_ts else None
        except Exception:
            last_activity = None
        if not last_activity:
            continue

        mandate_label = (c['role'] + ' \u2014 ' + c['client']) if c['role'] else ''

        # ── 2. Promised follow-up: most recent event is a promise-tag, and
        #     nothing has happened since, for longer than promise_hours ──
        is_promise = False
        if ev and ev['event_type'] == 'tag' and any(pt in (ev['detail'] or '') for pt in PROMISE_TAGS):
            hrs_since = (now - last_activity).total_seconds() / 3600
            if hrs_since >= promise_hours:
                matched_tag = next((pt for pt in PROMISE_TAGS if pt in ev['detail']), '')
                tasks.append({
                    'id': 'promise-' + str(c['id']), 'type': 'promise', 'ref_id': c['id'],
                    'candidate_id': c['id'], 'mandate_id': c['mandate_id'],
                    'title': c['name'] or 'Candidate',
                    'subtitle': 'Tagged "' + matched_tag + '" \u2014 no follow-up yet' + (' \u00b7 ' + mandate_label if mandate_label else ''),
                    'phone': c['phone'] or '',
                    'due_at': last_activity.isoformat(), 'section': 'today',
                })
                is_promise = True

        # ── 3. Stale candidate: no activity at all for stale_days ────────
        #     Only surface candidates still in an ACTIVE, worth-following-up stage.
        #     Dead/terminal-negative stages (Not Interested, Not Suitable, Client
        #     Rejected, Screened-Out, etc.) are excluded — no point chasing them.
        if not is_promise and (c['stage'] in STALE_FOLLOWUP_STAGES):
            days_since = (now - last_activity).total_seconds() / 86400
            if days_since >= stale_days:
                tasks.append({
                    'id': 'stale-' + str(c['id']), 'type': 'stale', 'ref_id': c['id'],
                    'candidate_id': c['id'], 'mandate_id': c['mandate_id'],
                    'title': c['name'] or 'Candidate',
                    'subtitle': 'No activity in ' + str(int(days_since)) + ' days' + (' \u00b7 ' + mandate_label if mandate_label else ''),
                    'phone': c['phone'] or '',
                    'due_at': last_activity.isoformat(), 'section': 'today',
                })

    # ── Interview follow-ups: day-of confirmation + next-day result chase ──
    iv_rows = conn.execute(
        "SELECT i.*, c.name, c.phone, m.role, m.client FROM interviews i "
        "LEFT JOIN candidates c ON c.id=i.candidate_id "
        "LEFT JOIN mandates m ON m.id=i.mandate_id "
        "WHERE i.owner_id=? AND i.status='scheduled'", (uid,)
    ).fetchall()
    for iv in iv_rows:
        snz = iv['task_snoozed_until']
        if snz:
            try:
                if datetime.datetime.fromisoformat(snz) > now: continue
            except Exception: pass
        try:
            sch = datetime.datetime.fromisoformat(iv['scheduled_at'])
        except Exception:
            continue
        mandate_label = (iv['role'] + ' \u2014 ' + iv['client']) if iv['role'] else ''
        nice = sch.strftime('%d %b, %I:%M %p')
        if sch.date() == today:
            # Interview is today — confirm candidate will attend
            tasks.append({
                'id': 'iv-day-' + str(iv['id']), 'type': 'interview', 'ref_id': iv['id'],
                'candidate_id': iv['candidate_id'], 'mandate_id': iv['mandate_id'],
                'title': iv['name'] or 'Candidate',
                'subtitle': iv['round_name'] + ' today at ' + nice + ' \u2014 confirm attendance' + (' \u00b7 ' + mandate_label if mandate_label else ''),
                'phone': iv['phone'] or '',
                'due_at': iv['scheduled_at'], 'section': 'today',
            })
        elif sch.date() < today:
            # Interview date passed, still 'scheduled' — chase the result from client
            tasks.append({
                'id': 'iv-result-' + str(iv['id']), 'type': 'interview', 'ref_id': iv['id'],
                'candidate_id': iv['candidate_id'], 'mandate_id': iv['mandate_id'],
                'title': iv['name'] or 'Candidate',
                'subtitle': iv['round_name'] + ' done (' + nice + ') \u2014 get result/feedback from client' + (' \u00b7 ' + mandate_label if mandate_label else ''),
                'phone': iv['phone'] or '',
                'due_at': iv['scheduled_at'], 'section': 'overdue',
            })
        elif sch.date() == today + datetime.timedelta(days=1):
            tasks.append({
                'id': 'iv-tom-' + str(iv['id']), 'type': 'interview', 'ref_id': iv['id'],
                'candidate_id': iv['candidate_id'], 'mandate_id': iv['mandate_id'],
                'title': iv['name'] or 'Candidate',
                'subtitle': iv['round_name'] + ' tomorrow at ' + nice + (' \u00b7 ' + mandate_label if mandate_label else ''),
                'phone': iv['phone'] or '',
                'due_at': iv['scheduled_at'], 'section': 'tomorrow',
            })

    # ── 4. Candidates who submitted an updated profile via self-update link ──
    upd_rows = conn.execute(
        "SELECT c.id, c.name, c.phone, c.mandate_id, c.update_submitted_at, "
        "m.role, m.client FROM candidates c LEFT JOIN mandates m ON m.id=c.mandate_id "
        "WHERE c.owner_id=? AND c.update_submitted_at!='' "
        "AND (c.task_snoozed_until IS NULL OR c.task_snoozed_until='' OR c.task_snoozed_until<?)",
        (uid, now.isoformat())
    ).fetchall()
    for c in upd_rows:
        mandate_label = (c['role'] + ' \u2014 ' + c['client']) if c['role'] else ''
        tasks.append({
            'id': 'updated-' + str(c['id']), 'type': 'updated', 'ref_id': c['id'],
            'candidate_id': c['id'], 'mandate_id': c['mandate_id'],
            'title': c['name'] or 'Candidate',
            'subtitle': 'Submitted updated profile \u2014 review now' + (' \u00b7 ' + mandate_label if mandate_label else ''),
            'phone': c['phone'] or '',
            'due_at': c['update_submitted_at'], 'section': 'overdue',
        })

    # ── 5. New submissions (not yet reviewed, not snoozed) ────────────────
    sub_rows = conn.execute(
        "SELECT * FROM submissions WHERE status='new' "
        "AND (task_snoozed_until IS NULL OR task_snoozed_until='' OR task_snoozed_until<?) "
        "ORDER BY created_at DESC", (now.isoformat(),)
    ).fetchall()
    for s in sub_rows:
        tasks.append({
            'id': 'submission-' + str(s['id']), 'type': 'submission', 'ref_id': s['id'],
            'candidate_id': None, 'mandate_id': None,
            'title': s['name'] or 'New applicant',
            'subtitle': (s['company'] or '') + (' \u00b7 ' + str(s['experience']) + 'y' if s['experience'] else ''),
            'due_at': s['created_at'], 'section': 'today',
        })

    conn.close()

    order = {'overdue': 0, 'today': 1, 'tomorrow': 2, 'upcoming': 3}
    tasks.sort(key=lambda t: (order.get(t['section'], 9), t['due_at']))
    counts = {'overdue': 0, 'today': 0, 'tomorrow': 0, 'upcoming': 0}
    for t in tasks:
        counts[t['section']] = counts.get(t['section'], 0) + 1
    counts['total'] = len(tasks)
    return jsonify({'ok': True, 'tasks': tasks, 'counts': counts})


@app.route('/api/tasks/snooze', methods=['POST'])
@login_required
def snooze_task():
    d = request.json or {}
    ttype = d.get('type')
    ref_id = d.get('ref_id')
    snoozed_until = (d.get('snoozed_until') or '').strip()
    if not ttype or not ref_id or not snoozed_until:
        return jsonify({'error': 'type, ref_id and snoozed_until required'}), 400
    conn = get_db()
    if ttype == 'reminder':
        conn.execute('UPDATE reminders SET due_at=? WHERE id=?', (snoozed_until, ref_id))
    elif ttype in ('stale', 'promise'):
        conn.execute('UPDATE candidates SET task_snoozed_until=? WHERE id=?', (snoozed_until, ref_id))
    elif ttype == 'interview':
        conn.execute('UPDATE interviews SET task_snoozed_until=? WHERE id=?', (snoozed_until, ref_id))
    elif ttype == 'submission':
        conn.execute('UPDATE submissions SET task_snoozed_until=? WHERE id=?', (snoozed_until, ref_id))
    else:
        conn.close(); return jsonify({'error': 'Unknown task type'}), 400
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/tasks/done', methods=['POST'])
@login_required
def task_done():
    d = request.json or {}
    ttype = d.get('type')
    ref_id = d.get('ref_id')
    if not ttype or not ref_id:
        return jsonify({'error': 'type and ref_id required'}), 400
    conn = get_db()
    if ttype == 'reminder':
        conn.execute('UPDATE reminders SET done=1 WHERE id=?', (ref_id,))
    elif ttype in ('stale', 'promise'):
        # Push the suppression window out by the relevant threshold so it
        # naturally resurfaces later if still untouched, rather than being
        # silenced forever.
        days = float(get_setting('stale_days', '7') or 7) if ttype == 'stale' else 1
        push_to = (_ist_now() + datetime.timedelta(days=days)).isoformat()
        conn.execute('UPDATE candidates SET task_snoozed_until=? WHERE id=?', (push_to, ref_id))
    elif ttype == 'updated':
        conn.execute("UPDATE candidates SET update_submitted_at='' WHERE id=?", (ref_id,))
    elif ttype == 'interview':
        conn.execute("UPDATE interviews SET status='completed' WHERE id=?", (ref_id,))
    elif ttype == 'submission':
        conn.execute("UPDATE submissions SET status='reviewed' WHERE id=?", (ref_id,))
    else:
        conn.close(); return jsonify({'error': 'Unknown task type'}), 400
    conn.commit(); conn.close()
    return jsonify({'ok': True})


def _save_wh_for(conn, cid, items):
    """Replace work_history for a candidate from extension data."""
    if not isinstance(items, list) or not items:
        return
    conn.execute('DELETE FROM work_history WHERE candidate_id=?', (cid,))
    for i, it in enumerate(items):
        conn.execute(
            'INSERT INTO work_history (candidate_id,company,designation,start_date,end_date,is_current,description,sort_order) '
            'VALUES (?,?,?,?,?,?,?,?)',
            (cid, (it.get('company') or '').strip(), (it.get('designation') or '').strip(),
             (it.get('start_date') or '').strip(), (it.get('end_date') or '').strip(),
             1 if it.get('is_current') else 0, (it.get('description') or '').strip(), i)
        )

@app.route('/api/extension/push', methods=['POST', 'OPTIONS'])
def extension_push():
    """Receive a candidate pushed from the Naukri Chrome extension.
    - Requires phone OR email present (locked profiles without contact are rejected by the extension).
    - If a candidate with the same phone exists in the SAME mandate -> UPDATE it.
    - Otherwise INSERT a new candidate into the chosen mandate at 'Screening' stage.
    """
    if request.method == 'OPTIONS':
        return ('', 204)

    if not session.get('user_id'):
        return jsonify({'error': 'auth_required', 'message': 'Please log into HireLab in this browser first.'}), 401

    d = request.json or {}
    mid = d.get('mandate_id')
    name = (d.get('name') or '').strip()
    phone = (d.get('phone') or '').strip()
    email = (d.get('email') or '').strip()

    if not mid:
        return jsonify({'error': 'Please select a mandate'}), 400

    # Verify the mandate belongs to the current (effective) user.
    # Freelancers are allowed if the mandate is ASSIGNED to them.
    _conn = get_db()
    _own = _conn.execute('SELECT owner_id FROM mandates WHERE id=?', (mid,)).fetchone()
    _is_freelancer_upload = False
    _cu = current_user()
    if _cu and _cu.get('role') == 'freelancer_sourcer':
        _is_freelancer_upload = True
        try:
            from modules.freelancer import freelancer_can_access_mandate
            ok_access = freelancer_can_access_mandate(_conn, real_user_id(), int(mid), effective_company_id())
        except Exception:
            ok_access = False
        _conn.close()
        if not ok_access:
            return jsonify({'error': 'This mandate is not assigned to you'}), 403
    else:
        _conn.close()
        if not _own or _own['owner_id'] != effective_user_id():
            return jsonify({'error': 'That mandate is not in your workspace'}), 403
    if not name:
        return jsonify({'error': 'Candidate name missing'}), 400
    if not phone and not email:
        return jsonify({'error': 'Profile appears locked (no phone/email). Unlock it on Naukri first.'}), 400

    def fnum(v):
        try: return float(v or 0)
        except: return 0.0
    def inum(v):
        try: return int(float(v or 0))
        except: return 0

    skills = d.get('key_skills') or []
    if isinstance(skills, str):
        skills = [s.strip() for s in skills.split(',') if s.strip()]
    skills_json = json.dumps(skills)

    conn = get_db(); c = conn.cursor()

    # ── Duplicate detection by EMAIL ──────────────────────────────────────
    # Same email in the SAME mandate  -> UPDATE the CV/profile fields, but
    #   NEVER touch the stage, journey (stage_history) or comments.
    # Same email in a DIFFERENT mandate -> still create a NEW entry here
    #   (each mandate has its own pipeline); we just tell the user it exists
    #   elsewhere so they have context.
    # Duplicate check. For FREELANCERS this is a hard block (per spec):
    # if the candidate already exists on this mandate (phone OR email match),
    # reject the upload entirely.
    if _is_freelancer_upload:
        import re as _re
        pd = _re.sub(r'[^0-9]', '', phone or '')
        dup = None
        if pd and len(pd) >= 10:
            dup = c.execute(
                "SELECT id, name FROM candidates WHERE mandate_id=? AND "
                "REPLACE(REPLACE(REPLACE(phone,' ',''),'-',''),'+','') LIKE ?",
                (mid, '%' + pd[-10:])).fetchone()
        if not dup and email:
            dup = c.execute('SELECT id, name FROM candidates WHERE mandate_id=? AND LOWER(email)=LOWER(?)',
                            (mid, email)).fetchone()
        if not dup and name:
            dup = c.execute('SELECT id, name FROM candidates WHERE mandate_id=? AND LOWER(name)=LOWER(?)',
                            (mid, name)).fetchone()
        if dup:
            conn.close()
            return jsonify({'error': 'duplicate',
                            'message': 'This candidate is already sourced on this mandate.',
                            'existing_name': dup['name']}), 409

    existing = None
    other_mandates = []
    if email:
        existing = c.execute(
            'SELECT * FROM candidates WHERE mandate_id=? AND LOWER(email)=LOWER(?) LIMIT 1',
            (mid, email)
        ).fetchone()
        # Find this person in OTHER mandates owned by the same user (for info)
        rows = c.execute(
            'SELECT c.id, m.role, m.client FROM candidates c '
            'JOIN mandates m ON m.id = c.mandate_id '
            'WHERE LOWER(c.email)=LOWER(?) AND c.mandate_id!=? AND m.owner_id=?',
            (email, mid, effective_user_id())
        ).fetchall()
        other_mandates = [ (r['role'] + ' @ ' + r['client']) for r in rows ]

    if existing:
        # UPDATE profile/CV fields ONLY. Do NOT modify stage, stage_history,
        # recruiter_feedback, client_feedback, general_comments, wa_response.
        c.execute(
            'UPDATE candidates SET name=?,company=?,designation=?,experience=?,ctc_current=?,'
            'ctc_expected=?,notice_period=?,location=?,phone=?,key_skills=?,updated_at=? WHERE id=?',
            (name, d.get('company',''), d.get('designation',''), fnum(d.get('experience')),
             fnum(d.get('ctc_current')), fnum(d.get('ctc_expected')), inum(d.get('notice_period')),
             d.get('location',''), phone or existing['phone'], skills_json, ts(), existing['id'])
        )
        _save_wh_for(conn, existing['id'], d.get('work_history'))
        conn.execute('UPDATE candidates SET qualification=?, specialization=?, preferred_location=?, linkedin_url=?, ai_insight_cv=? WHERE id=?',
                     (d.get('qualification',''), d.get('specialization',''), d.get('preferred_location',''),
                      d.get('linkedin_url',''), d.get('ai_insight_cv',''), existing['id']))
        conn.commit(); conn.close()
        return jsonify({'ok': True, 'action': 'updated', 'candidate_id': existing['id'],
                        'name': name, 'preserved': True,
                        'message': 'CV & details updated. Stage, journey and comments preserved.'})

    c.execute(
        'INSERT INTO candidates (mandate_id,name,company,designation,experience,ctc_current,'
        'ctc_expected,notice_period,location,phone,email,career_summary,key_skills,'
        'screening_decision,ai_reasoning,stage,created_at,updated_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
        (mid, name, d.get('company',''), d.get('designation',''), fnum(d.get('experience')),
         fnum(d.get('ctc_current')), fnum(d.get('ctc_expected')), inum(d.get('notice_period')),
         d.get('location',''), phone, email, d.get('career_summary',''), skills_json,
         'worth_opening', 'Pushed from Naukri', 'Screening', ts(), ts())
    )
    cid = c.lastrowid
    c.execute('UPDATE candidates SET qualification=?, specialization=?, preferred_location=?, linkedin_url=?, ai_insight_cv=? WHERE id=?',
              (d.get('qualification',''), d.get('specialization',''), d.get('preferred_location',''),
               d.get('linkedin_url',''), d.get('ai_insight_cv',''), cid))
    # If a freelancer sourced this candidate, stamp attribution
    if _is_freelancer_upload:
        c.execute('UPDATE candidates SET sourced_by=?, sourced_at=? WHERE id=?',
                  (real_user_id(), ts(), cid))
    c.execute('INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) VALUES (?,?,?,?,?)',
              (cid, '', 'Screening', 'Pushed from Naukri extension', ts()))
    _save_wh_for(conn, cid, d.get('work_history'))
    conn.commit(); conn.close()
    queue_embedding_job(cid)  # async: enqueue, background worker embeds (never blocks add)
    resp = {'ok': True, 'action': 'added', 'candidate_id': cid, 'name': name}
    if other_mandates:
        resp['also_in'] = other_mandates
        resp['message'] = 'Added here. This person also exists in: ' + ', '.join(other_mandates)
    return jsonify(resp)

@app.route('/api/extension/mandates', methods=['GET', 'OPTIONS'])
def extension_mandates():
    """Lightweight mandate list for the extension dropdown (active only).
    Login-aware: shows only the logged-in user's own active mandates."""
    if request.method == 'OPTIONS':
        return ('', 204)
    if not session.get('user_id'):
        return jsonify({'error': 'auth_required', 'message': 'Please log into HireLab in this browser first.'}), 401
    conn = get_db()
    # Freelancers only see mandates ASSIGNED to them; recruiters/admin see their own.
    cu = current_user()
    if cu and cu.get('role') == 'freelancer_sourcer':
        rows = conn.execute(
            "SELECT m.id, m.role, m.client, m.location FROM mandate_freelancers mf "
            "JOIN mandates m ON m.id=mf.mandate_id "
            "WHERE mf.freelancer_user_id=? AND mf.company_id=? AND mf.is_active=1 "
            "AND m.status='active' ORDER BY m.created_at DESC",
            (real_user_id(), effective_company_id())
        ).fetchall()
    else:
        rows = conn.execute(
            "SELECT id, role, client, location FROM mandates WHERE status='active' AND owner_id=? ORDER BY created_at DESC",
            (effective_user_id(),)
        ).fetchall()
    conn.close()
    return jsonify({'ok': True, 'mandates': [dict(r) for r in rows]})

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  AI INSIGHTS — Semantic Search (embeddings) + Stats (SQL + LLM summary)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
import math

GEMINI_EMBED_URL = ('https://generativelanguage.googleapis.com/v1beta/'
                    'models/gemini-embedding-001:embedContent')

# Settings that are PER-COMPANY (each tenant has their own). Everything else
# (billing config, pricing, central ids, API keys) stays global.
TENANT_SETTINGS = {
    'recruiter_name', 'company_name',
    'submission_cc_emails', 'company_website', 'submission_signature',
    'seller_gstin', 'seller_address', 'seller_udyam', 'seller_state', 'seller_state_code',
    'seller_reg_office', 'invoice_signatory', 'invoice_hsn', 'invoice_fy', 'imap_host',
    'cc_bank_cash', 'cc_monthly_fixed', 'cc_team_size', 'cc_year_target',
    'cc_funding_available', 'cc_target_total', 'cc_target_years', 'cc_notes', 'cc_last_plan', 'cc_task_prefs', 'cc_last_review',
    'embedding_api_key', 'embedding_base_url', 'embedding_model', 'rag_enabled',
    'seller_gstin', 'seller_address', 'seller_udyam', 'seller_state', 'seller_state_code',
    'seller_reg_office', 'invoice_signatory', 'invoice_hsn', 'invoice_prefix', 'seller_name',
    'template_msg1', 'template_fu1', 'template_fu2',
    'fu1_hours', 'fu2_hours',
    'workflow_mode',   # 'agency' (default) or 'corporate'
    'pipeline_stages',  # JSON array of custom stages before the fixed Placed/Joined
    'smtp_email', 'smtp_app_password', 'smtp_display_name',
    'imap_enabled', 'imap_last_uid',
    'email_templates',  # JSON array of {name, subject, body}
    'custom_status_tags',  # JSON array of user-created quick tags
    'stale_days', 'promise_hours',  # follow-up task detection thresholds
    'analytics_stale_days',  # dashboard stale-mandate threshold (separate)
    'bd_stale_days',  # BD command center: flag clients silent this many days
    'interview_template',  # default interview communication message
    'wa_templates',        # JSON: categorized WhatsApp message templates (27 defaults)
    'wa_inbound_token',    # per-company secret for the WhatsApp listener to post inbound msgs
    'wa_style_profile',    # DeepSeek-learned summary of how this recruiter writes (for AI drafts)
    'wa_followup_hours',   # hours to wait before AI suggests a follow-up (default 24)
    'wa_auto_categories',  # JSON list of situation categories the agent may auto-send (future)
}

def _safe_company_id():
    try:
        return effective_company_id() or 0
    except Exception:
        return 0

def get_setting(key, default=''):
    # Env var takes priority for sensitive keys (see _ENV_KEY_MAP)
    env_name = _ENV_KEY_MAP.get(key)
    if env_name:
        env_val = os.environ.get(env_name, '').strip()
        if env_val:
            return env_val
    conn = get_db()
    # Per-tenant keys: prefer this company's own value, else fall back to the
    # global row (which acts as the default seed).
    if key in TENANT_SETTINGS:
        cid = _safe_company_id()
        if cid:
            tr = conn.execute('SELECT value FROM tenant_settings WHERE company_id=? AND key=?',
                              (cid, key)).fetchone()
            if tr is not None:
                conn.close()
                return (tr['value'] or '') or default
    row = conn.execute('SELECT value FROM settings WHERE key=?', (key,)).fetchone()
    conn.close()
    return (row['value'] if row else '') or default

def set_setting(key, value):
    """Write a setting. Per-tenant keys go to this company's own row; global
    keys go to the shared settings table."""
    conn = get_db()
    if key in TENANT_SETTINGS:
        cid = _safe_company_id()
        if cid:
            conn.execute('INSERT OR REPLACE INTO tenant_settings (company_id,key,value) VALUES (?,?,?)',
                         (cid, key, str(value)))
            conn.commit(); conn.close()
            return
    conn.execute('INSERT OR REPLACE INTO settings (key,value) VALUES (?,?)', (key, str(value)))
    conn.commit(); conn.close()

def gemini_embed(text, api_key):
    """Return a list[float] embedding for the given text via Gemini, or None."""
    text = (text or '').strip()
    if not text:
        return None
    try:
        resp = requests.post(
            GEMINI_EMBED_URL + '?key=' + api_key,
            headers={'Content-Type': 'application/json'},
            json={'model': 'models/gemini-embedding-001',
                  'content': {'parts': [{'text': text[:8000]}]}},
            timeout=30)
        if resp.status_code != 200:
            return {'error': resp.json().get('error', {}).get('message', resp.text[:200])}
        return resp.json()['embedding']['values']
    except Exception as e:
        return {'error': str(e)}

def embed_one(text):
    """Embed ONE text with the configured provider (Jina by default — via
    embedding_api_key / embedding_base_url / embedding_model). Returns a
    list[float], or {'error': ...} so callers can log/skip. This is the single
    switch point: candidate, facet, mandate and QUERY embeddings all go through
    here, so query and documents always share the SAME vector space."""
    text = (text or '').strip()
    if not text:
        return None
    try:
        res = _embed_texts([text])
    except Exception as e:
        return {'error': str(e)}
    if not res or not res[0]:
        return {'error': 'embedding provider returned nothing — check the Jina API key / embedding settings'}
    return res[0]


def _row_get(c, key, default=''):
    """Safely read a column from a sqlite3.Row. Older/partial rows may lack a
    column added by a later migration; this returns `default` instead of
    raising KeyError, so embedding text never crashes on legacy data."""
    try:
        if key in c.keys():
            v = c[key]
            return default if v is None else v
    except Exception:
        pass
    return default


def _as_list_str(raw):
    """Normalise a field that may be a JSON list, a comma string, or plain text
    into a clean 'a, b, c' string. Empty/blank items are dropped."""
    if raw is None:
        return ''
    if isinstance(raw, list):
        items = raw
    else:
        s = str(raw).strip()
        if not s:
            return ''
        if s.startswith('['):
            try:
                items = json.loads(s)
            except Exception:
                items = [s]
        else:
            items = [p for p in s.replace(';', ',').split(',')]
    seen, out = set(), []
    for it in items:
        t = str(it).strip()
        if t and t.lower() not in seen:
            seen.add(t.lower())
            out.append(t)
    return ', '.join(out)


# Bumped whenever candidate_embed_text() changes shape, so we can tell which
# vectors were built with the old text vs the new structured text and reindex
# selectively.
#   v1 = legacy 6-field concat
#   v2 = structured labelled profile + Career History + full CV resume text
EMBED_TEXT_VERSION = 2

# ── Embedding pipeline metadata (Feature 2) ─────────────────────────────
# Stamped onto every vector so we can migrate embedding models in future
# without breaking existing vectors. Nothing here is hardcoded into search;
# it is descriptive metadata only.
EMBEDDING_MODEL    = 'jina-embeddings-v3'             # the model embed_one() -> _embed_texts() calls
EMBEDDING_VERSION  = 'v1'                             # our embedding-pipeline version
EMBED_TEXT_TEMPLATE = f'candidate-template-v{EMBED_TEXT_VERSION}'  # which text builder produced it

import array as _array
def _vec_to_blob(vec):
    """Pack a float vector into compact float32 bytes for the embedding_vec
    cache. ~5.6x smaller than the JSON form and ~1000x faster to read back."""
    try:
        return _array.array('f', vec).tobytes()
    except Exception:
        return None

# How much raw resume text to fold into the embedding. The structured labelled
# fields (high signal) always come first, so if the total exceeds the Gemini
# input cap it is the resume TAIL that gets clipped, never the structured data.
CV_EMBED_MAX_CHARS = 5000


def _candidate_cv_text(c, max_chars=CV_EMBED_MAX_CHARS):
    """Return plain text extracted from the candidate's stored CV file (Word or
    PDF), or '' if there is no file / it can't be read. Reuses the existing
    extract_text_from_file() so there is ONE resume-parsing code path.

    This is what lets the embedding capture detail the parsed columns miss —
    projects, tools, software, certifications, languages — straight from the
    resume the recruiter uploaded. Best-effort: never raises."""
    try:
        rel = str(_row_get(c, 'cv_path')).strip()
        if not rel:
            return ''
        fp = os.path.join(CV_DIR, rel)
        if not os.path.exists(fp):
            return ''
        # Prefer the original filename's extension (tells us pdf vs docx);
        # fall back to the stored name.
        name = str(_row_get(c, 'cv_original_name')).strip() or rel
        with open(fp, 'rb') as fh:
            data = fh.read()
        text, err = extract_text_from_file(data, name)
        if err or not text:
            return ''
        text = ' '.join(text.split())  # collapse whitespace/newlines
        return text[:max_chars]
    except Exception:
        return ''


def candidate_embed_text(c, conn=None):
    """Build the text blob we embed for a candidate.

    Mandate-agnostic on purpose so a candidate can surface for ANY role they
    fit. The output is a *labelled, sectioned* profile rather than a blind
    concatenation: clear field names ("Current Role:", "Domain Expertise:")
    give the embedding model structure to latch onto, which measurably improves
    semantic match quality over a bare bag-of-words.

    `conn` is optional and backward compatible: when supplied, previous
    employers are pulled from the work_history table to capture career depth.
    Existing callers that pass only `c` keep working unchanged.
    """
    lines = []

    def add(label, value):
        v = (value or '').strip() if isinstance(value, str) else value
        if v:
            lines.append(f'{label}: {v}')

    name = str(_row_get(c, 'name')).strip()
    if name:
        lines.append(name)

    # ── Current position ────────────────────────────────────────────────
    desig = str(_row_get(c, 'designation')).strip()
    company = str(_row_get(c, 'company')).strip()
    if desig and company:
        add('Current Role', f'{desig} at {company}')
    else:
        add('Current Designation', desig)
        add('Current Company', company)

    exp = _row_get(c, 'experience', 0)
    try:
        if float(exp) > 0:
            add('Total Experience', f'{exp} years')
    except Exception:
        pass

    add('Career Summary', str(_row_get(c, 'career_summary')))

    # ── Previous companies (career depth) ───────────────────────────────
    if conn is not None:
        try:
            wh = conn.execute(
                'SELECT company, designation FROM work_history '
                'WHERE candidate_id=? ORDER BY is_current DESC, sort_order ASC, id ASC',
                (_row_get(c, 'id', 0),)).fetchall()
            prev = []
            for w in wh:
                co = (w['company'] or '').strip()
                dg = (w['designation'] or '').strip()
                if co and dg:
                    prev.append(f'{dg} at {co}')
                elif co or dg:
                    prev.append(co or dg)
            if prev:
                add('Career History', '; '.join(prev[:8]))
        except Exception:
            pass

    # ── Skills, domain, products, functional expertise ──────────────────
    add('Key Skills',       _as_list_str(_row_get(c, 'key_skills')))
    add('Technical Skills', _as_list_str(_row_get(c, 'key_skill_tags')))
    add('Secondary Skills', _as_list_str(_row_get(c, 'secondary_skills')))
    add('Domain Expertise', _as_list_str(_row_get(c, 'domain_tags')))
    add('Industry',         _as_list_str(_row_get(c, 'industry_background')))
    add('Products',         _as_list_str(_row_get(c, 'product_handles')))
    add('Functional Expertise', _as_list_str(_row_get(c, 'function_tags')))

    # ── Education & location ────────────────────────────────────────────
    add('Education', str(_row_get(c, 'qualification')))
    loc = str(_row_get(c, 'location')).strip()
    pref = str(_row_get(c, 'preferred_location')).strip()
    if loc and pref and pref.lower() != loc.lower():
        add('Location', f'{loc} (open to {pref})')
    else:
        add('Location', loc or pref)

    # ── Full resume text (projects / tools / certs / languages the parsed
    #    columns don't capture). Appended LAST so structured fields survive
    #    any truncation at the embedding input cap. ──────────────────────
    cv_text = _candidate_cv_text(c)
    if cv_text:
        lines.append('Resume:\n' + cv_text)

    return '\n'.join(lines)

def cosine(a, b):
    if not a or not b or len(a) != len(b):
        return 0.0
    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x * x for x in a))
    nb = math.sqrt(sum(y * y for y in b))
    if na == 0 or nb == 0:
        return 0.0
    return dot / (na * nb)


@app.route('/api/candidates/<int:cid>/rate', methods=['POST'])
@login_required
def rate_candidate(cid):
    """Rate a candidate against their mandate's JD using DeepSeek.
    Returns AI suitability % + selection probability % + reasoning."""
    ds_key = get_setting('deepseek_api_key')
    if not ds_key:
        return jsonify({'error': 'DeepSeek API key not set. Add it in Settings.'}), 400

    conn = get_db()
    c = conn.execute('SELECT * FROM candidates WHERE id=? AND owner_id=?',
                     (cid, effective_company_id())).fetchone()
    if not c:
        conn.close()
        return jsonify({'error': 'Candidate not found'}), 404
    m = conn.execute('SELECT * FROM mandates WHERE id=?', (c['mandate_id'],)).fetchone()
    if not m:
        conn.close()
        return jsonify({'error': 'Mandate not found'}), 404

    try:
        skills = json.loads(c['key_skills'] or '[]')
        if isinstance(skills, list): skills = ', '.join(str(s) for s in skills)
    except Exception:
        skills = ''

    jd_text = html_to_text(m['jd']) if m['jd'] else ''
    role_ctx = (f"Role: {m['role']} at {m['client']}\n"
                f"Location: {m['location']}\n"
                f"CTC band: {m['ctc_min']}-{m['ctc_max']} LPA\n"
                + (f"Job Description:\n{jd_text}" if jd_text.strip() else "Job Description: (not provided)"))

    # Pull hidden client notes for this mandate (recruiter's private intel from client)
    client_notes = ''
    try:
        note_rows = conn.execute(
            'SELECT note, created_at FROM mandate_client_notes WHERE mandate_id=? AND is_active=1 '
            'ORDER BY created_at ASC', (c['mandate_id'],)).fetchall()
        if note_rows:
            client_notes = '\n'.join('- ' + (r['note'] or '') for r in note_rows if (r['note'] or '').strip())
    except Exception:
        client_notes = ''
    if client_notes.strip():
        role_ctx += ("\n\nIMPORTANT — Private client requirements & preferences "
                     "(shared confidentially by the client; weigh these heavily):\n" + client_notes)

    cand_ctx = (f"Name: {c['name']}\n"
                f"Current: {c['designation']} at {c['company']}\n"
                f"Experience: {c['experience']} years\n"
                f"Location: {c['location']}\n"
                f"Current CTC: {c['ctc_current']} LPA, Expected: {c['ctc_expected']} LPA\n"
                f"Notice period: {c['notice_period']} days\n"
                f"Skills: {skills}\n"
                f"Summary: {c['career_summary'] or ''}")

    prompt = ("You are an expert recruiter evaluating how well a candidate fits a role. "
              "Score strictly and realistically.\n\n"
              "=== ROLE ===\n" + role_ctx + "\n\n=== CANDIDATE ===\n" + cand_ctx + "\n\n"
              "Return ONLY a JSON object (no markdown, no extra text) with exactly these keys:\n"
              '{"suitability": <0-100 integer: how well candidate matches the role requirements>, '
              '"selection_probability": <0-100 integer: realistic chance of being shortlisted by the client>, '
              '"reasoning": "<2-3 concise sentences: key strengths and gaps for THIS role>"}')

    try:
        rr = call_deepseek(ds_key,
            {'model': 'deepseek-chat', 'temperature': 0.2, 'max_tokens': 400,
                  'messages': [{'role': 'user', 'content': prompt}]},
            timeout=60, endpoint='reasoning')
        if rr.status_code != 200:
            err = rr.json().get('error', {}).get('message', rr.text[:200])
            conn.close()
            return jsonify({'error': 'DeepSeek error: ' + err}), 500
        raw = rr.json()['choices'][0]['message']['content'].strip()
        # Strip code fences if present
        raw = re.sub(r'^```(json)?|```$', '', raw, flags=re.MULTILINE).strip()
        data = json.loads(raw)
    except json.JSONDecodeError:
        conn.close()
        return jsonify({'error': 'Could not parse AI response. Try again.'}), 500
    except Exception as e:
        conn.close()
        return jsonify({'error': str(e)}), 500

    suit = max(0, min(100, int(data.get('suitability', 0))))
    prob = max(0, min(100, int(data.get('selection_probability', 0))))
    reasoning = '[Rated vs JD] ' + (data.get('reasoning', '') or '')

    conn.execute('UPDATE candidates SET ai_score=?, ai_reasoning=?, updated_at=? WHERE id=?',
                 (suit, reasoning, ts(), cid))
    conn.commit(); conn.close()
    return jsonify({'ok': True, 'suitability': suit, 'selection_probability': prob,
                    'reasoning': data.get('reasoning', '')})

# ── Deep AI Analysis (capability-first hiring-manager evaluation) ─────────────
# A much richer evaluation than the quick /rate score. Judges ONLY demonstrated
# capability from the resume vs the JD, and returns a full markdown report:
# match score, skill table, strengths/gaps, transferable skills, interview
# questions and a final verdict. Result is cached in candidates.deep_analysis.
DEEP_ANALYSIS_PROMPT = """# ROLE

You are a world-class Hiring Manager, Technical Interviewer, and Recruitment Expert with 20+ years of experience hiring top performers across Engineering, Manufacturing, Software, Sales, Finance, Supply Chain, Operations, Semiconductor, Automotive, Data Center, Power Systems, Industrial Automation, AI, Cloud, and Enterprise Technology.

Your responsibility is NOT to filter candidates based on years of experience, salary, company brand, education, location, age, gender, notice period, or resume formatting.

Your only objective is to determine:

"Can this person successfully perform this job?"

You must think exactly like an experienced hiring manager.

# PRIMARY EVALUATION PRINCIPLE

Judge candidates ONLY on demonstrated capability. Capability is determined from evidence of: Technical Skills, Functional Skills, Domain Knowledge, Tools Used, Technologies Worked On, Project Complexity, Responsibilities, Ownership, Achievements, Problem Solving, Decision Making, Leadership (if applicable), Scale of Work, Hands-on Experience, Business Impact, Learning Ability, Adaptability.

Ignore years of experience unless they explain skill maturity. Never assume that someone with 15 years is better than someone with 4 years. Always prefer evidence over duration.

# STRICTLY IGNORE

Do NOT increase or decrease score because of: Years of experience, Current CTC, Expected CTC, Location, Nationality, Gender, College, Degree, University ranking, Company Brand, Notice Period, Resume Design, Grammar, English Fluency, Resume Length, Job Hopping, Career Gap. These are NOT evaluation criteria.

# INPUTS

You will receive: 1. Complete Job Description  2. Candidate Resume

# STEP 1 — Understand the Job (internal, do not display)
Extract: Primary Technical Skills, Secondary Skills, Domain Knowledge, Mandatory Technologies, Optional Technologies, Responsibilities, Expected Deliverables, Seniority Level, Business Problems to Solve.

# STEP 2 — Analyse the Resume
Extract evidence of: Technical Skills, Functional Expertise, Projects, Tools, Platforms, Industries, Products, Responsibilities, Leadership, Achievements, Business Impact, Complexity, Scale, Hands-on Work, Ownership, Cross-functional exposure.

# STEP 3 — Skill Matching
For every required skill determine: Strong Evidence / Moderate Evidence / Weak Evidence / No Evidence. Evidence must always come from the resume. Never hallucinate.

# STEP 4 — Transferable Skills
If exact technology is missing but equivalent experience exists, identify transferable skills (e.g. Allen Bradley PLC <-> Siemens PLC <-> Mitsubishi PLC; AWS <-> Azure <-> GCP; Oracle <-> SQL Server <-> PostgreSQL; React <-> Angular <-> Vue; AutoCAD <-> SolidWorks). Do not reject candidates because tools differ if underlying competency is similar.

# STEP 5 — Identify Missing Skills
List only skills that appear important in the JD but have no evidence.

# STEP 6 — Confidence Analysis
Estimate confidence that candidate can perform the job. Use evidence only.

# OUTPUT FORMAT (use this exact markdown structure and section order)

## Candidate Match Score

Overall Match: XX / 100

Recommendation: Strong Hire | Hire | Borderline | Reject

## Skill Match Table

| Required Skill | Match Level | Resume Evidence |
|---------------|------------|----------------|

## Strengths
(bullet list)

## Skill Gaps
(bullet list)

## Transferable Skills
(bullet list)

## Business Impact
Explain why this candidate could create value.

## Risks
Mention only genuine skill-related risks. Never mention salary, location, notice period or years.

## Hiring Confidence
High | Medium | Low  — explain why.

## Interview Focus Areas
List the areas that require validation during interview.

## Suggested Interview Questions

Generate 8-12 highly targeted interview questions based ONLY on the missing evidence or weak areas identified above. Every question must validate a specific skill; avoid generic HR questions; prefer scenario-based, troubleshooting, architecture/design, hands-on and decision-making questions; progressively increase in difficulty.

For EACH question output EXACTLY this three-line format (nothing else):

**Q1. <the exact interview question, phrased so the recruiter can read it out verbatim to the candidate>**
*Skill tested:* <the one specific skill or gap this question validates — one line>
**Indicative Best Answer:** <Write the ACTUAL model answer a strong candidate would give — concrete and technical, 3-6 lines. Name the real tools, steps, trade-offs and numbers expected (specific commands, config values, design choices, metrics). This is the benchmark the recruiter compares the candidate's live answer against. Do NOT write "the answer should demonstrate..." — write the answer itself.>

(Then Q2, Q3, ... in the same format.)

## Final Hiring Verdict
Summarize in 5-8 lines: Can the candidate do the job? Why? What evidence supports this? What remains unverified? Should the recruiter move forward? Base the verdict ONLY on demonstrated capability. Never use years, salary, location, company reputation, or education as deciding factors.

# GOLDEN RULE
The goal is NOT the most experienced candidate — it is the most CAPABLE candidate. Evidence beats assumptions. Skills beat tenure. Capability beats pedigree."""


def _deep_analysis_inputs(conn, c, m):
    """Build the (jd_text, resume_text) pair fed to the deep-analysis model.
    Prefers the candidate's FULL uploaded resume; falls back to structured
    profile fields when no CV file is available so the tool still works."""
    jd_text = html_to_text(m['jd']) if m['jd'] else ''
    jd_block = (f"Role: {m['role']} at {m['client']}\n"
                f"Location: {m['location']}\n"
                f"CTC band: {m['ctc_min']}-{m['ctc_max']} LPA\n\n"
                + (jd_text.strip() if jd_text.strip()
                   else "(No full JD text on file — evaluate against the role title above.)"))

    resume_text = _candidate_cv_text(c, max_chars=12000)
    if not resume_text.strip():
        # Fallback: assemble a structured resume from parsed fields.
        try:
            skills = json.loads(c['key_skills'] or '[]')
            if isinstance(skills, list): skills = ', '.join(str(s) for s in skills)
        except Exception:
            skills = ''
        try:
            sec = json.loads(c['secondary_skills'] or '[]')
            if isinstance(sec, list): sec = ', '.join(str(s) for s in sec)
        except Exception:
            sec = ''
        wh_lines = []
        try:
            for w in conn.execute(
                'SELECT company, designation, start_date, end_date FROM work_history '
                'WHERE candidate_id=? ORDER BY id ASC', (c['id'],)).fetchall():
                span = ' '.join(x for x in [(w['start_date'] or ''), '-', (w['end_date'] or '')] if x).strip(' -')
                wh_lines.append(f"- {w['designation'] or ''} at {w['company'] or ''} ({span})".strip())
        except Exception:
            pass
        resume_text = (
            f"Name: {c['name']}\n"
            f"Current: {c['designation']} at {c['company']} ({c['experience']} yrs total)\n"
            f"Qualification: {c['qualification']}\n"
            f"Industry: {c['industry_background']}\n"
            f"Primary skills: {skills}\n"
            f"Secondary skills: {sec}\n"
            + ("Career history:\n" + '\n'.join(wh_lines) + '\n' if wh_lines else '')
            + f"Summary: {c['career_summary'] or ''}\n"
            "(Note: full resume file not uploaded — evaluation based on the structured profile above.)")
    return jd_block, resume_text


@app.route('/api/candidates/<int:cid>/deep-analysis', methods=['GET'])
@login_required
def get_deep_analysis(cid):
    """Return the cached deep-analysis report for this candidate (if any)."""
    conn = get_db()
    try:
        c = conn.execute('SELECT deep_analysis FROM candidates WHERE id=? AND owner_id=?',
                         (cid, effective_company_id())).fetchone()
    except Exception:
        # Column not created yet (first boot with this feature) — treat as "no analysis".
        conn.close()
        return jsonify({'ok': True, 'cached': False})
    conn.close()
    if not c:
        return jsonify({'error': 'Candidate not found'}), 404
    raw = (c['deep_analysis'] or '').strip()
    if not raw:
        return jsonify({'ok': True, 'cached': False})
    try:
        data = json.loads(raw)
        return jsonify({'ok': True, 'cached': True, 'md': data.get('md', ''),
                        'at': data.get('at', ''), 'model': data.get('model', '')})
    except Exception:
        return jsonify({'ok': True, 'cached': False})


@app.route('/api/candidates/<int:cid>/deep-analysis', methods=['POST'])
@login_required
def run_deep_analysis(cid):
    """Generate a full capability-first hiring-manager analysis of the candidate
    against their mandate's JD using DeepSeek, cache it, and return the markdown."""
    ds_key = get_setting('deepseek_api_key')
    if not ds_key:
        return jsonify({'error': 'DeepSeek API key not set. Add it in Settings.'}), 400

    conn = None
    try:
        conn = get_db()
        # Self-heal: guarantee the cache column exists even if the boot-time
        # migration was skipped for any reason (idempotent, no-op if present).
        try:
            conn.execute('ALTER TABLE candidates ADD COLUMN deep_analysis TEXT DEFAULT ""')
            conn.commit()
        except Exception:
            pass
        c = conn.execute('SELECT * FROM candidates WHERE id=? AND owner_id=?',
                         (cid, effective_company_id())).fetchone()
        if not c:
            return jsonify({'error': 'Candidate not found'}), 404
        m = conn.execute('SELECT * FROM mandates WHERE id=?', (c['mandate_id'],)).fetchone()
        if not m:
            return jsonify({'error': 'Mandate not found (candidate has no linked role).'}), 404

        jd_block, resume_text = _deep_analysis_inputs(conn, c, m)
        user_msg = ("JOB DESCRIPTION:\n" + jd_block +
                    "\n\n----------------------------------------\n\n"
                    "CANDIDATE RESUME:\n" + resume_text)

        try:
            rr = call_deepseek(ds_key,
                {'model': 'deepseek-chat', 'temperature': 0.3, 'max_tokens': 3500,
                 'messages': [{'role': 'system', 'content': DEEP_ANALYSIS_PROMPT},
                              {'role': 'user', 'content': user_msg}]},
                timeout=180, endpoint='deep-analysis')
        except TokenCapError:
            return jsonify({'error': 'Monthly AI token cap reached.'}), 429
        except requests.exceptions.Timeout:
            return jsonify({'error': 'DeepSeek timed out (>180s). Resume/JD may be very long — try again.'}), 504
        except Exception as e:
            import traceback; traceback.print_exc()
            return jsonify({'error': f'Could not reach DeepSeek — {type(e).__name__}: {e}'}), 502

        if rr.status_code != 200:
            try:
                err = rr.json().get('error', {}).get('message', rr.text[:300])
            except Exception:
                err = rr.text[:300]
            return jsonify({'error': f'DeepSeek returned {rr.status_code}: {err}'}), 502

        try:
            md = rr.json()['choices'][0]['message']['content'].strip()
        except Exception as e:
            return jsonify({'error': f'Unexpected DeepSeek response: {type(e).__name__}: {e}'}), 502

        if not md:
            return jsonify({'error': 'AI returned an empty response. Please try again.'}), 502

        at = ts()
        blob = json.dumps({'md': md, 'at': at, 'model': 'deepseek-chat'})
        conn.execute('UPDATE candidates SET deep_analysis=?, updated_at=? WHERE id=?',
                     (blob, at, cid))
        conn.commit()
        try:
            log_candidate_event(cid, 'note', 'AI deep analysis generated')
        except Exception:
            pass
        return jsonify({'ok': True, 'md': md, 'at': at, 'model': 'deepseek-chat'})

    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': f'{type(e).__name__}: {e}'}), 500
    finally:
        if conn is not None:
            try:
                conn.close()
            except Exception:
                pass


# ── Recruiter Pitch (JD → candidate-attraction pitch for recruiters) ─────────
# Turns a mandate's JD + role details into a ready-to-use recruiter pitch:
# a hook, real selling points, target profile, objection handling and a full
# spoken call script. Cached in mandates.recruiter_pitch.
RECRUITER_PITCH_PROMPT = """# ROLE
You are an experienced agency recruiter writing a PHONE CALL PITCH SCRIPT that a
recruiter reads out (naturally, not robotically) when cold-calling a candidate
about a job. Write it in the flowing, spoken, first-person style of the flow below.

You will receive the role details (title, client company, location, CTC band), a
job description, and the recruiter's name and agency name for the introduction.

# HOW TO WRITE IT
Follow this exact flow and tone (spoken, warm, professional). Keep
"[Candidate Name]" as a LITERAL placeholder — do NOT invent a candidate name.

1. Opening: "Hi, am I speaking with [Candidate Name]?"
2. Intro + permission: "Hi [Candidate Name], this is <Recruiter Name> from
   <Agency Name>. Is this a good time for a quick 2-minute conversation?"
3. Reason for the call: name the exact role and the CLIENT company, plus ONE
   short credible line about the client's standing in its industry. Use only
   well-known, publicly-true facts about recognisable companies; for lesser-known
   clients keep it neutral (e.g. "a well-regarded company in the <sector> space").
   NEVER invent ownership, funding, parent companies or awards.
4. Why this candidate: "I came across your profile on Naukri, and it looks like
   you have strong experience in <the 1-2 most relevant skills from the JD>,
   which is why I wanted to connect with you."
5. Role basics: mention the LOCATION, the experience range they want (from the JD
   or role seniority), and the core work plus the key technologies/skills from
   the JD.
6. Depth: one line that elevates it beyond a typical role — the real ownership,
   architecture, leadership or impact from the JD ("It's much more than a typical
   <role> role. You'll be responsible for ...").
7. Close with discovery: "I'd love to understand your current role, the
   technologies/tools you're working on, your current CTC, expected CTC, and the
   kind of exposure you're looking for in your next opportunity."

# RULES
- Plain spoken text only. NO headings, NO bullets, NO markdown symbols, NO notes
  — just the script as flowing lines, exactly like the template style.
- Keep it honest and grounded in the JD. Do not fabricate facts.
- Adapt EVERY detail (role, client, skills, tech, responsibilities, experience,
  location) to the ACTUAL inputs — never copy the example's specific values.
- Output ONLY the call script. Nothing before or after it."""


def _recruiter_pitch_input(m, recruiter_name='', company_name=''):
    """Assemble the role details + JD + intro info fed to the pitch model."""
    jd_text = html_to_text(m['jd']) if m['jd'] else ''
    return ("Role / Job Title: " + str(m['role'] or '') + "\n"
            "Client company: " + str(m['client'] or '') + "\n"
            "Location: " + str(m['location'] or '') + "\n"
            "CTC band: " + str(m['ctc_min'] or '') + "-" + str(m['ctc_max'] or '') + " LPA\n"
            "Recruiter name (for the intro): " + (recruiter_name or 'the recruiter') + "\n"
            "Agency name (for the intro): " + (company_name or '') + "\n\n"
            "JOB DESCRIPTION:\n" + (jd_text.strip() if jd_text.strip()
                else "(No full JD text on file — write the call script from the role "
                     "title, client and location above.)"))


@app.route('/api/mandates/<int:mid>/recruiter-pitch', methods=['GET'])
@login_required
def get_recruiter_pitch(mid):
    """Return the cached recruiter pitch for this mandate (if any)."""
    conn = get_db()
    try:
        m = conn.execute('SELECT recruiter_pitch FROM mandates WHERE id=? AND owner_id=?',
                         (mid, effective_company_id())).fetchone()
    except Exception:
        conn.close()
        return jsonify({'ok': True, 'cached': False})
    conn.close()
    if not m:
        return jsonify({'error': 'Mandate not found'}), 404
    raw = (m['recruiter_pitch'] or '').strip()
    if not raw:
        return jsonify({'ok': True, 'cached': False})
    try:
        data = json.loads(raw)
        return jsonify({'ok': True, 'cached': True, 'md': data.get('md', ''),
                        'at': data.get('at', ''), 'model': data.get('model', '')})
    except Exception:
        return jsonify({'ok': True, 'cached': False})


@app.route('/api/mandates/<int:mid>/recruiter-pitch', methods=['POST'])
@login_required
def run_recruiter_pitch(mid):
    """Generate a recruiter pitch from the mandate's JD + role details using
    DeepSeek, cache it on the mandate, and return the markdown."""
    ds_key = get_setting('deepseek_api_key')
    if not ds_key:
        return jsonify({'error': 'DeepSeek API key not set. Add it in Settings.'}), 400

    conn = None
    try:
        conn = get_db()
        # Self-heal: guarantee the cache column exists (idempotent).
        try:
            conn.execute('ALTER TABLE mandates ADD COLUMN recruiter_pitch TEXT DEFAULT ""')
            conn.commit()
        except Exception:
            pass
        m = conn.execute('SELECT * FROM mandates WHERE id=? AND owner_id=?',
                         (mid, effective_company_id())).fetchone()
        if not m:
            return jsonify({'error': 'Mandate not found'}), 404

        # Pull the recruiter's name + company from their signature/profile.
        u = current_user()
        recruiter_name = ((u.get('display_name') or u.get('username') or '') if u else '') \
                         or get_setting('recruiter_name', '') or ''
        company_name = get_setting('company_name', '') \
                       or ((u.get('company_name') or '') if u else '') or ''
        user_msg = _recruiter_pitch_input(m, recruiter_name, company_name)

        try:
            rr = call_deepseek(ds_key,
                {'model': 'deepseek-chat', 'temperature': 0.6, 'max_tokens': 700,
                 'messages': [{'role': 'system', 'content': RECRUITER_PITCH_PROMPT},
                              {'role': 'user', 'content': user_msg}]},
                timeout=180, endpoint='recruiter-pitch')
        except TokenCapError:
            return jsonify({'error': 'Monthly AI token cap reached.'}), 429
        except requests.exceptions.Timeout:
            return jsonify({'error': 'DeepSeek timed out (>180s). Try again.'}), 504
        except Exception as e:
            import traceback; traceback.print_exc()
            return jsonify({'error': f'Could not reach DeepSeek — {type(e).__name__}: {e}'}), 502

        if rr.status_code != 200:
            try:
                err = rr.json().get('error', {}).get('message', rr.text[:300])
            except Exception:
                err = rr.text[:300]
            return jsonify({'error': f'DeepSeek returned {rr.status_code}: {err}'}), 502

        try:
            md = rr.json()['choices'][0]['message']['content'].strip()
        except Exception as e:
            return jsonify({'error': f'Unexpected DeepSeek response: {type(e).__name__}: {e}'}), 502

        if not md:
            return jsonify({'error': 'AI returned an empty response. Please try again.'}), 502

        at = ts()
        blob = json.dumps({'md': md, 'at': at, 'model': 'deepseek-chat'})
        conn.execute('UPDATE mandates SET recruiter_pitch=? WHERE id=?', (blob, mid))
        conn.commit()
        return jsonify({'ok': True, 'md': md, 'at': at, 'model': 'deepseek-chat'})

    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': f'{type(e).__name__}: {e}'}), 500
    finally:
        if conn is not None:
            try:
                conn.close()
            except Exception:
                pass


def _record_embedding(conn, cid, status, *, vec=None, txt=None, error='', duration_ms=0):
    """Single authoritative writer for a candidate's embedding + all metadata.
    Centralised so every code path (single add, batch reindex, future model
    migration) stamps identical metadata. Commits itself.

    status:
      'completed' -> store the vector, text and full metadata
      'missing'   -> no embeddable text; mark embedding='[]' so it is not
                     re-selected as pending, but flag it missing
      'failed'    -> Gemini/API error; leave embedding UNTOUCHED (stays '' so
                     the existing pending query retries it) and record the error
    """
    dim = len(vec) if isinstance(vec, list) else 0
    if status == 'completed' and isinstance(vec, list) and vec:
        conn.execute(
            'UPDATE candidates SET embedding=?, embedding_vec=?, embedding_text=?, embedded_at=?, '
            'embedding_model=?, embedding_version=?, embedding_dimension=?, '
            'embedding_status=?, embedding_error=?, embedding_duration_ms=?, '
            'embedding_text_version=? WHERE id=?',
            (json.dumps(vec), _vec_to_blob(vec), txt, ts(), EMBEDDING_MODEL, EMBEDDING_VERSION, dim,
             'completed', '', int(duration_ms or 0), EMBED_TEXT_TEMPLATE, cid))
    elif status == 'missing':
        conn.execute(
            "UPDATE candidates SET embedding='[]', embedded_at=?, embedding_dimension=0, "
            "embedding_model=?, embedding_version=?, embedding_status='missing', "
            "embedding_error=?, embedding_duration_ms=?, embedding_text_version=? WHERE id=?",
            (ts(), EMBEDDING_MODEL, EMBEDDING_VERSION, (error or 'no embeddable text')[:500],
             int(duration_ms or 0), EMBED_TEXT_TEMPLATE, cid))
    else:  # 'failed' — DO NOT touch the embedding column, so it stays retryable
        conn.execute(
            "UPDATE candidates SET embedding_status='failed', embedding_error=?, "
            "embedding_duration_ms=?, embedding_model=?, embedding_version=?, "
            "embedding_text_version=? WHERE id=?",
            ((error or 'unknown error')[:500], int(duration_ms or 0), EMBEDDING_MODEL,
             EMBEDDING_VERSION, EMBED_TEXT_TEMPLATE, cid))
    conn.commit()


def embed_candidate_row(conn, c, api_key):
    """Build the embed text for candidate row `c`, call Gemini (timed), and
    store the vector + metadata via _record_embedding. Returns a small status
    dict. Never raises. This is the ONE embedding code path used by both the
    single-add flow and the batch reindex."""
    import time as _time
    cid = c['id']
    txt = candidate_embed_text(c, conn)
    if not txt.strip():
        _record_embedding(conn, cid, 'missing', error='no embeddable text')
        print(f'[embed] cid={cid} missing (no embeddable text)')
        return {'status': 'missing', 'cid': cid}

    print(f'[embed] cid={cid} started model={EMBEDDING_MODEL} chars={len(txt)}')
    t0 = _time.perf_counter()
    vec = embed_one(txt)
    dur = int((_time.perf_counter() - t0) * 1000)

    if isinstance(vec, dict) and vec.get('error'):
        _record_embedding(conn, cid, 'failed', txt=txt, error=vec['error'], duration_ms=dur)
        print(f'[embed] cid={cid} FAILED in {dur}ms: {str(vec["error"])[:140]}')
        return {'status': 'failed', 'cid': cid, 'error': vec['error'], 'duration_ms': dur}
    if not vec:
        _record_embedding(conn, cid, 'failed', txt=txt, error='empty vector', duration_ms=dur)
        print(f'[embed] cid={cid} FAILED in {dur}ms: empty vector')
        return {'status': 'failed', 'cid': cid, 'error': 'empty vector', 'duration_ms': dur}

    _record_embedding(conn, cid, 'completed', vec=vec, txt=txt, duration_ms=dur)
    print(f'[embed] cid={cid} completed dim={len(vec)} in {dur}ms model={EMBEDDING_MODEL}')
    return {'status': 'completed', 'cid': cid, 'dimension': len(vec), 'duration_ms': dur}


# ══════════════════════════════════════════════════════════════════════
#  ASYNC EMBEDDING QUEUE (Sprint 3)
#  A lightweight, persistent job queue on top of the existing SQLite +
#  daemon-thread architecture (same pattern as the reminder scheduler).
#  Goals: candidate creation never waits on Gemini; failed embeds retry
#  with exponential backoff; jobs survive restarts; API bursts are throttled.
# ══════════════════════════════════════════════════════════════════════

# Tunables — all overridable at runtime via Settings (get_setting), so you can
# adjust throughput/retries without a redeploy. Defaults are safe for Render.
_EMBED_CFG_DEFAULTS = {
    'embed_max_retries':     5,     # attempts before a job is marked failed
    'embed_batch_per_cycle': 10,    # jobs processed per worker wake (throughput)
    'embed_interval_ms':     1200,  # gap between Gemini calls (anti-burst / rate limit)
    'embed_poll_sec':        5,     # how often the worker wakes to look for jobs
    'embed_backoff_base_sec': 10,   # backoff = base * 2^(retry-1)
    'embed_backoff_cap_sec': 600,   # backoff never exceeds this
    'embed_reconcile_cap':   200,   # max un-embedded candidates auto-queued per sweep
}

def _embed_cfg(key):
    """Read a queue tunable from Settings, falling back to the default. Values
    are coerced to the default's type so callers always get an int."""
    default = _EMBED_CFG_DEFAULTS[key]
    try:
        v = get_setting(key)
        if v is None or str(v).strip() == '':
            return default
        return type(default)(v)
    except Exception:
        return default


# Error substrings that mean "do NOT retry" — auth/permission/validation are
# permanent; retrying only wastes quota. Everything else (timeout, network,
# rate limit / RESOURCE_EXHAUSTED, 5xx, INTERNAL, DEADLINE) is transient.
_EMBED_PERMANENT_MARKERS = (
    'API KEY', 'API_KEY', 'PERMISSION', 'UNAUTHENTICATED',
    'INVALID_ARGUMENT', 'NOT_FOUND', 'FAILED_PRECONDITION',
)

def _is_retryable_embed_error(err):
    e = str(err or '').upper()
    return not any(m in e for m in _EMBED_PERMANENT_MARKERS)


def _embed_backoff_seconds(retry_count):
    base = _embed_cfg('embed_backoff_base_sec')
    cap = _embed_cfg('embed_backoff_cap_sec')
    return min(cap, base * (2 ** max(0, retry_count - 1)))


def queue_embedding_job(cid, conn=None):
    """Enqueue an embedding job for a candidate and return immediately. This is
    what candidate-creation calls instead of embedding inline — it is a single
    cheap INSERT, so the upload/add response is instant.

    Idempotent: if the candidate already has an ACTIVE job (pending/processing/
    retrying) we don't create a duplicate. `conn` optional (reused if given)."""
    own = False
    try:
        if conn is None:
            conn = get_db(); own = True
        active = conn.execute(
            "SELECT id FROM embedding_jobs WHERE candidate_id=? "
            "AND status IN ('pending','processing','retrying') LIMIT 1", (cid,)).fetchone()
        if active:
            return active['id']
        conn.execute(
            "INSERT INTO embedding_jobs (candidate_id, status, retry_count, max_retries, "
            "created_at, next_attempt_at) VALUES (?, 'pending', 0, ?, ?, ?)",
            (cid, _embed_cfg('embed_max_retries'), ts(), ts()))
        conn.commit()
        jid = conn.execute('SELECT last_insert_rowid() AS id').fetchone()['id']
        print(f'[embed-queue] job {jid} created for cid={cid}')
        return jid
    except Exception as e:
        print(f'[embed-queue] enqueue error cid={cid}: {e}')
        return None
    finally:
        if own and conn is not None:
            try: conn.close()
            except Exception: pass


def _claim_next_job(conn):
    """Atomically claim the next due job. The conditional UPDATE (…WHERE status
    IN pending/retrying) means even if two workers race — or a future multi-
    worker deploy runs several copies — only one claims each job."""
    now = ts()
    row = conn.execute(
        "SELECT id FROM embedding_jobs WHERE status IN ('pending','retrying') "
        "AND (next_attempt_at IS NULL OR next_attempt_at='' OR next_attempt_at<=?) "
        "ORDER BY created_at ASC, id ASC LIMIT 1", (now,)).fetchone()
    if not row:
        return None
    cur = conn.execute(
        "UPDATE embedding_jobs SET status='processing', started_at=? "
        "WHERE id=? AND status IN ('pending','retrying')", (now, row['id']))
    conn.commit()
    if cur.rowcount != 1:
        return None  # lost the race; another claimer got it
    return conn.execute('SELECT * FROM embedding_jobs WHERE id=?', (row['id'],)).fetchone()


def _process_job(conn, job, api_key):
    """Run one claimed job through the shared embed path and update its status.
    Reuses embed_candidate_row (Sprint 2) — there is still ONE embedding code
    path; the queue only decides WHEN it runs and whether to retry."""
    jid = job['id']; cid = job['candidate_id']
    c = conn.execute('SELECT * FROM candidates WHERE id=?', (cid,)).fetchone()
    if not c:
        conn.execute("UPDATE embedding_jobs SET status='cancelled', last_error='candidate deleted', "
                     "completed_at=? WHERE id=?", (ts(), jid))
        conn.commit()
        print(f'[embed-queue] job {jid} cancelled (cid={cid} gone)')
        return 'cancelled'

    res = embed_candidate_row(conn, c, api_key)   # writes vector + metadata
    st = res.get('status')
    dur = int(res.get('duration_ms', 0) or 0)

    if st in ('completed', 'missing'):
        conn.execute("UPDATE embedding_jobs SET status='completed', duration_ms=?, "
                     "completed_at=?, last_error='' WHERE id=?", (dur, ts(), jid))
        conn.commit()
        return 'completed'

    # failed — decide retry vs give up
    err = res.get('error', 'unknown error')
    if _is_retryable_embed_error(err) and (job['retry_count'] + 1) <= job['max_retries']:
        rc = job['retry_count'] + 1
        nxt = (_ist_now() + datetime.timedelta(seconds=_embed_backoff_seconds(rc))).isoformat(timespec='seconds')
        conn.execute("UPDATE embedding_jobs SET status='retrying', retry_count=?, last_error=?, "
                     "next_attempt_at=?, duration_ms=? WHERE id=?", (rc, str(err)[:500], nxt, dur, jid))
        conn.commit()
        print(f'[embed-queue] job {jid} retry {rc}/{job["max_retries"]} in '
              f'{_embed_backoff_seconds(rc)}s ({str(err)[:80]})')
        return 'retrying'
    else:
        reason = 'permanent' if not _is_retryable_embed_error(err) else 'max retries'
        conn.execute("UPDATE embedding_jobs SET status='failed', last_error=?, completed_at=?, "
                     "duration_ms=? WHERE id=?", (str(err)[:500], ts(), dur, jid))
        conn.commit()
        print(f'[embed-queue] job {jid} FAILED ({reason}): {str(err)[:100]}')
        return 'failed'


def _backfill_embedding_blobs(conn, cap):
    """Convert legacy JSON embeddings into the float32 embedding_vec cache in the
    background (no API calls — pure local). Runs a bounded batch per idle sweep
    so existing candidates gradually gain the fast search path."""
    try:
        rows = conn.execute(
            "SELECT id, embedding FROM candidates "
            "WHERE embedding_vec IS NULL AND embedding IS NOT NULL "
            "AND embedding NOT IN ('', '[]') LIMIT ?", (cap,)).fetchall()
        n = 0
        for r in rows:
            try:
                vec = json.loads(r['embedding'])
            except Exception:
                continue
            blob = _vec_to_blob(vec)
            if blob is not None:
                conn.execute("UPDATE candidates SET embedding_vec=? WHERE id=?", (blob, r['id']))
                n += 1
        if n:
            conn.commit()
            print(f'[embed-blob] backfilled {n} float32 vector-cache row(s)')
        return n
    except Exception as e:
        print(f'[embed-blob] backfill error: {e}')
        return 0


# ══════════════════════════════════════════════════════════════════════
#  MULTI-VECTOR (FACET) EMBEDDINGS (Phase 2 / Sprint 8)
#  Separate vectors for what a candidate can DO (skills), where they've BEEN
#  (experience) and what they've BUILT (projects). Lets a focused query match
#  the right facet instead of a diluted whole-profile vector. Generation is
#  gated OFF by default so no extra Gemini cost until deliberately enabled;
#  search is unchanged unless a request opts in with multivector=true.
# ══════════════════════════════════════════════════════════════════════
MV_FACETS = ('skills', 'experience', 'projects')
MV_TEXT_VERSION = 1
MV_TEXT_TEMPLATE = f'facet-template-v{MV_TEXT_VERSION}'

_MV_CFG_DEFAULTS = {
    'multivector_enabled':   0,      # generate facet vectors in the background?
    'multivector_batch':     20,     # candidates per background sweep
    'multivector_w_full':    0.6,    # blend weight: whole-profile vector
    'multivector_w_facet':   0.4,    # blend weight: best-matching facet
}
def _mv_cfg(key):
    dflt = _MV_CFG_DEFAULTS[key]
    try:
        v = get_setting(key)
        if v is None or str(v).strip() == '':
            return dflt
        return float(v) if isinstance(dflt, float) else int(float(v))
    except Exception:
        return dflt


def candidate_facet_text(c, facet, conn=None):
    """Build the text for one facet. Reuses the same tolerant field helpers as
    the full-profile builder so there is one parsing path. Returns '' when the
    facet has no meaningful content (so we skip embedding empty facets)."""
    if facet == 'skills':
        parts = []
        for label, col in [('Skills', 'key_skills'), ('Technical Skills', 'key_skill_tags'),
                           ('Secondary Skills', 'secondary_skills'), ('Domain Expertise', 'domain_tags'),
                           ('Products', 'product_handles'), ('Functional Expertise', 'function_tags')]:
            v = _as_list_str(_row_get(c, col))
            if v:
                parts.append(f'{label}: {v}')
        desig = str(_row_get(c, 'designation')).strip()
        if desig:
            parts.insert(0, f'Role: {desig}')
        return '\n'.join(parts)

    if facet == 'experience':
        lines = []
        desig = str(_row_get(c, 'designation')).strip()
        company = str(_row_get(c, 'company')).strip()
        if desig or company:
            lines.append(f'Current Role: {desig} at {company}'.strip())
        exp = _row_get(c, 'experience', 0)
        try:
            if float(exp) > 0:
                lines.append(f'Total Experience: {exp} years')
        except Exception:
            pass
        summ = str(_row_get(c, 'career_summary')).strip()
        if summ:
            lines.append(f'Summary: {summ}')
        ind = _as_list_str(_row_get(c, 'industry_background'))
        if ind:
            lines.append(f'Industry: {ind}')
        if conn is not None:
            try:
                wh = conn.execute(
                    'SELECT company, designation FROM work_history WHERE candidate_id=? '
                    'ORDER BY is_current DESC, sort_order ASC, id ASC', (_row_get(c, 'id', 0),)).fetchall()
                prev = [f"{(w['designation'] or '').strip()} at {(w['company'] or '').strip()}".strip(' at')
                        for w in wh if (w['company'] or w['designation'])]
                prev = [p for p in prev if p]
                if prev:
                    lines.append('Career History: ' + '; '.join(prev[:8]))
            except Exception:
                pass
        return '\n'.join(lines)

    if facet == 'projects':
        chunks = []
        summ = str(_row_get(c, 'career_summary')).strip()
        if summ:
            chunks.append(summ)
        if conn is not None:
            try:
                for w in conn.execute('SELECT description FROM work_history WHERE candidate_id=?',
                                      (_row_get(c, 'id', 0),)).fetchall():
                    dsc = (w['description'] or '').strip()
                    if dsc:
                        chunks.append(dsc)
            except Exception:
                pass
        cv = _candidate_cv_text(c, max_chars=3000)
        if cv:
            chunks.append(cv)
        return '\n'.join(chunks).strip()

    return ''


def _store_facet(conn, cid, facet, vec, txt, status, error='', duration_ms=0):
    dim = len(vec) if isinstance(vec, list) else 0
    blob = _vec_to_blob(vec) if (status == 'completed' and vec) else None
    conn.execute(
        "INSERT INTO candidate_vectors (candidate_id, facet, embedding_vec, embedding_text, "
        "embedding_model, embedding_version, embedding_dimension, embedding_text_version, status, embedded_at) "
        "VALUES (?,?,?,?,?,?,?,?,?,?) "
        "ON CONFLICT(candidate_id, facet) DO UPDATE SET embedding_vec=excluded.embedding_vec, "
        "embedding_text=excluded.embedding_text, embedding_model=excluded.embedding_model, "
        "embedding_version=excluded.embedding_version, embedding_dimension=excluded.embedding_dimension, "
        "embedding_text_version=excluded.embedding_text_version, status=excluded.status, "
        "embedded_at=excluded.embedded_at",
        (cid, facet, blob, (txt or '')[:20000], EMBEDDING_MODEL, EMBEDDING_VERSION, dim,
         MV_TEXT_TEMPLATE, status, ts()))
    conn.commit()


def embed_candidate_facets(conn, c, api_key, facets=MV_FACETS):
    """Generate + store the facet vectors for one candidate. Returns a per-facet
    status dict. Never raises."""
    cid = c['id']; out = {}
    import time as _time
    for facet in facets:
        try:
            txt = candidate_facet_text(c, facet, conn)
            if not txt.strip():
                _store_facet(conn, cid, facet, None, '', 'empty')
                out[facet] = 'empty'; continue
            t0 = _time.perf_counter()
            vec = embed_one(txt)
            dur = int((_time.perf_counter() - t0) * 1000)
            if isinstance(vec, dict) and vec.get('error'):
                _store_facet(conn, cid, facet, None, txt, 'failed', error=vec['error'], duration_ms=dur)
                out[facet] = 'failed'; continue
            if not vec:
                _store_facet(conn, cid, facet, None, txt, 'failed', error='empty vector', duration_ms=dur)
                out[facet] = 'failed'; continue
            _store_facet(conn, cid, facet, vec, txt, 'completed', duration_ms=dur)
            out[facet] = 'completed'
        except Exception as e:
            out[facet] = f'error:{e}'
    return out


def _backfill_candidate_facets(conn, cap, api_key):
    """Background pass (gated by multivector_enabled): generate facet vectors for
    candidates that have a completed full embedding but are missing current-
    template facet rows. Bounded per sweep so it never floods the API."""
    if not _mv_cfg('multivector_enabled') or not api_key:
        return 0
    try:
        need = _mv_cfg('multivector_batch')
        rows = conn.execute(
            "SELECT c.* FROM candidates c WHERE c.embedding_status='completed' AND c.id NOT IN ("
            "  SELECT candidate_id FROM candidate_vectors WHERE embedding_text_version=? "
            "  GROUP BY candidate_id HAVING COUNT(DISTINCT facet) >= ?"
            ") LIMIT ?", (MV_TEXT_TEMPLATE, len(MV_FACETS), min(cap, need))).fetchall()
        n = 0
        for c in rows:
            embed_candidate_facets(conn, c, api_key)
            n += 1
        if n:
            print(f'[embed-facet] generated facet vectors for {n} candidate(s)')
        return n
    except Exception as e:
        print(f'[embed-facet] backfill error: {e}')
        return 0


# ══════════════════════════════════════════════════════════════════════
#  PERSISTENT JD (JOB-DESCRIPTION) EMBEDDINGS (Phase 2 / Sprint 9)
#  Embed each mandate's JD ONCE and reuse it, so candidate<->JD matching
#  never re-embeds the JD. Mandates are few, so this runs on by default.
# ══════════════════════════════════════════════════════════════════════
JD_TEXT_VERSION = 1
JD_TEXT_TEMPLATE = f'jd-template-v{JD_TEXT_VERSION}'

def _jd_cfg_enabled():
    try:
        v = get_setting('jd_embed_enabled')
        return True if (v is None or str(v).strip() == '') else bool(int(float(v)))
    except Exception:
        return True


def mandate_jd_text(m):
    """Build the embeddable text for a mandate/JD: role, client, division,
    location, CTC band and the full JD body. Labelled for the model."""
    lines = []
    def add(label, val):
        v = (val or '').strip() if isinstance(val, str) else val
        if v:
            lines.append(f'{label}: {v}')
    add('Role', str(_row_get(m, 'role')))
    add('Client', str(_row_get(m, 'client')))
    add('Division', str(_row_get(m, 'division')))
    add('Location', str(_row_get(m, 'location')))
    cmin, cmax = _row_get(m, 'ctc_min', ''), _row_get(m, 'ctc_max', '')
    if cmin or cmax:
        add('CTC Range', f'{cmin}-{cmax} LPA')
    jd = str(_row_get(m, 'jd')).strip()
    if jd:
        lines.append('Job Description:\n' + jd[:6000])
    return '\n'.join(lines)


def _store_mandate_vec(conn, mid, vec, txt, status, duration_ms=0):
    dim = len(vec) if isinstance(vec, list) else 0
    blob = _vec_to_blob(vec) if (status == 'completed' and vec) else None
    conn.execute(
        "INSERT INTO mandate_vectors (mandate_id, embedding_vec, embedding_text, embedding_model, "
        "embedding_version, embedding_dimension, embedding_text_version, status, embedded_at) "
        "VALUES (?,?,?,?,?,?,?,?,?) "
        "ON CONFLICT(mandate_id) DO UPDATE SET embedding_vec=excluded.embedding_vec, "
        "embedding_text=excluded.embedding_text, embedding_model=excluded.embedding_model, "
        "embedding_version=excluded.embedding_version, embedding_dimension=excluded.embedding_dimension, "
        "embedding_text_version=excluded.embedding_text_version, status=excluded.status, "
        "embedded_at=excluded.embedded_at",
        (mid, blob, (txt or '')[:20000], EMBEDDING_MODEL, EMBEDDING_VERSION, dim,
         JD_TEXT_TEMPLATE, status, ts()))
    conn.commit()


def embed_mandate_jd(conn, m, api_key):
    """Generate + store the JD vector for one mandate. Returns status string."""
    import time as _time
    mid = m['id']
    txt = mandate_jd_text(m)
    if not txt.strip():
        _store_mandate_vec(conn, mid, None, '', 'empty')
        return 'empty'
    t0 = _time.perf_counter()
    vec = embed_one(txt)
    dur = int((_time.perf_counter() - t0) * 1000)
    if isinstance(vec, dict) and vec.get('error'):
        _store_mandate_vec(conn, mid, None, txt, 'failed')
        print(f'[embed-jd] mandate {mid} FAILED: {str(vec["error"])[:100]}')
        return 'failed'
    if not vec:
        _store_mandate_vec(conn, mid, None, txt, 'failed')
        return 'failed'
    _store_mandate_vec(conn, mid, vec, txt, 'completed', dur)
    print(f'[embed-jd] mandate {mid} embedded dim={len(vec)} in {dur}ms')
    return 'completed'


def _backfill_mandate_jd(conn, cap, api_key):
    """Background pass: embed mandates that have no current JD vector. Enabled by
    default (mandates are few); gate with jd_embed_enabled=0 to disable."""
    if not _jd_cfg_enabled() or not api_key:
        return 0
    try:
        rows = conn.execute(
            "SELECT m.* FROM mandates m WHERE m.id NOT IN ("
            "  SELECT mandate_id FROM mandate_vectors WHERE status='completed' AND embedding_text_version=?"
            ") LIMIT ?", (JD_TEXT_TEMPLATE, cap)).fetchall()
        n = 0
        for m in rows:
            embed_mandate_jd(conn, m, api_key)
            n += 1
        if n:
            print(f'[embed-jd] embedded {n} mandate JD(s)')
        return n
    except Exception as e:
        print(f'[embed-jd] backfill error: {e}')
        return 0


def _mandate_jd_vector(conn, mid):
    """Load a mandate's stored JD vector (or None). No re-embedding."""
    try:
        r = conn.execute("SELECT embedding_vec FROM mandate_vectors WHERE mandate_id=? AND status='completed'",
                         (mid,)).fetchone()
        if not r or not r['embedding_vec']:
            return None
        if _HAS_NUMPY:
            return list(_np.frombuffer(r['embedding_vec'], dtype=_np.float32))
        a = _array.array('f'); a.frombytes(r['embedding_vec']); return list(a)
    except Exception:
        return None


def _reconcile_missing_embeddings(conn, cap):
    """Self-healing: enqueue candidates that have no embedding and no active OR
    failed job. This is what covers the non-Naukri creation paths (manual add,
    import, public apply) without touching each endpoint — and it re-queues
    nothing that already failed permanently, so it can't loop."""
    try:
        rows = conn.execute(
            "SELECT id FROM candidates WHERE (embedding='' OR embedding IS NULL) "
            "AND id NOT IN (SELECT candidate_id FROM embedding_jobs "
            "               WHERE status IN ('pending','processing','retrying','failed')) "
            "ORDER BY id LIMIT ?", (cap,)).fetchall()
        n = 0
        for r in rows:
            if queue_embedding_job(r['id'], conn):
                n += 1
        if n:
            print(f'[embed-queue] reconciler queued {n} un-embedded candidate(s)')
        return n
    except Exception as e:
        print(f'[embed-queue] reconcile error: {e}')
        return 0


def _embedding_worker_loop():
    """Background daemon: drains the embedding_jobs queue with retry, backoff
    and anti-burst throttling. Mirrors the reminder scheduler's shape."""
    import time as _time
    _time.sleep(20)  # let the app finish booting

    # RESTART RECOVERY: any job left 'processing' when the process died is
    # requeued so nothing is stranded mid-flight.
    try:
        conn = get_db()
        n = conn.execute("UPDATE embedding_jobs SET status='pending', started_at='' "
                         "WHERE status='processing'").rowcount
        conn.commit(); conn.close()
        if n:
            print(f'[embed-worker] recovered {n} in-flight job(s) after restart')
    except Exception as e:
        print(f'[embed-worker] recovery error: {e}')

    reconcile_every = 6  # run the self-heal sweep roughly every 6 idle-ish cycles
    cycle = 0
    while True:
        cycle += 1
        try:
            api_key = get_setting('embedding_api_key')
            if not api_key:
                _time.sleep(30)  # nothing to do without a key
                continue

            interval = max(0, _embed_cfg('embed_interval_ms')) / 1000.0
            conn = get_db()
            processed = 0
            for _ in range(_embed_cfg('embed_batch_per_cycle')):
                job = _claim_next_job(conn)
                if not job:
                    break
                try:
                    _process_job(conn, job, api_key)
                except Exception as je:
                    # never let one bad job kill the worker
                    try:
                        nxt = (_ist_now() + datetime.timedelta(seconds=30)).isoformat(timespec='seconds')
                        conn.execute("UPDATE embedding_jobs SET status='retrying', last_error=?, "
                                     "next_attempt_at=? WHERE id=?", (str(je)[:500], nxt, job['id']))
                        conn.commit()
                    except Exception:
                        pass
                    print(f'[embed-worker] job {job["id"]} crashed: {je}')
                processed += 1
                _time.sleep(interval)  # throttle between API calls

            # If the queue was empty this cycle, periodically sweep for
            # un-embedded candidates and backfill the float32 vector cache.
            if processed == 0 and (cycle % reconcile_every == 0):
                _reconcile_missing_embeddings(conn, _embed_cfg('embed_reconcile_cap'))
                _backfill_embedding_blobs(conn, _embed_cfg('embed_reconcile_cap'))
                # Multi-vector facet generation (Sprint 8) — no-op unless
                # multivector_enabled is turned on in Settings.
                _backfill_candidate_facets(conn, _embed_cfg('embed_reconcile_cap'), api_key)
                # Persistent JD embeddings (Sprint 9) — mandates are few; on by default.
                _backfill_mandate_jd(conn, _embed_cfg('embed_reconcile_cap'), api_key)

            conn.close()
        except Exception as e:
            print(f'[embed-worker] loop error: {e}')
        _time.sleep(_embed_cfg('embed_poll_sec'))


_embedding_worker_started = False
def _start_embedding_worker():
    global _embedding_worker_started
    if _embedding_worker_started:
        return
    _embedding_worker_started = True
    import threading
    t = threading.Thread(target=_embedding_worker_loop, daemon=True)
    t.start()
    print('[embed-worker] background thread started')


def embed_candidate_async(cid):
    """Backward-compat shim: previously this embedded inline (blocking). It now
    just ENQUEUES a background job so callers return instantly. Kept so any
    existing caller keeps working unchanged."""
    return queue_embedding_job(cid)


@app.route('/api/ai/index-status', methods=['GET'])
@login_required
def ai_index_status():
    conn = get_db()
    total = conn.execute('SELECT COUNT(*) n FROM candidates').fetchone()['n']
    done = conn.execute("SELECT COUNT(*) n FROM candidates WHERE embedding!='' AND embedding IS NOT NULL").fetchone()['n']
    # Optional metadata breakdown (additive — existing keys unchanged).
    by_status, by_model = {}, {}
    try:
        for r in conn.execute(
                "SELECT COALESCE(NULLIF(embedding_status,''),'unknown') s, COUNT(*) n "
                "FROM candidates GROUP BY s").fetchall():
            by_status[r['s']] = r['n']
        for r in conn.execute(
                "SELECT COALESCE(NULLIF(embedding_model,''),'none') m, COUNT(*) n "
                "FROM candidates GROUP BY m").fetchall():
            by_model[r['m']] = r['n']
    except sqlite3.OperationalError:
        pass  # metadata columns not present yet (pre-migration)
    conn.close()
    has_key = bool(get_setting('gemini_api_key'))
    return jsonify({'ok': True, 'total': total, 'indexed': done,
                    'pending': total - done, 'has_gemini_key': has_key,
                    'by_status': by_status, 'by_model': by_model,
                    'text_template': EMBED_TEXT_TEMPLATE, 'model': EMBEDDING_MODEL})


@app.route('/api/ai/queue/status', methods=['GET'])
@login_required
def ai_queue_status():
    """Live embedding-queue metrics for monitoring."""
    conn = get_db()
    counts = {'pending': 0, 'processing': 0, 'retrying': 0,
              'completed': 0, 'failed': 0, 'cancelled': 0}
    try:
        for r in conn.execute("SELECT status, COUNT(*) n FROM embedding_jobs GROUP BY status").fetchall():
            counts[r['status']] = r['n']
        avg_row = conn.execute("SELECT AVG(duration_ms) a FROM embedding_jobs "
                               "WHERE status='completed' AND duration_ms>0").fetchone()
        avg_ms = int(avg_row['a']) if avg_row and avg_row['a'] else 0
        retried = conn.execute("SELECT COUNT(*) n FROM embedding_jobs WHERE retry_count>0").fetchone()['n']
        oldest = conn.execute("SELECT MIN(created_at) c FROM embedding_jobs "
                              "WHERE status IN ('pending','retrying')").fetchone()['c'] or ''
    except sqlite3.OperationalError:
        avg_ms = retried = 0; oldest = ''
    conn.close()
    waiting = counts['pending'] + counts['retrying']
    return jsonify({'ok': True,
                    'queue_length': waiting + counts['processing'],
                    'waiting': waiting,
                    'running': counts['processing'],
                    'failed': counts['failed'],
                    'retried': retried,
                    'avg_processing_ms': avg_ms,
                    'oldest_waiting_at': oldest,
                    'by_status': counts,
                    'worker_running': _embedding_worker_started,
                    'config': {k: _embed_cfg(k) for k in _EMBED_CFG_DEFAULTS}})


@app.route('/api/ai/queue/enqueue-missing', methods=['POST'])
@login_required
def ai_queue_enqueue_missing():
    """Queue every un-embedded candidate that has no active/failed job. This is
    the entry point for large imports ('500 resumes -> queue all'). Non-blocking:
    it only creates job rows; the background worker does the embedding."""
    d = request.json or {}
    cap = int(d.get('cap') or 5000)
    conn = get_db()
    n = _reconcile_missing_embeddings(conn, cap)
    conn.close()
    return jsonify({'ok': True, 'queued': n})


@app.route('/api/ai/queue/retry-failed', methods=['POST'])
@login_required
def ai_queue_retry_failed():
    """Reset failed jobs back to pending (e.g. after fixing an API key). The
    worker will pick them up on its next cycle."""
    conn = get_db()
    try:
        n = conn.execute("UPDATE embedding_jobs SET status='pending', retry_count=0, "
                         "last_error='', next_attempt_at=? WHERE status='failed'", (ts(),)).rowcount
        conn.commit()
    except sqlite3.OperationalError:
        n = 0
    conn.close()
    return jsonify({'ok': True, 'requeued': n})


@app.route('/api/ai/reindex', methods=['POST'])
@login_required
def ai_reindex():
    """Embed a BATCH of un-embedded candidates per call (default 25). The
    frontend calls this repeatedly until pending=0, so a single HTTP request
    never runs long enough to time out, and progress is visible."""
    d = request.json or {}
    force = bool(d.get('force'))
    batch = int(d.get('batch') or 25)
    api_key = get_setting('embedding_api_key')
    if not api_key:
        return jsonify({'error': 'Embedding (Jina) API key not set. Add your Jina key in Settings.'}), 400

    conn = get_db()
    if force:
        rows = conn.execute('SELECT * FROM candidates ORDER BY id LIMIT ?', (batch,)).fetchall()
        # For force, also clear so they re-embed; but simplest: process those
        # without embedding first; force re-embeds everything across calls by
        # clearing embeddings up front on the first force call.
        if d.get('reset'):
            conn.execute("UPDATE candidates SET embedding='', embedding_status='pending'")
            conn.commit()
            rows = conn.execute('SELECT * FROM candidates ORDER BY id LIMIT ?', (batch,)).fetchall()
    rows = conn.execute("SELECT * FROM candidates WHERE embedding='' OR embedding IS NULL ORDER BY id LIMIT ?", (batch,)).fetchall()

    done, failed, skipped = 0, 0, 0
    first_error = ''
    for c in rows:
        r = embed_candidate_row(conn, c, api_key)
        st = r['status']
        if st == 'missing':
            skipped += 1
        elif st == 'failed':
            failed += 1
            first_error = r.get('error', '')
            eu = str(first_error).upper()
            # Auth/permission errors won't fix themselves within this batch —
            # abort so we don't burn the whole batch (unchanged behaviour).
            if 'API KEY' in eu or 'PERMISSION' in eu or 'API_KEY' in eu:
                conn.close()
                return jsonify({'error': 'Gemini error: ' + first_error,
                                'indexed': done, 'failed': failed}), 400
        else:
            done += 1

    # remaining count
    pending = conn.execute("SELECT COUNT(*) n FROM candidates WHERE embedding='' OR embedding IS NULL").fetchone()['n']
    conn.close()
    return jsonify({'ok': True, 'indexed': done, 'failed': failed, 'skipped': skipped,
                    'pending': pending, 'first_error': first_error})


# ══════════════════════════════════════════════════════════════════════
#  SEMANTIC SEARCH — performance layer (Sprint 4)
#  Streaming chunked scan + heap top-N + optional numpy + query cache.
#  No new infrastructure: pure in-process, SQLite-backed.
# ══════════════════════════════════════════════════════════════════════
import heapq as _heapq
import collections as _collections
import threading as _search_threading

try:
    import numpy as _np
    _HAS_NUMPY = True
except Exception:
    _HAS_NUMPY = False

_SEARCH_CFG_DEFAULTS = {
    'search_max_candidates': 100000,  # hard cap on candidates evaluated
    'search_max_results':    200,     # ranked pool kept for pagination
    'search_min_score':      -1.0,    # min cosine (-1..1) to include; -1 = keep all
    'search_chunk_size':     2000,    # streaming chunk size
    'search_slow_ms':        1500,    # log searches slower than this
    'search_cache_ttl_sec':  120,     # query cache TTL
    'search_cache_max':      64,      # max cached queries
}

def _search_cfg(key):
    dflt = _SEARCH_CFG_DEFAULTS[key]
    try:
        v = get_setting(key)
        if v is None or str(v).strip() == '':
            return dflt
        return float(v) if isinstance(dflt, float) else int(float(v))
    except Exception:
        return dflt

# In-memory query cache + rolling metrics (bounded; no Redis).
_SEARCH_CACHE = {}
_SEARCH_CACHE_LOCK = _search_threading.Lock()
_SEARCH_METRICS = _collections.deque(maxlen=500)      # (duration_ms, scanned)
_SEARCH_METRICS_LOCK = _search_threading.Lock()

def _search_cache_get(key):
    import time as _time
    with _SEARCH_CACHE_LOCK:
        e = _SEARCH_CACHE.get(key)
        if not e:
            return None
        if _time.time() - e['ts'] > _search_cfg('search_cache_ttl_sec'):
            _SEARCH_CACHE.pop(key, None)
            return None
        return e

def _search_cache_put(key, payload):
    import time as _time
    with _SEARCH_CACHE_LOCK:
        payload = dict(payload); payload['ts'] = _time.time()
        _SEARCH_CACHE[key] = payload
        # evict oldest beyond cap
        cap = _search_cfg('search_cache_max')
        if len(_SEARCH_CACHE) > cap:
            for k in sorted(_SEARCH_CACHE, key=lambda k: _SEARCH_CACHE[k]['ts'])[:len(_SEARCH_CACHE) - cap]:
                _SEARCH_CACHE.pop(k, None)

def _search_metric_record(duration_ms, scanned):
    with _SEARCH_METRICS_LOCK:
        _SEARCH_METRICS.append((duration_ms, scanned))


def _search_rank(conn, qvec, min_score, max_cands, pool, owner_id):
    """Stream candidate vectors in chunks, score against qvec, keep a top-`pool`
    heap. Returns (ranked, scanned) where ranked is a descending list of
    (similarity, candidate_id). Memory stays bounded regardless of table size
    because we never hold all vectors at once and keep only `pool` results."""
    d = len(qvec)
    chunk = _search_cfg('search_chunk_size')
    q_norm = math.sqrt(sum(x * x for x in qvec)) or 1e-9
    q_np = _np.asarray(qvec, dtype=_np.float32) if _HAS_NUMPY else None

    heap = []   # min-heap of (sim, id)
    scanned = 0

    def _push(sim, cid):
        if sim < min_score:
            return
        if len(heap) < pool:
            _heapq.heappush(heap, (sim, cid))
        elif sim > heap[0][0]:
            _heapq.heappushpop(heap, (sim, cid))

    cur = conn.execute(
        "SELECT id, embedding_vec, embedding FROM candidates "
        "WHERE embedding IS NOT NULL AND embedding NOT IN ('', '[]') "
        "AND owner_id=?", (owner_id,))
    while True:
        batch = cur.fetchmany(chunk)
        if not batch:
            break
        ids, mats = [], []
        for r in batch:
            if scanned >= max_cands:
                break
            scanned += 1
            vec = None
            blob = r['embedding_vec']
            if blob:                                   # fast path: float32 bytes
                try:
                    if _HAS_NUMPY:
                        vec = _np.frombuffer(blob, dtype=_np.float32)
                    else:
                        vec = _array.array('f'); vec.frombytes(blob)
                except Exception:
                    vec = None
            if vec is None:                            # fallback: legacy JSON
                try:
                    vec = json.loads(r['embedding'])
                except Exception:
                    continue
            if vec is None or len(vec) != d:
                continue
            if _HAS_NUMPY:
                ids.append(r['id']); mats.append(vec)
            else:
                dot = 0.0; nb = 0.0
                for x, y in zip(qvec, vec):
                    dot += x * y; nb += y * y
                nb = math.sqrt(nb)
                if nb:
                    _push(dot / (nb * q_norm), r['id'])
        if _HAS_NUMPY and ids:
            m = _np.asarray(mats, dtype=_np.float32)          # (n, d)
            dots = m @ q_np                                    # (n,)
            norms = _np.sqrt((m * m).sum(axis=1))
            norms[norms == 0] = 1e-9
            sims = dots / (norms * q_norm)
            for cid, sim in zip(ids, sims):
                _push(float(sim), cid)
            del m, dots, norms, sims                           # release chunk memory
        if scanned >= max_cands:
            break

    ranked = sorted(heap, key=lambda x: x[0], reverse=True)
    return ranked, scanned


def _search_hydrate(conn, page_slice, owner_id=None):
    """Load full display fields (with mandate join) ONLY for the ranked page.
    Preserves ranked order and attaches the score. Same result shape as before."""
    if not page_slice:
        return []
    score_map = {cid: sim for sim, cid in page_slice}
    ids = [cid for _, cid in page_slice]
    ph = ','.join('?' * len(ids))
    where = f"c.id IN ({ph})"
    params = list(ids)
    if owner_id is not None:                 # defense-in-depth tenant scope
        where += " AND c.owner_id=?"
        params.append(owner_id)
    rows = conn.execute(
        "SELECT c.id, c.name, c.designation, c.company, c.experience, c.location, "
        "c.ctc_current, c.ctc_expected, c.notice_period, c.phone, c.email, c.key_skills, "
        "c.mandate_id, c.stage, m.role AS mandate_role, m.client AS mandate_client, "
        "m.status AS mandate_status "
        "FROM candidates c LEFT JOIN mandates m ON m.id=c.mandate_id "
        f"WHERE {where}", params).fetchall()
    rowmap = {r['id']: r for r in rows}
    results = []
    for cid in ids:  # preserve ranked order
        c = rowmap.get(cid)
        if not c:
            continue
        try:
            skills = json.loads(c['key_skills'] or '[]')
        except Exception:
            skills = []
        results.append({
            'id': c['id'], 'name': c['name'], 'designation': c['designation'],
            'company': c['company'], 'experience': c['experience'],
            'location': c['location'], 'ctc_current': c['ctc_current'],
            'ctc_expected': c['ctc_expected'], 'notice_period': c['notice_period'],
            'phone': c['phone'], 'email': c['email'],
            'key_skills': skills,
            'mandate_id': c['mandate_id'], 'mandate_role': c['mandate_role'],
            'mandate_client': c['mandate_client'], 'mandate_status': c['mandate_status'],
            'stage': c['stage'],
            'score': round(score_map[cid] * 100, 1),
        })
    return results


# ══════════════════════════════════════════════════════════════════════
#  HYBRID SEARCH ENGINE (Sprint 6)
#  Layers structured understanding + business-rule scoring on top of the
#  Sprint-4 semantic retrieval. Design principle: NL-extracted signals only
#  RE-RANK (never exclude) so recall is preserved and a filter-less query
#  scores exactly like pure semantic search (perfect backward compatibility).
#  Hard exclusion happens only for explicitly-supplied mandatory filters.
# ══════════════════════════════════════════════════════════════════════

_HYBRID_CFG_DEFAULTS = {
    'hybrid_enabled':      1,      # default-on; harmless for filter-less queries
    'hybrid_pool':         300,    # semantic top-N to re-rank structurally
    'hybrid_w_semantic':   0.55,
    'hybrid_w_skills':     0.20,
    'hybrid_w_industry':   0.06,
    'hybrid_w_company':    0.06,
    'hybrid_w_experience': 0.07,
    'hybrid_w_location':   0.06,
    'hybrid_w_recency':    0.0,    # off by default
    'hybrid_taxo_ttl_sec': 600,    # taxonomy cache lifetime
    'hybrid_taxo_sample':  20000,  # candidates sampled to build the vocabulary
}
def _hybrid_cfg(key):
    dflt = _HYBRID_CFG_DEFAULTS[key]
    try:
        v = get_setting(key)
        if v is None or str(v).strip() == '':
            return dflt
        return float(v) if isinstance(dflt, float) else int(float(v))
    except Exception:
        return dflt

# Noise tokens dropped from company/location vocab so "electric", "india" etc.
# don't cause spurious matches.
_HYBRID_STOP = {'electric', 'electricals', 'ltd', 'limited', 'pvt', 'private', 'india',
                'technologies', 'technology', 'solutions', 'systems', 'system', 'group',
                'co', 'company', 'corp', 'inc', 'llp', 'the', 'and', 'engineering',
                'services', 'industries', 'international', 'global'}

def _lc_tokens(s):
    return [t for t in re.findall(r'[a-z0-9]+', (s or '').lower()) if len(t) > 1]

def _parse_tag_list(raw):
    """Parse a candidate tag column (JSON list / comma string) into a lowercased
    phrase set. Reuses the same tolerant parsing as the embedder."""
    txt = _as_list_str(raw)  # 'a, b, c'
    out = set()
    for p in txt.split(','):
        p = p.strip().lower()
        if p:
            out.add(p)
    return out


_HYBRID_TAXO = {'ts': 0.0, 'data': None}
def _hybrid_taxonomy(conn):
    """Vocabulary built from EXISTING candidate data (the 'existing taxonomy'
    the brief refers to) — skill phrases, company tokens, location tokens,
    industry phrases. Cached with TTL so it isn't rebuilt every search."""
    import time as _time
    now = _time.time()
    if _HYBRID_TAXO['data'] is not None and now - _HYBRID_TAXO['ts'] < _hybrid_cfg('hybrid_taxo_ttl_sec'):
        return _HYBRID_TAXO['data']
    skills, companies, locations, industries = set(), set(), set(), set()
    try:
        sample = _hybrid_cfg('hybrid_taxo_sample')
        rows = conn.execute(
            "SELECT key_skill_tags, key_skills, secondary_skills, domain_tags, product_handles, "
            "function_tags, company, industry_background, location, preferred_location "
            "FROM candidates LIMIT ?", (sample,)).fetchall()
        for r in rows:
            for col in ('key_skill_tags', 'key_skills', 'secondary_skills', 'domain_tags',
                        'product_handles', 'function_tags'):
                skills |= _parse_tag_list(r[col])
            for tok in _lc_tokens(r['company']):
                if tok not in _HYBRID_STOP:
                    companies.add(tok)
            for col in ('location', 'preferred_location'):
                for tok in _lc_tokens(r[col]):
                    if tok not in _HYBRID_STOP:
                        locations.add(tok)
            industries |= _parse_tag_list(r['industry_background'])
        # previous employers too
        for r in conn.execute("SELECT DISTINCT company FROM work_history LIMIT ?", (sample,)).fetchall():
            for tok in _lc_tokens(r['company']):
                if tok not in _HYBRID_STOP:
                    companies.add(tok)
    except sqlite3.OperationalError:
        pass
    skills = {s for s in skills if len(s) > 1 and not s.isdigit()}
    data = {'skills': skills, 'companies': companies, 'locations': locations, 'industries': industries}
    _HYBRID_TAXO['data'] = data; _HYBRID_TAXO['ts'] = now
    return data


def _query_grams(query):
    toks = re.findall(r'[a-z0-9]+', query.lower())
    grams = set(toks)
    for n in (2, 3):
        for i in range(len(toks) - n + 1):
            grams.add(' '.join(toks[i:i + n]))
    return grams, set(toks)


def _extract_query_filters(query, taxo):
    """Detect structured signals from a natural-language recruiter query using
    the existing taxonomy + light regex. All soft (re-rank) unless promoted to
    mandatory by the caller. O(query length) via n-gram intersection, so it
    scales regardless of vocabulary size."""
    grams, toks = _query_grams(query)
    ql = ' ' + query.lower() + ' '
    detected = {'skills': [], 'companies': [], 'locations': [], 'industries': [],
                'min_experience': None, 'max_experience': None, 'education': [],
                'immediate': False, 'salary_max': None, 'employment_type': None}

    detected['skills']     = sorted(g for g in grams if g in taxo['skills'])
    detected['industries'] = sorted(g for g in grams if g in taxo['industries'])
    detected['companies']  = sorted(t for t in toks if t in taxo['companies'])
    detected['locations']  = sorted(t for t in toks if t in taxo['locations'])

    mrange = re.search(r'(\d+)\s*(?:-|to)\s*(\d+)\s*(?:\+?\s*)?(?:years?|yrs?)', ql)
    if mrange:
        detected['min_experience'] = float(mrange.group(1)); detected['max_experience'] = float(mrange.group(2))
    else:
        mexp = re.search(r'(\d+)\s*\+?\s*(?:years?|yrs?)', ql)
        if mexp:
            detected['min_experience'] = float(mexp.group(1))

    for pat, label in [(r'\bb\.?tech\b', 'btech'), (r'\bm\.?tech\b', 'mtech'),
                       (r'\bb\.?e\b', 'be'), (r'\bmba\b', 'mba'),
                       (r'\bdiploma\b', 'diploma'), (r'\bphd\b', 'phd'), (r'\bb\.?sc\b', 'bsc')]:
        if re.search(pat, ql):
            detected['education'].append(label)
    if re.search(r'\bimmediate(ly)?\b|\bimmediate joiner\b', ql):
        detected['immediate'] = True
    msal = re.search(r'(\d+(?:\.\d+)?)\s*(?:lpa|lakhs?|lac|l)\b', ql)
    if msal:
        detected['salary_max'] = float(msal.group(1))
    if re.search(r'\bcontract\b|\bcontractual\b', ql):
        detected['employment_type'] = 'contract'
    elif re.search(r'\bpermanent\b|\bfull[- ]?time\b', ql):
        detected['employment_type'] = 'permanent'
    elif re.search(r'\bfreelanc', ql):
        detected['employment_type'] = 'freelance'
    return detected


def _candidate_signal_sets(c, wh_companies):
    """Build the comparable sets for one candidate row (skills phrases, company
    tokens, location tokens, industry phrases)."""
    skills = set()
    for col in ('key_skill_tags', 'key_skills', 'secondary_skills', 'domain_tags',
                'product_handles', 'function_tags'):
        try:
            skills |= _parse_tag_list(c[col])
        except Exception:
            pass
    comp_toks = set(t for t in _lc_tokens(c['company']) if t not in _HYBRID_STOP)
    for co in (wh_companies or []):
        comp_toks |= set(t for t in _lc_tokens(co) if t not in _HYBRID_STOP)
    loc_toks = set(_lc_tokens(c['location'])) | set(_lc_tokens(c['preferred_location']))
    industries = set()
    try:
        industries |= _parse_tag_list(c['industry_background'])
    except Exception:
        pass
    return skills, comp_toks, loc_toks, industries


def _hybrid_score_one(sem, c, wh_companies, detected, weights, mandatory_skills):
    """Return (hybrid_score 0..1, structured_score 0..1, explain). Only signals
    actually present in the query contribute — so a filter-less query collapses
    to hybrid == semantic (identical to the old engine)."""
    cskills, ccomp, cloc, cind = _candidate_signal_sets(c, wh_companies)
    signals = {'semantic': max(0.0, min(1.0, sem))}
    active = {'semantic': weights['semantic']}
    explain = {'matched_skills': [], 'matched_related': [], 'matched_industries': [], 'matched_companies': [],
               'matched_experience': None, 'missing_mandatory_skills': []}

    core = detected.get('skills') or []
    exp = detected.get('skill_expand') or {}
    if core or exp:
        core_matched = [s for s in core if s in cskills]
        # related/adjacent skills present on the candidate (from the skill graph),
        # not already counted as an exact match — each adds a fractional boost.
        rel_hits = {s: w for s, w in exp.items() if s in cskills and s not in core}
        denom = len(core) if core else 1
        raw = len(core_matched) + sum(rel_hits.values())
        signals['skills'] = max(0.0, min(1.0, raw / denom))
        active['skills'] = weights['skills']
        explain['matched_skills'] = core_matched
        explain['matched_related'] = sorted(rel_hits.keys())
    if detected['industries']:
        matched = [s for s in detected['industries'] if s in cind]
        signals['industry'] = 1.0 if matched else 0.0
        active['industry'] = weights['industry']; explain['matched_industries'] = matched
    if detected['companies']:
        matched = [s for s in detected['companies'] if s in ccomp]
        signals['company'] = 1.0 if matched else 0.0
        active['company'] = weights['company']; explain['matched_companies'] = matched
    if detected['locations']:
        matched = [s for s in detected['locations'] if s in cloc]
        signals['location'] = 1.0 if matched else 0.0
        active['location'] = weights['location']
        explain['matched_location'] = matched
    if detected['min_experience'] is not None:
        try:
            exp = float(c['experience'] or 0)
        except Exception:
            exp = 0.0
        need = detected['min_experience']
        ok = exp >= need
        signals['experience'] = 1.0 if ok else max(0.0, exp / need if need else 1.0)
        active['experience'] = weights['experience']
        explain['matched_experience'] = {'candidate': exp, 'required_min': need, 'ok': ok}

    tot = sum(active.values()) or 1.0
    hybrid = sum(signals[k] * active[k] for k in active) / tot
    struct_w = {k: v for k, v in active.items() if k != 'semantic'}
    structured = (sum(signals[k] * struct_w[k] for k in struct_w) / (sum(struct_w.values()) or 1.0)) if struct_w else 0.0

    if mandatory_skills:
        explain['missing_mandatory_skills'] = [s for s in mandatory_skills if s not in cskills]
    return hybrid, structured, explain


def _hybrid_process(conn, query, ranked, d):
    """Re-rank the semantic pool with structured signals. Returns
    (scored_results, filters_detected). Each scored result is a full display
    dict (same shape as _search_hydrate) plus 'semantic_score' and 'explain'."""
    weights = {
        'semantic':   _hybrid_cfg('hybrid_w_semantic'),
        'skills':     _hybrid_cfg('hybrid_w_skills'),
        'industry':   _hybrid_cfg('hybrid_w_industry'),
        'company':    _hybrid_cfg('hybrid_w_company'),
        'experience': _hybrid_cfg('hybrid_w_experience'),
        'location':   _hybrid_cfg('hybrid_w_location'),
        'recency':    _hybrid_cfg('hybrid_w_recency'),
    }
    taxo = _hybrid_taxonomy(conn)
    detected = _extract_query_filters(query, taxo)

    # Skill-graph expansion — reuse the one computed in ai_search when present.
    _sg = d.get('_skill_sg')
    if _sg is None:
        try:
            _sg = _skill_expand(query, conn, effective_company_id())
        except Exception:
            _sg = {'expand': {}, 'recognized': [], 'related': []}
    detected['skill_expand'] = _sg['expand']
    detected['skill_recognized'] = _sg['recognized']
    detected['skill_related'] = _sg['related']

    # merge explicit filters (opt-in, can be mandatory/hard)
    exp_filters = d.get('filters') or {}
    if exp_filters.get('skills'):
        detected['skills'] = sorted(set(detected['skills']) | {s.lower() for s in exp_filters['skills']})
    if exp_filters.get('min_experience') is not None:
        detected['min_experience'] = float(exp_filters['min_experience'])
    if exp_filters.get('location'):
        detected['locations'] = sorted(set(detected['locations']) | set(_lc_tokens(exp_filters['location'])))
    mandatory_skills = [s.lower() for s in (d.get('mandatory_skills') or [])]
    hard_location = _lc_tokens(exp_filters.get('location', '')) if exp_filters.get('mandatory_location') else []
    hard_min_exp = float(exp_filters['min_experience']) if (exp_filters.get('min_experience') is not None
                        and exp_filters.get('mandatory_experience')) else None

    pool_ids = [cid for _, cid in ranked[:_hybrid_cfg('hybrid_pool')]]
    sem_map = {cid: sim for sim, cid in ranked}
    if not pool_ids:
        return [], detected

    ph = ','.join('?' * len(pool_ids))
    rows = conn.execute(
        "SELECT c.id, c.name, c.designation, c.company, c.experience, c.location, c.preferred_location, "
        "c.ctc_current, c.ctc_expected, c.notice_period, c.phone, c.email, c.key_skills, "
        "c.secondary_skills, c.key_skill_tags, c.domain_tags, c.product_handles, c.function_tags, "
        "c.industry_background, c.mandate_id, c.stage, "
        "m.role AS mandate_role, m.client AS mandate_client, m.status AS mandate_status "
        "FROM candidates c LEFT JOIN mandates m ON m.id=c.mandate_id "
        f"WHERE c.id IN ({ph})", pool_ids).fetchall()
    rowmap = {r['id']: r for r in rows}
    # batch previous-employer companies for the pool
    wh_map = {}
    try:
        for r in conn.execute(f"SELECT candidate_id, company FROM work_history WHERE candidate_id IN ({ph})",
                              pool_ids).fetchall():
            wh_map.setdefault(r['candidate_id'], []).append(r['company'])
    except sqlite3.OperationalError:
        pass

    scored = []
    for cid in pool_ids:
        c = rowmap.get(cid)
        if c is None:
            continue
        # hard (mandatory) exclusion — opt-in only
        if mandatory_skills or hard_location or hard_min_exp is not None:
            cskills, ccomp, cloc, _ = _candidate_signal_sets(c, wh_map.get(cid))
            if mandatory_skills and any(s not in cskills for s in mandatory_skills):
                continue
            if hard_location and not (set(hard_location) & cloc):
                continue
            if hard_min_exp is not None:
                try:
                    if float(c['experience'] or 0) < hard_min_exp:
                        continue
                except Exception:
                    continue
        sem = sem_map.get(cid, 0.0)
        hyb, structured, explain = _hybrid_score_one(sem, c, wh_map.get(cid), detected, weights, mandatory_skills)
        try:
            skills = json.loads(c['key_skills'] or '[]')
        except Exception:
            skills = []
        scored.append({
            'id': c['id'], 'name': c['name'], 'designation': c['designation'],
            'company': c['company'], 'experience': c['experience'], 'location': c['location'],
            'ctc_current': c['ctc_current'], 'ctc_expected': c['ctc_expected'],
            'notice_period': c['notice_period'], 'phone': c['phone'], 'email': c['email'],
            'key_skills': skills, 'mandate_id': c['mandate_id'], 'mandate_role': c['mandate_role'],
            'mandate_client': c['mandate_client'], 'mandate_status': c['mandate_status'],
            'stage': c['stage'],
            'score': round(hyb * 100, 1),
            'semantic_score': round(sem * 100, 1),
            'structured_score': round(structured * 100, 1),
            'explain': explain,
            '_sem': sem,
        })
    # sort by hybrid score desc, tie-break on semantic
    scored.sort(key=lambda x: (x['score'], x['_sem']), reverse=True)
    for x in scored:
        x.pop('_sem', None)
    return scored, detected


def _rme_capture_search(query, detected, result_count, user):
    """Log the search into the Recruitment Memory Engine as a search_activity
    event (Sprint 5). Best-effort, never blocks or breaks search."""
    try:
        conn = get_db()
        rme_add_event(conn, 'search_activity', 'recruiter',
                      str((user or {}).get('id', 'system')),
                      actor=str((user or {}).get('username', 'system')),
                      summary=query[:200],
                      data={'skills': detected.get('skills', []),
                            'companies': detected.get('companies', []),
                            'locations': detected.get('locations', []),
                            'industries': detected.get('industries', []),
                            'min_experience': detected.get('min_experience'),
                            'results': result_count})
        conn.close()
    except Exception:
        pass


# ── Multi-vector (facet) re-rank (Sprint 8) ────────────────────────────
def _facet_vectors_for(conn, ids):
    """Bulk-load completed facet vectors for candidate ids.
    Returns {candidate_id: [vector, ...]}."""
    if not ids:
        return {}
    ph = ','.join('?' * len(ids))
    out = {}
    try:
        for r in conn.execute(
                f"SELECT candidate_id, embedding_vec FROM candidate_vectors "
                f"WHERE status='completed' AND embedding_vec IS NOT NULL AND candidate_id IN ({ph})",
                ids).fetchall():
            blob = r['embedding_vec']
            if not blob:
                continue
            try:
                if _HAS_NUMPY:
                    v = _np.frombuffer(blob, dtype=_np.float32)
                else:
                    v = _array.array('f'); v.frombytes(blob)
            except Exception:
                continue
            out.setdefault(r['candidate_id'], []).append(v)
    except sqlite3.OperationalError:
        pass
    return out


def _multivector_rerank(conn, qvec, ranked):
    """Opt-in: blend the whole-profile similarity with the BEST-matching facet
    (skills/experience/projects) so a focused query surfaces a candidate whose
    relevant facet matches strongly even if their full-profile vector is diluted
    by unrelated resume text. Only the semantic pool is touched; candidates with
    no facet vectors keep their original score. Returns a re-sorted (sim, id) list."""
    pool_n = _hybrid_cfg('hybrid_pool')
    pool = ranked[:pool_n]
    ids = [cid for _, cid in pool]
    fmap = _facet_vectors_for(conn, ids)
    if not fmap:
        return ranked  # no facet vectors yet -> unchanged (backward compatible)
    wf = _mv_cfg('multivector_w_full'); wx = _mv_cfg('multivector_w_facet')
    d = len(qvec)
    qn = math.sqrt(sum(x * x for x in qvec)) or 1e-9
    q_np = _np.asarray(qvec, dtype=_np.float32) if _HAS_NUMPY else None

    def _cos(v):
        if len(v) != d:
            return 0.0
        if _HAS_NUMPY:
            nb = float(_np.linalg.norm(v)) or 1e-9
            return float(_np.dot(q_np, v)) / (nb * qn)
        dot = 0.0; nb = 0.0
        for x, y in zip(qvec, v):
            dot += x * y; nb += y * y
        nb = math.sqrt(nb) or 1e-9
        return dot / (nb * qn)

    boosted = []
    for sim, cid in pool:
        facets = fmap.get(cid)
        if facets:
            best = max(_cos(v) for v in facets)
            boosted.append((wf * sim + wx * best, cid))
        else:
            boosted.append((sim, cid))
    boosted.sort(key=lambda x: x[0], reverse=True)
    return boosted + ranked[pool_n:]  # re-ranked pool + untouched tail


# ── LLM re-ranking (Sprint 10) ─────────────────────────────────────────
_RERANK_CFG_DEFAULTS = {'llm_rerank_enabled': 0, 'llm_rerank_top_n': 10}
def _rerank_cfg(key):
    dflt = _RERANK_CFG_DEFAULTS[key]
    try:
        v = get_setting(key)
        return dflt if (v is None or str(v).strip() == '') else int(float(v))
    except Exception:
        return dflt


def _llm_rerank(results, query, top_n, api_key):
    """Second-stage AI re-rank over the top-N hybrid results. Sends compact
    candidate summaries to DeepSeek and asks for a fit-ordered list with short
    reasons. Only the top-N slice is reordered; on ANY failure the original
    order is returned unchanged (never worse than hybrid)."""
    subset = results[:top_n]
    if len(subset) < 2 or not api_key:
        return results, False
    lines = []
    for r in subset:
        skills = ', '.join((r.get('key_skills') or [])[:8])
        lines.append(f"id={r['id']} | {r.get('name','')} | {r.get('designation','')} at "
                     f"{r.get('company','')} | {r.get('experience',0)}y | {r.get('location','')} | "
                     f"skills: {skills}")
    prompt = ("You are an expert technical recruiter. Rank these candidates by fit for the "
              f"search query: \"{query}\".\n\nCandidates:\n" + '\n'.join(lines) +
              "\n\nReturn ONLY a JSON array, best fit first, each item exactly: "
              '{"id": <candidate id>, "fit": <0-100 integer>, "reason": "<max 12 words>"}. '
              "No prose, no markdown.")
    try:
        resp = call_deepseek(api_key,
            {'model': 'deepseek-chat', 'temperature': 0.2, 'max_tokens': 700,
             'messages': [{'role': 'user', 'content': prompt}]},
            timeout=60, endpoint='rerank')
        if resp.status_code != 200:
            return results, False
        arr = parse_json(resp.json()['choices'][0]['message']['content'])
        if not isinstance(arr, list) or not arr:
            return results, False
    except Exception:
        return results, False

    by_id = {r['id']: r for r in subset}
    ordered, seen = [], set()
    for item in arr:
        try:
            cid = int(item.get('id'))
        except Exception:
            continue
        r = by_id.get(cid)
        if r and cid not in seen:
            r = dict(r)
            r['llm_score'] = item.get('fit')
            r['llm_reason'] = str(item.get('reason', ''))[:160]
            ordered.append(r); seen.add(cid)
    # any top-N candidates the model omitted keep their original relative order
    for r in subset:
        if r['id'] not in seen:
            ordered.append(r)
    return ordered + results[top_n:], True


@app.route('/api/ai/search', methods=['POST'])
@login_required
def ai_search():
    """Semantic talent search across ALL candidates (ignores mandate
    boundaries). Sprint 4: streams candidate vectors in chunks (bounded
    memory), keeps a top-N heap, loads full display fields only for the
    returned page, and caches query embeddings + ranked results."""
    import time as _time
    t_start = _time.perf_counter()
    d = request.json or {}
    query = (d.get('query') or '').strip()
    if not query:
        return jsonify({'error': 'Empty query'}), 400

    # ── Params / config ────────────────────────────────────────────────
    top_k       = int(d.get('top_k') or 10)
    page        = max(1, int(d.get('page') or 1))
    page_size   = int(d.get('page_size') or top_k)
    max_cands   = int(d.get('max_candidates') or _search_cfg('search_max_candidates'))
    min_score   = float(d.get('min_score') if d.get('min_score') is not None else _search_cfg('search_min_score'))
    pool        = max(_search_cfg('search_max_results'), page * page_size)
    no_cache    = bool(d.get('no_cache'))
    offset      = (page - 1) * page_size

    api_key = get_setting('gemini_api_key')
    if not api_key:
        return jsonify({'error': 'Gemini API key not set. Add it in Settings.'}), 400

    timing = {}
    _oid = effective_company_id()
    # Cache key MUST include the tenant id, else two agencies searching the same
    # phrase could share cached rankings (cross-tenant leak).
    ckey = f'{_oid}|{query.lower()}|{min_score}|{max_cands}|{pool}'

    # ── 1) Query embedding (cache) ─────────────────────────────────────
    t0 = _time.perf_counter()
    cached = None if no_cache else _search_cache_get(ckey)
    if cached:
        qvec = cached['qvec']; ranked = cached['ranked']; scanned = cached['scanned']
        timing['embed_ms'] = 0; timing['scan_ms'] = 0; from_cache = True
    else:
        qvec = embed_one(query)
        if isinstance(qvec, dict) and qvec.get('error'):
            return jsonify({'error': 'Gemini error: ' + qvec['error']}), 400
        if not qvec:
            return jsonify({'error': 'Could not embed query'}), 400
        timing['embed_ms'] = int((_time.perf_counter() - t0) * 1000)
        from_cache = False

    conn = get_db()

    # ── 2) Stream + score + rank (only when not cached) ────────────────
    if not from_cache:
        t0 = _time.perf_counter()
        ranked, scanned = _search_rank(conn, qvec, min_score, max_cands, pool, _oid)
        timing['scan_ms'] = int((_time.perf_counter() - t0) * 1000)
        if not no_cache:
            _search_cache_put(ckey, {'qvec': qvec, 'ranked': ranked, 'scanned': scanned})

    total = len(ranked)

    # ── 2a) Multi-vector re-rank (Sprint 8) — opt-in, off by default ────
    # Blends best-matching facet vector into the ranking. Returns ranked
    # unchanged if no facet vectors exist, so it's safe even before generation.
    mv_on = bool(d.get('multivector')) if ('multivector' in d) else bool(_mv_cfg('multivector_enabled'))
    if mv_on:
        t0 = _time.perf_counter()
        ranked = _multivector_rerank(conn, qvec, ranked)
        timing['multivector_ms'] = int((_time.perf_counter() - t0) * 1000)
        total = len(ranked)

    # ── 2a½) Skill-graph expansion (Intelligence Layer · Phase 1) ──────
    # Computed ONCE here, independent of hybrid on/off, so the "recognised
    # skills" banner always renders and the hybrid scorer can reuse it.
    try:
        _sg = _skill_expand(query, conn, _oid)
    except Exception as _e:
        print('[skill-graph] expand error:', _e)
        _sg = {'expand': {}, 'recognized': [], 'related': []}
    d['_skill_sg'] = _sg

    # ── 2b) Hybrid re-rank (Sprint 6) ──────────────────────────────────
    # Default-on but soft: with no detectable structured signals it collapses
    # to pure semantic order (identical to the old engine). Opt out with hybrid=false.
    hybrid_on = bool(d.get('hybrid')) if ('hybrid' in d) else bool(_hybrid_cfg('hybrid_enabled'))
    filters_detected = {}
    if hybrid_on:
        t0 = _time.perf_counter()
        scored_pool, filters_detected = _hybrid_process(conn, query, ranked, d)
        timing['hybrid_ms'] = int((_time.perf_counter() - t0) * 1000)
        total = len(scored_pool)
        page_items = scored_pool[offset:offset + page_size]
        results = page_items
        timing['hydrate_ms'] = 0
        page_slice = [(it['semantic_score'] / 100.0, it['id']) for it in page_items]
    else:
        # ── 3) Two-phase: full fields for THIS page only (semantic-only) ──
        t0 = _time.perf_counter()
        page_slice = ranked[offset:offset + page_size]
        results = _search_hydrate(conn, page_slice, _oid)
        timing['hydrate_ms'] = int((_time.perf_counter() - t0) * 1000)

    # Skill recognition is shown even on the semantic-only path.
    if not isinstance(filters_detected, dict):
        filters_detected = {}
    filters_detected['skill_recognized'] = _sg['recognized']
    filters_detected['skill_related'] = _sg['related']

    # ── 3b) LLM re-rank (Sprint 10) — opt-in, page 1 only ──────────────
    # A second AI pass reorders the top results and attaches fit reasons.
    # Off by default; one DeepSeek call, and it never worsens hybrid order.
    rerank_on = bool(d.get('rerank')) if ('rerank' in d) else bool(_rerank_cfg('llm_rerank_enabled'))
    reranked = False
    if rerank_on and page == 1 and len(results) >= 2:
        ds_key = get_setting('deepseek_api_key')
        if ds_key:
            t0 = _time.perf_counter()
            results, reranked = _llm_rerank(results, query, _rerank_cfg('llm_rerank_top_n'), ds_key)
            timing['rerank_ms'] = int((_time.perf_counter() - t0) * 1000)

    # ── 4) Optional AI reasoning over the page (unchanged behaviour) ───
    reasoning = ''
    if bool(d.get('explain')) and results:
        ds_key = get_setting('deepseek_api_key')
        if ds_key:
            cand_lines = []
            for r in results[:6]:
                cand_lines.append(f"- {r['name']} ({r['designation']} at {r['company']}, "
                                  f"{r['experience']}y, {r['location']}, skills: {', '.join(r['key_skills'][:6])}) "
                                  f"[currently in mandate: {r['mandate_role'] or 'N/A'}]")
            prompt = ("A recruiter searched their candidate pool for: \"" + query + "\".\n\n"
                      "Here are the top matches:\n" + '\n'.join(cand_lines) + "\n\n"
                      "In 3-5 short bullet points, explain which candidates fit best and WHY "
                      "(note if someone saved for a different role is still a strong fit). "
                      "Be concise and practical. Plain text, start each line with '- '.")
            try:
                rr = call_deepseek(ds_key,
                    {'model': 'deepseek-chat', 'temperature': 0.3, 'max_tokens': 400,
                          'messages': [{'role': 'user', 'content': prompt}]},
                    timeout=60, endpoint='reasoning')
                if rr.status_code == 200:
                    reasoning = rr.json()['choices'][0]['message']['content'].strip()
            except Exception:
                pass

    conn.close()

    total_ms = int((_time.perf_counter() - t_start) * 1000)
    timing['total_ms'] = total_ms
    avg_sim = round(sum(s for s, _ in page_slice) / len(page_slice), 4) if page_slice else 0.0
    _search_metric_record(total_ms, scanned)

    # Capture the search into the Recruitment Memory Engine — only on a fresh
    # first page, so scrolling/pagination doesn't spam duplicate events.
    if hybrid_on and page == 1 and not from_cache:
        _rme_capture_search(query, filters_detected, total, current_user())

    # slow-search + per-search logging
    if total_ms >= _search_cfg('search_slow_ms'):
        print(f'[search] SLOW {total_ms}ms q="{query[:40]}" scanned={scanned} '
              f'returned={len(results)} hybrid={hybrid_on} cache={"hit" if from_cache else "miss"} {timing}')
    else:
        print(f'[search] {total_ms}ms q="{query[:40]}" scanned={scanned} '
              f'returned={len(results)} hybrid={hybrid_on} cache={"hit" if from_cache else "miss"}')

    return jsonify({
        'ok': True,
        'results': results,
        'reasoning': reasoning,
        'searched': scanned,       # backward-compatible key (candidates scanned)
        'total': total,            # total ranked (pagination pool)
        'page': page, 'page_size': page_size,
        'has_more': offset + page_size < total,
        'avg_similarity': avg_sim,
        'from_cache': from_cache,
        'hybrid': hybrid_on,
        'multivector': mv_on,
        'reranked': reranked,
        'filters_detected': filters_detected,
        'timing': timing,
    })


@app.route('/api/skill-graph', methods=['GET'])
@login_required
def skill_graph_list():
    """List the domain skill graph (shared seed + this tenant's own nodes)."""
    conn = get_db()
    oid = effective_company_id()
    rows = conn.execute(
        "SELECT id,owner_id,canonical,display,category,aliases,parents,related FROM skill_graph "
        "WHERE owner_id=0 OR owner_id=? ORDER BY category, canonical", (oid,)).fetchall()
    conn.close()
    def _jl(v):
        try: return json.loads(v or '[]')
        except Exception: return []
    out = []
    for r in rows:
        out.append({'id': r['id'], 'canonical': r['canonical'], 'display': r['display'],
                    'category': r['category'], 'aliases': _jl(r['aliases']),
                    'parents': _jl(r['parents']), 'related': _jl(r['related']),
                    'custom': (r['owner_id'] == oid and oid != 0)})
    return jsonify({'ok': True, 'nodes': out, 'count': len(out)})


@app.route('/api/skill-graph', methods=['POST'])
@login_required
def skill_graph_add():
    """Add or extend a skill node for THIS tenant (owner_id = company). Merges
    aliases/parents/related into any existing tenant node with the same canonical.
    Lets Nitin grow the graph as he learns his market — this is the moat compounding."""
    d = request.json or {}
    can = (d.get('canonical') or '').strip().lower()
    if not can:
        return jsonify({'error': 'canonical required'}), 400
    def _norm(lst):
        return sorted({str(x).strip().lower() for x in (lst or []) if str(x).strip()})
    aliases, parents, related = _norm(d.get('aliases')), _norm(d.get('parents')), _norm(d.get('related'))
    display = (d.get('display') or can).strip()
    category = (d.get('category') or '').strip()
    conn = get_db(); oid = effective_company_id()
    existing = conn.execute("SELECT aliases,parents,related FROM skill_graph WHERE owner_id=? AND canonical=?",
                            (oid, can)).fetchone()
    def _jl(v):
        try: return json.loads(v or '[]')
        except Exception: return []
    if existing:
        aliases = sorted(set(aliases) | set(_jl(existing['aliases'])))
        parents = sorted(set(parents) | set(_jl(existing['parents'])))
        related = sorted(set(related) | set(_jl(existing['related'])))
        conn.execute("UPDATE skill_graph SET display=?, category=?, aliases=?, parents=?, related=?, updated_at=? "
                     "WHERE owner_id=? AND canonical=?",
                     (display, category, json.dumps(aliases), json.dumps(parents), json.dumps(related),
                      ts(), oid, can))
    else:
        conn.execute("INSERT INTO skill_graph (owner_id,canonical,display,category,aliases,parents,related,created_at,updated_at) "
                     "VALUES (?,?,?,?,?,?,?,?,?)",
                     (oid, can, display, category, json.dumps(aliases), json.dumps(parents), json.dumps(related),
                      ts(), ts()))
    conn.commit(); conn.close()
    _SKILL_GRAPH_CACHE['nodes'] = None  # invalidate cache so search picks it up
    return jsonify({'ok': True, 'canonical': can})


@app.route('/api/skill-graph/expand', methods=['GET'])
@login_required
def skill_graph_expand_preview():
    """Preview what a query expands to — a quick way to test the graph."""
    q = (request.args.get('q') or '').strip()
    if not q:
        return jsonify({'ok': True, 'recognized': [], 'related': [], 'expand': {}})
    conn = get_db()
    res = _skill_expand(q, conn, effective_company_id())
    conn.close()
    return jsonify({'ok': True, **res})


def _jd_pool(conn, oid, terms, location=''):
    """Candidates in this tenant whose skill/tag columns mention ANY of `terms`.
    Loose LIKE match across every tag column so the pool estimate is generous."""
    if not terms:
        return []
    terms = list(terms)[:30]
    tagcols = ['key_skill_tags', 'key_skills', 'secondary_skills', 'domain_tags', 'product_handles']
    ors, params = [], []
    for t in terms:
        lt = '%' + t.lower() + '%'
        for col in tagcols:
            ors.append(f"LOWER(COALESCE({col},'')) LIKE ?")
            params.append(lt)
    sql = ("SELECT id,name,company,location,experience FROM candidates "
           "WHERE owner_id=? AND (" + " OR ".join(ors) + ")")
    args = [oid] + params
    if location:
        sql += " AND LOWER(COALESCE(location,'')) LIKE ?"
        args.append('%' + location.lower() + '%')
    try:
        return conn.execute(sql, args).fetchall()
    except Exception:
        return []


@app.route('/api/jd/analyze', methods=['POST'])
@login_required
def jd_analyze():
    """Intelligence Layer · Phase 2 — JD Intelligence. Turn a pasted JD into a
    sourcing plan: recognised + adjacent skills (skill graph), a structured
    requirement (DeepSeek), and — from YOUR OWN database — the talent-pool size,
    the top source companies, and where that talent sits."""
    d = request.json or {}
    jd = (d.get('jd') or d.get('jd_text') or '').strip()
    if not jd:
        return jsonify({'error': 'Paste a JD first'}), 400
    conn = get_db(); oid = effective_company_id()

    # 1) Skill graph — recognise skills + expand to adjacent/related/child
    sg = _skill_expand(jd, conn, oid)
    g = _load_skill_graph(conn, oid)
    lookup = g['lookup']
    grams, _t = _query_grams(jd)
    core_terms = sorted({lookup[gr] for gr in grams if gr in lookup})
    exp_terms = sorted((sg.get('expand') or {}).keys())
    search_terms = sorted(set(core_terms) | set(exp_terms))

    # 2) DeepSeek — structured requirement (best-effort, degrades gracefully)
    structured = {}
    ds_key = get_setting('deepseek_api_key')
    if ds_key:
        try:
            sysmsg = ("Extract the hiring requirement from the job description. "
                      "Reply with ONLY compact JSON, no prose, no code fences. Keys: "
                      "title (string), must_have_skills (array of short strings), "
                      "good_to_have_skills (array of short strings), "
                      "min_experience (number or null), max_experience (number or null), "
                      "location (string), industry (string), seniority (string).")
            rr = call_deepseek(ds_key,
                {'model': 'deepseek-chat', 'temperature': 0.1, 'max_tokens': 600,
                 'messages': [{'role': 'system', 'content': sysmsg},
                              {'role': 'user', 'content': jd[:9000]}]},
                timeout=60, endpoint='jd_intel')
            if rr.status_code == 200:
                raw = (rr.json()['choices'][0]['message']['content'] or '').strip()
                a, b = raw.find('{'), raw.rfind('}')
                if a >= 0 and b > a:
                    structured = json.loads(raw[a:b + 1])
        except TokenCapError:
            structured = {'_error': 'token cap reached'}
        except Exception as e:
            structured = {'_error': str(e)[:140]}

    # 3) From YOUR DB — pool size, top source companies, top locations
    pool = _jd_pool(conn, oid, search_terms)
    from collections import Counter
    comp, loc = Counter(), Counter()
    for r in pool:
        cc = (r['company'] or '').strip()
        if cc:
            comp[cc] += 1
        ll = (r['location'] or '').strip()
        if ll:
            loc[ll] += 1
    conn.close()
    return jsonify({
        'ok': True,
        'structured': structured,
        'recognized_skills': sg.get('recognized', []),
        'adjacent_skills': sg.get('related', []),
        'search_terms': search_terms,
        'pool_size': len(pool),
        'top_companies': [{'name': k, 'count': v} for k, v in comp.most_common(8)],
        'top_locations': [{'name': k, 'count': v} for k, v in loc.most_common(8)],
    })


@app.route('/api/ai/search/metrics', methods=['GET'])
@login_required
def ai_search_metrics():
    """Rolling search performance metrics (in-memory, last N searches)."""
    with _SEARCH_METRICS_LOCK:
        durs = sorted(x[0] for x in _SEARCH_METRICS)
        scans = [x[1] for x in _SEARCH_METRICS]
    if not durs:
        return jsonify({'ok': True, 'samples': 0})
    def _pct(a, p):
        if not a: return 0
        i = min(len(a) - 1, int(round((p / 100.0) * (len(a) - 1))))
        return a[i]
    return jsonify({'ok': True, 'samples': len(durs),
                    'avg_ms': int(sum(durs) / len(durs)),
                    'p95_ms': _pct(durs, 95),
                    'max_ms': durs[-1],
                    'avg_candidates_scanned': int(sum(scans) / len(scans)) if scans else 0})


@app.route('/api/ai/vectors/status', methods=['GET'])
@login_required
def ai_vectors_status():
    """Multi-vector (facet) coverage + config. Confirms the architecture is live."""
    conn = get_db()
    by_facet = {}
    try:
        for r in conn.execute("SELECT facet, status, COUNT(*) n FROM candidate_vectors GROUP BY facet, status"):
            by_facet.setdefault(r['facet'], {})[r['status']] = r['n']
        fully = conn.execute(
            "SELECT COUNT(*) n FROM (SELECT candidate_id FROM candidate_vectors "
            "WHERE status='completed' AND embedding_text_version=? "
            "GROUP BY candidate_id HAVING COUNT(DISTINCT facet) >= ?)",
            (MV_TEXT_TEMPLATE, len(MV_FACETS))).fetchone()['n']
        total_c = conn.execute("SELECT COUNT(*) n FROM candidates WHERE embedding_status='completed'").fetchone()['n']
    except sqlite3.OperationalError:
        fully = total_c = 0
    conn.close()
    return jsonify({'ok': True, 'enabled': bool(_mv_cfg('multivector_enabled')),
                    'facets': MV_FACETS, 'template': MV_TEXT_TEMPLATE,
                    'candidates_fully_faceted': fully, 'candidates_embedded': total_c,
                    'by_facet': by_facet,
                    'weights': {'full': _mv_cfg('multivector_w_full'),
                                'facet': _mv_cfg('multivector_w_facet')}})


@app.route('/api/ai/vectors/rebuild', methods=['POST'])
@login_required
def ai_vectors_rebuild():
    """Manually generate facet vectors for candidates that need them (bounded).
    Explicit intent, so it runs even if multivector_enabled is off — but it still
    needs a Gemini key. Use `cap` to limit how many candidates per call."""
    d = request.json or {}
    cap = min(int(d.get('cap') or 50), 500)
    api_key = get_setting('gemini_api_key')
    if not api_key:
        return jsonify({'error': 'Gemini API key not set. Add it in Settings.'}), 400
    conn = get_db()
    try:
        rows = conn.execute(
            "SELECT c.* FROM candidates c WHERE c.embedding_status='completed' AND c.id NOT IN ("
            "  SELECT candidate_id FROM candidate_vectors WHERE embedding_text_version=? "
            "  GROUP BY candidate_id HAVING COUNT(DISTINCT facet) >= ?"
            ") LIMIT ?", (MV_TEXT_TEMPLATE, len(MV_FACETS), cap)).fetchall()
        done = 0
        for c in rows:
            embed_candidate_facets(conn, c, api_key)
            done += 1
    finally:
        conn.close()
    return jsonify({'ok': True, 'processed': done})


@app.route('/api/mandates/<int:mid>/match', methods=['GET'])
@login_required
def mandate_match_candidates(mid):
    """Rank candidates for this mandate using its STORED JD vector — no
    re-embedding. Reuses the Sprint-4 retrieval + Sprint-6 hydrate."""
    conn = get_db()
    if not _tenant_owns_mandate(conn, mid):
        conn.close(); return jsonify({'ok': False, 'error': 'Not found'}), 404
    jdvec = _mandate_jd_vector(conn, mid)
    if jdvec is None:
        conn.close()
        return jsonify({'ok': False, 'error': 'JD not embedded yet. It will be generated shortly.',
                        'pending': True, 'results': []})
    top_k = min(int(request.args.get('top_k') or 20), 100)
    min_score = float(request.args.get('min_score') or -1.0)
    pool = max(_search_cfg('search_max_results'), top_k)
    ranked, scanned = _search_rank(conn, jdvec, min_score, int(_search_cfg('search_max_candidates')), pool, effective_company_id())
    results = _search_hydrate(conn, ranked[:top_k], effective_company_id())
    conn.close()
    return jsonify({'ok': True, 'mandate_id': mid, 'searched': scanned,
                    'total': len(ranked), 'results': results})


@app.route('/api/candidates/<int:cid>/match-mandates', methods=['GET'])
@login_required
def candidate_match_mandates(cid):
    """Reverse match: given a candidate's vector, rank the open mandates whose
    stored JD vector best fits. Mandates are few, so this is a trivial scan."""
    conn = get_db()
    cr = conn.execute("SELECT embedding_vec FROM candidates WHERE id=? AND owner_id=?",
                      (cid, effective_company_id())).fetchone()
    if not cr or not cr['embedding_vec']:
        conn.close(); return jsonify({'ok': False, 'error': 'Candidate not embedded yet', 'results': []})
    try:
        cvec = _np.frombuffer(cr['embedding_vec'], dtype=_np.float32) if _HAS_NUMPY else \
               (lambda a: (a.frombytes(cr['embedding_vec']) or a))(_array.array('f'))
    except Exception:
        conn.close(); return jsonify({'ok': False, 'error': 'Bad candidate vector', 'results': []})
    rows = conn.execute(
        "SELECT mv.mandate_id, mv.embedding_vec, m.role, m.client, m.location, m.status "
        "FROM mandate_vectors mv JOIN mandates m ON m.id=mv.mandate_id "
        "WHERE mv.status='completed' AND mv.embedding_vec IS NOT NULL AND m.owner_id=?",
        (effective_company_id(),)).fetchall()
    out = []
    for r in rows:
        try:
            jv = _np.frombuffer(r['embedding_vec'], dtype=_np.float32) if _HAS_NUMPY else \
                 (lambda a: (a.frombytes(r['embedding_vec']) or a))(_array.array('f'))
            sim = cosine(list(cvec), list(jv))
        except Exception:
            continue
        out.append({'mandate_id': r['mandate_id'], 'role': r['role'], 'client': r['client'],
                    'location': r['location'], 'status': r['status'], 'score': round(sim * 100, 1)})
    out.sort(key=lambda x: x['score'], reverse=True)
    conn.close()
    return jsonify({'ok': True, 'candidate_id': cid, 'results': out[:20]})


@app.route('/api/ai/jd/status', methods=['GET'])
@login_required
def ai_jd_status():
    conn = get_db()
    by_status = {}
    try:
        for r in conn.execute("SELECT status, COUNT(*) n FROM mandate_vectors GROUP BY status"):
            by_status[r['status']] = r['n']
        total_m = conn.execute("SELECT COUNT(*) n FROM mandates").fetchone()['n']
        done = conn.execute("SELECT COUNT(*) n FROM mandate_vectors WHERE status='completed' "
                            "AND embedding_text_version=?", (JD_TEXT_TEMPLATE,)).fetchone()['n']
    except sqlite3.OperationalError:
        total_m = done = 0
    conn.close()
    return jsonify({'ok': True, 'enabled': _jd_cfg_enabled(), 'template': JD_TEXT_TEMPLATE,
                    'mandates_total': total_m, 'mandates_embedded': done, 'by_status': by_status})


# ══════════════════════════════════════════════════════════════════════
#  RECRUITMENT MEMORY ENGINE (Sprint 5) — primitives + thin API.
#  These are the reusable write/read functions future sprints will call to
#  attach memories, log events and record relationships on any entity.
#  Nothing here auto-generates data; existing flows are untouched.
# ══════════════════════════════════════════════════════════════════════

# Canonical vocabularies (documentation + /status breakdown + light validation).
# New values are allowed (extensibility) — these are the known-good set.
RME_ENTITY_TYPES = ('candidate', 'job', 'client', 'company', 'recruiter', 'placement',
                    'interview', 'submission', 'communication', 'skill', 'technology', 'industry')
RME_MEMORY_TYPES = ('fact', 'event', 'relationship', 'observation', 'preference',
                    'ai_insight', 'recruiter_note', 'system_note', 'interaction')
RME_EVENT_TYPES = ('candidate_created', 'resume_uploaded', 'candidate_updated',
                   'interview_scheduled', 'interview_feedback', 'submission',
                   'offer_released', 'offer_accepted', 'placement', 'communication',
                   'status_change', 'search_activity')
RME_VISIBILITY = ('internal', 'team', 'private', 'candidate_visible')
RME_STATUS = ('active', 'archived', 'deleted')


def _rme_json(v, default):
    """Coerce a value to a JSON string for storage (accepts list/dict/str)."""
    if v is None:
        return default
    if isinstance(v, (list, dict)):
        try:
            return json.dumps(v)
        except Exception:
            return default
    return str(v)


def rme_add_memory(conn, entity_type, entity_id, memory_type='fact', title='', content='',
                   source='system', created_by='', visibility='internal', confidence=1.0,
                   importance=0, tags=None, metadata=None, status='active'):
    """Attach a memory to any entity. Returns the new memory id. This is the
    core primitive future sprints (AI insights, recruiter notes) build on."""
    now = ts()
    cur = conn.execute(
        "INSERT INTO rme_memories (entity_type, entity_id, memory_type, title, content, source, "
        "created_by, created_at, updated_at, visibility, confidence, importance, tags, metadata, status) "
        "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (str(entity_type), str(entity_id), memory_type, title, content, source, str(created_by),
         now, now, visibility, float(confidence), int(importance),
         _rme_json(tags, '[]'), _rme_json(metadata, '{}'), status))
    conn.commit()
    return cur.lastrowid


def rme_add_event(conn, event_type, entity_type, entity_id, actor='system', summary='',
                  data=None, related_entity_type='', related_entity_id=''):
    """Append an immutable event to the RME timeline. Returns the event id."""
    cur = conn.execute(
        "INSERT INTO rme_events (event_type, entity_type, entity_id, actor, summary, data, "
        "related_entity_type, related_entity_id, created_at) VALUES (?,?,?,?,?,?,?,?,?)",
        (str(event_type), str(entity_type), str(entity_id), str(actor), summary,
         _rme_json(data, '{}'), str(related_entity_type), str(related_entity_id), ts()))
    conn.commit()
    return cur.lastrowid


def rme_set_relationship(conn, from_type, from_id, to_type, to_id, rel_type,
                         weight=1.0, confidence=1.0, source='system', metadata=None):
    """Upsert a relationship edge (unique per from/to/rel_type). Returns row id."""
    now = ts()
    conn.execute(
        "INSERT INTO rme_relationships (from_type, from_id, to_type, to_id, rel_type, weight, "
        "confidence, source, metadata, created_at, updated_at) VALUES (?,?,?,?,?,?,?,?,?,?,?) "
        "ON CONFLICT(from_type, from_id, to_type, to_id, rel_type) DO UPDATE SET "
        "weight=excluded.weight, confidence=excluded.confidence, source=excluded.source, "
        "metadata=excluded.metadata, updated_at=excluded.updated_at, status='active'",
        (str(from_type), str(from_id), str(to_type), str(to_id), str(rel_type),
         float(weight), float(confidence), source, _rme_json(metadata, '{}'), now, now))
    conn.commit()
    row = conn.execute(
        "SELECT id FROM rme_relationships WHERE from_type=? AND from_id=? AND to_type=? "
        "AND to_id=? AND rel_type=?",
        (str(from_type), str(from_id), str(to_type), str(to_id), str(rel_type))).fetchone()
    return row['id'] if row else None


def _rme_row(r):
    return {k: r[k] for k in r.keys()}


@app.route('/api/rme/memory', methods=['GET', 'POST'])
@login_required
def rme_memory():
    if request.method == 'POST':
        d = request.json or {}
        et, eid = (d.get('entity_type') or '').strip(), str(d.get('entity_id') or '').strip()
        if not et or not eid:
            return jsonify({'error': 'entity_type and entity_id are required'}), 400
        conn = get_db()
        mid = rme_add_memory(
            conn, et, eid,
            memory_type=(d.get('memory_type') or 'fact'),
            title=d.get('title', ''), content=d.get('content', ''),
            source=d.get('source', 'recruiter'),
            created_by=str((current_user() or {}).get('username', '')),
            visibility=d.get('visibility', 'internal'),
            confidence=d.get('confidence', 1.0), importance=d.get('importance', 0),
            tags=d.get('tags'), metadata=d.get('metadata'))
        conn.close()
        return jsonify({'ok': True, 'id': mid})
    # GET — list memories for an entity
    et = (request.args.get('entity_type') or '').strip()
    eid = str(request.args.get('entity_id') or '').strip()
    if not et or not eid:
        return jsonify({'error': 'entity_type and entity_id are required'}), 400
    limit = min(int(request.args.get('limit') or 100), 500)
    conn = get_db()
    rows = conn.execute(
        "SELECT * FROM rme_memories WHERE entity_type=? AND entity_id=? AND status!='deleted' "
        "ORDER BY importance DESC, created_at DESC LIMIT ?", (et, eid, limit)).fetchall()
    conn.close()
    return jsonify({'ok': True, 'memories': [_rme_row(r) for r in rows]})


@app.route('/api/rme/memory/<int:mid>/archive', methods=['POST'])
@login_required
def rme_memory_archive(mid):
    conn = get_db()
    conn.execute("UPDATE rme_memories SET status='archived', updated_at=? WHERE id=?", (ts(), mid))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/rme/event', methods=['GET', 'POST'])
@login_required
def rme_event():
    if request.method == 'POST':
        d = request.json or {}
        etype = (d.get('event_type') or '').strip()
        et, eid = (d.get('entity_type') or '').strip(), str(d.get('entity_id') or '').strip()
        if not etype or not et or not eid:
            return jsonify({'error': 'event_type, entity_type and entity_id are required'}), 400
        conn = get_db()
        evid = rme_add_event(conn, etype, et, eid,
                             actor=str((current_user() or {}).get('username', 'system')),
                             summary=d.get('summary', ''), data=d.get('data'),
                             related_entity_type=d.get('related_entity_type', ''),
                             related_entity_id=d.get('related_entity_id', ''))
        conn.close()
        return jsonify({'ok': True, 'id': evid})
    et = (request.args.get('entity_type') or '').strip()
    eid = str(request.args.get('entity_id') or '').strip()
    limit = min(int(request.args.get('limit') or 100), 500)
    conn = get_db()
    if et and eid:
        rows = conn.execute("SELECT * FROM rme_events WHERE entity_type=? AND entity_id=? "
                            "ORDER BY created_at DESC, id DESC LIMIT ?", (et, eid, limit)).fetchall()
    else:
        rows = conn.execute("SELECT * FROM rme_events ORDER BY created_at DESC, id DESC LIMIT ?",
                            (limit,)).fetchall()
    conn.close()
    return jsonify({'ok': True, 'events': [_rme_row(r) for r in rows]})


@app.route('/api/rme/relationship', methods=['GET', 'POST'])
@login_required
def rme_relationship():
    if request.method == 'POST':
        d = request.json or {}
        req = ['from_type', 'from_id', 'to_type', 'to_id', 'rel_type']
        if not all(str(d.get(k) or '').strip() for k in req):
            return jsonify({'error': 'from_type, from_id, to_type, to_id, rel_type are required'}), 400
        conn = get_db()
        rid = rme_set_relationship(conn, d['from_type'], d['from_id'], d['to_type'], d['to_id'],
                                   d['rel_type'], weight=d.get('weight', 1.0),
                                   confidence=d.get('confidence', 1.0),
                                   source=d.get('source', 'recruiter'), metadata=d.get('metadata'))
        conn.close()
        return jsonify({'ok': True, 'id': rid})
    ft = (request.args.get('from_type') or '').strip()
    fid = str(request.args.get('from_id') or '').strip()
    limit = min(int(request.args.get('limit') or 200), 1000)
    conn = get_db()
    if ft and fid:
        rows = conn.execute("SELECT * FROM rme_relationships WHERE from_type=? AND from_id=? "
                            "AND status='active' ORDER BY weight DESC LIMIT ?", (ft, fid, limit)).fetchall()
    else:
        rows = conn.execute("SELECT * FROM rme_relationships WHERE status='active' "
                            "ORDER BY id DESC LIMIT ?", (limit,)).fetchall()
    conn.close()
    return jsonify({'ok': True, 'relationships': [_rme_row(r) for r in rows]})


@app.route('/api/rme/status', methods=['GET'])
@login_required
def rme_status():
    """Confirms the RME architecture is live and reports counts. Useful to
    verify the foundation without any memories having been generated yet."""
    conn = get_db()
    def _count(t):
        try:
            return conn.execute(f'SELECT COUNT(*) n FROM {t}').fetchone()['n']
        except sqlite3.OperationalError:
            return None
    mem_by_type, evt_by_type = {}, {}
    try:
        for r in conn.execute("SELECT memory_type m, COUNT(*) n FROM rme_memories GROUP BY m"):
            mem_by_type[r['m']] = r['n']
        for r in conn.execute("SELECT event_type e, COUNT(*) n FROM rme_events GROUP BY e"):
            evt_by_type[r['e']] = r['n']
    except sqlite3.OperationalError:
        pass
    out = {'ok': True,
           'memories': _count('rme_memories'),
           'events': _count('rme_events'),
           'relationships': _count('rme_relationships'),
           'memories_by_type': mem_by_type,
           'events_by_type': evt_by_type,
           'vocab': {'entity_types': RME_ENTITY_TYPES, 'memory_types': RME_MEMORY_TYPES,
                     'event_types': RME_EVENT_TYPES}}
    conn.close()
    return jsonify(out)


# ══════════════════════════════════════════════════════════════════════
#  RECRUITMENT KNOWLEDGE GRAPH (Sprint 7) — repository layer + interfaces.
#  Everything the rest of RecruitOS touches goes through these rkg_* helpers,
#  never raw SQL — that abstraction is what lets a graph DB replace SQLite
#  later without changing callers. Architecture only: no extraction, no
#  reasoning, no recommendations.
# ══════════════════════════════════════════════════════════════════════

RKG_ENTITY_TYPES = ('candidate', 'company', 'job', 'client', 'recruiter', 'skill', 'technology',
                    'certification', 'industry', 'product', 'location', 'education',
                    'institution', 'project', 'role', 'designation')
RKG_REL_TYPES = (
    # candidate
    'WORKED_AT', 'HAS_SKILL', 'HAS_CERTIFICATION', 'KNOWS_TECHNOLOGY', 'WORKED_ON',
    'LOCATED_IN', 'HAS_EDUCATION', 'REPORTS_TO', 'REFERRED_BY',
    # company
    'OPERATES_IN', 'COMPETES_WITH', 'USES_TECHNOLOGY', 'MANUFACTURES', 'HIRES_FOR',
    'SUPPLIES', 'PARTNERS_WITH',
    # job
    'REQUIRES_SKILL', 'REQUIRES_CERTIFICATION', 'REQUIRES_TECHNOLOGY',
    'BELONGS_TO_INDUSTRY', 'POSTED_BY',
    # recruiter
    'MANAGES', 'PLACED', 'WORKS_WITH', 'SEARCHED', 'CONTACTED',
    # generic concept graph
    'RELATED_TO',
)

# Seed alias map — small on purpose (architecture, not a big dictionary).
# Maps a canonical normalized key -> known surface variants. Extend freely.
_RKG_ALIAS_SEED = {
    'motorcontrolcentre': ['mcc', 'motor control center', 'motor control centre', 'intelligent mcc'],
    'powercontrolcentre': ['pcc', 'power control center', 'power control centre'],
    'lowvoltage':         ['lv', 'l v', 'low voltage'],
    'iec61439':           ['iec 61439', 'iec61439'],
}


def rkg_normalize(name):
    """Canonical dedup key: lowercase, punctuation/separators removed, spaces
    stripped. Collapses spacing/punctuation variants of code-like terms, e.g.
    'IEC 61439' / 'I.E.C-61439' / 'iec61439' -> 'iec61439'. Abbreviation and
    spelling variants (MCC, centre/center) are handled by the alias layer."""
    s = (name or '').lower().strip()
    s = re.sub(r'[^a-z0-9]+', '', s)
    return s


def _rkg_canonical_norm(name):
    """Resolve a surface form to its canonical normalized key via the seed alias
    map, falling back to the plain normalized form. Returns (canonical, norm)."""
    norm = rkg_normalize(name)
    if not norm:
        return '', ''
    if norm in _RKG_ALIAS_SEED:
        return norm, norm
    for canon, variants in _RKG_ALIAS_SEED.items():
        if norm == canon or any(rkg_normalize(v) == norm for v in variants):
            return canon, norm
    return norm, norm


def _rkg_register_alias(conn, entity_id, entity_type, surface):
    """Record a surface form -> entity mapping so later lookups of this alias
    resolve to the same canonical node. Idempotent; ignores conflicts."""
    na = rkg_normalize(surface)
    if not na:
        return
    try:
        conn.execute("INSERT OR IGNORE INTO rkg_aliases (entity_id, entity_type, normalized_alias, created_at) "
                     "VALUES (?,?,?,?)", (entity_id, entity_type, na, ts()))
    except sqlite3.OperationalError:
        pass


def _rkg_lookup(conn, entity_type, name):
    """Resolve a surface form to an existing entity id, trying (1) canonical
    normalized_name, then (2) the learned alias table. Returns id or None."""
    canon, norm = _rkg_canonical_norm(name)
    if not canon:
        return None
    row = conn.execute("SELECT id FROM rkg_entities WHERE entity_type=? AND normalized_name=?",
                       (entity_type, canon)).fetchone()
    if row:
        return row['id']
    try:
        arow = conn.execute("SELECT entity_id FROM rkg_aliases WHERE entity_type=? AND normalized_alias=?",
                            (entity_type, norm)).fetchone()
        if arow:
            return arow['entity_id']
    except sqlite3.OperationalError:
        pass
    return None


def rkg_get_or_create_entity(conn, entity_type, name, description='', metadata=None, source='system'):
    """Idempotently resolve/create a canonical entity. Same concept under a
    different surface form returns the SAME id (and records the new alias).
    This is THE entry point for putting anything into the graph."""
    display = (name or '').strip()
    if not display:
        return None
    canon, norm = _rkg_canonical_norm(display)
    if not canon:
        return None
    eid = _rkg_lookup(conn, entity_type, display)
    if eid:
        # record a newly-seen surface form (on the entity + alias table)
        row = conn.execute("SELECT aliases FROM rkg_entities WHERE id=?", (eid,)).fetchone()
        try:
            al = json.loads(row['aliases'] or '[]') if row else []
        except Exception:
            al = []
        if display.lower() not in [a.lower() for a in al]:
            al.append(display)
            conn.execute("UPDATE rkg_entities SET aliases=?, updated_at=? WHERE id=?",
                         (json.dumps(al), ts(), eid))
        _rkg_register_alias(conn, eid, entity_type, display)
        conn.commit()
        return eid
    now = ts()
    cur = conn.execute(
        "INSERT INTO rkg_entities (entity_type, display_name, normalized_name, aliases, description, "
        "metadata, created_at, updated_at) VALUES (?,?,?,?,?,?,?,?)",
        (entity_type, display, canon, json.dumps([display]), description,
         _rme_json(metadata, '{}'), now, now))
    eid = cur.lastrowid
    # register both the canonical form and the surface form as aliases
    _rkg_register_alias(conn, eid, entity_type, display)
    _rkg_register_alias(conn, eid, entity_type, canon)
    conn.commit()
    return eid


def rkg_add_alias(conn, entity_id, alias):
    row = conn.execute("SELECT entity_type, aliases FROM rkg_entities WHERE id=?", (entity_id,)).fetchone()
    if not row:
        return False
    try:
        al = json.loads(row['aliases'] or '[]')
    except Exception:
        al = []
    if alias and alias.lower() not in [a.lower() for a in al]:
        al.append(alias)
        conn.execute("UPDATE rkg_entities SET aliases=?, updated_at=? WHERE id=?",
                     (json.dumps(al), ts(), entity_id))
    # register in the resolution table so this alias now resolves to the node
    _rkg_register_alias(conn, entity_id, row['entity_type'], alias)
    conn.commit()
    return True


def rkg_find_entity(conn, entity_type, name):
    eid = _rkg_lookup(conn, entity_type, name)
    if not eid:
        return None
    row = conn.execute("SELECT * FROM rkg_entities WHERE id=?", (eid,)).fetchone()
    return dict(row) if row else None


def rkg_link(conn, source_id, target_id, rel_type, confidence=1.0, source='system', metadata=None):
    """Upsert a directed edge between two entities (unique per src/tgt/rel_type)."""
    now = ts()
    conn.execute(
        "INSERT INTO rkg_edges (source_id, target_id, rel_type, confidence, source, metadata, "
        "created_at, updated_at) VALUES (?,?,?,?,?,?,?,?) "
        "ON CONFLICT(source_id, target_id, rel_type) DO UPDATE SET "
        "confidence=excluded.confidence, source=excluded.source, metadata=excluded.metadata, "
        "updated_at=excluded.updated_at, status='active'",
        (source_id, target_id, str(rel_type), float(confidence), source,
         _rme_json(metadata, '{}'), now, now))
    conn.commit()
    row = conn.execute("SELECT id FROM rkg_edges WHERE source_id=? AND target_id=? AND rel_type=?",
                       (source_id, target_id, str(rel_type))).fetchone()
    return row['id'] if row else None


def rkg_neighbors(conn, entity_id, rel_type=None, direction='out', limit=200):
    """Traverse one hop. direction: 'out' (entity is source), 'in' (target),
    'both'. Returns connected entities with the edge type + confidence."""
    out = []
    def _q(join_on, other):
        sql = (f"SELECT e.*, k.rel_type AS _rel, k.confidence AS _conf, k.id AS _edge_id "
               f"FROM rkg_edges k JOIN rkg_entities e ON e.id=k.{other} "
               f"WHERE k.{join_on}=? AND k.status='active'")
        args = [entity_id]
        if rel_type:
            sql += " AND k.rel_type=?"; args.append(rel_type)
        sql += " LIMIT ?"; args.append(limit)
        for r in conn.execute(sql, args).fetchall():
            d = dict(r); d['_direction'] = ('out' if join_on == 'source_id' else 'in')
            out.append(d)
    if direction in ('out', 'both'):
        _q('source_id', 'target_id')
    if direction in ('in', 'both'):
        _q('target_id', 'source_id')
    return out


# ── Memory Engine integration ──────────────────────────────────────────
# A graph node can carry a memory timeline via the existing RME (Sprint 5),
# addressed as entity_type='rkg_entity'. No new memory storage needed.
def rkg_attach_memory(conn, entity_id, memory_type='fact', title='', content='',
                      source='system', created_by='', confidence=1.0, importance=0,
                      tags=None, metadata=None):
    return rme_add_memory(conn, 'rkg_entity', str(entity_id), memory_type=memory_type,
                          title=title, content=content, source=source, created_by=created_by,
                          confidence=confidence, importance=importance, tags=tags, metadata=metadata)

def rkg_memories(conn, entity_id, limit=100):
    rows = conn.execute("SELECT * FROM rme_memories WHERE entity_type='rkg_entity' AND entity_id=? "
                        "AND status!='deleted' ORDER BY importance DESC, created_at DESC LIMIT ?",
                        (str(entity_id), limit)).fetchall()
    return [_rme_row(r) for r in rows]


# ── Knowledge-extraction interface (DESIGN ONLY — no implementation) ────
# Future sprints register extractors that turn raw text (resume, JD, notes,
# email, WhatsApp) into (entity, relationship) candidates. The registry +
# signature are defined now so callers and extractors can be built against a
# stable contract; the default extractor intentionally returns nothing.
_RKG_EXTRACTORS = {}
def rkg_register_extractor(source_type, fn):
    """Register a callable fn(text, context)->{'entities':[...],'edges':[...]}."""
    _RKG_EXTRACTORS[source_type] = fn

def rkg_extract(source_type, text, context=None):
    """Dispatch to a registered extractor. No-op until extractors are added in a
    later sprint (architecture only)."""
    fn = _RKG_EXTRACTORS.get(source_type)
    if not fn:
        return {'entities': [], 'edges': [], 'implemented': False}
    try:
        return fn(text, context or {})
    except Exception as e:
        return {'entities': [], 'edges': [], 'error': str(e)}


# ── Search integration interface (exposed, NOT wired into search yet) ───
def rkg_resolve_terms(conn, terms):
    """Map query terms to canonical entities + their aliases. Hybrid search can
    call this later for synonym expansion; search behaviour is unchanged now."""
    resolved = []
    for t in (terms or []):
        # resolve type-agnostically: direct canonical match, then learned alias
        canon, norm = _rkg_canonical_norm(t)
        row = conn.execute("SELECT id, entity_type, display_name, aliases FROM rkg_entities "
                           "WHERE normalized_name=? LIMIT 1", (canon,)).fetchone()
        if not row:
            arow = conn.execute("SELECT entity_id FROM rkg_aliases WHERE normalized_alias=? LIMIT 1",
                               (norm,)).fetchone()
            if arow:
                row = conn.execute("SELECT id, entity_type, display_name, aliases FROM rkg_entities "
                                   "WHERE id=?", (arow['entity_id'],)).fetchone()
        if row:
            try:
                al = json.loads(row['aliases'] or '[]')
            except Exception:
                al = []
            resolved.append({'term': t, 'entity_id': row['id'], 'entity_type': row['entity_type'],
                             'canonical': row['display_name'], 'aliases': al})
        else:
            resolved.append({'term': t, 'entity_id': None, 'canonical': None, 'aliases': []})
    return resolved


# ── Thin API (additive, login_required) ────────────────────────────────
@app.route('/api/rkg/entity', methods=['GET', 'POST'])
@login_required
def rkg_entity():
    if request.method == 'POST':
        d = request.json or {}
        et, name = (d.get('entity_type') or '').strip(), (d.get('name') or '').strip()
        if not et or not name:
            return jsonify({'error': 'entity_type and name are required'}), 400
        conn = get_db()
        eid = rkg_get_or_create_entity(conn, et, name, description=d.get('description', ''),
                                       metadata=d.get('metadata'), source=d.get('source', 'recruiter'))
        row = conn.execute("SELECT * FROM rkg_entities WHERE id=?", (eid,)).fetchone()
        conn.close()
        return jsonify({'ok': True, 'id': eid, 'entity': dict(row) if row else None})
    conn = get_db()
    if request.args.get('id'):
        row = conn.execute("SELECT * FROM rkg_entities WHERE id=?", (request.args.get('id'),)).fetchone()
        conn.close()
        return jsonify({'ok': True, 'entity': dict(row) if row else None})
    et = (request.args.get('entity_type') or '').strip()
    name = (request.args.get('name') or '').strip()
    if et and name:
        ent = rkg_find_entity(conn, et, name); conn.close()
        return jsonify({'ok': True, 'entity': ent})
    q = "SELECT * FROM rkg_entities WHERE status='active'"
    args = []
    if et:
        q += " AND entity_type=?"; args.append(et)
    q += " ORDER BY id DESC LIMIT ?"; args.append(min(int(request.args.get('limit') or 100), 500))
    rows = conn.execute(q, args).fetchall(); conn.close()
    return jsonify({'ok': True, 'entities': [dict(r) for r in rows]})


@app.route('/api/rkg/entity/<int:eid>/alias', methods=['POST'])
@login_required
def rkg_entity_alias(eid):
    d = request.json or {}
    conn = get_db(); ok = rkg_add_alias(conn, eid, (d.get('alias') or '').strip()); conn.close()
    return jsonify({'ok': ok})


@app.route('/api/rkg/edge', methods=['POST'])
@login_required
def rkg_edge():
    d = request.json or {}
    for k in ('source_id', 'target_id', 'rel_type'):
        if not str(d.get(k) or '').strip():
            return jsonify({'error': 'source_id, target_id and rel_type are required'}), 400
    conn = get_db()
    rid = rkg_link(conn, int(d['source_id']), int(d['target_id']), d['rel_type'],
                   confidence=d.get('confidence', 1.0), source=d.get('source', 'recruiter'),
                   metadata=d.get('metadata'))
    conn.close()
    return jsonify({'ok': True, 'id': rid})


@app.route('/api/rkg/neighbors', methods=['GET'])
@login_required
def rkg_neighbors_api():
    eid = request.args.get('entity_id')
    if not eid:
        return jsonify({'error': 'entity_id is required'}), 400
    conn = get_db()
    nb = rkg_neighbors(conn, int(eid), rel_type=request.args.get('rel_type'),
                       direction=request.args.get('direction', 'out'))
    conn.close()
    return jsonify({'ok': True, 'neighbors': nb})


@app.route('/api/rkg/status', methods=['GET'])
@login_required
def rkg_status():
    conn = get_db()
    def _c(t):
        try:
            return conn.execute(f'SELECT COUNT(*) n FROM {t}').fetchone()['n']
        except sqlite3.OperationalError:
            return None
    by_type = {}
    try:
        for r in conn.execute("SELECT entity_type t, COUNT(*) n FROM rkg_entities GROUP BY t"):
            by_type[r['t']] = r['n']
    except sqlite3.OperationalError:
        pass
    out = {'ok': True, 'entities': _c('rkg_entities'), 'edges': _c('rkg_edges'),
           'entities_by_type': by_type,
           'vocab': {'entity_types': RKG_ENTITY_TYPES, 'rel_types': RKG_REL_TYPES}}
    conn.close()
    return jsonify(out)


@app.route('/api/ai/stats', methods=['POST'])
@login_required
def ai_stats():
    """Approach B: answer counting/analytics questions. We pull a compact,
    structured snapshot of all candidates+mandates and let DeepSeek reason
    over it (no embeddings needed). Good for 'how many', 'which', 'average'."""
    d = request.json or {}
    question = (d.get('question') or '').strip()
    if not question:
        return jsonify({'error': 'Empty question'}), 400
    ds_key = get_setting('deepseek_api_key')
    if not ds_key:
        return jsonify({'error': 'DeepSeek API key not set. Add it in Settings.'}), 400

    conn = get_db()
    mandates = conn.execute('SELECT id, role, client, location, status, ctc_min, ctc_max FROM mandates').fetchall()
    cands = conn.execute('SELECT name, company, designation, experience, ctc_current, '
                         'ctc_expected, notice_period, location, key_skills, stage, mandate_id '
                         'FROM candidates').fetchall()
    conn.close()

    m_lines = []
    for m in mandates:
        m_lines.append(f"Mandate#{m['id']}: {m['role']} @ {m['client']} | loc:{m['location']} | "
                       f"status:{m['status']} | CTC {m['ctc_min']}-{m['ctc_max']}L")
    c_lines = []
    for c in cands:
        try:
            sk = ', '.join(json.loads(c['key_skills'] or '[]')[:6])
        except Exception:
            sk = ''
        c_lines.append(f"{c['name']} | {c['designation']} @ {c['company']} | exp:{c['experience']}y | "
                       f"CTC cur:{c['ctc_current']} exp:{c['ctc_expected']} | NP:{c['notice_period']}d | "
                       f"loc:{c['location']} | stage:{c['stage']} | mandate:{c['mandate_id']} | skills:{sk}")

    # Keep prompt within limits; if too many candidates, note truncation
    MAX = 400
    truncated = len(c_lines) > MAX
    snapshot = ("MANDATES:\n" + '\n'.join(m_lines) + "\n\nCANDIDATES (" +
                str(len(c_lines)) + " total" + (", showing first 400" if truncated else "") + "):\n" +
                '\n'.join(c_lines[:MAX]))

    prompt = ("You are a recruitment data analyst for HireLab. Answer the recruiter's "
              "question using ONLY the data snapshot below. Be precise with numbers, "
              "and list relevant candidate names when useful. If the data is truncated "
              "and you can't be exact, say so. Keep it concise and practical.\n\n"
              "DATA:\n" + snapshot + "\n\nQUESTION: " + question + "\n\nANSWER:")

    try:
        rr = call_deepseek(ds_key,
            {'model': 'deepseek-chat', 'temperature': 0.2, 'max_tokens': 800,
                  'messages': [{'role': 'user', 'content': prompt}]},
            timeout=90, endpoint='analysis')
        if rr.status_code != 200:
            err = rr.json().get('error', {}).get('message', rr.text[:200])
            return jsonify({'error': 'DeepSeek error: ' + err}), 500
        answer = rr.json()['choices'][0]['message']['content'].strip()
        return jsonify({'ok': True, 'answer': answer, 'candidate_count': len(c_lines),
                        'truncated': truncated})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/tags/<tag_type>')
@login_required
def get_tag_suggestions(tag_type):
    """Return distinct tags used across all candidates for autocomplete.
    tag_type: 'product' (Product Handles) or 'function' (Function)
    """
    col_map = {'product': 'product_handles', 'function': 'function_tags', 'status': 'status_tags'}
    col = col_map.get(tag_type)
    if not col:
        return jsonify({'ok': False, 'error': 'Invalid tag type'}), 400

    defaults_map = {
        'product': [
            'Electrical Wires & Cables', 'Switches & Sockets', 'Low Voltage Switchgear',
            'Medium Voltage Switchgear', 'Circuit Breakers', 'Distribution Boards',
            'Lighting', 'Solar Inverters', 'Solar Panels', 'Energy Storage / BESS',
            'Transformers', 'Motors & Drives', 'Automation & Controls', 'Cable Management',
            'HVAC', 'Building Management Systems', 'EV Charging', 'Renewable Energy',
            'Industrial Automation', 'Power Distribution', 'Wiring Devices', 'MCCBs', 'ACBs',
            'Busbar Systems', 'UPS Systems', 'Genset / DG Sets'
        ],
        'function': [
            'Sales', 'Marketing', 'Business Development', 'Channel Sales',
            'Key Account Management', 'Product Management', 'Pre-Sales',
            'Technical Sales', 'Operations', 'Supply Chain', 'Procurement',
            'Project Management', 'Engineering', 'R&D', 'Quality',
            'Service & Support', 'After Sales Service', 'Finance', 'HR',
            'General Management', 'Strategy', 'Application Engineering', 'Design'
        ]
    }

    conn = get_db()
    rows = conn.execute(f'SELECT {col} FROM candidates WHERE owner_id=? AND {col} IS NOT NULL AND {col} != "" AND {col} != "[]"', (effective_company_id(),)).fetchall()
    conn.close()

    used = set()
    for r in rows:
        try:
            for t in json.loads(r[col] or '[]'):
                if t and t.strip():
                    used.add(t.strip())
        except Exception:
            pass

    all_tags = sorted(used) + [d for d in defaults_map.get(tag_type, []) if d not in used]
    return jsonify({'ok': True, 'tags': all_tags})


@app.route('/api/candidates/<int:cid>/tags', methods=['POST'])
@login_required
def save_candidate_tags(cid):
    """Save Product Handles or Function tags for a candidate."""
    d = request.json or {}
    tag_type = d.get('tag_type')
    tags = d.get('tags', [])
    col_map = {'product': 'product_handles', 'function': 'function_tags', 'status': 'status_tags'}
    col = col_map.get(tag_type)
    if not col:
        return jsonify({'ok': False, 'error': 'Invalid tag type'}), 400
    if not isinstance(tags, list):
        return jsonify({'ok': False, 'error': 'tags must be a list'}), 400
    tags = [str(t).strip() for t in tags if str(t).strip()]
    conn = get_db()
    if not _tenant_owns_candidate(conn, cid):
        conn.close(); return jsonify({'ok': False, 'error': 'Not found'}), 404
    conn.execute(f'UPDATE candidates SET {col}=?, updated_at=? WHERE id=?', (json.dumps(tags), ts(), cid))
    conn.commit(); conn.close()
    if tag_type == 'status' and tags:
        log_candidate_event(cid, 'tag', 'Tag updated — ' + ', '.join(tags))
    return jsonify({'ok': True, 'tags': tags})

@app.route('/api/health')
def health():
    return jsonify({'ok': True, 'data_dir': DATA_DIR, 'db': DB_PATH})

# Settings
@app.route('/api/settings', methods=['GET'])
def get_settings():
    conn = get_db()
    rows = conn.execute('SELECT key,value FROM settings').fetchall()
    out = {r['key']: r['value'] for r in rows}
    # Overlay this company's own per-tenant settings on top of the global defaults.
    cid = _safe_company_id()
    if cid:
        trows = conn.execute('SELECT key,value FROM tenant_settings WHERE company_id=?', (cid,)).fetchall()
        for r in trows:
            out[r['key']] = r['value']
        # company_name must be THIS tenant's, never the platform/global one.
        try:
            crow = conn.execute('SELECT name FROM companies WHERE id=?', (cid,)).fetchone()
            if crow and crow['name']:
                out['company_name'] = crow['name']
        except Exception:
            pass
    conn.close()
    # SECURITY: never expose platform secrets (AI keys, passwords, tokens) to any tenant UI.
    def _is_secret(k):
        kl = (k or '').lower()
        if kl.endswith('_api_key') or kl.endswith('_apikey') or kl.endswith('api_key'):
            return True
        if 'secret' in kl or 'password' in kl or kl.endswith('_pass') or kl.endswith('_token'):
            return True
        return kl in ('claude_api_key','deepseek_api_key','groq_api_key','openai_api_key',
                      'anthropic_api_key','smtp_pass','smtp_password','wa_token',
                      'wa_access_token','verify_token','flask_secret_key','secret_key')
    for k in list(out.keys()):
        if _is_secret(k):
            # Tell the UI whether a key is configured, without leaking its value.
            out[k] = '__set__' if (out.get(k) or '').strip() else ''
    # Ensure workflow_mode is always present so the UI can branch on it.
    if 'workflow_mode' not in out or not out.get('workflow_mode'):
        out['workflow_mode'] = 'agency'
    return jsonify(out)

@app.route('/api/settings', methods=['POST'])
@login_required
def save_settings():
    admin = is_admin()
    for k, v in (request.json or {}).items():
        # Global billing config can only be changed by the platform owner,
        # so an agency cannot set its own price/GST.
        if k.startswith('billing_') and not admin:
            continue
        set_setting(k, v)   # routes per-tenant keys automatically
    return jsonify({'ok': True})

# Mandates
@app.route('/api/mandates', methods=['GET'])
@login_required
def list_mandates():
    conn = get_db()
    if is_company_admin():
        rows = conn.execute('SELECT * FROM mandates WHERE owner_id=? ORDER BY created_at DESC',
                            (effective_company_id(),)).fetchall()
    else:
        # Recruiter: only mandates assigned to them within their company.
        rows = conn.execute('SELECT * FROM mandates WHERE owner_id=? AND assigned_user_id=? ORDER BY created_at DESC',
                            (effective_company_id(), real_user_id())).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/mandates/<int:mid>/client-notes', methods=['GET'])
@login_required
def get_client_notes(mid):
    conn = get_db()
    m = conn.execute('SELECT owner_id FROM mandates WHERE id=?', (mid,)).fetchone()
    if not m or m['owner_id'] != effective_company_id():
        conn.close(); return jsonify({'error': 'Mandate not found'}), 404
    rows = conn.execute(
        'SELECT n.id, n.note, n.created_at, n.created_by, u.display_name AS author '
        'FROM mandate_client_notes n LEFT JOIN users u ON u.id=n.created_by '
        'WHERE n.mandate_id=? AND n.is_active=1 ORDER BY n.created_at DESC',
        (mid,)).fetchall()
    conn.close()
    return jsonify({'ok': True, 'notes': [dict(r) for r in rows]})


@app.route('/api/mandates/<int:mid>/client-notes', methods=['POST'])
@login_required
def add_client_note(mid):
    d = request.json or {}
    note = (d.get('note') or '').strip()
    if not note:
        return jsonify({'error': 'Note text required'}), 400
    conn = get_db()
    m = conn.execute('SELECT owner_id FROM mandates WHERE id=?', (mid,)).fetchone()
    if not m or m['owner_id'] != effective_company_id():
        conn.close(); return jsonify({'error': 'Mandate not found'}), 404
    conn.execute(
        'INSERT INTO mandate_client_notes (mandate_id, owner_id, note, created_by, created_at, is_active) '
        'VALUES (?,?,?,?,?,1)',
        (mid, effective_company_id(), note, real_user_id(), ts()))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/mandates/<int:mid>/client-notes/<int:nid>', methods=['DELETE'])
@login_required
def delete_client_note(mid, nid):
    conn = get_db()
    m = conn.execute('SELECT owner_id FROM mandates WHERE id=?', (mid,)).fetchone()
    if not m or m['owner_id'] != effective_company_id():
        conn.close(); return jsonify({'error': 'Mandate not found'}), 404
    conn.execute('UPDATE mandate_client_notes SET is_active=0 WHERE id=? AND mandate_id=?', (nid, mid))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/mandates', methods=['POST'])
@login_required
def create_mandate():
    d = request.json or {}
    if not d.get('client') or not d.get('role'):
        return jsonify({'error': 'Client and Role required'}), 400
    conn = get_db(); c = conn.cursor()
    # Resolve CRM client link (Option B). If a crm_client_id is passed, use it.
    # Otherwise try to auto-match by normalised name so existing CRM clients link up.
    crm_client_id = int(d.get('crm_client_id') or 0)
    if not crm_client_id and d.get('client'):
        try:
            crm_client_id = _match_crm_client_by_name(conn, effective_company_id(), d['client'])
        except Exception:
            crm_client_id = 0
    c.execute('INSERT INTO mandates (client,role,location,division,ctc_min,ctc_max,jd,status,created_at,owner_id,assigned_user_id,crm_client_id) VALUES (?,?,?,?,?,?,?,?,?,?,?,?)',
              (d['client'], d['role'], d.get('location',''), d.get('division',''),
               float(d.get('ctc_min', 0)), float(d.get('ctc_max', 0)), d.get('jd',''), 'active', ts(), effective_company_id(), real_user_id(), crm_client_id))
    mid = c.lastrowid; conn.commit(); conn.close()
    log_activity('create_mandate', d['role'] + ' @ ' + d['client'])
    return jsonify({'ok': True, 'id': mid, 'crm_client_id': crm_client_id})


def _norm_client_name(name):
    """Normalise a company name for fuzzy matching (mirror of CRM logic)."""
    import re as _re
    s = (name or '').lower().strip()
    # strip common suffixes and punctuation
    s = _re.sub(r'[.,&\-/()]', ' ', s)
    for suffix in ['private limited', 'pvt ltd', 'pvt. ltd', 'private ltd',
                   'limited', 'ltd', 'llp', 'inc', 'incorporated', 'corporation',
                   'corp', 'technologies', 'technology', 'solutions', 'services',
                   'india', 'pvt', 'and', 'company', 'co']:
        s = _re.sub(r'\b' + _re.escape(suffix) + r'\b', ' ', s)
    s = _re.sub(r'\s+', ' ', s).strip()
    return s


def _match_crm_client_by_name(conn, company_id, client_name):
    """Return crm_clients.id whose normalised name matches, else 0."""
    target = _norm_client_name(client_name)
    if not target:
        return 0
    rows = conn.execute(
        'SELECT id, name FROM crm_clients WHERE company_id=? AND is_active=1',
        (company_id,)).fetchall()
    for r in rows:
        if _norm_client_name(r['name']) == target:
            return r['id']
    return 0


@app.route('/api/mandates/<int:mid>/link-client', methods=['POST'])
@login_required
def link_mandate_client(mid):
    """Manually set (or clear) the CRM client a mandate belongs to."""
    d = request.json or {}
    crm_client_id = int(d.get('crm_client_id') or 0)
    conn = get_db()
    m = conn.execute('SELECT id, owner_id FROM mandates WHERE id=?', (mid,)).fetchone()
    if not m or m['owner_id'] != effective_company_id():
        conn.close(); return jsonify({'error': 'Mandate not found'}), 404
    # Validate the client belongs to this tenant (if non-zero)
    if crm_client_id:
        cl = conn.execute('SELECT id FROM crm_clients WHERE id=? AND company_id=? AND is_active=1',
                          (crm_client_id, effective_company_id())).fetchone()
        if not cl:
            conn.close(); return jsonify({'error': 'CRM client not found'}), 404
    conn.execute('UPDATE mandates SET crm_client_id=? WHERE id=?', (crm_client_id, mid))
    conn.commit(); conn.close()
    return jsonify({'ok': True, 'crm_client_id': crm_client_id})


@app.route('/api/crm-link/auto-map', methods=['POST'])
@login_required
def auto_map_mandates():
    """One-click: link every unlinked mandate to a CRM client by name match.
    Returns how many were mapped and which couldn't be matched."""
    conn = get_db()
    company_id = effective_company_id()
    mandates = conn.execute(
        'SELECT id, client, crm_client_id FROM mandates WHERE owner_id=?',
        (company_id,)).fetchall()
    mapped, unmatched = 0, []
    for m in mandates:
        if m['crm_client_id']:
            continue  # already linked
        cid = _match_crm_client_by_name(conn, company_id, m['client'])
        if cid:
            conn.execute('UPDATE mandates SET crm_client_id=? WHERE id=?', (cid, m['id']))
            mapped += 1
        else:
            if m['client'] and m['client'] not in unmatched:
                unmatched.append(m['client'])
    conn.commit(); conn.close()
    return jsonify({'ok': True, 'mapped': mapped, 'unmatched': unmatched})


@app.route('/api/crm-link/client/<int:crm_client_id>/mandates', methods=['GET'])
@login_required
def client_mandates(crm_client_id):
    """List all mandates linked to a CRM client (for the client detail page)."""
    conn = get_db()
    company_id = effective_company_id()
    rows = conn.execute(
        'SELECT id, role, location, status, created_at FROM mandates '
        'WHERE owner_id=? AND crm_client_id=? ORDER BY created_at DESC',
        (company_id, crm_client_id)).fetchall()
    # Attach candidate counts
    out = []
    for r in rows:
        d = dict(r)
        d['candidate_count'] = conn.execute(
            'SELECT COUNT(*) n FROM candidates WHERE mandate_id=?', (r['id'],)).fetchone()['n']
        out.append(d)
    conn.close()
    return jsonify({'ok': True, 'mandates': out})

# ══════════════════════════════════════════════════════════════════════════
#  CRM — Email thread + AI-assisted drafting (Sprint 1)
#  Emails live as crm_activities rows (activity_type='email'); the meta JSON
#  carries direction / to / from / sent so the thread can render richly.
#  All three endpoints are additive and only read/write the existing table.
# ══════════════════════════════════════════════════════════════════════════
def _crm_client_or_none(conn, company_id, client_id):
    return conn.execute(
        'SELECT id, name, industry, website FROM crm_clients '
        'WHERE id=? AND company_id=? AND is_active=1',
        (client_id, company_id)).fetchone()

@app.route('/api/crm/clients/<int:cid>/emails', methods=['GET'])
@login_required
def crm_client_emails(cid):
    """Return the email thread for a client (optionally scoped to one contact)."""
    conn = get_db()
    company_id = effective_company_id()
    client = _crm_client_or_none(conn, company_id, cid)
    if not client:
        conn.close()
        return jsonify({'error': 'Client not found'}), 404
    contact_id = request.args.get('contact_id', type=int)
    sql = ("SELECT * FROM crm_activities WHERE company_id=? AND client_id=? "
           "AND activity_type='email' AND is_active=1")
    params = [company_id, cid]
    if contact_id:
        sql += ' AND contact_id=?'
        params.append(contact_id)
    sql += ' ORDER BY created_at ASC, id ASC'
    rows = conn.execute(sql, params).fetchall()
    conn.close()
    out = []
    for r in rows:
        d = dict(r)
        try:
            meta = json.loads(d.get('meta') or '{}')
        except Exception:
            meta = {}
        out.append({
            'id': d['id'], 'contact_id': d['contact_id'],
            'subject': d.get('subject', ''), 'body': d.get('body', ''),
            'created_at': d.get('created_at', ''),
            'direction': meta.get('direction', 'outgoing'),
            'to': meta.get('to', ''), 'from': meta.get('from', ''),
            'sent': bool(meta.get('sent', False)),
        })
    return jsonify({'ok': True, 'emails': out})


@app.route('/api/crm/clients/<int:cid>/email', methods=['POST'])
@login_required
def crm_client_email_send(cid):
    """Log an outgoing email on the client/contact timeline, optionally sending
    it via the tenant's configured SMTP. Stored as a crm_activities email row."""
    conn = get_db()
    company_id = effective_company_id()
    client = _crm_client_or_none(conn, company_id, cid)
    if not client:
        conn.close()
        return jsonify({'error': 'Client not found'}), 404
    d = request.get_json(force=True) or {}
    to_email = (d.get('to') or '').strip()
    subject = (d.get('subject') or '').strip()
    body = (d.get('body') or '').strip()
    contact_id = int(d.get('contact_id') or 0)
    do_send = bool(d.get('send'))
    if not subject and not body:
        conn.close()
        return jsonify({'error': 'Subject or body is required'}), 400

    sent_ok = False
    if do_send:
        if not to_email:
            conn.close()
            return jsonify({'error': 'A "To" email address is required to send.'}), 400
        sent_ok, send_err = _smtp_send(to_email, subject, body)
        if not sent_ok:
            conn.close()
            return jsonify({'error': send_err or 'Failed to send email'}), 400

    now = ts()
    uid = real_user_id()
    meta = json.dumps({'direction': 'outgoing', 'to': to_email,
                       'from': get_setting('smtp_email', ''), 'sent': sent_ok})
    conn.execute(
        "INSERT INTO crm_activities (company_id,client_id,contact_id,activity_type,"
        "subject,body,outcome,due_at,completed_at,status,owner_user_id,meta,"
        "is_active,created_by,updated_by,created_at,updated_at) "
        "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (company_id, cid, contact_id, 'email', subject, body, '', '', now,
         'done', uid, meta, 1, uid, uid, now, now))
    new_id = conn.execute('SELECT last_insert_rowid() AS id').fetchone()['id']
    conn.commit()
    conn.close()
    try:
        log_activity('email_' + ('sent' if sent_ok else 'logged'),
                     detail=subject, entity_type='crm_client', entity_id=cid)
    except Exception:
        pass
    return jsonify({'ok': True, 'id': new_id, 'sent': sent_ok})


@app.route('/api/crm/ai/draft-email', methods=['POST'])
@login_required
def crm_ai_draft_email():
    """AI-assisted email drafting for the CRM (tenant's DeepSeek key).
    Returns a suggested subject + body from the client/contact context and the
    recruiter's intent. Nothing is sent or stored here — draft only."""
    conn = get_db()
    company_id = effective_company_id()
    d = request.get_json(force=True) or {}
    client_id = int(d.get('client_id') or 0)
    contact_id = int(d.get('contact_id') or 0)
    intent = (d.get('intent') or '').strip()
    tone = (d.get('tone') or 'professional').strip()
    if not intent:
        conn.close()
        return jsonify({'error': 'Please describe what the email should say.'}), 400

    client = _crm_client_or_none(conn, company_id, client_id) if client_id else None
    contact = None
    if contact_id:
        contact = conn.execute(
            'SELECT name, designation, email FROM crm_contacts '
            'WHERE id=? AND company_id=? AND is_active=1',
            (contact_id, company_id)).fetchone()
    thread_ctx = ''
    if client_id:
        prev = conn.execute(
            "SELECT subject, body FROM crm_activities WHERE company_id=? AND client_id=? "
            "AND activity_type='email' AND is_active=1 ORDER BY id DESC LIMIT 3",
            (company_id, client_id)).fetchall()
        if prev:
            thread_ctx = '\n\n'.join(
                '- ' + (p['subject'] or '') + ': ' + (p['body'] or '')[:300] for p in prev)
    conn.close()

    ds_key = get_setting('deepseek_api_key')
    if not ds_key:
        return jsonify({'error': 'DeepSeek API key not configured. Add it in Settings.'}), 400

    sender_name = get_setting('smtp_display_name', '') or 'the recruitment team'
    ctx_lines = []
    if client:
        ctx_lines.append('Client company: ' + (client['name'] or '') +
                         ((' (' + client['industry'] + ')') if client['industry'] else ''))
    if contact:
        ctx_lines.append('Recipient: ' + (contact['name'] or '') +
                         ((', ' + contact['designation']) if contact['designation'] else ''))
    ctx = '\n'.join(ctx_lines) or 'A business client contact.'

    system_msg = (
        "You are an assistant that writes clear, concise, professional B2B "
        "recruitment emails for an Indian staffing agency. Write in natural "
        "business English. Keep it short and specific. Do NOT invent facts, "
        "names, numbers, or commitments that were not provided. "
        "Return ONLY a JSON object of the form {\"subject\": \"...\", \"body\": \"...\"}. "
        "The body must be ready to send and signed off as " + sender_name + "."
    )
    user_msg = (
        'Context:\n' + ctx + '\n\n'
        'Desired tone: ' + tone + '\n'
        'What the email should accomplish:\n' + intent
        + (('\n\nRecent email history for reference:\n' + thread_ctx) if thread_ctx else '')
    )
    try:
        resp = call_deepseek(ds_key,
            {'model': 'deepseek-chat', 'temperature': 0.5, 'max_tokens': 700,
             'messages': [{'role': 'system', 'content': system_msg},
                          {'role': 'user', 'content': user_msg[:8000]}]},
            timeout=60, endpoint='crm_email_draft')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    if resp.status_code != 200:
        try: err = resp.json().get('error', {}).get('message', 'DeepSeek API error')
        except Exception: err = resp.text[:200]
        return jsonify({'error': err}), 500
    content = resp.json()['choices'][0]['message']['content'].strip()
    parsed = parse_json(content) or {}
    subject = (parsed.get('subject') or '').strip()
    body = (parsed.get('body') or '').strip()
    if not body:
        body = content
    return jsonify({'ok': True, 'subject': subject, 'body': body})


# ══════════════════════════════════════════════════════════════════════════
#  CRM — Attachments (Sprint 2). Files attached to a company or a contact,
#  stored on the persistent disk (CRM_FILES_DIR) and tracked in crm_attachments.
#  Additive; reuses the same storage approach as candidate CVs.
# ══════════════════════════════════════════════════════════════════════════
CRM_ATT_CATEGORIES = {'NDA', 'Agreement', 'PO', 'Requirement', 'Proposal',
                      'Presentation', 'Invoice', 'Resume', 'BusinessCard', 'Other'}
ALLOWED_ATT_EXT = {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx',
                   '.png', '.jpg', '.jpeg', '.txt', '.csv', '.zip'}

def _crm_resolve_client_id(conn, company_id, entity_type, entity_id):
    """Every attachment is anchored to a client_id (a contact resolves to its
    parent client). Returns 0 if the entity does not belong to the tenant."""
    if entity_type == 'contact':
        row = conn.execute(
            'SELECT client_id FROM crm_contacts WHERE id=? AND company_id=? AND is_active=1',
            (entity_id, company_id)).fetchone()
        return row['client_id'] if row else 0
    row = conn.execute(
        'SELECT id FROM crm_clients WHERE id=? AND company_id=? AND is_active=1',
        (entity_id, company_id)).fetchone()
    return row['id'] if row else 0

def _att_public(r):
    d = dict(r)
    d.pop('stored_name', None)  # never expose the on-disk filename
    return d

@app.route('/api/crm/<entity_type>/<int:entity_id>/attachments', methods=['GET'])
@login_required
def crm_list_attachments(entity_type, entity_id):
    if entity_type not in ('client', 'contact'):
        return jsonify({'error': 'Invalid entity type'}), 400
    conn = get_db()
    company_id = effective_company_id()
    rows = conn.execute(
        'SELECT * FROM crm_attachments WHERE company_id=? AND entity_type=? AND entity_id=? '
        'AND is_active=1 ORDER BY created_at DESC, id DESC',
        (company_id, entity_type, entity_id)).fetchall()
    conn.close()
    return jsonify({'ok': True, 'attachments': [_att_public(r) for r in rows]})

@app.route('/api/crm/<entity_type>/<int:entity_id>/attachments', methods=['POST'])
@login_required
def crm_upload_attachment(entity_type, entity_id):
    if entity_type not in ('client', 'contact'):
        return jsonify({'error': 'Invalid entity type'}), 400
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    f = request.files['file']
    if not f or not f.filename:
        return jsonify({'error': 'No file selected'}), 400
    ext = Path(f.filename).suffix.lower()
    if ext not in ALLOWED_ATT_EXT:
        return jsonify({'error': 'Unsupported file type: ' + (ext or 'unknown')}), 400
    category = (request.form.get('category') or 'Other').strip()
    if category not in CRM_ATT_CATEGORIES:
        category = 'Other'
    conn = get_db()
    company_id = effective_company_id()
    client_id = _crm_resolve_client_id(conn, company_id, entity_type, entity_id)
    if not client_id:
        conn.close()
        return jsonify({'error': 'Record not found'}), 404
    now = ts()
    uid = real_user_id()
    stored = ('att_' + entity_type + str(entity_id) + '_'
              + datetime.datetime.now().strftime('%Y%m%d%H%M%S%f') + ext)
    try:
        f.save(os.path.join(CRM_FILES_DIR, stored))
    except Exception as e:
        conn.close()
        return jsonify({'error': 'Failed to save file: ' + str(e)}), 500
    try:
        size = os.path.getsize(os.path.join(CRM_FILES_DIR, stored))
    except Exception:
        size = 0
    conn.execute(
        'INSERT INTO crm_attachments (company_id,client_id,entity_type,entity_id,category,'
        'original_name,stored_name,size_bytes,mime,uploaded_by,is_active,created_at) '
        'VALUES (?,?,?,?,?,?,?,?,?,?,?,?)',
        (company_id, client_id, entity_type, entity_id, category, f.filename, stored,
         size, f.mimetype or '', uid, 1, now))
    aid = conn.execute('SELECT last_insert_rowid() AS id').fetchone()['id']
    conn.commit()
    row = conn.execute('SELECT * FROM crm_attachments WHERE id=?', (aid,)).fetchone()
    conn.close()
    try:
        log_activity('attachment_added', detail=category + ': ' + f.filename,
                     entity_type='crm_' + entity_type, entity_id=entity_id)
    except Exception:
        pass
    return jsonify({'ok': True, 'attachment': _att_public(row)})

@app.route('/api/crm/attachments/<int:aid>/download', methods=['GET'])
@login_required
def crm_download_attachment(aid):
    conn = get_db()
    company_id = effective_company_id()
    r = conn.execute('SELECT * FROM crm_attachments WHERE id=? AND company_id=? AND is_active=1',
                     (aid, company_id)).fetchone()
    conn.close()
    if not r:
        return jsonify({'error': 'Not found'}), 404
    fp = os.path.join(CRM_FILES_DIR, r['stored_name'])
    if not os.path.exists(fp):
        return jsonify({'error': 'File missing on disk'}), 404
    try:
        return send_file(fp, as_attachment=True,
                         download_name=(r['original_name'] or r['stored_name']))
    except TypeError:
        # Older Flask uses attachment_filename instead of download_name
        return send_file(fp, as_attachment=True,
                         attachment_filename=(r['original_name'] or r['stored_name']))

@app.route('/api/crm/attachments/<int:aid>', methods=['DELETE'])
@login_required
def crm_delete_attachment(aid):
    conn = get_db()
    company_id = effective_company_id()
    r = conn.execute('SELECT * FROM crm_attachments WHERE id=? AND company_id=? AND is_active=1',
                     (aid, company_id)).fetchone()
    if not r:
        conn.close()
        return jsonify({'error': 'Not found'}), 404
    conn.execute('UPDATE crm_attachments SET is_active=0 WHERE id=?', (aid,))
    conn.commit()
    conn.close()
    try:
        log_activity('attachment_removed', detail=r['original_name'],
                     entity_type='crm_' + r['entity_type'], entity_id=r['entity_id'])
    except Exception:
        pass
    return jsonify({'ok': True})


@app.route('/api/mandates/<int:mid>', methods=['GET'])
@login_required
def get_mandate(mid):
    conn = get_db()
    r = conn.execute('SELECT * FROM mandates WHERE id=? AND owner_id=?', (mid, effective_company_id())).fetchone()
    conn.close()
    if not r:
        return jsonify({'error': 'Not found'}), 404
    # Recruiters can only open mandates assigned to them.
    if not is_company_admin() and r['assigned_user_id'] != real_user_id():
        return jsonify({'error': 'Not found'}), 404
    return jsonify(dict(r))


@app.route('/api/my-profile', methods=['GET'])
@login_required
def get_my_profile():
    u = current_user()
    return jsonify({'ok': True, 'profile': {
        'display_name': u.get('display_name', ''),
        'profile_phone': u.get('profile_phone', ''),
        'profile_designation': u.get('profile_designation', ''),
        'profile_email': u.get('profile_email', ''),
    }})

@app.route('/api/my-profile', methods=['POST'])
@login_required
def update_my_profile():
    d = request.json or {}
    conn = get_db()
    uid = session.get('user_id')
    for field in ('display_name', 'profile_phone', 'profile_designation', 'profile_email'):
        if field in d:
            conn.execute(f'UPDATE users SET {field}=? WHERE id=?', (d[field], uid))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/my-team', methods=['GET'])
@login_required
def my_team():
    """Recruiters in the current company (for the assign-to dropdown).
    Company-admin only."""
    if not is_company_admin():
        return jsonify({'error': 'Not allowed'}), 403
    conn = get_db()
    rows = conn.execute('''SELECT id, username, display_name, is_company_admin
                           FROM users WHERE company_id=? AND status='approved'
                           ORDER BY is_company_admin DESC, id''',
                        (effective_company_id(),)).fetchall()
    conn.close()
    return jsonify({'ok': True, 'team': [dict(r) for r in rows]})


@app.route('/api/mandates/<int:mid>/assign', methods=['POST'])
@login_required
def assign_mandate(mid):
    """Company-admin assigns a mandate to a recruiter in the same company."""
    if not is_company_admin():
        return jsonify({'error': 'Only an admin can assign mandates'}), 403
    d = request.json or {}
    target = d.get('user_id')
    conn = get_db()
    m = conn.execute('SELECT id, role, client FROM mandates WHERE id=? AND owner_id=?',
                     (mid, effective_company_id())).fetchone()
    if not m:
        conn.close(); return jsonify({'error': 'Mandate not found'}), 404
    # Target must be a user in the same company.
    u = conn.execute('SELECT id, display_name, username FROM users WHERE id=? AND company_id=?',
                     (target, effective_company_id())).fetchone()
    if not u:
        conn.close(); return jsonify({'error': 'Recruiter not in your company'}), 400
    conn.execute('UPDATE mandates SET assigned_user_id=? WHERE id=?', (target, mid))
    conn.commit(); conn.close()
    log_activity('assign_mandate', f"{m['role']} @ {m['client']} → {u['display_name'] or u['username']}")
    return jsonify({'ok': True})


@app.route('/api/mandates/<int:mid>', methods=['DELETE'])
@login_required
def delete_mandate(mid):
    """Company-admin deletes a job (mandate). Its candidates are NOT deleted —
    they are detached and kept in the company's Central Database."""
    if not is_company_admin():
        return jsonify({'error': 'Only an admin can delete a job'}), 403
    conn = get_db()
    m = conn.execute('SELECT id, role, client FROM mandates WHERE id=? AND owner_id=?',
                     (mid, effective_company_id())).fetchone()
    if not m:
        conn.close(); return jsonify({'error': 'Mandate not found'}), 404
    # Move this mandate's candidates to the company's central pool so they
    # survive in the Central Database (owner_id already = company, so they stay
    # visible there). We only repoint mandate_id to avoid a dangling reference.
    central_mid = get_or_create_central_mandate()
    kept = conn.execute('SELECT COUNT(*) n FROM candidates WHERE mandate_id=? AND owner_id=?',
                        (mid, effective_company_id())).fetchone()['n']
    conn.execute('UPDATE candidates SET mandate_id=? WHERE mandate_id=? AND owner_id=?',
                 (central_mid, mid, effective_company_id()))
    conn.execute('DELETE FROM mandates WHERE id=?', (mid,))
    conn.commit(); conn.close()
    log_activity('delete_mandate', f"{m['role']} @ {m['client']} (kept {kept} candidates in Central DB)")
    return jsonify({'ok': True, 'candidates_kept': kept})


@app.route('/api/mandates/<int:mid>', methods=['PUT'])
@login_required
def update_mandate(mid):
    d = request.json or {}
    conn = get_db()
    own = conn.execute('SELECT owner_id FROM mandates WHERE id=?', (mid,)).fetchone()
    if not own or own['owner_id'] != effective_user_id():
        conn.close(); return jsonify({'error': 'Not found'}), 404
    conn.execute('UPDATE mandates SET client=?,role=?,location=?,division=?,ctc_min=?,ctc_max=?,experience=?,jd=?,status=? WHERE id=?',
                 (d.get('client',''), d.get('role',''), d.get('location',''), d.get('division',''),
                  float(d.get('ctc_min', 0)), float(d.get('ctc_max', 0)), d.get('experience',''), d.get('jd',''), d.get('status','active'), mid))
    conn.commit(); conn.close()
    return jsonify({'ok': True})

@app.route('/api/mandates/<int:mid>/candidates')
@login_required
def list_candidates(mid):
    check_timers()
    conn = get_db()
    if not _tenant_owns_mandate(conn, mid):
        conn.close(); return jsonify({'error': 'Not found'}), 404
    rows = conn.execute('SELECT * FROM candidates WHERE mandate_id=? ORDER BY created_at DESC', (mid,)).fetchall()
    conn.close()
    out = []
    for r in rows:
        d = _cand_public(r)   # drop embedding / embedding_text / embedding_vec
        try: d['key_skills'] = json.loads(d['key_skills'] or '[]')
        except: d['key_skills'] = []
        try: d['secondary_skills'] = json.loads(d['secondary_skills'] or '[]')
        except: d['secondary_skills'] = []
        out.append(d)
    return jsonify(out)

@app.route('/api/candidates/<int:cid>/ai-compose', methods=['POST'])
@login_required
def ai_compose_email(cid):
    """Use DeepSeek to draft/refine an email for a candidate based on user's command."""
    d = request.json or {}
    command = (d.get('command') or '').strip()
    context = (d.get('context') or '').strip()
    current_draft = (d.get('current_draft') or '').strip()
    if not command:
        return jsonify({'error': 'Please give a command'}), 400
    key = get_setting('deepseek_api_key', '') or os.environ.get('DEEPSEEK_API_KEY', '')
    if not key:
        return jsonify({'error': 'DeepSeek API key not configured. Add it in Settings.'}), 400

    # Get candidate info for context
    conn = get_db()
    c = conn.execute('SELECT name,company,designation,email,ctc_current,experience,location,mandate_id FROM candidates WHERE id=?', (cid,)).fetchone()
    cand_info = dict(c) if c else {}

    # Get mandate JD for context
    jd_text = ''
    if cand_info.get('mandate_id'):
        mandate = conn.execute('SELECT client,role,jd,location,division,ctc_min,ctc_max FROM mandates WHERE id=?', (cand_info['mandate_id'],)).fetchone()
        if mandate:
            jd_text = mandate['jd'] or ''
            cand_info['mandate_role'] = mandate['role'] or ''
            cand_info['mandate_client'] = mandate['client'] or ''
            cand_info['mandate_location'] = mandate['location'] or ''
            cand_info['ctc_range'] = f"{mandate['ctc_min']}-{mandate['ctc_max']} LPA"

    # Get recruiter profile for signature
    u = current_user()
    recruiter_name = (u.get('display_name') or u.get('username') or '') if u else ''
    recruiter_phone = (u.get('profile_phone') or '') if u else ''
    recruiter_email_addr = (u.get('profile_email') or get_setting('smtp_email', '')) if u else ''
    recruiter_designation = (u.get('profile_designation') or '') if u else ''
    company_name = get_setting('company_name', '') or get_setting('recruiter_name', '')
    conn.close()

    signature_block = f"{recruiter_name}"
    if recruiter_designation: signature_block += f"\n{recruiter_designation}"
    if company_name: signature_block += f"\n{company_name}"
    if recruiter_phone: signature_block += f"\nPhone: {recruiter_phone}"
    if recruiter_email_addr: signature_block += f"\nEmail: {recruiter_email_addr}"

    system_prompt = f"""You are an expert Talent Acquisition and Recruitment Communication Specialist for an Indian Executive Search and Recruitment firm.
Your responsibility is to generate highly professional, personalized recruitment emails that encourage candidates to respond.
Always write naturally like an experienced recruiter, never like AI.
---------------------------------------------------
AVAILABLE DATA
Candidate Details
- Name: {cand_info.get('name','')}
- Current Company: {cand_info.get('company','')}
- Current Designation: {cand_info.get('designation','')}
- Experience: {cand_info.get('experience','')}
- Current Location: {cand_info.get('location','')}
Job Details
- Role: {cand_info.get('mandate_role','')}
- Hiring Company / Client: {cand_info.get('mandate_client','')}
- Job Location: {cand_info.get('mandate_location','')}
- Complete Job Description:
{jd_text or '(not provided)'}
Recruiter Signature
{signature_block}
{('Previous email context:\n' + context) if context else ''}
{('Current draft in compose box (improve or continue from this):\n' + current_draft) if current_draft else ''}
---------------------------------------------------
GENERAL WRITING STYLE
Use:
\u2022 Professional
\u2022 Warm
\u2022 Personalized
\u2022 Easy to read
\u2022 Natural recruiter language
\u2022 Indian business communication style
Avoid:
\u2022 Robotic writing
\u2022 Marketing language
\u2022 AI sounding text
\u2022 Over excitement
\u2022 Emoji
\u2022 ALL CAPS
---------------------------------------------------
EMAIL STRUCTURE
Start with:
Dear {cand_info.get('name','Candidate')},
Introduce yourself in 1-2 lines.
Briefly explain why you are reaching out.
Then generate the requested content.
Always end with:
Regards,
{signature_block}
---------------------------------------------------
IF COMMAND = "create JD"
Generate a complete recruitment email including:
1. Opening paragraph
Mention:
\u2022 Candidate's current role
\u2022 Current company (if available)
\u2022 Why the profile appears relevant
2. About the Opportunity
Short paragraph introducing:
\u2022 Role
\u2022 Client
\u2022 Location
3. Key Responsibilities
Use HTML unordered list.
Only include responsibilities that exist in the provided Job Description.
Do NOT invent responsibilities.
4. Desired Skills & Experience
Use HTML unordered list.
Extract only from JD.
5. Why Consider This Opportunity
Summarize important highlights from JD such as:
\u2022 Industry
\u2022 Technologies
\u2022 Growth
\u2022 Projects
\u2022 Team
\u2022 Leadership
\u2022 Exposure
Only if mentioned.
6. Closing Paragraph
Invite candidate to share:
\u2022 Updated Resume
\u2022 Availability
\u2022 Interest
---------------------------------------------------
IF COMMAND = "follow up"
Generate a short polite follow-up email.
Maximum 120 words.
Mention that you're checking whether the candidate had a chance to review the earlier email.
Invite them to respond if interested.
---------------------------------------------------
IMPORTANT RULES
Never mention:
\u2022 Salary
\u2022 CTC
\u2022 Compensation
\u2022 Budget
unless explicitly present in the prompt AND specifically requested.
Never fabricate information.
If any information is unavailable, simply omit it.
Never write placeholders like:
[Company]
[TBD]
Not Available
---------------------------------------------------
HTML FORMAT
Body must be valid HTML.
Allowed tags:
<p>
<strong>
<ul>
<li>
<br>
No CSS.
No tables.
---------------------------------------------------
SUBJECT LINE
Generate an engaging subject.
Examples:
Opportunity for Senior Electrical Engineer | Mumbai
Business Development Opportunity | Data Centre Industry
Exciting Career Opportunity \u2013 Project Sales | Delhi NCR
Do not use clickbait.
---------------------------------------------------
OUTPUT FORMAT
Return ONLY valid JSON.
{{"subject":"...","body":"<p>...</p>"}}
Do not include markdown.
Do not include explanations.
Do not include additional text outside JSON."""

    text = ''
    try:
        resp = call_deepseek(key, {
            'model': 'deepseek-chat',
            'max_tokens': 2400,
            'messages': [
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': command}
            ]
        }, timeout=150, endpoint='ai-compose')
        try:
            data = resp.json()
        except Exception:
            return jsonify({'error': 'AI service returned an invalid response. Check your DeepSeek API key in Settings.'}), 502
        if isinstance(data, dict) and data.get('error'):
            emsg = data['error'].get('message', 'Unknown error') if isinstance(data['error'], dict) else str(data['error'])
            return jsonify({'error': 'DeepSeek error: ' + emsg}), 502
        choices = data.get('choices') or []
        if not choices:
            return jsonify({'error': 'AI did not return any content. Please try again.'}), 502
        text = (choices[0].get('message', {}).get('content', '') or '').strip()
        text = text.replace('```json', '').replace('```', '').strip()
        subject = ''; body = ''; parsed = None
        try:
            parsed = json.loads(text)
        except Exception:
            try:
                s = text.find('{'); e = text.rfind('}')
                if s >= 0 and e > s:
                    parsed = json.loads(text[s:e+1])
            except Exception:
                parsed = None
        if isinstance(parsed, dict) and (parsed.get('body') or parsed.get('subject')):
            subject = parsed.get('subject', '') or ''
            body = parsed.get('body', '') or ''
        else:
            body = text  # not usable JSON — use whole text as body
        if not (body or '').strip():
            return jsonify({'error': 'AI returned an empty email. Please rephrase your command and try again.'}), 502
        return jsonify({'ok': True, 'subject': subject, 'body': body})
    except TokenCapError:
        return jsonify({'error': 'Monthly AI token cap reached.'}), 429
    except Exception as e:
        return jsonify({'error': f'AI compose failed: {str(e)}'}), 500


@app.route('/api/mandates/<int:mid>/submission-excel')
@login_required
def submission_excel(mid):
    """Generate a client-submission Excel for a mandate's candidates,
    matching the standard submission format (yellow bold headers, borders)."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    from openpyxl.utils import get_column_letter

    conn = get_db()
    m = conn.execute('SELECT role, client FROM mandates WHERE id=? AND owner_id=?',
                     (mid, effective_company_id())).fetchone()
    if not m:
        conn.close(); return jsonify({'error': 'Mandate not found'}), 404
    stage = (request.args.get('stage') or '').strip()
    if stage:
        rows = conn.execute(
            'SELECT * FROM candidates WHERE mandate_id=? AND stage=? ORDER BY name', (mid, stage)
        ).fetchall()
    else:
        rows = conn.execute(
            'SELECT * FROM candidates WHERE mandate_id=? ORDER BY name', (mid,)
        ).fetchall()
    conn.close()

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Submission'

    headers = ['Candidate Name', 'Contact Number', 'Email ID', 'Educational Qualification',
               'Current Company', 'Total Experience', 'Current CTC', 'Expected CTC',
               'Current Location', 'Preferred Location', 'Notice Period']
    widths = [26.5, 16, 31.5, 25.7, 31, 16.5, 27.7, 32.5, 16.3, 18.3, 28.3]

    header_fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
    header_font = Font(bold=True, size=10, color='222222')
    thin = Side(style='thin', color='BBBBBB')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal='left', vertical='center', wrap_text=True)

    for i, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=i, value=h)
        cell.fill = header_fill; cell.font = header_font
        cell.border = border; cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        ws.column_dimensions[get_column_letter(i)].width = widths[i-1]
    ws.row_dimensions[1].height = 28

    def fmt_exp(v):
        try:
            v = float(v or 0)
            return f"{int(v)} Years" if v == int(v) else f"{v} Years"
        except Exception:
            return ''
    def fmt_ctc(v):
        try:
            v = float(v or 0)
            return f"{int(v)} LPA" if v == int(v) else f"{v} LPA"
        except Exception:
            return ''
    def fmt_notice(v):
        try:
            v = int(v or 0)
            return f"{v} Days" if v else ''
        except Exception:
            return ''

    r = 2
    for c in rows:
        d = dict(c)
        vals = [
            d.get('name', ''), d.get('phone', ''), d.get('email', ''),
            d.get('qualification', ''), d.get('company', ''),
            fmt_exp(d.get('experience')), fmt_ctc(d.get('ctc_current')),
            fmt_ctc(d.get('ctc_expected')), d.get('location', ''),
            d.get('preferred_location', ''), fmt_notice(d.get('notice_period')),
        ]
        for i, val in enumerate(vals, 1):
            cell = ws.cell(row=r, column=i, value=val)
            cell.border = border; cell.alignment = center; cell.font = Font(size=10)
        r += 1

    bio = io.BytesIO()
    wb.save(bio); bio.seek(0)
    safe_role = re.sub(r'[^A-Za-z0-9_-]+', '_', (m['role'] or 'Submission'))[:40]
    safe_stage = ('_' + re.sub(r'[^A-Za-z0-9_-]+', '_', stage)) if stage else ''
    fname = f"Submission_{safe_role}{safe_stage}.xlsx"
    return send_file(bio, as_attachment=True, download_name=fname,
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


# ── Share to Client (submission email: table + resume attachments) ───────────
def _sub_esc(s):
    return (str(s if s is not None else '')
            .replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))

def _sub_fmt_exp(v):
    try:
        v = float(v or 0)
        if v <= 0: return ''
        return f"{int(v)} Years" if v == int(v) else f"{v} Years"
    except Exception:
        return ''

def _sub_fmt_ctc(v):
    try:
        v = float(v or 0)
        if v <= 0: return ''
        return f"{int(v)} LPA" if v == int(v) else f"{v} LPA"
    except Exception:
        return ''

def _sub_fmt_notice(v):
    try:
        v = int(v or 0)
        return f"{v} Days" if v else ''
    except Exception:
        return ''

# Columns for the submission table (matches the standard email format).
_SUB_COLS = ['Position', 'Position Location', 'Candidate Name', 'Contact Number',
             'Email ID', 'Educational Qualification', 'Current Company',
             'Total Experience', 'Current CTC', 'Expected CTC', 'Current Location',
             'Preferred Location', 'Notice Period']

def _sub_row_values(m, c):
    d = dict(c)
    return [
        m['role'] or '', m['location'] or '',
        d.get('name', ''), d.get('phone', ''), d.get('email', ''),
        d.get('qualification', ''), d.get('company', ''),
        _sub_fmt_exp(d.get('experience')), _sub_fmt_ctc(d.get('ctc_current')),
        _sub_fmt_ctc(d.get('ctc_expected')), d.get('location', ''),
        d.get('preferred_location', ''), _sub_fmt_notice(d.get('notice_period')),
    ]

def _submission_table_html(m, rows):
    th = ('<th style="background:#FFFF00;border:1px solid #999;padding:6px 8px;'
          'font-size:12px;font-weight:bold;text-align:left;color:#222">')
    td = '<td style="border:1px solid #bbb;padding:6px 8px;font-size:12px;color:#222">'
    head = ''.join(th + _sub_esc(h) + '</th>' for h in _SUB_COLS)
    body = ''
    for c in rows:
        vals = _sub_row_values(m, c)
        body += '<tr>' + ''.join(td + _sub_esc(v) + '</td>' for v in vals) + '</tr>'
    return ('<table style="border-collapse:collapse;border:1px solid #999;'
            'font-family:Calibri,Arial,sans-serif"><thead><tr>' + head
            + '</tr></thead><tbody>' + body + '</tbody></table>')

def _submission_signature_html():
    """Signature for the submission email. Uses the exact override setting if
    provided, else auto-builds from the recruiter's profile + company."""
    override = (get_setting('submission_signature', '') or '').strip()
    if override:
        return '<div style="font-family:Calibri,Arial,sans-serif;font-size:13px;color:#333;white-space:pre-wrap">' \
               + _sub_esc(override) + '</div>'
    u = current_user() or {}
    name = (u.get('display_name') or u.get('username') or get_setting('recruiter_name', '') or '').strip()
    desig = (u.get('profile_designation') or '').strip()
    company = (get_setting('company_name', '') or (u.get('company_name') or '')).strip()
    email = (u.get('profile_email') or get_setting('smtp_email', '') or '').strip()
    phone = (u.get('profile_phone') or '').strip()
    web = (get_setting('company_website', '') or '').strip()
    line2 = ' &nbsp;|&nbsp; '.join([x for x in [desig, company, 'Talent Acquisition'] if x])
    parts = ['<div style="font-family:Calibri,Arial,sans-serif;font-size:13px;color:#333;line-height:1.5;margin-top:14px">']
    if name: parts.append('<div style="font-weight:bold;color:#1a2a6c;font-size:14px">' + _sub_esc(name) + '</div>')
    if line2: parts.append('<div style="color:#666">' + line2 + '</div>')
    contact = []
    if email: contact.append('E: <a href="mailto:' + _sub_esc(email) + '">' + _sub_esc(email) + '</a>')
    if phone: contact.append('M: ' + _sub_esc(phone))
    if web: contact.append('W: ' + _sub_esc(web))
    if contact: parts.append('<div style="margin-top:4px;color:#444">' + ' &nbsp; '.join(contact) + '</div>')
    parts.append('</div>')
    return ''.join(parts)

def _mandate_client_contacts(conn, company_id, crm_client_id):
    """Return [{name, email}] for the mandate's linked CRM client (with emails)."""
    if not crm_client_id:
        return []
    try:
        rows = conn.execute(
            'SELECT name, email FROM crm_contacts '
            'WHERE client_id=? AND company_id=? AND is_active=1 AND email != "" '
            'ORDER BY id ASC', (crm_client_id, company_id)).fetchall()
        return [{'name': r['name'] or '', 'email': r['email'] or ''} for r in rows if (r['email'] or '').strip()]
    except Exception:
        return []


@app.route('/api/mandates/<int:mid>/spoc-options', methods=['GET'])
@login_required
def mandate_spoc_options(mid):
    """Contacts of the mandate's linked CRM client (for choosing a SPOC), plus
    the currently selected SPOC and whether the client is CRM-linked."""
    conn = get_db()
    company_id = effective_company_id()
    m = conn.execute('SELECT * FROM mandates WHERE id=? AND owner_id=?',
                     (mid, company_id)).fetchone()
    if not m:
        conn.close()
        return jsonify({'error': 'Mandate not found'}), 404
    crm_id = m['crm_client_id'] or 0
    contacts = []
    if crm_id:
        try:
            rows = conn.execute(
                'SELECT id, name, designation, email FROM crm_contacts '
                'WHERE client_id=? AND company_id=? AND is_active=1 ORDER BY id ASC',
                (crm_id, company_id)).fetchall()
            contacts = [{'id': r['id'], 'name': r['name'] or '',
                         'designation': r['designation'] or '', 'email': r['email'] or ''}
                        for r in rows]
        except Exception:
            contacts = []
    try:
        spoc_id = m['spoc_contact_id'] or 0
    except Exception:
        spoc_id = 0
    conn.close()
    return jsonify({'ok': True, 'crm_linked': bool(crm_id), 'crm_client_id': crm_id,
                    'client': m['client'], 'spoc_contact_id': spoc_id, 'contacts': contacts})


@app.route('/api/mandates/<int:mid>/spoc', methods=['POST'])
@login_required
def set_mandate_spoc(mid):
    """Set the mandate's SPOC to a contact that belongs to its linked CRM client."""
    d = request.json or {}
    spoc_id = int(d.get('spoc_contact_id') or 0)
    conn = None
    try:
        conn = get_db()
        company_id = effective_company_id()
        try:
            conn.execute('ALTER TABLE mandates ADD COLUMN spoc_contact_id INTEGER DEFAULT 0')
            conn.commit()
        except Exception:
            pass
        m = conn.execute('SELECT * FROM mandates WHERE id=? AND owner_id=?',
                         (mid, company_id)).fetchone()
        if not m:
            return jsonify({'error': 'Mandate not found'}), 404
        if spoc_id:
            ok = conn.execute(
                'SELECT id FROM crm_contacts WHERE id=? AND client_id=? AND company_id=? AND is_active=1',
                (spoc_id, m['crm_client_id'] or 0, company_id)).fetchone()
            if not ok:
                return jsonify({'error': 'SPOC must be a contact created under this client in CRM.'}), 400
        conn.execute('UPDATE mandates SET spoc_contact_id=? WHERE id=?', (spoc_id, mid))
        conn.commit()
        return jsonify({'ok': True, 'spoc_contact_id': spoc_id})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': f'{type(e).__name__}: {e}'}), 500
    finally:
        if conn is not None:
            try: conn.close()
            except Exception: pass


@app.route('/api/mandates/<int:mid>/cc-options', methods=['GET'])
@login_required
def mandate_cc_options(mid):
    """External CC candidates (client's CRM contacts, minus SPOC) + Internal CC
    candidates (team members with an email), plus the current selections."""
    conn = get_db()
    company_id = effective_company_id()
    m = conn.execute('SELECT * FROM mandates WHERE id=? AND owner_id=?',
                     (mid, company_id)).fetchone()
    if not m:
        conn.close()
        return jsonify({'error': 'Mandate not found'}), 404
    crm_id = m['crm_client_id'] or 0
    try: spoc_id = m['spoc_contact_id'] or 0
    except Exception: spoc_id = 0
    external = []
    if crm_id:
        try:
            rows = conn.execute(
                'SELECT id, name, designation, email FROM crm_contacts '
                'WHERE client_id=? AND company_id=? AND is_active=1 AND email != "" ORDER BY id',
                (crm_id, company_id)).fetchall()
            external = [{'id': r['id'], 'name': r['name'] or '', 'designation': r['designation'] or '',
                         'email': r['email'] or ''} for r in rows if r['id'] != spoc_id]
        except Exception:
            external = []
    internal = []
    try:
        urows = conn.execute(
            "SELECT id, username, display_name, profile_email FROM users "
            "WHERE company_id=? AND status='approved' ORDER BY id", (company_id,)).fetchall()
        for u in urows:
            d = dict(u)
            em = (d.get('profile_email') or '').strip()
            if not em and '@' in (d.get('username') or ''):
                em = d['username'].strip()
            if em:
                internal.append({'name': d.get('display_name') or d.get('username') or em, 'email': em})
    except Exception:
        internal = []
    try: sel_ext = json.loads(m['cc_external_ids'] or '[]')
    except Exception: sel_ext = []
    try: sel_int = json.loads(m['cc_internal_emails'] or '[]')
    except Exception: sel_int = []
    conn.close()
    return jsonify({'ok': True, 'crm_linked': bool(crm_id), 'external': external,
                    'internal': internal, 'selected_external_ids': sel_ext,
                    'selected_internal_emails': sel_int})


@app.route('/api/mandates/<int:mid>/cc-config', methods=['POST'])
@login_required
def set_mandate_cc(mid):
    """Save the mandate's External (CRM contact ids) + Internal (emails) CC lists."""
    d = request.json or {}
    ext_ids = [int(x) for x in (d.get('external_ids') or []) if str(x).isdigit()]
    int_emails = [e.strip() for e in (d.get('internal_emails') or []) if e and e.strip()]
    conn = None
    try:
        conn = get_db()
        company_id = effective_company_id()
        for col, typ in [('cc_external_ids', "TEXT DEFAULT '[]'"), ('cc_internal_emails', "TEXT DEFAULT '[]'")]:
            try:
                conn.execute(f'ALTER TABLE mandates ADD COLUMN {col} {typ}'); conn.commit()
            except Exception:
                pass
        m = conn.execute('SELECT id FROM mandates WHERE id=? AND owner_id=?', (mid, company_id)).fetchone()
        if not m:
            return jsonify({'error': 'Mandate not found'}), 404
        conn.execute('UPDATE mandates SET cc_external_ids=?, cc_internal_emails=? WHERE id=?',
                     (json.dumps(ext_ids), json.dumps(int_emails), mid))
        conn.commit()
        return jsonify({'ok': True})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': f'{type(e).__name__}: {e}'}), 500
    finally:
        if conn is not None:
            try: conn.close()
            except Exception: pass


@app.route('/api/mandates/<int:mid>/submission-prefill', methods=['GET'])
@login_required
def submission_prefill(mid):
    """Everything the Share-to-Client modal needs: candidate list (grouped by
    stage), auto-filled To (client contacts) + CC (settings), subject, greeting
    and intro."""
    conn = get_db()
    company_id = effective_company_id()
    m = conn.execute('SELECT * FROM mandates WHERE id=? AND owner_id=?',
                     (mid, company_id)).fetchone()
    if not m:
        conn.close()
        return jsonify({'error': 'Mandate not found'}), 404
    rows = conn.execute(
        'SELECT * FROM candidates WHERE mandate_id=? AND stage != "Screened-Out" '
        'ORDER BY stage, name', (mid,)).fetchall()
    contacts = _mandate_client_contacts(conn, company_id, m['crm_client_id'])
    # If a SPOC is set on the mandate, put them first (primary To + greeting).
    spoc = None
    try:
        sid = m['spoc_contact_id'] or 0
    except Exception:
        sid = 0
    if sid:
        try:
            sc = conn.execute('SELECT id, name, email FROM crm_contacts WHERE id=? AND company_id=? AND is_active=1',
                              (sid, company_id)).fetchone()
            if sc and (sc['email'] or '').strip():
                spoc = {'name': sc['name'] or '', 'email': sc['email'] or ''}
                contacts = [spoc] + [c for c in contacts if c['email'] != spoc['email']]
        except Exception:
            spoc = None

    # CC: per-mandate External (CRM contact ids) + Internal (emails); if none
    # configured, fall back to the global submission_cc_emails setting.
    cc_list = []
    try:
        ext_ids = json.loads(m['cc_external_ids'] or '[]')
    except Exception:
        ext_ids = []
    try:
        int_emails = json.loads(m['cc_internal_emails'] or '[]')
    except Exception:
        int_emails = []
    if ext_ids:
        try:
            q = 'SELECT email FROM crm_contacts WHERE company_id=? AND id IN (%s)' % ','.join('?' * len(ext_ids))
            for r in conn.execute(q, tuple([company_id] + ext_ids)).fetchall():
                if (r['email'] or '').strip():
                    cc_list.append(r['email'].strip())
        except Exception:
            pass
    for e in int_emails:
        if e and e.strip():
            cc_list.append(e.strip())
    conn.close()
    if not cc_list:
        cc_raw = get_setting('submission_cc_emails', '') or ''
        cc_list = [e.strip() for e in re.split(r'[,\n;]+', cc_raw) if e.strip()]
    # de-dup, drop any that equal a To recipient
    _seen = set()
    cc_list = [e for e in cc_list if not (e in _seen or _seen.add(e))]

    cands = []
    for c in rows:
        d = dict(c)
        cands.append({
            'id': d['id'], 'name': d.get('name', ''), 'stage': d.get('stage', ''),
            'company': d.get('company', ''), 'experience': d.get('experience'),
            'has_cv': bool(d.get('cv_path')),
        })
    # To = the SPOC only (when set). Other client contacts must NOT auto-fill
    # To — the ones chosen as external CC belong in CC, not To.
    to_list = [spoc['email']] if spoc else [c['email'] for c in contacts]
    primary_name = (contacts[0]['name'].split()[0] if contacts and contacts[0]['name'] else '')
    cc_list = [e for e in cc_list if e not in to_list]
    n_default = len([c for c in cands if c['stage'] == 'Shared with Client']) or len(cands)
    subject = f"{m['role']} — {n_default} profile{'s' if n_default != 1 else ''} | {get_setting('company_name','') or 'HireLab'}".strip(' |')
    greeting = f"Dear {primary_name}," if primary_name else "Dear Team,"
    intro = f"Please find attached {n_default} profile{'s' if n_default != 1 else ''} for {m['role']}:"

    return jsonify({'ok': True, 'role': m['role'], 'client': m['client'],
                    'to': to_list, 'cc': cc_list, 'contacts': contacts,
                    'subject': subject, 'greeting': greeting, 'intro': intro,
                    'candidates': cands})


@app.route('/api/mandates/<int:mid>/submission-preview', methods=['POST'])
@login_required
def submission_preview(mid):
    """Return the exact HTML email body (greeting + intro + table + signature)
    for the given greeting/intro/candidate selection — for live preview."""
    d = request.json or {}
    greeting = (d.get('greeting') or '').strip()
    intro = (d.get('intro') or '').strip()
    cand_ids = [int(x) for x in (d.get('candidate_ids') or []) if str(x).isdigit()]
    conn = get_db()
    company_id = effective_company_id()
    m = conn.execute('SELECT * FROM mandates WHERE id=? AND owner_id=?',
                     (mid, company_id)).fetchone()
    if not m:
        conn.close()
        return jsonify({'error': 'Mandate not found'}), 404
    rows = []
    if cand_ids:
        rows = conn.execute(
            'SELECT * FROM candidates WHERE mandate_id=? AND owner_id=? AND id IN (%s) '
            'ORDER BY name' % ','.join('?' * len(cand_ids)),
            tuple([mid, company_id] + cand_ids)).fetchall()
    table = _submission_table_html(m, rows) if rows else \
        '<div style="color:#999;font-size:12px;padding:8px 0">(No candidates selected yet)</div>'
    sig = _submission_signature_html()
    conn.close()
    html_body = (
        '<p style="font-family:Calibri,Arial,sans-serif;font-size:14px">' + _sub_esc(greeting) + '</p>'
        + '<p style="font-family:Calibri,Arial,sans-serif;font-size:14px">' + _sub_esc(intro) + '</p>'
        + table + '<br>' + sig)
    return jsonify({'ok': True, 'html': html_body, 'count': len(rows)})


def _send_submission_email(to_list, cc_list, subject, html_body, attach_cids, conn, company_id):
    """Build a MIME email (HTML body + resume PDF attachments) and SMTP-send it
    to To + CC. Returns (ok, error)."""
    from email.mime.application import MIMEApplication
    smtp_email = get_setting('smtp_email', '')
    smtp_pass = get_setting('smtp_app_password', '')
    smtp_name = get_setting('smtp_display_name', '') or get_setting('company_name', '') or smtp_email
    if not smtp_email or not smtp_pass:
        return False, 'Email not configured. Go to Settings → Email Configuration.'

    msg = MIMEMultipart('mixed')
    msg['From'] = f'{smtp_name} <{smtp_email}>' if smtp_name else smtp_email
    msg['To'] = ', '.join(to_list)
    if cc_list:
        msg['Cc'] = ', '.join(cc_list)
    msg['Subject'] = subject
    import email.utils as _eut
    domain = smtp_email.split('@')[-1] if '@' in smtp_email else 'hirelab.local'
    msg['Message-ID'] = _eut.make_msgid(domain=domain)
    msg['Date'] = _eut.formatdate(localtime=True)

    alt = MIMEMultipart('alternative')
    alt.attach(MIMEText(re.sub('<[^<]+?>', '', html_body), 'plain', 'utf-8'))
    alt.attach(MIMEText('<div style="font-family:Calibri,Arial,sans-serif;font-size:14px">'
                        + html_body + '</div>', 'html', 'utf-8'))
    msg.attach(alt)

    # Attach each candidate's CV file
    for cid in attach_cids:
        try:
            c = conn.execute('SELECT name, cv_path, cv_original_name FROM candidates '
                             'WHERE id=? AND owner_id=?', (cid, company_id)).fetchone()
            if not c or not c['cv_path']:
                continue
            fp = os.path.join(CV_DIR, str(c['cv_path']))
            if not os.path.exists(fp):
                continue
            with open(fp, 'rb') as fh:
                data = fh.read()
            # Attachment name is ALWAYS the candidate's name (keep original extension).
            ext = os.path.splitext(c['cv_original_name'] or c['cv_path'] or '')[1] or '.pdf'
            safe_name = re.sub(r'[^A-Za-z0-9 _.-]+', '', (c['name'] or 'Candidate')).strip() or 'Candidate'
            fname = safe_name + ext
            part = MIMEApplication(data, Name=fname)
            part['Content-Disposition'] = f'attachment; filename="{fname}"'
            msg.attach(part)
        except Exception:
            continue

    if '@gmail' in smtp_email.lower() or '@googlemail' in smtp_email.lower():
        smtp_host, smtp_port = 'smtp.gmail.com', 587
    elif '@outlook' in smtp_email.lower() or '@hotmail' in smtp_email.lower() or '@live' in smtp_email.lower():
        smtp_host, smtp_port = 'smtp-mail.outlook.com', 587
    elif '@yahoo' in smtp_email.lower():
        smtp_host, smtp_port = 'smtp.mail.yahoo.com', 587
    else:
        smtp_host, smtp_port = 'smtp.gmail.com', 587

    all_rcpts = list(dict.fromkeys([e for e in (to_list + cc_list) if e]))
    try:
        server = smtplib.SMTP(smtp_host, smtp_port, timeout=30)
        server.starttls()
        server.login(smtp_email, smtp_pass)
        server.sendmail(smtp_email, all_rcpts, msg.as_string())
        server.quit()
        return True, ''
    except smtplib.SMTPAuthenticationError:
        return False, 'Email authentication failed. Check email + app password in Settings.'
    except Exception as e:
        return False, f'Failed to send: {e}'


@app.route('/api/mandates/<int:mid>/submit-to-client', methods=['POST'])
@login_required
def submit_to_client(mid):
    """Send (or save as draft) a client-submission email: greeting + intro +
    candidate table + signature, with the selected candidates' resumes attached."""
    d = request.json or {}
    action = (d.get('action') or 'send').strip()
    to_list = [e.strip() for e in (d.get('to') or []) if e and e.strip()]
    cc_list = [e.strip() for e in (d.get('cc') or []) if e and e.strip()]
    subject = (d.get('subject') or '').strip()
    greeting = (d.get('greeting') or '').strip()
    intro = (d.get('intro') or '').strip()
    body_html = (d.get('body_html') or '').strip()
    cand_ids = [int(x) for x in (d.get('candidate_ids') or []) if str(x).isdigit()]

    conn = None
    try:
        conn = get_db()
        company_id = effective_company_id()
        m = conn.execute('SELECT * FROM mandates WHERE id=? AND owner_id=?',
                         (mid, company_id)).fetchone()
        if not m:
            return jsonify({'error': 'Mandate not found'}), 404

        # Save-as-draft: persist and return (no validation of recipients needed).
        if action == 'draft':
            now = ts()
            existing = d.get('draft_id')
            if existing:
                conn.execute(
                    'UPDATE submission_drafts SET to_emails=?, cc_emails=?, subject=?, '
                    'greeting=?, intro=?, candidate_ids=?, body_html=?, updated_at=? WHERE id=? AND owner_id=?',
                    (', '.join(to_list), ', '.join(cc_list), subject, greeting, intro,
                     json.dumps(cand_ids), body_html, now, int(existing), company_id))
                did = int(existing)
            else:
                cur = conn.execute(
                    'INSERT INTO submission_drafts (owner_id, mandate_id, to_emails, cc_emails, '
                    'subject, greeting, intro, candidate_ids, body_html, created_at, updated_at) '
                    'VALUES (?,?,?,?,?,?,?,?,?,?,?)',
                    (company_id, mid, ', '.join(to_list), ', '.join(cc_list), subject,
                     greeting, intro, json.dumps(cand_ids), body_html, now, now))
                did = cur.lastrowid
            conn.commit()
            return jsonify({'ok': True, 'draft_id': did, 'saved': True})

        # Send: validate + build + SMTP send.
        if not to_list:
            return jsonify({'error': 'Add at least one recipient (To).'}), 400
        if not cand_ids:
            return jsonify({'error': 'Select at least one candidate.'}), 400
        if not subject:
            return jsonify({'error': 'Subject is required.'}), 400

        rows = conn.execute(
            'SELECT * FROM candidates WHERE mandate_id=? AND owner_id=? AND id IN (%s) '
            'ORDER BY name' % ','.join('?' * len(cand_ids)),
            tuple([mid, company_id] + cand_ids)).fetchall()
        if not rows:
            return jsonify({'error': 'Selected candidates not found.'}), 404

        # Use the recruiter's edited body if provided, else build the standard one.
        if body_html:
            final_html = body_html
        else:
            table = _submission_table_html(m, rows)
            sig = _submission_signature_html()
            final_html = (
                '<p style="font-family:Calibri,Arial,sans-serif;font-size:14px">' + _sub_esc(greeting) + '</p>'
                + '<p style="font-family:Calibri,Arial,sans-serif;font-size:14px">' + _sub_esc(intro) + '</p>'
                + table + '<br>' + sig)

        ok, err = _send_submission_email(to_list, cc_list, subject, final_html,
                                         [c['id'] for c in rows], conn, company_id)
        if not ok:
            return jsonify({'error': err}), 502

        # Log to each candidate's journey.
        for c in rows:
            try:
                log_candidate_event(c['id'], 'note',
                                    'Profile shared with client: ' + ', '.join(to_list))
            except Exception:
                pass
        return jsonify({'ok': True, 'sent': True, 'count': len(rows),
                        'to': to_list, 'cc': cc_list})

    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': f'{type(e).__name__}: {e}'}), 500
    finally:
        if conn is not None:
            try:
                conn.close()
            except Exception:
                pass


@app.route('/api/submission-drafts', methods=['GET'])
@login_required
def list_submission_drafts():
    conn = get_db()
    company_id = effective_company_id()
    mid = request.args.get('mandate_id', type=int)
    try:
        if mid:
            rows = conn.execute('SELECT * FROM submission_drafts WHERE owner_id=? AND mandate_id=? '
                                'ORDER BY updated_at DESC', (company_id, mid)).fetchall()
        else:
            rows = conn.execute('SELECT * FROM submission_drafts WHERE owner_id=? '
                                'ORDER BY updated_at DESC', (company_id,)).fetchall()
    except Exception:
        conn.close(); return jsonify({'ok': True, 'drafts': []})
    conn.close()
    out = []
    for r in rows:
        d = dict(r)
        try: cids = json.loads(d.get('candidate_ids') or '[]')
        except Exception: cids = []
        out.append({'id': d['id'], 'mandate_id': d['mandate_id'], 'subject': d.get('subject', ''),
                    'to': d.get('to_emails', ''), 'cc': d.get('cc_emails', ''),
                    'greeting': d.get('greeting', ''), 'intro': d.get('intro', ''),
                    'body_html': d.get('body_html', ''),
                    'candidate_ids': cids, 'updated_at': d.get('updated_at', '')})
    return jsonify({'ok': True, 'drafts': out})


@app.route('/api/submission-drafts/<int:did>', methods=['DELETE'])
@login_required
def delete_submission_draft(did):
    conn = get_db()
    conn.execute('DELETE FROM submission_drafts WHERE id=? AND owner_id=?',
                 (did, effective_company_id()))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


# ── Invoicing (GST tax invoices + expenses) ──────────────────────────────────
def _seller_dict():
    return {
        'name': get_setting('company_name', '') or 'HireLab Talent Resource',
        'address': get_setting('seller_address', ''),
        'udyam': get_setting('seller_udyam', ''),
        'gstin': get_setting('seller_gstin', ''),
        'state': get_setting('seller_state', ''),
        'state_code': get_setting('seller_state_code', ''),
        'reg_office': get_setting('seller_reg_office', ''),
    }

def _invoice_engine_dict(r):
    d = dict(r)
    try: extra = json.loads(d.get('extra_lines') or '[]')
    except Exception: extra = []
    buyer = {'name': d['buyer_name'], 'address': d['buyer_address'], 'gstin': d['buyer_gstin'],
             'state': d['buyer_state'], 'state_code': d['buyer_state_code']}
    if d.get('consignee_same'):
        con = buyer
    else:
        con = {'name': d['con_name'] or d['buyer_name'], 'address': d['con_address'],
               'gstin': d['con_gstin'], 'state': d['con_state'], 'state_code': d['con_state_code']}
    return {
        'seller': _seller_dict(), 'buyer': buyer, 'consignee': con,
        'invoice_no': d['invoice_no'], 'invoice_date': d['invoice_date'],
        'ref_no': d['ref_no'], 'other_ref': d['other_ref'],
        'order_no': d['order_no'], 'order_date': d['order_date'],
        'place_of_supply': d['place_of_supply'] or d['buyer_state'],
        'description': d['description'] or 'Charge Towards Recruitment Services',
        'candidate_name': d['candidate_name'], 'role': d['role'], 'extra_lines': extra,
        'hsn': d['hsn'], 'quantity': d['quantity'], 'rate': d['rate'], 'per': d['per'],
        'total_qty': d['total_qty'], 'amount': d['amount'], 'gst_rate': d['gst_rate'],
        'signatory_name': get_setting('invoice_signatory', ''),
    }

def _next_gst_invoice_no(conn, owner_id):
    fy = get_setting('invoice_fy', '') or '2026-27'
    row = conn.execute('SELECT MAX(seq) AS m FROM invoices WHERE owner_id=? AND fy=?',
                       (owner_id, fy)).fetchone()
    seq = (row['m'] or 0) + 1
    return fy, seq, f"{fy}/{seq:04d}"


@app.route('/api/invoices', methods=['POST'])
@login_required
def create_invoice():
    d = request.json or {}
    conn = None
    try:
        conn = get_db(); oid = effective_company_id(); now = ts()
        fy, seq, inv_no = _next_gst_invoice_no(conn, oid)
        if (d.get('invoice_no') or '').strip():
            inv_no = d['invoice_no'].strip()
        due = (d.get('due_date') or '').strip()
        if not due:
            import datetime as _dt
            due = (_dt.date.today() + _dt.timedelta(days=45)).isoformat()
        cur = conn.execute("""INSERT INTO invoices
            (owner_id, invoice_no, invoice_date, fy, seq, buyer_name, buyer_address, buyer_gstin,
             buyer_state, buyer_state_code, consignee_same, con_name, con_address, con_gstin,
             con_state, con_state_code, candidate_name, role, description, extra_lines, hsn, quantity,
             rate, per, total_qty, amount, gst_rate, place_of_supply, ref_no, other_ref, order_no,
             order_date, status, due_date, client_id, candidate_id, mandate_id, created_at, updated_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (oid, inv_no, d.get('invoice_date',''), fy, seq, d.get('buyer_name',''), d.get('buyer_address',''),
             d.get('buyer_gstin',''), d.get('buyer_state',''), d.get('buyer_state_code',''),
             1 if d.get('consignee_same', True) else 0, d.get('con_name',''), d.get('con_address',''),
             d.get('con_gstin',''), d.get('con_state',''), d.get('con_state_code',''),
             d.get('candidate_name',''), d.get('role',''),
             d.get('description','') or 'Charge Towards Recruitment Services',
             json.dumps(d.get('extra_lines') or []), d.get('hsn','') or get_setting('invoice_hsn','998512'),
             str(d.get('quantity','')), str(d.get('rate','')), d.get('per','CTC') or 'CTC',
             d.get('total_qty',''), float(d.get('amount',0) or 0), float(d.get('gst_rate',18) or 18),
             d.get('place_of_supply',''), d.get('ref_no',''), d.get('other_ref',''), d.get('order_no',''),
             d.get('order_date',''), d.get('status','Sent') or 'Sent', due,
             int(d.get('client_id',0) or 0), int(d.get('candidate_id',0) or 0), int(d.get('mandate_id',0) or 0),
             now, now))
        conn.commit()
        # Remember this client's billing details on the CRM record (fill-once).
        cli_id = int(d.get('client_id', 0) or 0)
        if cli_id:
            try:
                conn.execute('UPDATE crm_clients SET gstin=?, bill_address=?, bill_state=?, bill_state_code=? WHERE id=? AND company_id=?',
                             (d.get('buyer_gstin', ''), d.get('buyer_address', ''), d.get('buyer_state', ''),
                              d.get('buyer_state_code', ''), cli_id, oid))
                conn.commit()
            except Exception:
                pass
        return jsonify({'ok': True, 'id': cur.lastrowid, 'invoice_no': inv_no})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': f'{type(e).__name__}: {e}'}), 500
    finally:
        if conn is not None:
            try: conn.close()
            except Exception: pass


@app.route('/api/invoices/<int:iid>', methods=['PUT'])
@login_required
def update_invoice(iid):
    d = request.json or {}
    conn = get_db(); oid = effective_company_id()
    cols = ['invoice_date','buyer_name','buyer_address','buyer_gstin','buyer_state','buyer_state_code',
            'con_name','con_address','con_gstin','con_state','con_state_code','candidate_name','role',
            'description','hsn','quantity','rate','per','total_qty','place_of_supply','ref_no','other_ref',
            'order_no','order_date','status','due_date','received_date']
    sets, vals = [], []
    for k in cols:
        if k in d:
            sets.append(f'{k}=?'); vals.append(str(d[k]) if d[k] is not None else '')
    if 'received_amount' in d:
        sets.append('received_amount=?'); vals.append(float(d['received_amount'] or 0))
    if 'consignee_same' in d:
        sets.append('consignee_same=?'); vals.append(1 if d['consignee_same'] else 0)
    if 'amount' in d:
        sets.append('amount=?'); vals.append(float(d['amount'] or 0))
    if 'gst_rate' in d:
        sets.append('gst_rate=?'); vals.append(float(d['gst_rate'] or 18))
    if 'extra_lines' in d:
        sets.append('extra_lines=?'); vals.append(json.dumps(d['extra_lines'] or []))
    sets.append('updated_at=?'); vals.append(ts())
    vals += [iid, oid]
    conn.execute(f'UPDATE invoices SET {",".join(sets)} WHERE id=? AND owner_id=?', tuple(vals))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/invoices', methods=['GET'])
@login_required
def list_invoices():
    conn = get_db(); oid = effective_company_id()
    rows = conn.execute('SELECT * FROM invoices WHERE owner_id=? ORDER BY id DESC', (oid,)).fetchall()
    conn.close()
    out = []
    for r in rows:
        d = dict(r)
        taxable = d['amount'] or 0
        gross = round(taxable * (1 + (d['gst_rate'] or 18)/100))
        out.append({'id': d['id'], 'invoice_no': d['invoice_no'], 'invoice_date': d['invoice_date'],
                    'buyer_name': d['buyer_name'], 'candidate_name': d['candidate_name'],
                    'amount': taxable, 'total': gross, 'gst_rate': d['gst_rate'],
                    'status': d['status'], 'due_date': d['due_date'],
                    'received_amount': d['received_amount'], 'received_date': d['received_date']})
    return jsonify({'ok': True, 'invoices': out})


@app.route('/api/invoices/<int:iid>', methods=['GET'])
@login_required
def get_invoice(iid):
    conn = get_db()
    r = conn.execute('SELECT * FROM invoices WHERE id=? AND owner_id=?', (iid, effective_company_id())).fetchone()
    conn.close()
    if not r:
        return jsonify({'error': 'Invoice not found'}), 404
    d = dict(r)
    try: d['extra_lines'] = json.loads(d.get('extra_lines') or '[]')
    except Exception: d['extra_lines'] = []
    return jsonify({'ok': True, 'invoice': d})


@app.route('/api/invoices/<int:iid>/mark-paid', methods=['POST'])
@login_required
def mark_invoice_paid(iid):
    d = request.json or {}
    conn = get_db()
    conn.execute('UPDATE invoices SET status=?, received_date=?, received_amount=?, updated_at=? WHERE id=? AND owner_id=?',
                 ('paid', d.get('received_date', '') or ts()[:10], float(d.get('received_amount', 0) or 0),
                  ts(), iid, effective_company_id()))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/invoices/<int:iid>', methods=['DELETE'])
@login_required
def delete_invoice(iid):
    conn = get_db()
    conn.execute('DELETE FROM invoices WHERE id=? AND owner_id=?', (iid, effective_company_id()))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/invoices/<int:iid>/print', methods=['GET'])
@login_required
def print_invoice(iid):
    if not HAS_INVOICE:
        return Response('Invoice engine not deployed (invoice_engine.py missing).', mimetype='text/plain')
    conn = get_db()
    r = conn.execute('SELECT * FROM invoices WHERE id=? AND owner_id=?', (iid, effective_company_id())).fetchone()
    conn.close()
    if not r:
        return Response('Invoice not found', status=404, mimetype='text/plain')
    html = invoice_engine.build_invoice_html(_invoice_engine_dict(r), for_print=True)
    if request.args.get('auto') == '1':
        html = html.replace('</body>', '<script>window.addEventListener("load",function(){setTimeout(function(){window.print();},350);});</script></body>')
    return Response(html, mimetype='text/html')


# Expenses
@app.route('/api/expenses', methods=['POST'])
@login_required
def create_expense():
    d = request.json or {}
    conn = get_db()
    conn.execute('INSERT INTO expenses (owner_id, date, category, payee, amount, note, invoice_id, created_at) VALUES (?,?,?,?,?,?,?,?)',
                 (effective_company_id(), d.get('date','') or ts()[:10], d.get('category',''), d.get('payee',''),
                  float(d.get('amount',0) or 0), d.get('note',''), int(d.get('invoice_id',0) or 0), ts()))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/expenses', methods=['GET'])
@login_required
def list_expenses():
    conn = get_db()
    rows = conn.execute('SELECT * FROM expenses WHERE owner_id=? ORDER BY date DESC, id DESC',
                        (effective_company_id(),)).fetchall()
    conn.close()
    return jsonify({'ok': True, 'expenses': [dict(r) for r in rows]})


@app.route('/api/expenses/<int:eid>', methods=['DELETE'])
@login_required
def delete_expense(eid):
    conn = get_db()
    conn.execute('DELETE FROM expenses WHERE id=? AND owner_id=?', (eid, effective_company_id()))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/invoices/next-number', methods=['GET'])
@login_required
def invoice_next_number():
    conn = get_db()
    _, _, inv_no = _next_gst_invoice_no(conn, effective_company_id())
    conn.close()
    return jsonify({'ok': True, 'invoice_no': inv_no})


@app.route('/api/invoices/summary', methods=['GET'])
@login_required
def invoicing_summary():
    import datetime as _dt
    conn = get_db(); oid = effective_company_id()
    today = _dt.date.today().isoformat()
    invs = conn.execute('SELECT invoice_no, buyer_name, amount, gst_rate, status, received_amount, due_date FROM invoices WHERE owner_id=?', (oid,)).fetchall()
    exps = conn.execute('SELECT category, amount FROM expenses WHERE owner_id=?', (oid,)).fetchall()
    invoiced = received = outstanding = overdue_amt = 0.0
    n_paid = unpaid_count = overdue_count = 0
    outstanding_list = []
    for r in invs:
        gross = round((r['amount'] or 0) * (1 + (r['gst_rate'] or 18) / 100))
        invoiced += gross
        if (r['status'] or '').lower() == 'paid':
            received += (r['received_amount'] or gross); n_paid += 1
        else:
            outstanding += gross; unpaid_count += 1
            due = (r['due_date'] or '')[:10]
            is_over = bool(due and due < today)
            if is_over:
                overdue_amt += gross; overdue_count += 1
            outstanding_list.append({'invoice_no': r['invoice_no'], 'buyer_name': r['buyer_name'],
                                     'total': gross, 'due_date': due, 'overdue': is_over})
    outstanding_list.sort(key=lambda x: (not x['overdue'], x['due_date'] or '9999'))
    exp_by_cat = {}; total_exp = 0.0
    for e in exps:
        exp_by_cat[e['category']] = exp_by_cat.get(e['category'], 0) + (e['amount'] or 0)
        total_exp += (e['amount'] or 0)
    # Revenue projection from candidates' offered CTC / placement fee
    confirmed = committed = probable = 0.0
    try:
        offc = conn.execute("SELECT stage,offered_ctc,fee_percent,placement_fee,joining_date FROM candidates WHERE owner_id=? AND (offered_ctc>0 OR placement_fee>0)", (oid,)).fetchall()
        def _fee(r):
            if r['offered_ctc'] and r['offered_ctc'] > 0:
                return r['offered_ctc'] * (r['fee_percent'] or 8.33) / 100.0
            return r['placement_fee'] or 0
        def _jd(r):
            return (r['joining_date'] or '')[:10]
        for r in offc:
            jd = _jd(r); fee = _fee(r)
            if (jd and jd <= today) or (r['stage'] == 'Joined' and not (jd and jd > today)):
                confirmed += fee
            elif jd and jd > today:
                committed += fee
            elif r['stage'] in ('Shared with Client', 'Interview Inprocess'):
                probable += fee
    except Exception:
        pass
    conn.close()
    return jsonify({'ok': True, 'invoiced': round(invoiced), 'received': round(received),
                    'outstanding': round(outstanding), 'overdue_amount': round(overdue_amt),
                    'count': len(invs), 'n_paid': n_paid, 'unpaid_count': unpaid_count, 'overdue_count': overdue_count,
                    'outstanding_list': outstanding_list,
                    'total_expenses': round(total_exp), 'expenses_by_category': exp_by_cat,
                    'net_profit': round(received - total_exp),
                    'projection': {'confirmed': round(confirmed), 'committed': round(committed),
                                   'probable': round(probable),
                                   'total_potential': round(confirmed + committed + probable)}})


# ── Email Box (IMAP inbox + sent, for Command Center signal) ─────────────────
def _email_decode(s):
    if not s:
        return ''
    try:
        import email.header
        out = ''
        for txt, enc in email.header.decode_header(s):
            if isinstance(txt, bytes):
                try: out += txt.decode(enc or 'utf-8', 'ignore')
                except Exception: out += txt.decode('utf-8', 'ignore')
            else:
                out += txt
        return out
    except Exception:
        return str(s)

def _email_bodies(msg):
    """Return (clean_text, html) — original HTML preserved for proper display."""
    import html as _htmllib
    plain = ''; htmlbody = ''
    try:
        if msg.is_multipart():
            for part in msg.walk():
                ct = part.get_content_type(); cd = str(part.get('Content-Disposition') or '')
                if 'attachment' in cd:
                    continue
                if ct == 'text/plain' and not plain:
                    p = part.get_payload(decode=True)
                    if p: plain = p.decode(part.get_content_charset() or 'utf-8', 'ignore')
                elif ct == 'text/html' and not htmlbody:
                    p = part.get_payload(decode=True)
                    if p: htmlbody = p.decode(part.get_content_charset() or 'utf-8', 'ignore')
        else:
            p = msg.get_payload(decode=True)
            if p:
                txt = p.decode(msg.get_content_charset() or 'utf-8', 'ignore')
                if msg.get_content_type() == 'text/html':
                    htmlbody = txt
                else:
                    plain = txt
    except Exception:
        pass
    text = plain
    if not text and htmlbody:
        t = re.sub(r'(?is)<(script|style)[^>]*>.*?</\1>', ' ', htmlbody)
        t = re.sub(r'(?i)<br\s*/?>', '\n', t)
        t = re.sub(r'(?i)</(p|div|tr|li|h[1-6])>', '\n', t)
        t = re.sub(r'<[^>]+>', ' ', t)
        t = _htmllib.unescape(t)
        t = re.sub(r'[ \t\r\f]+', ' ', t)
        t = re.sub(r'\n[ \t]*\n[ \t]*\n+', '\n\n', t)
        text = t.strip()
    return (text or '').strip(), (htmlbody or '')

def _find_sent_folder(M):
    try:
        typ, data = M.list()
        if typ == 'OK':
            for line in data:
                s = line.decode('utf-8', 'ignore') if isinstance(line, bytes) else str(line)
                if '\\Sent' in s or 'Sent Mail' in s:
                    m = re.search(r'"([^"]+)"\s*$', s)
                    if m: return m.group(1)
    except Exception:
        pass
    return '[Gmail]/Sent Mail'


@app.route('/api/emailbox/sync', methods=['POST'])
@login_required
def emailbox_sync():
    import imaplib, email as emaillib, email.utils, datetime as _dt
    host = get_setting('imap_host', 'imap.gmail.com') or 'imap.gmail.com'
    user = get_setting('smtp_email', ''); pw = get_setting('smtp_app_password', '')
    if not user or not pw:
        return jsonify({'error': 'Email not configured. Settings → Email Configuration.'}), 400
    oid = effective_company_id()
    since = (_dt.date.today() - _dt.timedelta(days=30)).strftime('%d-%b-%Y')
    try:
        M = imaplib.IMAP4_SSL(host)
        M.login(user, pw)
    except imaplib.IMAP4.error as e:
        return jsonify({'error': f'IMAP login failed: {e}. Gmail mein IMAP enable karo aur app password check karo.'}), 502
    except Exception as e:
        return jsonify({'error': f'Could not connect ({host}): {e}'}), 502

    folders = [('INBOX', 'Inbox'), ('"' + _find_sent_folder(M) + '"', 'Sent')]
    added = 0
    conn = get_db()
    try:
        try:
            conn.execute("ALTER TABLE emails ADD COLUMN body_html TEXT DEFAULT ''"); conn.commit()
        except Exception:
            pass
        for fexpr, fname in folders:
            try:
                typ, _ = M.select(fexpr, readonly=True)
                if typ != 'OK': continue
                typ, data = M.search(None, f'(SINCE {since})')
                if typ != 'OK' or not data or not data[0]: continue
                ids = data[0].split()[-400:]
                for num in ids:
                    try:
                        typ, md = M.fetch(num, '(RFC822)')
                        if typ != 'OK' or not md or not md[0]: continue
                        msg = emaillib.message_from_bytes(md[0][1])
                        mid = (msg.get('Message-ID') or '').strip() or (fname + ':' + num.decode())
                        nm, addr = email.utils.parseaddr(_email_decode(msg.get('From')))
                        dt = msg.get('Date') or ''
                        try: dts = email.utils.parsedate_to_datetime(dt).timestamp()
                        except Exception: dts = 0
                        text, htmlbody = _email_bodies(msg)
                        text = text[:20000]; htmlbody = (htmlbody or '')[:400000]
                        snip = re.sub(r'\s+', ' ', text)[:220]
                        ex = conn.execute('SELECT id FROM emails WHERE owner_id=? AND msg_id=?', (oid, mid)).fetchone()
                        if ex:
                            # backfill/refresh body so previously-synced emails display correctly
                            conn.execute('UPDATE emails SET body=?, body_html=?, snippet=? WHERE id=?',
                                         (text, htmlbody, snip, ex['id']))
                            continue
                        conn.execute('INSERT INTO emails (owner_id,msg_id,folder,from_addr,from_name,to_addr,subject,date_str,date_ts,snippet,body,body_html,created_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)',
                            (oid, mid, fname, addr, nm or addr, _email_decode(msg.get('To')),
                             _email_decode(msg.get('Subject')), dt, dts, snip, text, htmlbody, ts()))
                        added += 1
                    except Exception:
                        continue
            except Exception:
                continue
        conn.commit()
    finally:
        conn.close()
        try: M.logout()
        except Exception: pass
    return jsonify({'ok': True, 'added': added})


@app.route('/api/emailbox/list', methods=['GET'])
@login_required
def emailbox_list():
    conn = get_db(); folder = request.args.get('folder', '')
    q = 'SELECT id,folder,from_addr,from_name,to_addr,subject,date_str,date_ts,snippet,is_read FROM emails WHERE owner_id=?'
    params = [effective_company_id()]
    if folder in ('Inbox', 'Sent'):
        q += ' AND folder=?'; params.append(folder)
    q += ' ORDER BY date_ts DESC LIMIT 500'
    rows = conn.execute(q, tuple(params)).fetchall()
    conn.close()
    return jsonify({'ok': True, 'emails': [dict(r) for r in rows]})


@app.route('/api/emailbox/<int:eid>', methods=['GET'])
@login_required
def emailbox_get(eid):
    conn = get_db()
    r = conn.execute('SELECT * FROM emails WHERE id=? AND owner_id=?', (eid, effective_company_id())).fetchone()
    if r:
        conn.execute('UPDATE emails SET is_read=1 WHERE id=?', (eid,)); conn.commit()
    conn.close()
    if not r:
        return jsonify({'error': 'Not found'}), 404
    return jsonify({'ok': True, 'email': dict(r)})


# ── CEO Command Center (strategic brain toward the ₹100 Cr goal) ─────────────
CEO_BRAIN_PROMPT = """You are not an AI assistant. You are the Executive Leadership Team of HireLab, acting simultaneously as CEO, COO, CRO, CFO, Head of Recruitment, Delivery Manager, Account Director, and Business Strategist. Your only responsibility is to maximize the long-term enterprise value of HireLab. Ignore vanity metrics. Every recommendation must increase one or more of: Revenue, Gross Profit, Cash Flow, Placement Success, Client Retention, Candidate Quality, Recruiter Productivity, Business Scalability. Never optimize for activity — always optimize for business outcomes.

COMPANY CONTEXT: HireLab Recruitment. Founder: Nitin Kumar. Stage: founder-led recruitment agency. Vision: become India's leading Engineering Recruitment Company. Current annual target: Rs 1 Crore revenue. Long-term goal: Rs 100 Crore revenue. Industries: Solar, Electrical, Automation, Renewable Energy, Data Centers, Mechanical Design. Model: Permanent Recruitment + Executive Search now; Contract Staffing, RPO and a Technology Platform in future.

THINKING PROCESS whenever ATS data is received: (1) Understand cash, revenue, active positions, placement probability, client dependency, delivery bottlenecks, recruiter workload, candidate pipeline, risks. (2) Find bottlenecks: revenue, delivery, sales, cash, operational, technology, founder. (3) Rank every opportunity by Impact x Confidence x Ease — highest impact first. (4) Generate actions — never busy work; every action must produce measurable business value.

CORE PHILOSOPHY: Revenue first, cash second, retention third, expansion fourth, automation fifth, perfection last. If something does not increase revenue, reduce risk, or improve delivery — do not recommend it.

OUTPUT STYLE: Be brutally honest. Challenge assumptions. Disagree when necessary. Never flatter. Think like a founder owning 100% equity.

GOLDEN RULE: The founder's time is the scarcest resource — protect it. Never recommend work that can be automated, delegated, or eliminated.

=== BOARD MEETING (STRATEGIC PLAN) ===
Conduct a BOARD MEETING using the ATS data provided. Participants: CEO, COO, CFO, CRO, Head of Delivery, Head of Recruitment. Each executive independently evaluates the data, then they debate and disagree where necessary, then produce ONE unified decision.
Output in markdown with these exact sections:
## Executive Summary
## Revenue Forecast
## Cash Forecast
## Placements Forecast
## Biggest Risk
## Fastest Revenue Opportunity
## Highest-ROI Client
## Highest-ROI Position
## Founder Focus
## What to Stop Doing
## What to Delegate
## The One Decision That Changes Everything
## Weekly KPIs
## Monthly KPIs
## Probability of Achieving Rs 1 Crore Target
## Probability of Achieving Rs 100 Crore Vision
Be specific and numeric. Explain WHY for every forecast and probability. No flattery."""


# ══════════════════════════════════════════════════════════════════════════
#  VECTOR / RAG ENGINE  (pluggable embeddings, vectors stored in SQLite)
# ══════════════════════════════════════════════════════════════════════════
def _embed_texts(texts):
    """Embed a list of texts via an OpenAI-compatible /embeddings endpoint.
    Provider/key/model are configurable in settings; returns None if not configured
    so the app gracefully falls back to the structured brief."""
    key = get_setting('embedding_api_key', '')
    if not key or not texts:
        return None
    base = (get_setting('embedding_base_url', 'https://api.jina.ai/v1') or 'https://api.jina.ai/v1').rstrip('/')
    model = get_setting('embedding_model', 'jina-embeddings-v3') or 'jina-embeddings-v3'
    import requests as _rq
    out = []
    B = 64
    for i in range(0, len(texts), B):
        batch = [ (t or '')[:6000] for t in texts[i:i+B] ]
        try:
            r = _rq.post(base + '/embeddings',
                         headers={'Authorization': 'Bearer ' + key, 'Content-Type': 'application/json'},
                         json={'model': model, 'input': batch}, timeout=90)
            if r.status_code != 200:
                return None
            data = r.json().get('data', [])
            out.extend([d.get('embedding') for d in data])
        except Exception:
            return None
    return out


def _cosine(a, b):
    if not a or not b or len(a) != len(b):
        return -1.0
    import math
    dot = s1 = s2 = 0.0
    for x, y in zip(a, b):
        dot += x * y; s1 += x * x; s2 += y * y
    if s1 == 0 or s2 == 0:
        return -1.0
    return dot / (math.sqrt(s1) * math.sqrt(s2))


def _chash(text):
    import hashlib
    return hashlib.md5((text or '').encode('utf-8')).hexdigest()


def _cand_chunk_text(d):
    parts = []
    def add(label, v):
        if v and str(v).strip():
            parts.append(f"{label}: {str(v).strip()}")
    add("Candidate", d.get('name'))
    add("Current company", d.get('company'))
    add("Designation", d.get('designation'))
    add("Experience", d.get('experience'))
    add("Location", d.get('location'))
    add("Qualification", d.get('qualification'))
    add("Skills", d.get('key_skills'))
    add("Summary", d.get('career_summary'))
    add("Industry", d.get('industry_background'))
    add("Stage", d.get('stage'))
    add("Recruiter feedback", d.get('recruiter_feedback'))
    add("Client feedback", d.get('client_feedback'))
    add("Notes", d.get('general_comments'))
    return '\n'.join(parts)


def _collect_index_items(conn, oid):
    """Return list of (source_type, source_id, text) for everything worth embedding."""
    items = []
    try:
        for r in conn.execute("SELECT * FROM candidates WHERE owner_id=?", (oid,)).fetchall():
            d = dict(r); txt = _cand_chunk_text(d)
            if txt.strip():
                items.append(('candidate', d['id'], txt))
    except Exception:
        pass
    try:
        for m in conn.execute("SELECT id,role,client,location,jd,status FROM mandates WHERE owner_id=?", (oid,)).fetchall():
            d = dict(m)
            jd = re.sub('<[^>]+>', ' ', d.get('jd') or '')
            txt = f"Position: {d.get('role','')} @ {d.get('client','')} ({d.get('location','')}) [{d.get('status','active')}]\n{jd[:2500]}"
            items.append(('mandate', d['id'], txt.strip()))
    except Exception:
        pass
    try:
        for e in conn.execute("SELECT id,folder,from_name,from_addr,subject,snippet,body FROM emails WHERE owner_id=?", (oid,)).fetchall():
            d = dict(e)
            body = (d.get('body') or d.get('snippet') or '')[:2500]
            txt = f"Email [{d.get('folder','')}] from {d.get('from_name','')} <{d.get('from_addr','')}> — {d.get('subject','')}\n{body}"
            items.append(('email', d['id'], txt.strip()))
    except Exception:
        pass
    return items


def _reindex(conn, oid, force=False):
    """Incrementally embed changed/new items. Returns {embedded, skipped, total, error}."""
    if not get_setting('embedding_api_key', ''):
        return {'error': 'No embedding API key set. Command Center → Knowledge → set key.'}
    items = _collect_index_items(conn, oid)
    existing = {}
    for row in conn.execute("SELECT source_type, source_id, chash FROM vec_chunks WHERE owner_id=?", (oid,)).fetchall():
        existing[(row['source_type'], row['source_id'])] = row['chash']
    to_embed = []
    for st, sid, txt in items:
        h = _chash(txt)
        if not force and existing.get((st, sid)) == h:
            continue
        to_embed.append((st, sid, txt, h))
    embedded = 0
    B = 64
    for i in range(0, len(to_embed), B):
        chunk = to_embed[i:i+B]
        vecs = _embed_texts([c[2] for c in chunk])
        if not vecs or len(vecs) != len(chunk):
            return {'error': 'Embedding API call failed. Check key/model/credit.', 'embedded': embedded}
        for (st, sid, txt, h), v in zip(chunk, vecs):
            conn.execute("DELETE FROM vec_chunks WHERE owner_id=? AND source_type=? AND source_id=?", (oid, st, sid))
            conn.execute("INSERT INTO vec_chunks (owner_id,source_type,source_id,chash,text,embedding,updated_at) VALUES (?,?,?,?,?,?,?)",
                         (oid, st, sid, h, txt, json.dumps(v), ts()))
            embedded += 1
        conn.commit()
    return {'embedded': embedded, 'skipped': len(items) - len(to_embed), 'total': len(items)}


def _vector_search(conn, oid, query, k=8):
    """Semantic search over indexed ATS records. Returns list of text snippets."""
    try:
        if get_setting('rag_enabled', '1') != '1' or not get_setting('embedding_api_key', ''):
            return []
        qv = _embed_texts([query])
        if not qv:
            return []
        qv = qv[0]
        rows = conn.execute("SELECT source_type, text, embedding FROM vec_chunks WHERE owner_id=?", (oid,)).fetchall()
        scored = []
        for r in rows:
            try:
                emb = json.loads(r['embedding'])
            except Exception:
                continue
            scored.append((_cosine(qv, emb), r['source_type'], r['text']))
        scored.sort(key=lambda x: -x[0])
        return [f"[{st}] {txt}" for sc, st, txt in scored[:k] if sc > 0.15]
    except Exception:
        return []


@app.route('/api/vector/status', methods=['GET'])
@login_required
def vector_status():
    conn = get_db(); oid = effective_company_id()
    try:
        by = {}
        for r in conn.execute("SELECT source_type, COUNT(*) n FROM vec_chunks WHERE owner_id=? GROUP BY source_type", (oid,)).fetchall():
            by[r['source_type']] = r['n']
        total = sum(by.values())
    finally:
        conn.close()
    return jsonify({'ok': True, 'total': total, 'by_type': by,
                    'has_key': bool(get_setting('embedding_api_key', '')),
                    'model': get_setting('embedding_model', ''),
                    'base_url': get_setting('embedding_base_url', ''),
                    'rag_enabled': get_setting('rag_enabled', '1') == '1'})


@app.route('/api/vector/config', methods=['POST'])
@login_required
def vector_config():
    d = request.json or {}
    if 'embedding_api_key' in d and (d['embedding_api_key'] or '').strip():
        set_setting('embedding_api_key', d['embedding_api_key'].strip())
    if 'embedding_base_url' in d:
        set_setting('embedding_base_url', (d['embedding_base_url'] or 'https://api.openai.com/v1').strip())
    if 'embedding_model' in d:
        set_setting('embedding_model', (d['embedding_model'] or 'text-embedding-3-small').strip())
    if 'rag_enabled' in d:
        set_setting('rag_enabled', '1' if d['rag_enabled'] else '0')
    return jsonify({'ok': True})


@app.route('/api/vector/reindex', methods=['POST'])
@login_required
def vector_reindex():
    force = bool((request.json or {}).get('force'))
    conn = get_db()
    try:
        res = _reindex(conn, effective_company_id(), force=force)
    finally:
        conn.close()
    if res.get('error'):
        return jsonify(res), 400
    return jsonify({'ok': True, **res})


@app.route('/api/vector/search', methods=['POST'])
@login_required
def vector_search_ep():
    q = ((request.json or {}).get('query') or '').strip()
    if not q:
        return jsonify({'error': 'Empty query'}), 400
    conn = get_db()
    try:
        hits = _vector_search(conn, effective_company_id(), q, int((request.json or {}).get('k', 8)))
    finally:
        conn.close()
    return jsonify({'ok': True, 'results': hits})


def _command_overview(conn, oid):
    d = {}
    # Revenue from invoices
    invs = conn.execute('SELECT amount, gst_rate, status, received_amount, invoice_date FROM invoices WHERE owner_id=?', (oid,)).fetchall()
    invoiced = received = outstanding = 0.0
    for r in invs:
        gross = round((r['amount'] or 0) * (1 + (r['gst_rate'] or 18)/100))
        invoiced += gross
        if (r['status'] or '').lower() == 'paid':
            received += (r['received_amount'] or gross)
        else:
            outstanding += gross
    # Expenses
    exps = conn.execute('SELECT category, amount FROM expenses WHERE owner_id=?', (oid,)).fetchall()
    total_exp = sum((e['amount'] or 0) for e in exps)
    exp_by_cat = {}
    for e in exps:
        exp_by_cat[e['category']] = exp_by_cat.get(e['category'], 0) + (e['amount'] or 0)
    # Pipeline
    try:
        mand_rows = conn.execute('SELECT id, status FROM mandates WHERE owner_id=?', (oid,)).fetchall()
        open_mandates = sum(1 for m in mand_rows if (m['status'] or 'active').lower() in ('active', 'open', ''))
    except Exception:
        open_mandates = 0
    try:
        cand_rows = conn.execute('SELECT stage FROM candidates WHERE owner_id=?', (oid,)).fetchall()
        total_cand = len(cand_rows)
        by_stage = {}
        for c in cand_rows:
            s = c['stage'] or 'Unknown'; by_stage[s] = by_stage.get(s, 0) + 1
        placed = conn.execute("SELECT COUNT(*) n FROM candidates WHERE owner_id=? AND " + PLACED_SQL, (oid,)).fetchone()['n']
        shared = by_stage.get('Shared with Client', 0)
    except Exception:
        total_cand = placed = shared = 0; by_stage = {}
    # Snapshot inputs
    def _num(k):
        try: return float(get_setting(k, '') or 0)
        except Exception: return 0.0
    bank_cash = _num('cc_bank_cash'); monthly_fixed = _num('cc_monthly_fixed')
    year_target = _num('cc_year_target'); funding = _num('cc_funding_available')
    target_total = _num('cc_target_total') or 1000000000
    target_years = _num('cc_target_years') or 3
    runway = round(bank_cash / monthly_fixed, 1) if monthly_fixed else 0
    d.update({
        'invoiced': round(invoiced), 'received': round(received), 'outstanding': round(outstanding),
        'total_expenses': round(total_exp), 'net_profit': round(received - total_exp),
        'expenses_by_category': exp_by_cat,
        'open_mandates': open_mandates, 'total_candidates': total_cand, 'placed': placed,
        'shared_with_client': shared, 'by_stage': by_stage,
        'bank_cash': bank_cash, 'monthly_fixed': monthly_fixed, 'runway_months': runway,
        'year_target': year_target, 'funding_available': funding,
        'target_total': target_total, 'target_years': target_years, 'team_size': get_setting('cc_team_size', '1'),
    })
    return d


CEO_CHAT_PROMPT = """You are not an AI assistant. You are the Executive Leadership Team of HireLab, acting simultaneously as CEO, COO, CRO, CFO, Head of Recruitment, Delivery Manager, Account Director, and Business Strategist. Your only responsibility is to maximize the long-term enterprise value of HireLab. Ignore vanity metrics. Every recommendation must increase one or more of: Revenue, Gross Profit, Cash Flow, Placement Success, Client Retention, Candidate Quality, Recruiter Productivity, Business Scalability. Never optimize for activity — always optimize for business outcomes.

COMPANY CONTEXT: HireLab Recruitment. Founder: Nitin Kumar. Stage: founder-led recruitment agency. Vision: become India's leading Engineering Recruitment Company. Current annual target: Rs 1 Crore revenue. Long-term goal: Rs 100 Crore revenue. Industries: Solar, Electrical, Automation, Renewable Energy, Data Centers, Mechanical Design. Model: Permanent Recruitment + Executive Search now; Contract Staffing, RPO and a Technology Platform in future.

THINKING PROCESS whenever ATS data is received: (1) Understand cash, revenue, active positions, placement probability, client dependency, delivery bottlenecks, recruiter workload, candidate pipeline, risks. (2) Find bottlenecks: revenue, delivery, sales, cash, operational, technology, founder. (3) Rank every opportunity by Impact x Confidence x Ease — highest impact first. (4) Generate actions — never busy work; every action must produce measurable business value.

CORE PHILOSOPHY: Revenue first, cash second, retention third, expansion fourth, automation fifth, perfection last. If something does not increase revenue, reduce risk, or improve delivery — do not recommend it.

OUTPUT STYLE: Be brutally honest. Challenge assumptions. Disagree when necessary. Never flatter. Think like a founder owning 100% equity.

GOLDEN RULE: The founder's time is the scarcest resource — protect it. Never recommend work that can be automated, delegated, or eliminated.

=== FOUNDER CHAT / DAILY DASHBOARD ===
You are talking directly with founder Nitin Kumar as his executive team. INTERPRET the ATS — don't just summarize. When he asks about the business, answer: what changed and why, is the business healthier, are we growing, where are we losing money, which client/recruiter/position/invoice/candidate needs immediate attention, and what decision a world-class CEO would make today. Rank opportunities and risks. Treat facts he tells you as current truth.

Before answering ANY question, run it through these filters — does it (1) increase revenue, (2) improve cash flow, (3) protect existing placements, (4) improve client relationships, (5) improve recruiter productivity, (6) reduce founder workload, (7) can be automated, (8) can be delegated, (9) align with the Rs 1 Crore annual goal, (10) align with the Rs 100 Crore vision. If the answer is NO to all — do not recommend it.
Keep replies tight and practical, in Nitin's Hinglish/English mix."""


def _command_chat_history(conn, oid, limit=24):
    rows = conn.execute('SELECT role, content FROM command_chat WHERE owner_id=? ORDER BY id DESC LIMIT ?',
                        (oid, limit)).fetchall()
    return [{'role': r['role'], 'content': r['content']} for r in reversed(rows)]


@app.route('/api/command/chat', methods=['GET'])
@login_required
def command_chat_history():
    conn = get_db()
    try:
        hist = _command_chat_history(conn, effective_company_id(), 60)
    finally:
        conn.close()
    return jsonify({'ok': True, 'messages': hist})


@app.route('/api/command/chat', methods=['DELETE'])
@login_required
def command_chat_clear():
    conn = get_db()
    conn.execute('DELETE FROM command_chat WHERE owner_id=?', (effective_company_id(),))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/command/chat', methods=['POST'])
@login_required
def command_chat():
    ds_key = get_setting('deepseek_api_key')
    if not ds_key:
        return jsonify({'error': 'DeepSeek API key not set. Add it in Settings.'}), 400
    msg = ((request.json or {}).get('message') or '').strip()
    if not msg:
        return jsonify({'error': 'Empty message'}), 400
    conn = get_db()
    try:
        oid = effective_company_id()
        o = _command_overview(conn, oid)
        hist = _command_chat_history(conn, oid, 24)
        try:
            brief = _work_status_brief(conn, oid)
        except Exception:
            brief = ''
        def r(n): return f"₹{int(n or 0):,}"
        live = ("FULL ATS SCAN (use this to answer):\n" + brief +
                f"\n\nSNAPSHOT: bank cash {r(o['bank_cash'])}, monthly fixed {r(o['monthly_fixed'])}, "
                f"runway {o['runway_months']} months, this-year target {r(o['year_target'])}.")
        rag = _vector_search(conn, oid, msg, 8)
        if rag:
            live += "\n\nRELEVANT RECORDS (semantic search of candidates/positions/emails for this question):\n" + '\n---\n'.join(rag)
        messages = [{'role': 'system', 'content': CEO_CHAT_PROMPT + "\n\n" + live}]
        messages += hist
        messages.append({'role': 'user', 'content': msg})

        try:
            rr = call_deepseek(ds_key,
                {'model': 'deepseek-chat', 'temperature': 0.5, 'max_tokens': 900, 'messages': messages},
                timeout=120, endpoint='command-chat')
        except TokenCapError:
            return jsonify({'error': 'Monthly AI token cap reached.'}), 429
        except Exception as e:
            return jsonify({'error': f'Could not reach DeepSeek — {type(e).__name__}: {e}'}), 502
        if rr.status_code != 200:
            try: err = rr.json().get('error', {}).get('message', rr.text[:300])
            except Exception: err = rr.text[:300]
            return jsonify({'error': f'DeepSeek returned {rr.status_code}: {err}'}), 502
        try:
            reply = rr.json()['choices'][0]['message']['content'].strip()
        except Exception as e:
            return jsonify({'error': f'Unexpected DeepSeek response: {e}'}), 502

        now = ts()
        conn.execute('INSERT INTO command_chat (owner_id, role, content, created_at) VALUES (?,?,?,?)', (oid, 'user', msg, now))
        conn.execute('INSERT INTO command_chat (owner_id, role, content, created_at) VALUES (?,?,?,?)', (oid, 'assistant', reply, now))
        conn.commit()
        return jsonify({'ok': True, 'reply': reply})
    finally:
        conn.close()


def _work_status_brief(conn, oid):
    """Full-ATS snapshot of what's actionable right now — feeds daily task generation."""
    import datetime as _dt
    today = _dt.date.today().isoformat()
    soon = (_dt.date.today() + _dt.timedelta(days=15)).isoformat()
    stale_cut = (_dt.date.today() - _dt.timedelta(days=4)).isoformat()
    ACTIVE = {'Screening', 'Follow Up 1', 'Follow Up 2', 'Not Contacted', 'Called',
              'Interested', 'Updated CV awaited', 'Shared with Client', 'Interview Inprocess'}
    lines = []

    # 1) MONEY position + TARGET GAP
    try:
        import datetime as _dt2
        o = _command_overview(conn, oid)
        def _r(n): return f"₹{int(n or 0):,}"
        lines.append(f"MONEY: received {_r(o['received'])}, outstanding {_r(o['outstanding'])}, "
                     f"net profit {_r(o['net_profit'])}, runway {o['runway_months']} months.")
        yt = o.get('year_target') or 0
        if yt:
            gap = yt - (o['received'] or 0)
            # months left in this financial/calendar year
            months_left = max(1, 12 - _dt2.date.today().month + 1)
            per_month = gap / months_left if months_left else gap
            # active clients + open mandates for growth context
            try:
                nclients = conn.execute("SELECT COUNT(*) n FROM crm_clients WHERE company_id=? AND is_active=1", (oid,)).fetchone()['n']
            except Exception:
                nclients = 0
            lines.append(f"TARGET & GAP (think strategically about this every day): this-year target {_r(yt)}, "
                         f"achieved {_r(o['received'])} ({int((o['received'] or 0)/yt*100)}%), GAP REMAINING {_r(gap)} "
                         f"over ~{months_left} months = {_r(per_month)}/month needed. Active clients: {nclients}, open mandates: {o.get('open_mandates',0)}.")
    except Exception:
        pass
    # 2) Invoices — overdue + unpaid
    try:
        rows = conn.execute("SELECT invoice_no, buyer_name, due_date FROM invoices WHERE owner_id=? AND lower(status)!='paid'", (oid,)).fetchall()
        overdue = [r for r in rows if r['due_date'] and r['due_date'][:10] < today]
        if overdue:
            lines.append("OVERDUE PAYMENTS (chase today): " + '; '.join(f"{r['invoice_no']} {r['buyer_name']} (due {r['due_date'][:10]})" for r in overdue[:8]))
        pend = [r for r in rows if not (r['due_date'] and r['due_date'][:10] < today)]
        if pend:
            lines.append("UNPAID (not yet due): " + '; '.join(f"{r['invoice_no']} {r['buyer_name']}" for r in pend[:6]))
    except Exception:
        pass
    # 3) Candidate lifecycle — invoicing, guarantees, then PIPELINE BY POSITION
    try:
        cands = conn.execute("SELECT id,name,stage,joining_date,guarantee_days,updated_at,phone,email,mandate_id FROM candidates WHERE owner_id=?", (oid,)).fetchall()
        inv_cids = set(r['candidate_id'] for r in conn.execute("SELECT candidate_id FROM invoices WHERE owner_id=? AND candidate_id>0", (oid,)).fetchall())
        def _joined_date(c):
            return (c['joining_date'] or '')[:10]
        # JOINED = joining date has actually arrived (<= today), OR stage 'Joined' with no future date.
        awaiting_inv = [c for c in cands
                        if c['id'] not in inv_cids
                        and ((_joined_date(c) and _joined_date(c) <= today)
                             or (c['stage'] == 'Joined' and not (_joined_date(c) and _joined_date(c) > today)))]
        if awaiting_inv:
            lines.append("JOINED (joining date has passed), NO INVOICE YET — raise invoice now: " + '; '.join(f"{c['name']} [ref cand:{c['mandate_id']}:{c['id']}]" for c in awaiting_inv[:8]))
        # UPCOMING JOININGS = future joining date. Do NOT invoice yet.
        upcoming = [c for c in cands if _joined_date(c) and _joined_date(c) > today]
        if upcoming:
            lines.append("UPCOMING JOININGS (candidate confirmed but has NOT joined yet — DO NOT create any invoice task for these until their joining date arrives): "
                         + '; '.join(f"{c['name']} joining {_joined_date(c)} [ref cand:{c['mandate_id']}:{c['id']}]" for c in upcoming[:8]))
        guar = []
        for c in cands:
            if c['joining_date']:
                try:
                    gend = (_dt.date.fromisoformat(c['joining_date'][:10]) + _dt.timedelta(days=int(c['guarantee_days'] or 90))).isoformat()
                    if today <= gend <= soon:
                        guar.append(f"{c['name']} (ends {gend}) [ref cand:{c['mandate_id']}:{c['id']}]")
                except Exception:
                    pass
        if guar:
            lines.append("GUARANTEE ENDING SOON (confirm still on job): " + '; '.join(guar[:6]))
        # Map mandate id -> role @ client  and  id -> status
        mroles = {}; mstatus = {}
        for m in conn.execute("SELECT id,role,client,status FROM mandates WHERE owner_id=?", (oid,)).fetchall():
            mroles[m['id']] = f"{m['role']} @ {m['client']}"
            mstatus[m['id']] = (m['status'] or 'active').lower()
        active_mids = {mid for mid, st in mstatus.items() if st in ('active', 'open', '')}
        closed_positions = [f"{mroles[mid]}" for mid, st in mstatus.items() if st == 'closed']
        hold_positions = [f"{mroles[mid]}" for mid, st in mstatus.items() if st == 'hold']
        if closed_positions or hold_positions:
            note = "POSITIONS THAT ARE NOT OPEN — do NOT create sourcing/follow-up/feedback tasks for these: "
            if closed_positions: note += "CLOSED: " + '; '.join(closed_positions[:10]) + ". "
            if hold_positions: note += "ON HOLD: " + '; '.join(hold_positions[:10]) + "."
            lines.append(note)
        ACTIVE_STAGES = ['Not Contacted', 'Called', 'Screening', 'Interested', 'Follow Up 1',
                         'Follow Up 2', 'Updated CV awaited', 'Shared with Client', 'Interview Inprocess']
        by_pos = {}
        for c in cands:
            # ONLY candidates whose position is still OPEN belong in active pipeline tasks
            if c['stage'] in ACTIVE_STAGES and c['mandate_id'] in active_mids:
                by_pos.setdefault(c['mandate_id'], []).append(c)
        if by_pos:
            lines.append("PIPELINE BY POSITION (only OPEN positions — make SEPARATE tasks per position, never mix candidates of different positions in one task):")
            for mid, cl in by_pos.items():
                pos = mroles.get(mid, f"Mandate#{mid}")
                sg = {}
                for c in cl:
                    sg.setdefault(c['stage'], []).append(c)
                parts = []
                for st in ACTIVE_STAGES:
                    if sg.get(st):
                        parts.append(f"{st}: " + ', '.join(f"{c['name']} [ref cand:{mid}:{c['id']}]" for c in sg[st][:6]))
                lines.append(f"  • {pos} [ref mandate:{mid}] — " + ' | '.join(parts))
        # stale: active stage, OPEN position, not touched 4+ days
        stale = [c for c in cands if (c['stage'] in ACTIVE_STAGES) and (c['mandate_id'] in active_mids) and (not c['updated_at'] or c['updated_at'][:10] < stale_cut)]
        if stale:
            lines.append(f"STALE FOLLOW-UPS ({len(stale)}, open positions, no activity 4+ days): " + '; '.join(f"{c['name']}[{c['stage']}, {mroles.get(c['mandate_id'],'?')}] [ref cand:{c['mandate_id']}:{c['id']}]" for c in stale[:8]))
    except Exception:
        pass
    # 4) Mandates — sourcing gaps
    try:
        mand = conn.execute("SELECT id,role,client FROM mandates WHERE owner_id=? AND lower(coalesce(status,'active')) IN ('active','open','')", (oid,)).fetchall()
        counts = {}
        for r in conn.execute("SELECT mandate_id, COUNT(*) n FROM candidates WHERE owner_id=? GROUP BY mandate_id", (oid,)).fetchall():
            counts[r['mandate_id']] = r['n']
        if mand:
            thin = [m for m in mand if counts.get(m['id'], 0) < 3]
            if thin:
                lines.append("NEEDS SOURCING (thin pipeline <3 candidates): " + '; '.join(f"{m['role']} @ {m['client']} [ref mandate:{m['id']}]" for m in thin[:6]))
    except Exception:
        pass
    # 4b) Client concentration, candidate funnel, this-month activity, pipeline forecast
    try:
        import datetime as _d3
        month_start = _d3.date.today().replace(day=1).isoformat()
        crows = conn.execute("SELECT buyer_name, SUM(amount*(1+COALESCE(gst_rate,18)/100.0)) billed FROM invoices WHERE owner_id=? AND buyer_name!='' GROUP BY buyer_name ORDER BY billed DESC", (oid,)).fetchall()
        if crows:
            total_billed = sum((r['billed'] or 0) for r in crows) or 1
            conc = int((crows[0]['billed'] or 0) / total_billed * 100)
            lines.append("CLIENT CONCENTRATION: " + '; '.join(f"{r['buyer_name']} Rs {int(r['billed'] or 0):,}" for r in crows[:5])
                         + f". Top client = {conc}% of billing" + (" (HIGH dependency risk — diversify)" if conc >= 50 else ""))
        allc = conn.execute("SELECT stage, COUNT(*) n FROM candidates WHERE owner_id=? GROUP BY stage", (oid,)).fetchall()
        if allc:
            lines.append("CANDIDATE FUNNEL (all stages): " + '; '.join(f"{r['stage'] or '?'}: {r['n']}" for r in allc))
        jm = conn.execute("SELECT COUNT(*) n FROM candidates WHERE owner_id=? AND substr(joining_date,1,10)>=?", (oid, month_start)).fetchone()['n']
        interviews = conn.execute("SELECT COUNT(*) n FROM candidates WHERE owner_id=? AND stage='Interview Inprocess'", (oid,)).fetchone()['n']
        lines.append(f"THIS MONTH: {jm} joined, {interviews} currently in interview stage.")
        fc = conn.execute("SELECT ctc_max FROM mandates WHERE owner_id=? AND lower(coalesce(status,'active')) IN ('active','open','')", (oid,)).fetchall()
        potential = sum(((r['ctc_max'] or 0) * 100000 * 0.0833) for r in fc if (r['ctc_max'] or 0) > 0)
        if potential:
            lines.append(f"PIPELINE FORECAST: {len(fc)} open mandates, ceiling fee value ~Rs {int(potential):,} if all filled (8.33% of CTC).")
        # Revenue projection from ACTUAL offers (offered_ctc * fee%)
        try:
            offc = conn.execute("SELECT name,stage,offered_ctc,fee_percent,placement_fee,joining_date FROM candidates WHERE owner_id=? AND (offered_ctc>0 OR placement_fee>0)", (oid,)).fetchall()
            def _fee(r):
                if r['offered_ctc'] and r['offered_ctc'] > 0:
                    return r['offered_ctc'] * (r['fee_percent'] or 8.33) / 100.0
                return r['placement_fee'] or 0
            def _jd(r):
                return (r['joining_date'] or '')[:10]
            confirmed = sum(_fee(r) for r in offc if (_jd(r) and _jd(r) <= today) or (r['stage'] == 'Joined' and not (_jd(r) and _jd(r) > today)))
            committed = sum(_fee(r) for r in offc if _jd(r) and _jd(r) > today)
            probable = sum(_fee(r) for r in offc if r['stage'] in ('Shared with Client', 'Interview Inprocess') and not _jd(r))
            if confirmed or committed or probable:
                lines.append(f"REVENUE PROJECTION (from actual offered CTCs): "
                             f"CONFIRMED (already joined, fee due NOW) ~Rs {int(confirmed):,}; "
                             f"COMMITTED (offer accepted, joining on a future date — fee only AFTER joining, do not invoice yet) ~Rs {int(committed):,}; "
                             f"PROBABLE (still in interview/shared) ~Rs {int(probable):,}.")
        except Exception:
            pass
    except Exception:
        pass
    # 5) Email signal — unread inbox
    try:
        em = conn.execute("SELECT from_name, subject FROM emails WHERE owner_id=? AND folder='Inbox' AND is_read=0 ORDER BY date_ts DESC LIMIT 8", (oid,)).fetchall()
        if em:
            lines.append("UNREAD INBOX (may need reply): " + ' | '.join(f"{e['from_name']}: {e['subject']}" for e in em))
    except Exception:
        pass
    # 6) Founder notes from chat
    try:
        notes = [m['content'][:160] for m in _command_chat_history(conn, oid, 8) if m['role'] == 'user']
        if notes:
            lines.append("FOUNDER NOTES (recent, treat as current facts): " + ' | '.join(notes[:5]))
    except Exception:
        pass
    return '\n'.join(lines) or 'No active pipeline/billing data yet.'


TASK_GEN_PROMPT = """You are not an AI assistant. You are the Executive Leadership Team of HireLab, acting simultaneously as CEO, COO, CRO, CFO, Head of Recruitment, Delivery Manager, Account Director, and Business Strategist. Your only responsibility is to maximize the long-term enterprise value of HireLab. Ignore vanity metrics. Every recommendation must increase one or more of: Revenue, Gross Profit, Cash Flow, Placement Success, Client Retention, Candidate Quality, Recruiter Productivity, Business Scalability. Never optimize for activity — always optimize for business outcomes.

COMPANY CONTEXT: HireLab Recruitment. Founder: Nitin Kumar. Stage: founder-led recruitment agency. Vision: become India's leading Engineering Recruitment Company. Current annual target: Rs 1 Crore revenue. Long-term goal: Rs 100 Crore revenue. Industries: Solar, Electrical, Automation, Renewable Energy, Data Centers, Mechanical Design. Model: Permanent Recruitment + Executive Search now; Contract Staffing, RPO and a Technology Platform in future.

THINKING PROCESS whenever ATS data is received: (1) Understand cash, revenue, active positions, placement probability, client dependency, delivery bottlenecks, recruiter workload, candidate pipeline, risks. (2) Find bottlenecks: revenue, delivery, sales, cash, operational, technology, founder. (3) Rank every opportunity by Impact x Confidence x Ease — highest impact first. (4) Generate actions — never busy work; every action must produce measurable business value.

CORE PHILOSOPHY: Revenue first, cash second, retention third, expansion fourth, automation fifth, perfection last. If something does not increase revenue, reduce risk, or improve delivery — do not recommend it.

OUTPUT STYLE: Be brutally honest. Challenge assumptions. Disagree when necessary. Never flatter. Think like a founder owning 100% equity.

GOLDEN RULE: The founder's time is the scarcest resource — protect it. Never recommend work that can be automated, delegated, or eliminated.

=== TODAY'S TASK GENERATION ===
You have complete access to the ATS. Your objective is NOT to create tasks — it is to maximize TODAY'S business outcome.
Before creating tasks, hunt for: (1) revenue leaks, (2) placements at risk, (3) clients losing attention, (4) delayed invoices, (5) guarantees ending, (6) positions that can close fastest, (7) candidates likely to drop, (8) inactive recruiters, (9) follow-ups affecting revenue, (10) founder bottlenecks.
Then generate ONLY the highest-ROI actions. Maximum 12 tasks. If no high-value task exists, return fewer — never pad with low-value admin work.

HARD RULES:
- POSITION-WISE: every pipeline/follow-up task is for ONE position (role @ client) only — never mix candidates of different positions in one task; name the position and the specific candidate(s).
- RESPECT STATUS: never create sourcing/follow-up/feedback tasks for positions marked CLOSED or ON HOLD (only exception: collecting invoice/payment for an already-placed candidate).
- STRATEGIC: always include 2-3 tasks that move Nitin toward the Rs 1 Crore target — account expansion of existing paying clients (cheapest revenue), targeted new BD, or closing the monthly gap — tied to the TARGET & GAP data.

Return ONLY a JSON array (no prose, no markdown fences), each item exactly:
{"text": "one imperative action naming the position/candidate/client/invoice",
 "category": "Payment|Invoice|Placement|Follow-up|Sourcing|Client|Email|Admin|Growth",
 "priority": "high|medium|low",
 "reason": "pack WHY NOW + business impact + estimated revenue impact + whether it can be delegated (and to whom) into one tight line",
 "ref": "entity to open, copied EXACTLY from the [ref ...] tag in the data (e.g. cand:40:786, mandate:40, or invoice); empty string if none"}"""


@app.route('/api/command/tasks', methods=['GET'])
@login_required
def command_tasks_list():
    import datetime as _dt
    today = _dt.date.today().isoformat()
    conn = get_db()
    rows = conn.execute("SELECT id,text,category,priority,done,source,reason,ref,snooze_until,task_date FROM command_tasks WHERE owner_id=? ORDER BY done ASC, CASE priority WHEN 'high' THEN 0 WHEN 'medium' THEN 1 ELSE 2 END, id DESC",
                        (effective_company_id(),)).fetchall()
    conn.close()
    active, snoozed = [], 0
    for r in rows:
        d = dict(r)
        if d.get('snooze_until') and d['snooze_until'][:10] > today and not d['done']:
            snoozed += 1
            continue
        active.append(d)
    return jsonify({'ok': True, 'tasks': active, 'snoozed_count': snoozed})


@app.route('/api/command/tasks', methods=['POST'])
@login_required
def command_tasks_add():
    d = request.json or {}
    txt = (d.get('text') or '').strip()
    if not txt:
        return jsonify({'error': 'Empty task'}), 400
    conn = get_db()
    conn.execute("INSERT INTO command_tasks (owner_id,text,category,priority,done,source,task_date,created_at) VALUES (?,?,?,?,0,'manual',?,?)",
                 (effective_company_id(), txt, d.get('category', ''), d.get('priority', 'medium'), ts()[:10], ts()))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/command/tasks/<int:tid>', methods=['PUT'])
@login_required
def command_tasks_update(tid):
    d = request.json or {}
    conn = get_db(); oid = effective_company_id()
    sets, vals = [], []
    if 'done' in d: sets.append('done=?'); vals.append(1 if d['done'] else 0)
    if 'text' in d: sets.append('text=?'); vals.append((d['text'] or '').strip())
    if 'priority' in d: sets.append('priority=?'); vals.append(d['priority'])
    if 'snooze_until' in d: sets.append('snooze_until=?'); vals.append((d['snooze_until'] or '')[:10])
    if sets:
        vals += [tid, oid]
        conn.execute(f'UPDATE command_tasks SET {",".join(sets)} WHERE id=? AND owner_id=?', tuple(vals))
        conn.commit()
    conn.close()
    return jsonify({'ok': True})


@app.route('/api/command/tasks/<int:tid>', methods=['DELETE'])
@login_required
def command_tasks_delete(tid):
    conn = get_db()
    conn.execute('DELETE FROM command_tasks WHERE id=? AND owner_id=?', (tid, effective_company_id()))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


def _attach_task_refs(conn, oid, tasks):
    """Reliably attach deep-link refs by matching candidate/mandate names in each task text
    (don't depend on the AI echoing the [ref ...] tags)."""
    try:
        cands = conn.execute("SELECT id,name,mandate_id FROM candidates WHERE owner_id=? AND name!=''", (oid,)).fetchall()
        mands = conn.execute("SELECT id,role,client FROM mandates WHERE owner_id=?", (oid,)).fetchall()
    except Exception:
        return tasks
    cand_list = sorted([(c['name'].strip(), c['id'], c['mandate_id']) for c in cands if len((c['name'] or '').strip()) >= 4],
                       key=lambda x: -len(x[0]))
    mand_list = [(m['id'], (m['role'] or '').strip(), (m['client'] or '').strip()) for m in mands]
    for t in tasks:
        if not isinstance(t, dict):
            continue
        ref = (t.get('ref') or '').strip()
        if ref.startswith('cand:') or ref.startswith('mandate:') or ref == 'invoice':
            continue
        txt = (t.get('text') or '').lower()
        found = ''
        for nm, cid, mid in cand_list:
            if nm.lower() in txt:
                found = f"cand:{mid}:{cid}"; break
        if not found:
            for mid, role, client in mand_list:
                if (role and role.lower() in txt) or (client and len(client) >= 4 and client.lower() in txt):
                    found = f"mandate:{mid}"; break
        if not found and any(w in txt for w in ('invoice', 'payment', 'overdue', 'gst')):
            found = 'invoice'
        t['ref'] = found
    return tasks


def _run_task_generation(conn, oid, refine_instruction='', current_tasks=None):
    """Generate (or revise) the daily AI task list. Honors persisted task prefs.
    Returns (tasks_list, error_json_tuple_or_None)."""
    ds_key = get_setting('deepseek_api_key')
    if not ds_key:
        return None, (jsonify({'error': 'DeepSeek API key not set. Add it in Settings.'}), 400)
    brief = _work_status_brief(conn, oid)
    prefs = (get_setting('cc_task_prefs', '') or '').strip()
    sys = TASK_GEN_PROMPT
    if prefs:
        sys += "\n\nSTANDING PREFERENCES from the founder (always follow these): " + prefs
    user = 'FULL ATS WORK STATUS:\n' + brief
    if current_tasks:
        user += '\n\nCURRENT TASK LIST:\n' + '\n'.join('- ' + t for t in current_tasks)
    if refine_instruction:
        user += ('\n\nThe founder wants you to REVISE the task list with this instruction: "'
                 + refine_instruction + '"\nReturn the FULL revised task list (not just the changes), as a JSON array.')
    else:
        user += "\n\nGive today's task list as a JSON array."
    try:
        rr = call_deepseek(ds_key,
            {'model': 'deepseek-chat', 'temperature': 0.4, 'max_tokens': 1600,
             'messages': [{'role': 'system', 'content': sys}, {'role': 'user', 'content': user}]},
            timeout=150, endpoint='task-gen')
    except TokenCapError:
        return None, (jsonify({'error': 'Monthly AI token cap reached.'}), 429)
    except Exception as e:
        return None, (jsonify({'error': f'Could not reach DeepSeek — {type(e).__name__}: {e}'}), 502)
    if rr.status_code != 200:
        try: err = rr.json().get('error', {}).get('message', rr.text[:200])
        except Exception: err = rr.text[:200]
        return None, (jsonify({'error': f'DeepSeek returned {rr.status_code}: {err}'}), 502)
    text = rr.json()['choices'][0]['message']['content'].strip()
    text = re.sub(r'^```[a-zA-Z]*\n?|```$', '', text).strip()
    tasks = None
    try:
        tasks = json.loads(text)
    except Exception:
        s = text.find('['); e = text.rfind(']')
        if s >= 0 and e > s:
            try: tasks = json.loads(text[s:e+1])
            except Exception: tasks = None
    if not isinstance(tasks, list) or not tasks:
        return None, (jsonify({'error': 'AI could not produce a task list. Try again.'}), 502)
    tasks = _attach_task_refs(conn, oid, tasks)
    conn.execute("DELETE FROM command_tasks WHERE owner_id=? AND source='ai' AND done=0", (oid,))
    now = ts(); td = now[:10]
    for t in tasks[:12]:
        if isinstance(t, dict) and (t.get('text') or '').strip():
            conn.execute("INSERT INTO command_tasks (owner_id,text,category,priority,done,source,reason,ref,task_date,created_at) VALUES (?,?,?,?,0,'ai',?,?,?,?)",
                         (oid, t['text'].strip()[:400], (t.get('category') or '')[:40], (t.get('priority') or 'medium')[:10],
                          (t.get('reason') or '')[:300], (t.get('ref') or '')[:40], td, now))
    conn.commit()
    rows = conn.execute("SELECT id,text,category,priority,done,source,reason,ref,snooze_until,task_date FROM command_tasks WHERE owner_id=? ORDER BY done ASC, CASE priority WHEN 'high' THEN 0 WHEN 'medium' THEN 1 ELSE 2 END, id DESC", (oid,)).fetchall()
    return [dict(r) for r in rows], None


@app.route('/api/command/tasks/generate', methods=['POST'])
@login_required
def command_tasks_generate():
    conn = get_db()
    try:
        tasks, err = _run_task_generation(conn, effective_company_id())
        if err:
            return err
        return jsonify({'ok': True, 'tasks': tasks})
    finally:
        conn.close()


WEEKLY_REVIEW_PROMPT = """You are not an AI assistant. You are the Executive Leadership Team of HireLab, acting simultaneously as CEO, COO, CRO, CFO, Head of Recruitment, Delivery Manager, Account Director, and Business Strategist. Your only responsibility is to maximize the long-term enterprise value of HireLab. Ignore vanity metrics. Every recommendation must increase one or more of: Revenue, Gross Profit, Cash Flow, Placement Success, Client Retention, Candidate Quality, Recruiter Productivity, Business Scalability. Never optimize for activity — always optimize for business outcomes.

COMPANY CONTEXT: HireLab Recruitment. Founder: Nitin Kumar. Stage: founder-led recruitment agency. Vision: become India's leading Engineering Recruitment Company. Current annual target: Rs 1 Crore revenue. Long-term goal: Rs 100 Crore revenue. Industries: Solar, Electrical, Automation, Renewable Energy, Data Centers, Mechanical Design. Model: Permanent Recruitment + Executive Search now; Contract Staffing, RPO and a Technology Platform in future.

THINKING PROCESS whenever ATS data is received: (1) Understand cash, revenue, active positions, placement probability, client dependency, delivery bottlenecks, recruiter workload, candidate pipeline, risks. (2) Find bottlenecks: revenue, delivery, sales, cash, operational, technology, founder. (3) Rank every opportunity by Impact x Confidence x Ease — highest impact first. (4) Generate actions — never busy work; every action must produce measurable business value.

CORE PHILOSOPHY: Revenue first, cash second, retention third, expansion fourth, automation fifth, perfection last. If something does not increase revenue, reduce risk, or improve delivery — do not recommend it.

OUTPUT STYLE: Be brutally honest. Challenge assumptions. Disagree when necessary. Never flatter. Think like a founder owning 100% equity.

GOLDEN RULE: The founder's time is the scarcest resource — protect it. Never recommend work that can be automated, delegated, or eliminated.

=== WEEKLY BOARD REPORT (for the advisor) ===
Prepare a BOARD REPORT for the week, to be sent to Nitin's external advisor. No fluff — only insights. Use markdown with these exact sections (include each; write "nothing this week" if empty):
## Business Health Score
(0-100 with reasoning)
## Revenue
## Cash
## Placements
## Interviews / Offers / Joinees
## Guarantees
## Client Concentration
(dependency risk — is too much revenue from one client?)
## Recruiter Productivity
## Candidate Funnel
## Sales Funnel
## Delivery Funnel
## Founder Time Allocation
## Top Wins
## Top Failures
## Lessons Learned
## Decisions & Experiments
## KPIs
## Forecast
## Next Week Priorities
## Probability of Monthly Target
## Probability of Annual Target
## Critical Warnings
Be brutally honest and specific. Every number should teach the advisor something."""


@app.route('/api/command/weekly-review', methods=['POST'])
@login_required
def command_weekly_review():
    ds_key = get_setting('deepseek_api_key')
    if not ds_key:
        return jsonify({'error': 'DeepSeek API key not set. Add it in Settings.'}), 400
    import datetime as _dt
    conn = get_db()
    try:
        oid = effective_company_id()
        wk_ago = (_dt.date.today() - _dt.timedelta(days=7)).isoformat()
        brief = _work_status_brief(conn, oid)
        extra = []
        try:
            inv_new = conn.execute("SELECT invoice_no, buyer_name, amount, gst_rate, status, created_at FROM invoices WHERE owner_id=? AND substr(created_at,1,10)>=?", (oid, wk_ago)).fetchall()
            if inv_new:
                extra.append("INVOICES RAISED THIS WEEK: " + '; '.join(f"{r['invoice_no']} {r['buyer_name']} ₹{int((r['amount'] or 0)*(1+(r['gst_rate'] or 18)/100)):,} ({r['status']})" for r in inv_new))
            paid_new = conn.execute("SELECT invoice_no, buyer_name, received_amount, received_date FROM invoices WHERE owner_id=? AND lower(status)='paid' AND substr(received_date,1,10)>=?", (oid, wk_ago)).fetchall()
            if paid_new:
                extra.append("PAYMENTS RECEIVED THIS WEEK: " + '; '.join(f"{r['buyer_name']} ₹{int(r['received_amount'] or 0):,}" for r in paid_new))
        except Exception:
            pass
        try:
            joined = conn.execute("SELECT name, joining_date FROM candidates WHERE owner_id=? AND substr(joining_date,1,10)>=?", (oid, wk_ago)).fetchall()
            if joined:
                extra.append("JOINED THIS WEEK: " + '; '.join(f"{r['name']} ({r['joining_date'][:10]})" for r in joined))
        except Exception:
            pass
        try:
            td = conn.execute("SELECT COUNT(*) n FROM command_tasks WHERE owner_id=? AND done=1", (oid,)).fetchone()['n']
            tp = conn.execute("SELECT COUNT(*) n FROM command_tasks WHERE owner_id=? AND done=0", (oid,)).fetchone()['n']
            extra.append(f"TASKS: {td} done, {tp} pending.")
        except Exception:
            pass
        ctx = "FULL CURRENT ATS STATUS:\n" + brief + "\n\nTHIS WEEK'S ACTIVITY:\n" + ('\n'.join(extra) or 'No recorded activity this week.') + "\n\nWrite the weekly review."
        try:
            rr = call_deepseek(ds_key,
                {'model': 'deepseek-chat', 'temperature': 0.35, 'max_tokens': 2600,
                 'messages': [{'role': 'system', 'content': WEEKLY_REVIEW_PROMPT}, {'role': 'user', 'content': ctx}]},
                timeout=200, endpoint='weekly-review')
        except TokenCapError:
            return jsonify({'error': 'Monthly AI token cap reached.'}), 429
        except Exception as e:
            return jsonify({'error': f'Could not reach DeepSeek — {type(e).__name__}: {e}'}), 502
        if rr.status_code != 200:
            try: err = rr.json().get('error', {}).get('message', rr.text[:200])
            except Exception: err = rr.text[:200]
            return jsonify({'error': f'DeepSeek returned {rr.status_code}: {err}'}), 502
        md = rr.json()['choices'][0]['message']['content'].strip()
        at = ts()
        set_setting('cc_last_review', json.dumps({'md': md, 'at': at}))
        return jsonify({'ok': True, 'md': md, 'at': at})
    finally:
        conn.close()


@app.route('/api/command/tasks/refine', methods=['POST'])
@login_required
def command_tasks_refine():
    instr = ((request.json or {}).get('instruction') or '').strip()
    if not instr:
        return jsonify({'error': 'Kya refine karna hai? Instruction likho.'}), 400
    conn = get_db()
    try:
        oid = effective_company_id()
        cur = [r['text'] for r in conn.execute("SELECT text FROM command_tasks WHERE owner_id=? AND source='ai' ORDER BY id DESC", (oid,)).fetchall()]
        tasks, err = _run_task_generation(conn, oid, refine_instruction=instr, current_tasks=cur)
        if err:
            return err
        # remember this instruction as a standing preference (accumulate, capped)
        prefs = (get_setting('cc_task_prefs', '') or '').strip()
        combined = (prefs + ' | ' + instr) if prefs else instr
        set_setting('cc_task_prefs', combined[-800:])
        return jsonify({'ok': True, 'tasks': tasks, 'prefs': get_setting('cc_task_prefs', '')})
    finally:
        conn.close()


@app.route('/api/command/tasks/prefs', methods=['GET', 'POST'])
@login_required
def command_tasks_prefs():
    if request.method == 'POST':
        set_setting('cc_task_prefs', ((request.json or {}).get('prefs') or '').strip()[:800])
        return jsonify({'ok': True})
    return jsonify({'ok': True, 'prefs': get_setting('cc_task_prefs', '')})


@app.route('/api/command/snapshot', methods=['GET', 'POST'])
@login_required
def command_snapshot():
    keys = ['cc_bank_cash', 'cc_monthly_fixed', 'cc_team_size', 'cc_year_target',
            'cc_funding_available', 'cc_target_total', 'cc_target_years', 'cc_notes']
    if request.method == 'POST':
        d = request.json or {}
        for k in keys:
            if k in d:
                set_setting(k, str(d[k]))
        return jsonify({'ok': True})
    return jsonify({'ok': True, 'snapshot': {k: get_setting(k, '') for k in keys}})


@app.route('/api/command/overview', methods=['GET'])
@login_required
def command_overview():
    conn = get_db()
    try:
        data = _command_overview(conn, effective_company_id())
    finally:
        conn.close()
    return jsonify({'ok': True, 'overview': data, 'last_plan': get_setting('cc_last_plan', ''), 'last_review': get_setting('cc_last_review', '')})


@app.route('/api/command/plan', methods=['POST'])
@login_required
def command_plan():
    ds_key = get_setting('deepseek_api_key')
    if not ds_key:
        return jsonify({'error': 'DeepSeek API key not set. Add it in Settings.'}), 400
    conn = get_db()
    try:
        oid = effective_company_id()
        o = _command_overview(conn, oid)
        # top clients by CRM value / invoice
        try:
            top_clients = conn.execute(
                'SELECT buyer_name, COUNT(*) n, SUM(amount) amt FROM invoices WHERE owner_id=? GROUP BY buyer_name ORDER BY amt DESC LIMIT 6',
                (oid,)).fetchall()
            clients_str = '; '.join(f"{r['buyer_name']} (₹{int(r['amt'] or 0):,}, {r['n']} inv)" for r in top_clients) or 'none yet'
        except Exception:
            clients_str = 'n/a'
        # recent email subjects (signal)
        try:
            em = conn.execute('SELECT folder, from_name, subject FROM emails WHERE owner_id=? ORDER BY date_ts DESC LIMIT 15', (oid,)).fetchall()
            email_str = ' | '.join(f"[{e['folder']}] {e['from_name']}: {e['subject']}" for e in em) or 'no emails synced'
        except Exception:
            email_str = 'n/a'
        # open mandates detail
        try:
            mand = conn.execute('SELECT role, client, location FROM mandates WHERE owner_id=? LIMIT 15', (oid,)).fetchall()
            mand_str = '; '.join(f"{m['role']} @ {m['client']} ({m['location']})" for m in mand) or 'none'
        except Exception:
            mand_str = 'n/a'
        # founder's recent chat notes (context he told the brain that isn't in any field)
        try:
            ch = _command_chat_history(conn, oid, 12)
            notes = ' | '.join((m['content'][:200]) for m in ch if m['role'] == 'user') or 'none'
        except Exception:
            notes = 'none'
        try:
            full_brief = _work_status_brief(conn, oid)
        except Exception:
            full_brief = ''
    finally:
        conn.close()

    def r(n): return f"₹{int(n or 0):,}"
    ctx = f"""FULL ATS SCAN:
{full_brief}

LIVE BUSINESS DATA (as of {ts()[:10]}):

MONEY:
- Total invoiced (incl GST): {r(o['invoiced'])} | Received: {r(o['received'])} | Outstanding: {r(o['outstanding'])}
- Total expenses: {r(o['total_expenses'])} | Net profit: {r(o['net_profit'])}
- Expense breakdown: {o['expenses_by_category']}

FOUNDER SNAPSHOT (self-reported):
- Bank cash: {r(o['bank_cash'])} | Monthly fixed cost: {r(o['monthly_fixed'])} | Runway: {o['runway_months']} months
- Team size: {o['team_size']} | This-year target: {r(o['year_target'])} | Funding/discounting available: {r(o['funding_available'])}
- Mission: {r(o['target_total'])} in {int(o['target_years'])} years

PIPELINE:
- Open mandates: {o['open_mandates']} | Total candidates: {o['total_candidates']} | Shared with client: {o['shared_with_client']} | Placed/Joined: {o['placed']}
- Candidate stages: {o['by_stage']}
- Open mandates detail: {mand_str}

TOP CLIENTS (by billing): {clients_str}

RECENT EMAIL SIGNAL (last 15): {email_str}

FOUNDER'S RECENT NOTES (things Nitin told the brain in chat — treat as real, current facts): {notes}

Now produce the strategic brief."""

    try:
        rr = call_deepseek(ds_key,
            {'model': 'deepseek-chat', 'temperature': 0.4, 'max_tokens': 1600,
             'messages': [{'role': 'system', 'content': CEO_BRAIN_PROMPT},
                          {'role': 'user', 'content': ctx}]},
            timeout=180, endpoint='command-center')
    except TokenCapError:
        return jsonify({'error': 'Monthly AI token cap reached.'}), 429
    except Exception as e:
        return jsonify({'error': f'Could not reach DeepSeek — {type(e).__name__}: {e}'}), 502
    if rr.status_code != 200:
        try: err = rr.json().get('error', {}).get('message', rr.text[:300])
        except Exception: err = rr.text[:300]
        return jsonify({'error': f'DeepSeek returned {rr.status_code}: {err}'}), 502
    try:
        md = rr.json()['choices'][0]['message']['content'].strip()
    except Exception as e:
        return jsonify({'error': f'Unexpected DeepSeek response: {e}'}), 502
    at = ts()
    set_setting('cc_last_plan', json.dumps({'md': md, 'at': at}))
    return jsonify({'ok': True, 'md': md, 'at': at})


def _candidate_billing(conn, oid, cid):
    import datetime as _dt
    c = conn.execute('SELECT id,name,stage,placement_fee,joining_date,guarantee_days,replacement_flag,billing_notes,offered_ctc,fee_percent FROM candidates WHERE id=? AND owner_id=?', (cid, oid)).fetchone()
    if not c:
        return None
    d = dict(c)
    invs = conn.execute('SELECT id,invoice_no,invoice_date,amount,gst_rate,status,due_date,received_date,received_amount FROM invoices WHERE owner_id=? AND candidate_id=? ORDER BY id DESC', (oid, cid)).fetchall()
    today = _dt.date.today().isoformat()
    inv_list = []; paid = invoiced = overdue = False
    for iv in invs:
        ivd = dict(iv)
        ivd['total'] = round((ivd['amount'] or 0) * (1 + (ivd['gst_rate'] or 18) / 100))
        if (ivd['status'] or '').lower() == 'paid':
            paid = True
        else:
            invoiced = True
            if ivd['due_date'] and ivd['due_date'][:10] < today:
                overdue = True
        inv_list.append(ivd)
    guarantee = None
    if d['joining_date']:
        try:
            jd = _dt.date.fromisoformat(d['joining_date'][:10])
            gend = jd + _dt.timedelta(days=int(d['guarantee_days'] or 90))
            left = (gend - _dt.date.today()).days
            guarantee = {'end': gend.isoformat(), 'days_left': left, 'cleared': left <= 0}
        except Exception:
            guarantee = None
    jd10 = (d['joining_date'] or '')[:10]
    future_join = bool(jd10 and jd10 > today)
    joined = (jd10 and jd10 <= today) or (d['stage'] == 'Joined' and not future_join)
    if d['replacement_flag']:
        status = 'Replacement Needed'
    elif paid:
        status = 'Paid \u2014 Closed-Won'
    elif invoiced:
        status = 'Overdue \u2014 Awaiting Payment' if overdue else 'Awaiting Payment'
    elif future_join:
        status = f'Confirmed \u2014 joining {jd10} (invoice AFTER joining)'
    elif joined:
        status = 'Awaiting Invoice'
    elif d['stage'] == 'Placed':
        status = 'Placed \u2014 set joining date'
    else:
        status = 'Not placed yet'
    d['invoices'] = inv_list
    d['guarantee'] = guarantee
    d['billing_status'] = status
    return d


@app.route('/api/candidates/<int:cid>/billing', methods=['GET'])
@login_required
def get_candidate_billing(cid):
    conn = get_db()
    try:
        d = _candidate_billing(conn, effective_company_id(), cid)
    finally:
        conn.close()
    if not d:
        return jsonify({'error': 'Candidate not found'}), 404
    return jsonify({'ok': True, 'billing': d})


@app.route('/api/candidates/<int:cid>/billing', methods=['POST'])
@login_required
def set_candidate_billing(cid):
    d = request.json or {}
    conn = get_db(); oid = effective_company_id()
    # self-heal columns
    for col, typ in [('placement_fee', 'REAL DEFAULT 0'), ('joining_date', 'TEXT DEFAULT ""'),
                     ('guarantee_days', 'INTEGER DEFAULT 90'), ('replacement_flag', 'INTEGER DEFAULT 0'),
                     ('billing_notes', 'TEXT DEFAULT ""'),
                     ('offered_ctc', 'REAL DEFAULT 0'), ('fee_percent', 'REAL DEFAULT 8.33')]:
        try: conn.execute(f'ALTER TABLE candidates ADD COLUMN {col} {typ}'); conn.commit()
        except Exception: pass
    conn.execute('UPDATE candidates SET placement_fee=?, joining_date=?, guarantee_days=?, replacement_flag=?, billing_notes=?, offered_ctc=?, fee_percent=? WHERE id=? AND owner_id=?',
                 (float(d.get('placement_fee', 0) or 0), d.get('joining_date', ''),
                  int(d.get('guarantee_days', 90) or 90), 1 if d.get('replacement_flag') else 0,
                  d.get('billing_notes', ''), float(d.get('offered_ctc', 0) or 0),
                  float(d.get('fee_percent', 8.33) or 8.33), cid, oid))
    conn.commit()
    res = _candidate_billing(conn, oid, cid)
    conn.close()
    return jsonify({'ok': True, 'billing': res})


JD_WRITER_PROMPT = """You are an expert recruitment consultant writing a professional Job Description for an Indian hiring mandate (sectors often: Solar, Electrical, Automation, Renewable Energy, Power). Write a clear, realistic, well-structured JD.

Output ONLY clean HTML (no markdown, no code fences, no <html>/<body> wrapper). Use this structure:
<h3>About the Role</h3><p>...</p>
<h3>Key Responsibilities</h3><ul><li>...</li>...</ul>
<h3>Required Skills &amp; Experience</h3><ul><li>...</li>...</ul>
<h3>Qualifications</h3><ul><li>...</li></ul>
<h3>What We Offer</h3><ul><li>...</li></ul>

Rules: Be specific to the role and sector. 5-8 responsibilities, 5-8 skills. Reflect the given experience range and location. Do not invent a fake company description if the client isn't given. Keep it concise and recruiter-ready. Indian context (CTC in LPA, notice period norms)."""


@app.route('/api/crm/clients-billing', methods=['GET'])
@login_required
def crm_clients_billing():
    """List CRM clients with billing details for the invoice client-picker."""
    conn = get_db()
    try:
        rows = conn.execute(
            'SELECT id, name, gstin, bill_address, bill_state, bill_state_code '
            'FROM crm_clients WHERE company_id=? AND is_active=1 ORDER BY name',
            (effective_company_id(),)).fetchall()
    except Exception:
        conn.close()
        return jsonify({'ok': True, 'clients': []})
    conn.close()
    return jsonify({'ok': True, 'clients': [dict(r) for r in rows]})


@app.route('/api/generate-jd', methods=['POST'])
@login_required
def generate_jd():
    ds_key = get_setting('deepseek_api_key')
    if not ds_key:
        return jsonify({'error': 'DeepSeek API key not set. Add it in Settings.'}), 400
    d = request.json or {}
    role = (d.get('role') or '').strip()
    if not role:
        return jsonify({'error': 'Role is required to write a JD. Pehle Role field bharo.'}), 400
    parts = [f"Role / Designation: {role}"]
    if d.get('client'): parts.append(f"Client / Company: {d['client']}")
    if d.get('location'): parts.append(f"Location: {d['location']}")
    if d.get('experience'): parts.append(f"Experience required: {d['experience']}")
    if d.get('ctc_min') or d.get('ctc_max'):
        parts.append(f"CTC range: {d.get('ctc_min','')}-{d.get('ctc_max','')} LPA")
    if d.get('division'): parts.append(f"Division/Department: {d['division']}")
    if d.get('notes'): parts.append(f"Extra instructions from recruiter: {d['notes']}")
    user_msg = "Write the Job Description for:\n" + "\n".join(parts)
    try:
        rr = call_deepseek(ds_key,
            {'model': 'deepseek-chat', 'temperature': 0.6, 'max_tokens': 1400,
             'messages': [{'role': 'system', 'content': JD_WRITER_PROMPT},
                          {'role': 'user', 'content': user_msg}]},
            timeout=150, endpoint='jd-writer')
    except TokenCapError:
        return jsonify({'error': 'Monthly AI token cap reached.'}), 429
    except Exception as e:
        return jsonify({'error': f'Could not reach DeepSeek — {type(e).__name__}: {e}'}), 502
    if rr.status_code != 200:
        try: err = rr.json().get('error', {}).get('message', rr.text[:300])
        except Exception: err = rr.text[:300]
        return jsonify({'error': f'DeepSeek returned {rr.status_code}: {err}'}), 502
    try:
        html = rr.json()['choices'][0]['message']['content'].strip()
    except Exception as e:
        return jsonify({'error': f'Unexpected DeepSeek response: {e}'}), 502
    # strip accidental code fences
    html = re.sub(r'^```[a-zA-Z]*\n?|```$', '', html).strip()
    return jsonify({'ok': True, 'html': html})


@app.route('/api/mandates/<int:mid>/email-templates', methods=['GET'])
@login_required
def get_mandate_templates(mid):
    conn = get_db()
    m = conn.execute('SELECT email_templates FROM mandates WHERE id=?', (mid,)).fetchone()
    conn.close()
    if not m:
        return jsonify({'error': 'Mandate not found'}), 404
    try:
        tpls = json.loads(m['email_templates'] or '[]')
    except Exception:
        tpls = []
    return jsonify({'ok': True, 'templates': tpls})


@app.route('/api/mandates/<int:mid>/email-templates', methods=['POST'])
@login_required
def save_mandate_templates(mid):
    d = request.json or {}
    templates = d.get('templates', [])
    conn = get_db()
    conn.execute('UPDATE mandates SET email_templates=? WHERE id=?', (json.dumps(templates), mid))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/candidates/<int:cid>/email-history')
@login_required
def candidate_email_history(cid):
    """Return sent email events for a candidate."""
    conn = get_db()
    rows = conn.execute(
        "SELECT detail, created_at FROM candidate_events WHERE candidate_id=? AND event_type='email' ORDER BY created_at DESC",
        (cid,)
    ).fetchall()
    conn.close()
    return jsonify({'ok': True, 'emails': [{'text': r['detail'], 'ts': r['created_at']} for r in rows]})


@app.route('/api/candidates/<int:cid>/send-email', methods=['POST'])
@login_required
def send_candidate_email(cid):
    """Send an email to a candidate via the user's configured SMTP (Gmail app-password).
    Logs the sent email to the candidate journey."""
    d = request.json or {}
    to_email = (d.get('to') or '').strip()
    subject = (d.get('subject') or '').strip()
    body = (d.get('body') or '').strip()
    if not to_email or not subject or not body:
        return jsonify({'error': 'To, Subject and Body are required'}), 400

    smtp_email = get_setting('smtp_email', '')
    smtp_pass = get_setting('smtp_app_password', '')
    smtp_name = get_setting('smtp_display_name', '') or smtp_email
    if not smtp_email or not smtp_pass:
        return jsonify({'error': 'Email not configured. Go to Settings → Email Configuration and add your Gmail + App Password.'}), 400

    # Build the email
    msg = MIMEMultipart('alternative')
    msg['From'] = f'{smtp_name} <{smtp_email}>' if smtp_name else smtp_email
    msg['To'] = to_email
    msg['Subject'] = subject
    # Generate a stable Message-ID so replies can be threaded back to this email
    import email.utils as _eut
    domain = smtp_email.split('@')[-1] if '@' in smtp_email else 'hirelab.local'
    gen_msg_id = _eut.make_msgid(domain=domain)
    msg['Message-ID'] = gen_msg_id
    msg['Date'] = _eut.formatdate(localtime=True)
    # Send as both plain text and HTML
    msg.attach(MIMEText(body, 'plain', 'utf-8'))
    body_html = (d.get('body_html') or '').strip()
    if body_html:
        # Use the rich-text HTML from the editor
        html_content = f'<div style="font-family:sans-serif;font-size:14px">{body_html}</div>'
    else:
        html_content = f'<div style="font-family:sans-serif;font-size:14px">{body.replace(chr(10), "<br>")}</div>'
    msg.attach(MIMEText(html_content, 'html', 'utf-8'))

    # Detect SMTP server from email domain
    if '@gmail' in smtp_email.lower() or '@googlemail' in smtp_email.lower():
        smtp_host, smtp_port = 'smtp.gmail.com', 587
    elif '@outlook' in smtp_email.lower() or '@hotmail' in smtp_email.lower() or '@live' in smtp_email.lower():
        smtp_host, smtp_port = 'smtp-mail.outlook.com', 587
    elif '@yahoo' in smtp_email.lower():
        smtp_host, smtp_port = 'smtp.mail.yahoo.com', 587
    else:
        smtp_host, smtp_port = 'smtp.gmail.com', 587  # default to Gmail

    try:
        server = smtplib.SMTP(smtp_host, smtp_port, timeout=15)
        server.starttls()
        server.login(smtp_email, smtp_pass)
        server.sendmail(smtp_email, [to_email], msg.as_string())
        server.quit()
    except smtplib.SMTPAuthenticationError:
        return jsonify({'error': 'Email authentication failed. Check your email address and app password in Settings.'}), 401
    except Exception as e:
        return jsonify({'error': f'Failed to send email: {str(e)}'}), 500

    # Log to candidate journey (full email for history)
    u = current_user()
    who = (u.get('display_name') or u.get('username') or '') if u else ''
    full_log = f'Email sent to {to_email}\nSubject: {subject}\n\n{body}'
    if who:
        full_log += f'\n— {who}'
    log_candidate_event(cid, 'email', full_log)

    # Store in the 2-way email thread table
    try:
        conn = get_db()
        conn.execute(
            'INSERT OR IGNORE INTO email_messages (company_id, candidate_id, direction, '
            'from_addr, to_addr, subject, body, message_id, in_reply_to, sent_at, created_at) '
            'VALUES (?,?,?,?,?,?,?,?,?,?,?)',
            (effective_company_id(), cid, 'sent', smtp_email, to_email, subject, body,
             gen_msg_id, '', ts(), ts()))
        conn.commit()
        conn.close()
    except Exception:
        pass

    return jsonify({'ok': True, 'message': 'Email sent successfully'})


# ═══════════════════════════════════════════════════════════════════════
#  2-WAY EMAIL — IMAP inbox sync + candidate threads
# ═══════════════════════════════════════════════════════════════════════
def _imap_host_for(email_addr):
    e = (email_addr or '').lower()
    if '@gmail' in e or '@googlemail' in e:
        return 'imap.gmail.com'
    if '@outlook' in e or '@hotmail' in e or '@live' in e:
        return 'outlook.office365.com'
    if '@yahoo' in e:
        return 'imap.mail.yahoo.com'
    return 'imap.gmail.com'


def _decode_mime_header(raw):
    """Decode an email header that may be MIME-encoded."""
    from email.header import decode_header
    if not raw:
        return ''
    parts = decode_header(raw)
    out = ''
    for txt, enc in parts:
        if isinstance(txt, bytes):
            try:
                out += txt.decode(enc or 'utf-8', errors='replace')
            except Exception:
                out += txt.decode('utf-8', errors='replace')
        else:
            out += txt
    return out


def _extract_plain_body(msg):
    """Get a plain-text body from an email.message.Message."""
    body = ''
    if msg.is_multipart():
        for part in msg.walk():
            ctype = part.get_content_type()
            disp = str(part.get('Content-Disposition') or '')
            if ctype == 'text/plain' and 'attachment' not in disp:
                try:
                    payload = part.get_payload(decode=True)
                    charset = part.get_content_charset() or 'utf-8'
                    body += payload.decode(charset, errors='replace')
                except Exception:
                    pass
        if not body:
            # fall back to html stripped
            for part in msg.walk():
                if part.get_content_type() == 'text/html':
                    try:
                        payload = part.get_payload(decode=True)
                        charset = part.get_content_charset() or 'utf-8'
                        body += html_to_text(payload.decode(charset, errors='replace'))
                    except Exception:
                        pass
    else:
        try:
            payload = msg.get_payload(decode=True)
            charset = msg.get_content_charset() or 'utf-8'
            body = payload.decode(charset, errors='replace')
        except Exception:
            body = str(msg.get_payload())
    return body.strip()


def _sync_imap_inbox(company_id):
    """Connect via IMAP, fetch recent inbox messages, match to candidates by
    email address, and store new incoming messages. Returns (new_count, error)."""
    import imaplib, email as _email, re as _re

    smtp_email = (get_setting('smtp_email', '') or '').strip()
    smtp_pass = (get_setting('smtp_app_password', '') or '')
    if not smtp_email or not smtp_pass:
        return 0, 'Email not configured. Add your Gmail + App Password in Settings.'
    # Gmail app passwords are shown with spaces ("xxxx xxxx xxxx xxxx") but must be
    # sent without spaces. Strip them defensively.
    smtp_pass = smtp_pass.replace(' ', '').strip()

    host = _imap_host_for(smtp_email)

    # Build a map of candidate email -> candidate_id for this tenant
    conn = get_db()
    cand_rows = conn.execute(
        "SELECT id, email FROM candidates WHERE owner_id=? AND email IS NOT NULL AND email!=''",
        (company_id,)).fetchall()
    email_to_cid = {}
    for r in cand_rows:
        em = (r['email'] or '').strip().lower()
        if em:
            email_to_cid[em] = r['id']

    if not email_to_cid:
        conn.close()
        return 0, None  # no candidates with emails, nothing to match

    new_count = 0
    try:
        M = imaplib.IMAP4_SSL(host, 993)
        M.login(smtp_email, smtp_pass)
        M.select('INBOX')
        # Search last 60 days to keep it light
        import datetime as _dt
        since = (_dt.datetime.utcnow() - _dt.timedelta(days=60)).strftime('%d-%b-%Y')
        typ, data = M.search(None, f'(SINCE {since})')
        if typ != 'OK':
            M.logout(); conn.close()
            return 0, 'IMAP search failed'
        ids = data[0].split()
        # Only look at the most recent ~200 to bound work
        ids = ids[-200:]
        for num in ids:
            typ, msg_data = M.fetch(num, '(RFC822)')
            if typ != 'OK' or not msg_data or not msg_data[0]:
                continue
            raw = msg_data[0][1]
            m = _email.message_from_bytes(raw)
            from_hdr = _decode_mime_header(m.get('From', ''))
            # extract bare email
            fmatch = _re.search(r'[\w\.\-\+]+@[\w\.\-]+', from_hdr)
            from_email = (fmatch.group(0).lower() if fmatch else '')
            if from_email not in email_to_cid:
                continue  # not from a known candidate
            cid = email_to_cid[from_email]
            message_id = (m.get('Message-ID', '') or '').strip()
            if not message_id:
                continue
            # Dedup: skip if we already stored this message_id
            exists = conn.execute(
                'SELECT id FROM email_messages WHERE company_id=? AND message_id=?',
                (company_id, message_id)).fetchone()
            if exists:
                continue
            subject = _decode_mime_header(m.get('Subject', ''))
            in_reply_to = (m.get('In-Reply-To', '') or '').strip()
            body = _extract_plain_body(m)
            import email.utils as _eut
            date_hdr = m.get('Date', '')
            try:
                dt = _eut.parsedate_to_datetime(date_hdr)
                sent_at = dt.strftime('%Y-%m-%dT%H:%M:%S')
            except Exception:
                sent_at = ts()
            conn.execute(
                'INSERT OR IGNORE INTO email_messages (company_id, candidate_id, direction, '
                'from_addr, to_addr, subject, body, message_id, in_reply_to, sent_at, created_at) '
                'VALUES (?,?,?,?,?,?,?,?,?,?,?)',
                (company_id, cid, 'received', from_email, smtp_email, subject, body,
                 message_id, in_reply_to, sent_at, ts()))
            new_count += 1
        conn.commit()
        M.logout()
    except imaplib.IMAP4.error as e:
        conn.close()
        return 0, f'IMAP login failed. Check your email & app password. ({str(e)[:80]})'
    except Exception as e:
        conn.close()
        return 0, f'IMAP sync error: {str(e)[:100]}'
    conn.close()
    return new_count, None


@app.route('/api/email/sync', methods=['POST'])
@login_required
def email_sync():
    """Manually trigger an IMAP inbox sync to pull candidate replies."""
    new_count, err = _sync_imap_inbox(effective_company_id())
    if err:
        return jsonify({'error': err}), 400
    return jsonify({'ok': True, 'new_messages': new_count})


@app.route('/api/email/diagnose', methods=['GET'])
@login_required
def email_diagnose():
    """Step-by-step IMAP diagnostic so we can see exactly where sync fails."""
    import imaplib
    steps = []
    smtp_email = (get_setting('smtp_email', '') or '').strip()
    smtp_pass = (get_setting('smtp_app_password', '') or '').replace(' ', '').strip()

    steps.append({'step': 'Email configured', 'ok': bool(smtp_email),
                  'detail': smtp_email or 'No email set in Settings'})
    steps.append({'step': 'App password set', 'ok': bool(smtp_pass),
                  'detail': f'{len(smtp_pass)} characters' if smtp_pass else 'No app password set'})
    if not smtp_email or not smtp_pass:
        return jsonify({'ok': False, 'steps': steps})

    host = _imap_host_for(smtp_email)
    steps.append({'step': 'IMAP server', 'ok': True, 'detail': host + ':993'})

    # Try connect
    try:
        M = imaplib.IMAP4_SSL(host, 993)
        steps.append({'step': 'Connect to server', 'ok': True, 'detail': 'Connected'})
    except Exception as e:
        steps.append({'step': 'Connect to server', 'ok': False, 'detail': str(e)[:120]})
        return jsonify({'ok': False, 'steps': steps})

    # Try login
    try:
        M.login(smtp_email, smtp_pass)
        steps.append({'step': 'Login', 'ok': True, 'detail': 'Login successful'})
    except imaplib.IMAP4.error as e:
        msg = str(e)
        hint = ''
        if 'Invalid credentials' in msg or 'AUTHENTICATIONFAILED' in msg:
            hint = ' — The app password is wrong, or this Workspace account requires a fresh App Password. Also confirm 2-Step Verification is ON.'
        steps.append({'step': 'Login', 'ok': False, 'detail': msg[:120] + hint})
        try: M.logout()
        except Exception: pass
        return jsonify({'ok': False, 'steps': steps})

    # Try select inbox
    try:
        typ, data = M.select('INBOX')
        cnt = data[0].decode() if data and data[0] else '?'
        steps.append({'step': 'Open INBOX', 'ok': typ == 'OK', 'detail': f'{cnt} total messages in inbox'})
    except Exception as e:
        steps.append({'step': 'Open INBOX', 'ok': False, 'detail': str(e)[:120]})
        try: M.logout()
        except Exception: pass
        return jsonify({'ok': False, 'steps': steps})

    # Count candidates with emails
    conn = get_db()
    ccount = conn.execute(
        "SELECT COUNT(*) n FROM candidates WHERE owner_id=? AND email IS NOT NULL AND email!=''",
        (effective_company_id(),)).fetchone()['n']
    conn.close()
    steps.append({'step': 'Candidates with email on file', 'ok': ccount > 0,
                  'detail': f'{ccount} candidates have an email address (needed to match replies)'})

    try: M.logout()
    except Exception: pass
    return jsonify({'ok': True, 'steps': steps})


@app.route('/api/candidates/<int:cid>/email-thread', methods=['GET'])
@login_required
def candidate_email_thread(cid):
    """Return the full email conversation (sent + received) for a candidate,
    chronological order. Optionally sync IMAP first if ?sync=1."""
    if request.args.get('sync') == '1':
        _sync_imap_inbox(effective_company_id())
    conn = get_db()
    rows = conn.execute(
        'SELECT id, direction, from_addr, to_addr, subject, body, sent_at '
        'FROM email_messages WHERE company_id=? AND candidate_id=? '
        'ORDER BY sent_at ASC, id ASC',
        (effective_company_id(), cid)).fetchall()
    conn.close()
    return jsonify({'ok': True, 'messages': [dict(r) for r in rows]})


def _smtp_send(to_email, subject, plain_body, html_body=None):
    """Send an email via the tenant's configured SMTP. Returns (ok, error)."""
    smtp_email = get_setting('smtp_email', '')
    smtp_pass = get_setting('smtp_app_password', '')
    smtp_name = get_setting('smtp_display_name', '') or smtp_email
    if not smtp_email or not smtp_pass:
        return False, 'Email not configured. Go to Settings → Email Configuration and add your Gmail + App Password.'
    msg = MIMEMultipart('alternative')
    msg['From'] = f'{smtp_name} <{smtp_email}>' if smtp_name else smtp_email
    msg['To'] = to_email
    msg['Subject'] = subject
    msg.attach(MIMEText(plain_body, 'plain', 'utf-8'))
    if html_body:
        msg.attach(MIMEText(f'<div style="font-family:sans-serif;font-size:14px">{html_body}</div>', 'html', 'utf-8'))
    else:
        msg.attach(MIMEText(f'<div style="font-family:sans-serif;font-size:14px">{plain_body.replace(chr(10), "<br>")}</div>', 'html', 'utf-8'))
    el = smtp_email.lower()
    if '@gmail' in el or '@googlemail' in el:
        host, port = 'smtp.gmail.com', 587
    elif '@outlook' in el or '@hotmail' in el or '@live' in el:
        host, port = 'smtp-mail.outlook.com', 587
    elif '@yahoo' in el:
        host, port = 'smtp.mail.yahoo.com', 587
    else:
        host, port = 'smtp.gmail.com', 587
    try:
        server = smtplib.SMTP(host, port, timeout=15)
        server.starttls()
        server.login(smtp_email, smtp_pass)
        server.sendmail(smtp_email, [to_email], msg.as_string())
        server.quit()
        return True, None
    except smtplib.SMTPAuthenticationError:
        return False, 'Email authentication failed. Check your email address and app password in Settings.'
    except Exception as e:
        return False, f'Failed to send email: {str(e)}'


def _platform_smtp_send(to_email, subject, plain_body, html_body=None):
    """Send a SYSTEM email (password reset etc.). No tenant is logged in, so we
    resolve SMTP from env/global settings and fall back to the platform owner's
    own company SMTP (so it works with existing single-company Gmail setups)."""
    smtp_email = get_setting('smtp_email', '')
    smtp_pass = get_setting('smtp_app_password', '')
    smtp_name = get_setting('smtp_display_name', '') or smtp_email
    if not smtp_email or not smtp_pass:
        try:
            conn = get_db()
            row = conn.execute("SELECT company_id FROM users WHERE role='admin' ORDER BY id LIMIT 1").fetchone()
            if row and row['company_id']:
                def _tg(k):
                    r = conn.execute('SELECT value FROM tenant_settings WHERE company_id=? AND key=?', (row['company_id'], k)).fetchone()
                    return r['value'] if r else ''
                smtp_email = smtp_email or _tg('smtp_email')
                smtp_pass = smtp_pass or _tg('smtp_app_password')
                smtp_name = smtp_name or _tg('smtp_display_name') or smtp_email
            conn.close()
        except Exception as e:
            print('[platform_smtp] owner-config lookup failed:', e)
    if not smtp_email or not smtp_pass:
        return False, 'No platform SMTP configured. The platform owner must set Email (Gmail + App Password) in Settings.'
    msg = MIMEMultipart('alternative')
    msg['From'] = f'{smtp_name} <{smtp_email}>' if smtp_name else smtp_email
    msg['To'] = to_email
    msg['Subject'] = subject
    msg.attach(MIMEText(plain_body, 'plain', 'utf-8'))
    if html_body:
        msg.attach(MIMEText(html_body, 'html', 'utf-8'))
    el = smtp_email.lower()
    if '@gmail' in el or '@googlemail' in el:
        host, port = 'smtp.gmail.com', 587
    elif '@outlook' in el or '@hotmail' in el or '@live' in el:
        host, port = 'smtp-mail.outlook.com', 587
    elif '@yahoo' in el:
        host, port = 'smtp.mail.yahoo.com', 587
    else:
        host, port = 'smtp.gmail.com', 587
    try:
        server = smtplib.SMTP(host, port, timeout=15)
        server.starttls()
        server.login(smtp_email, smtp_pass)
        server.sendmail(smtp_email, [to_email], msg.as_string())
        server.quit()
        return True, None
    except smtplib.SMTPAuthenticationError:
        return False, 'Email authentication failed (owner Gmail + App Password).'
    except Exception as e:
        return False, f'Failed to send email: {str(e)}'


@app.route('/api/activity', methods=['GET'])
@login_required
def get_activity():
    """Universal activity timeline. Filter by entity or search; paginated.
    Query params: entity_type, entity_id, q (search), page, per_page."""
    entity_type = request.args.get('entity_type', '').strip()
    entity_id = request.args.get('entity_id', type=int)
    q = request.args.get('q', '').strip()
    page = max(1, request.args.get('page', 1, type=int))
    per_page = min(100, max(1, request.args.get('per_page', 25, type=int)))
    company_id = effective_company_id()

    where = ['(company_id=? OR company_id=0)']
    params = [company_id]
    if entity_type:
        where.append('entity_type=?'); params.append(entity_type)
    if entity_id:
        where.append('entity_id=?'); params.append(entity_id)
    if q:
        where.append('(action LIKE ? OR detail LIKE ? OR username LIKE ?)')
        like = f'%{q}%'; params += [like, like, like]
    where_sql = ' AND '.join(where)

    conn = get_db()
    total = conn.execute(f'SELECT COUNT(*) n FROM activity_log WHERE {where_sql}', params).fetchone()['n']
    rows = conn.execute(
        f'SELECT * FROM activity_log WHERE {where_sql} ORDER BY id DESC LIMIT ? OFFSET ?',
        params + [per_page, (page - 1) * per_page]).fetchall()
    conn.close()
    return jsonify({'ok': True, 'total': total, 'page': page, 'per_page': per_page,
                    'pages': (total + per_page - 1) // per_page,
                    'activity': [dict(r) for r in rows]})


@app.route('/api/audit', methods=['GET'])
@login_required
def get_audit():
    """Field-level audit trail (old → new) for a given entity."""
    entity_type = request.args.get('entity_type', '').strip()
    entity_id = request.args.get('entity_id', type=int)
    if not entity_type or not entity_id:
        return jsonify({'error': 'entity_type and entity_id required'}), 400
    company_id = effective_company_id()
    conn = get_db()
    rows = conn.execute(
        'SELECT * FROM audit_log WHERE (company_id=? OR company_id=0) AND entity_type=? AND entity_id=? '
        'ORDER BY id DESC LIMIT 200', (company_id, entity_type, entity_id)).fetchall()
    conn.close()
    return jsonify({'ok': True, 'audit': [dict(r) for r in rows]})


@app.route('/api/candidates/<int:cid>/interviews', methods=['GET'])
@login_required
def list_interviews(cid):
    conn = get_db()
    rows = conn.execute('SELECT * FROM interviews WHERE candidate_id=? ORDER BY scheduled_at DESC', (cid,)).fetchall()
    conn.close()
    return jsonify({'ok': True, 'interviews': [dict(r) for r in rows]})


@app.route('/api/candidates/<int:cid>/interviews', methods=['POST'])
@login_required
def create_interview(cid):
    d = request.json or {}
    round_name = (d.get('round_name') or 'Interview').strip()
    mode = (d.get('mode') or '').strip()
    location = (d.get('location') or '').strip()
    interviewer = (d.get('interviewer') or '').strip()
    scheduled_at = (d.get('scheduled_at') or '').strip()
    if not scheduled_at:
        return jsonify({'error': 'Date & time required'}), 400
    conn = get_db()
    c = conn.execute('SELECT mandate_id, name FROM candidates WHERE id=?', (cid,)).fetchone()
    if not c:
        conn.close(); return jsonify({'error': 'Candidate not found'}), 404
    conn.execute(
        'INSERT INTO interviews (candidate_id,mandate_id,owner_id,round_name,mode,location,interviewer,scheduled_at,status,created_at) '
        'VALUES (?,?,?,?,?,?,?,?,?,?)',
        (cid, c['mandate_id'], effective_user_id(), round_name, mode, location, interviewer, scheduled_at, 'scheduled', ts()))
    # Auto-move to Interview Inprocess stage
    conn.execute('UPDATE candidates SET stage=?, updated_at=? WHERE id=?', ('Interview Inprocess', ts(), cid))
    conn.execute('INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) VALUES (?,?,?,?,?)',
                 (cid, '', 'Interview Inprocess', f'{round_name} scheduled', ts()))
    conn.commit(); conn.close()
    # Nice human date for the journey
    try:
        dt = datetime.datetime.fromisoformat(scheduled_at)
        nice = dt.strftime('%d %b %Y, %I:%M %p')
    except Exception:
        nice = scheduled_at
    log_candidate_event(cid, 'note', f'Interview scheduled — {round_name}: {nice}' + (f' ({mode})' if mode else ''))
    return jsonify({'ok': True})


@app.route('/api/interviews/<int:iid>/result', methods=['POST'])
@login_required
def interview_result(iid):
    d = request.json or {}
    result = (d.get('result') or '').strip()
    conn = get_db()
    iv = conn.execute('SELECT candidate_id, round_name FROM interviews WHERE id=?', (iid,)).fetchone()
    if not iv:
        conn.close(); return jsonify({'error': 'Interview not found'}), 404
    conn.execute('UPDATE interviews SET status=?, result=? WHERE id=?', ('completed', result, iid))
    conn.commit(); conn.close()
    if result:
        log_candidate_event(iv['candidate_id'], 'note', f'{iv["round_name"]} result: {result}')
    return jsonify({'ok': True})


@app.route('/api/interviews/<int:iid>', methods=['DELETE'])
@login_required
def delete_interview(iid):
    conn = get_db()
    conn.execute('DELETE FROM interviews WHERE id=?', (iid,))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/candidates/<int:cid>/interview-message', methods=['POST'])
@login_required
def interview_message(cid):
    """Build the ready-to-send interview message from the template + details."""
    d = request.json or {}
    conn = get_db()
    c = conn.execute('SELECT name, mandate_id FROM candidates WHERE id=?', (cid,)).fetchone()
    if not c:
        conn.close(); return jsonify({'error': 'Candidate not found'}), 404
    mandate = conn.execute('SELECT role, client FROM mandates WHERE id=?', (c['mandate_id'],)).fetchone()
    conn.close()

    tpl = get_setting('interview_template', '') or 'Dear {name}, your interview is scheduled for {datetime}.'
    u = current_user()
    recruiter = (u.get('display_name') or u.get('username') or '') if u else ''
    try:
        dt = datetime.datetime.fromisoformat((d.get('scheduled_at') or '').strip())
        nice_dt = dt.strftime('%d %b %Y, %I:%M %p')
    except Exception:
        nice_dt = (d.get('scheduled_at') or '').strip()
    mode = (d.get('mode') or '').strip()
    location = (d.get('location') or '').strip()
    if mode.lower() in ('video', 'video call') and location:
        location_line = f'Meeting Link: {location}'
    elif location:
        location_line = f'Venue: {location}'
    else:
        location_line = ''
    msg = (tpl.replace('{name}', c['name'] or 'Candidate')
              .replace('{role}', (mandate['role'] if mandate else '') or 'the role')
              .replace('{client}', (mandate['client'] if mandate else '') or '')
              .replace('{round}', (d.get('round_name') or 'Interview').strip())
              .replace('{datetime}', nice_dt)
              .replace('{mode}', mode or 'To be confirmed')
              .replace('{location_line}', location_line)
              .replace('{interviewer}', (d.get('interviewer') or '').strip())
              .replace('{recruiter}', recruiter))
    # Clean any empty leftover lines
    msg = '\n'.join([ln for ln in msg.split('\n') if ln.strip() != ''] ) if False else msg
    return jsonify({'ok': True, 'message': msg})


@app.route('/api/candidates/<int:cid>/request-update', methods=['POST'])
@login_required
def request_candidate_update(cid):
    """Generate a secure self-update link and email it to the candidate."""
    import secrets as _secrets
    conn = get_db()
    c = conn.execute('SELECT * FROM candidates WHERE id=? AND owner_id=?',
                     (cid, effective_company_id())).fetchone()
    if not c:
        conn.close(); return jsonify({'error': 'Candidate not found'}), 404
    if not (c['email'] or '').strip():
        conn.close(); return jsonify({'error': 'Candidate ka email nahi hai. Pehle email add karein.'}), 400

    token = _secrets.token_urlsafe(24)
    conn.execute('UPDATE candidates SET update_token=?, update_requested_at=?, update_submitted_at=? WHERE id=?',
                 (token, ts(), '', cid))
    # Mandate + recruiter context for the email
    mandate = conn.execute('SELECT role, client FROM mandates WHERE id=?', (c['mandate_id'],)).fetchone()
    conn.commit(); conn.close()

    role = mandate['role'] if mandate else 'a role'
    u = current_user()
    recruiter_name = (u.get('display_name') or u.get('username') or 'Recruiter') if u else 'Recruiter'
    company = get_setting('company_name', '') or 'our team'

    base = request.host_url.rstrip('/')
    link = f'{base}/update-profile?token={token}'

    subject = f'Please share your updated profile — {role}'
    plain = (f"Dear {c['name'] or 'Candidate'},\n\n"
             f"Thank you for your interest in the {role} position"
             + (f" at {mandate['client']}" if mandate and mandate['client'] else '') + ".\n\n"
             f"To move ahead, please review and update your details and upload your latest resume "
             f"using the secure link below:\n\n{link}\n\n"
             f"This link is personal to you. It will take just 2 minutes.\n\n"
             f"Regards,\n{recruiter_name}\n{company}")
    html = (f"Dear {esc_html(c['name'] or 'Candidate')},<br><br>"
            f"Thank you for your interest in the <b>{esc_html(role)}</b> position"
            + (f" at <b>{esc_html(mandate['client'])}</b>" if mandate and mandate['client'] else '') + ".<br><br>"
            f"To move ahead, please review and update your details and upload your latest resume "
            f"using the secure link below:<br><br>"
            f'<a href="{link}" style="display:inline-block;background:#1D9E75;color:#fff;padding:11px 22px;border-radius:8px;text-decoration:none;font-weight:600">Update My Profile</a><br><br>'
            f'<span style="font-size:12px;color:#666">Or copy this link: {link}</span><br><br>'
            f"This link is personal to you. It will take just 2 minutes.<br><br>"
            f"Regards,<br><b>{esc_html(recruiter_name)}</b><br>{esc_html(company)}")

    ok, err = _smtp_send(c['email'], subject, plain, html)
    if not ok:
        return jsonify({'error': err}), 400
    log_candidate_event(cid, 'note', f'Requested updated resume — link emailed to {c["email"]}')
    return jsonify({'ok': True, 'message': 'Update request email sent!'})


@app.route('/update-profile')
def update_profile_page():
    return send_file('update-profile.html')


@app.route('/api/public/candidate-update/<token>', methods=['GET'])
def public_get_candidate(token):
    """Return the candidate's editable fields for the self-update page."""
    if not token or len(token) < 10:
        return jsonify({'error': 'Invalid link'}), 400
    conn = get_db()
    c = conn.execute('SELECT * FROM candidates WHERE update_token=?', (token,)).fetchone()
    conn.close()
    if not c:
        return jsonify({'error': 'This link is invalid or has expired.'}), 404
    # Expiry: 14 days from request
    try:
        req_at = datetime.datetime.fromisoformat(c['update_requested_at'])
        if (datetime.datetime.now() - req_at).days > 14:
            return jsonify({'error': 'This link has expired. Please ask your recruiter for a new one.'}), 410
    except Exception:
        pass
    try:
        skills = json.loads(c['key_skills'] or '[]')
    except Exception:
        skills = []
    return jsonify({'ok': True, 'candidate': {
        'name': c['name'] or '', 'phone': c['phone'] or '', 'email': c['email'] or '',
        'company': c['company'] or '', 'designation': c['designation'] or '',
        'experience': c['experience'] or '', 'ctc_current': c['ctc_current'] or '',
        'ctc_expected': c['ctc_expected'] or '', 'notice_period': c['notice_period'] or '',
        'location': c['location'] or '', 'preferred_location': c['preferred_location'] or '',
        'qualification': c['qualification'] or '', 'key_skills': skills,
        'already_submitted': bool(c['update_submitted_at']),
    }})


@app.route('/api/public/candidate-update/<token>', methods=['POST'])
def public_save_candidate(token):
    """Candidate submits their updated details (+ optional resume) via the link."""
    if not token or len(token) < 10:
        return jsonify({'error': 'Invalid link'}), 400
    conn = get_db()
    c = conn.execute('SELECT * FROM candidates WHERE update_token=?', (token,)).fetchone()
    if not c:
        conn.close(); return jsonify({'error': 'This link is invalid or has expired.'}), 404
    cid = c['id']

    d = request.form if request.form else (request.json or {})
    fields = ['name', 'phone', 'email', 'company', 'designation', 'location',
              'preferred_location', 'qualification']
    num_fields = ['experience', 'ctc_current', 'ctc_expected', 'notice_period']
    sets, vals = [], []
    for f in fields:
        if f in d:
            sets.append(f'{f}=?'); vals.append(str(d.get(f) or '').strip())
    for f in num_fields:
        if f in d:
            try:
                sets.append(f'{f}=?'); vals.append(float(d.get(f) or 0))
            except Exception:
                pass
    if 'key_skills' in d:
        ks = d.get('key_skills')
        if isinstance(ks, str):
            try: ks = json.loads(ks)
            except Exception: ks = [s.strip() for s in ks.split(',') if s.strip()]
        sets.append('key_skills=?'); vals.append(json.dumps(ks or []))

    # Optional resume upload
    resume_saved = False
    if 'resume' in request.files:
        f = request.files['resume']
        if f and f.filename:
            ext = Path(f.filename).suffix.lower()
            if ext in ['.pdf', '.doc', '.docx']:
                fname = 'c' + str(cid) + '_' + datetime.datetime.now().strftime('%Y%m%d%H%M%S') + ext
                f.save(os.path.join(CV_DIR, fname))
                sets.append('cv_path=?'); vals.append(fname)
                sets.append('cv_original_name=?'); vals.append(f.filename)
                resume_saved = True

    if sets:
        vals += [ts(), cid]
        conn.execute('UPDATE candidates SET ' + ','.join(sets) + ',updated_at=? WHERE id=?', vals)
    conn.execute('UPDATE candidates SET update_submitted_at=? WHERE id=?', (ts(), cid))
    conn.commit(); conn.close()

    log_candidate_event(cid, 'update', 'Candidate submitted updated profile via self-update link'
                        + (' (with new resume)' if resume_saved else ''))
    return jsonify({'ok': True, 'message': 'Thank you! Your details have been updated.'})


@app.route('/api/candidates/<int:cid>/note', methods=['POST'])
@login_required
def add_candidate_note(cid):
    """Add a free-text comment to the candidate's journey."""
    text = (request.json or {}).get('text', '').strip()
    if not text:
        return jsonify({'error': 'Empty comment'}), 400
    u = current_user()
    who = (u.get('display_name') or u.get('username') or '') if u else ''
    detail = text + (' — ' + who if who else '')
    log_candidate_event(cid, 'note', detail)
    return jsonify({'ok': True})


@app.route('/api/candidates/<int:cid>/journey')
@login_required
def candidate_journey(cid):
    """Aggregate a candidate's full journey from every real event source,
    newest first. Each event: {ts, type, text, icon, color}."""
    conn = get_db()
    c = conn.execute('SELECT * FROM candidates WHERE id=? AND owner_id=?',
                     (cid, effective_company_id())).fetchone()
    if not c:
        conn.close(); return jsonify({'error': 'Not found'}), 404
    ev = []
    def add(t, text, icon, color):
        if t:
            ev.append({'ts': t, 'text': text, 'icon': icon, 'color': color})

    # Sourced / created
    add(c['created_at'], 'Candidate added to pipeline', 'user-plus', 'gray')
    # Stage changes
    for h in conn.execute('SELECT * FROM stage_history WHERE candidate_id=? ORDER BY created_at', (cid,)).fetchall():
        frm = h['from_stage'] or '—'
        add(h['created_at'], f"Stage changed — {frm} to {h['to_stage']}", 'arrow-right', 'purple')
    # WhatsApp sends
    add(c['msg1_sent_at'], 'WhatsApp intro sent', 'brand-whatsapp', 'green')
    add(c['fu1_sent_at'], 'WhatsApp follow up 1 sent', 'brand-whatsapp', 'green')
    add(c['fu2_sent_at'], 'WhatsApp follow up 2 sent', 'brand-whatsapp', 'green')
    # WhatsApp / call response
    if c['wa_response']:
        rmap = {'interested': 'Interested', 'callback': 'Callback', 'not_interested': 'Not interested', 'no_reply': 'No reply'}
        add(c['wa_response_at'] or c['updated_at'], 'Response logged — ' + rmap.get(c['wa_response'], c['wa_response']), 'message-dots', 'teal')
    # Reminders
    for r in conn.execute('SELECT * FROM reminders WHERE candidate_id=? ORDER BY created_at', (cid,)).fetchall():
        note = (r['note'] or 'Reminder')
        add(r['created_at'], 'Reminder set — ' + note, 'bell', 'amber')
        if r['done']:
            add(r['due_at'], 'Reminder completed — ' + note, 'check', 'teal')
    # Logged events (tags added, call analysed, etc.)
    for e in conn.execute('SELECT * FROM candidate_events WHERE candidate_id=? ORDER BY created_at', (cid,)).fetchall():
        icon = {'tag': 'tag', 'call': 'phone', 'note': 'note', 'email': 'mail', 'edit': 'edit'}.get(e['event_type'], 'point')
        color = {'tag': 'gray', 'call': 'teal', 'note': 'blue', 'email': 'purple', 'edit': 'amber'}.get(e['event_type'], 'gray')
        add(e['created_at'], e['detail'], icon, color)
    conn.close()

    ev = [x for x in ev if x['ts']]
    ev.sort(key=lambda x: x['ts'], reverse=True)
    return jsonify({'ok': True, 'events': ev})


@app.route('/api/candidates/<int:cid>')
@login_required
def get_candidate(cid):
    conn = get_db()
    r = conn.execute('SELECT * FROM candidates WHERE id=? AND owner_id=?',
                     (cid, effective_company_id())).fetchone()
    if not r: conn.close(); return jsonify({'error': 'Not found'}), 404
    d = _cand_public(r)   # drop embedding / embedding_text / embedding_vec
    try: d['key_skills'] = json.loads(d['key_skills'] or '[]')
    except: d['key_skills'] = []
    try: d['secondary_skills'] = json.loads(d['secondary_skills'] or '[]')
    except: d['secondary_skills'] = []
    hist = conn.execute('SELECT * FROM stage_history WHERE candidate_id=? ORDER BY created_at', (cid,)).fetchall()
    d['history'] = [dict(h) for h in hist]
    wh = conn.execute('SELECT * FROM work_history WHERE candidate_id=? ORDER BY is_current DESC, sort_order ASC, id ASC', (cid,)).fetchall()
    d['work_history'] = [dict(w) for w in wh]
    conn.close()
    return jsonify(d)

@app.route('/api/candidates/<int:cid>/move', methods=['POST'])
@login_required
def move_candidate(cid):
    """Move a candidate to a different mandate (within the same tenant)."""
    d = request.json or {}
    target_mid = d.get('mandate_id')
    if not target_mid:
        return jsonify({'error': 'Target mandate required'}), 400
    conn = get_db()
    cand = conn.execute('SELECT mandate_id, name FROM candidates WHERE id=?', (cid,)).fetchone()
    if not cand:
        conn.close(); return jsonify({'error': 'Candidate not found'}), 404
    # Verify target mandate belongs to this tenant
    tgt = conn.execute('SELECT id, role, client, owner_id FROM mandates WHERE id=?', (target_mid,)).fetchone()
    if not tgt or tgt['owner_id'] != effective_company_id():
        conn.close(); return jsonify({'error': 'Target mandate not found'}), 404
    old = conn.execute('SELECT role FROM mandates WHERE id=?', (cand['mandate_id'],)).fetchone()
    old_label = old['role'] if old else 'previous mandate'
    conn.execute('UPDATE candidates SET mandate_id=?, updated_at=? WHERE id=?', (target_mid, ts(), cid))
    conn.execute('INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) VALUES (?,?,?,?,?)',
                 (cid, '', '', f'Moved from "{old_label}" to "{tgt["role"]}"', ts()))
    conn.commit(); conn.close()
    log_candidate_event(cid, 'note', f'Moved to mandate: {tgt["role"]} ({tgt["client"]})')
    return jsonify({'ok': True, 'mandate_id': target_mid})


@app.route('/api/candidates/<int:cid>', methods=['PUT'])
@login_required
def update_candidate(cid):
    d = request.json or {}
    conn = get_db()
    c = conn.execute('SELECT * FROM candidates WHERE id=? AND owner_id=?',
                     (cid, effective_company_id())).fetchone()
    if not c: conn.close(); return jsonify({'error': 'Not found'}), 404

    fields = ['name','company','designation','experience','ctc_current','ctc_expected',
              'notice_period','location','preferred_location','phone','email','qualification','specialization','career_summary',
              'key_skills','secondary_skills','recruiter_feedback','client_feedback','general_comments',
              'linkedin_url','ai_insight_cv']
    sets = []; vals = []
    for f in fields:
        if f in d:
            sets.append(f + '=?')
            val = d[f]
            if isinstance(val, (list, dict)): val = json.dumps(val)
            vals.append(val)

    if sets:
        vals += [ts(), cid]
        conn.execute('UPDATE candidates SET ' + ','.join(sets) + ',updated_at=? WHERE id=?', vals)

        # Build a human-readable list of what changed, for the journey
        labels = {
            'name':'Name','company':'Company','designation':'Designation',
            'experience':'Experience','ctc_current':'Current CTC','ctc_expected':'Expected CTC',
            'notice_period':'Notice period','location':'Location','preferred_location':'Preferred location','phone':'Phone','email':'Email',
            'qualification':'Qualification','specialization':'Specialization','career_summary':'Summary',
            'linkedin_url':'LinkedIn URL','ai_insight_cv':'AI Insight (CV)'
        }
        changes = []
        for f, lbl in labels.items():
            if f in d:
                old_v = c[f] if f in c.keys() else ''
                new_v = d[f]
                if str(old_v or '') != str(new_v or ''):
                    if new_v not in (None, '', 0):
                        changes.append(f"{lbl}: {old_v or '—'} \u2192 {new_v}")
        # Skills change
        if 'key_skills' in d:
            try:
                new_skills = d['key_skills'] if isinstance(d['key_skills'], list) else json.loads(d['key_skills'] or '[]')
                old_skills = json.loads(c['key_skills'] or '[]')
                if set(new_skills) != set(old_skills):
                    changes.append('Skills updated')
            except Exception:
                pass

        notes = []
        if 'recruiter_feedback' in d and d['recruiter_feedback'] and d['recruiter_feedback'] != (c['recruiter_feedback'] or ''):
            notes.append('Recruiter feedback updated')
        if 'client_feedback' in d and d['client_feedback'] and d['client_feedback'] != (c['client_feedback'] or ''):
            notes.append('Client feedback updated')

        conn.commit(); conn.close()

        # Log each change to the journey as an edit event
        if changes:
            u = current_user()
            who = (u.get('display_name') or u.get('username') or '') if u else ''
            detail = 'Profile updated \u2014 ' + '; '.join(changes) + (f' (by {who})' if who else '')
            log_candidate_event(cid, 'edit', detail)
        for note in notes:
            log_candidate_event(cid, 'note', note)
        return jsonify({'ok': True})
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/candidates/<int:cid>/stage', methods=['POST'])
@login_required
def move_stage(cid):
    # Freelancers cannot change stages
    try:
        from modules.freelancer import block_if_freelancer
        blocked = block_if_freelancer()
        if blocked:
            return blocked
    except Exception:
        pass
    d = request.json or {}
    conn = get_db()
    r = conn.execute('SELECT stage FROM candidates WHERE id=?', (cid,)).fetchone()
    if not r: conn.close(); return jsonify({'error': 'Not found'}), 404
    old_stage = r['stage']
    # keep_stage=true means just add a note to history without changing stage
    keep_stage = d.get('keep_stage', False)
    new_stage = old_stage if keep_stage else d.get('stage', old_stage)
    if not keep_stage:
        conn.execute('UPDATE candidates SET stage=?,updated_at=? WHERE id=?', (new_stage, ts(), cid))
    conn.execute('INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) VALUES (?,?,?,?,?)',
                 (cid, old_stage, new_stage, d.get('note',''), ts()))
    conn.commit(); conn.close()
    return jsonify({'ok': True})

@app.route('/api/mandates/<int:mid>/candidates/manual', methods=['POST'])
@login_required
def add_manual(mid):
    d = request.json or {}
    if not d.get('name') or not d.get('company'):
        return jsonify({'error': 'Name and Company are required'}), 400
    conn = get_db(); c = conn.cursor()
    if not _tenant_owns_mandate(conn, mid):
        conn.close(); return jsonify({'error': 'Mandate not found'}), 404
    c.execute(
        'INSERT INTO candidates (mandate_id,name,company,designation,experience,ctc_current,'
        'ctc_expected,notice_period,location,phone,email,career_summary,key_skills,'
        'screening_decision,ai_reasoning,stage,created_at,updated_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
        (mid, d['name'], d['company'], d.get('designation',''), float(d.get('experience') or 0),
         float(d.get('ctc_current') or 0), float(d.get('ctc_expected') or 0), int(d.get('notice_period') or 0),
         d.get('location',''), d.get('phone',''), d.get('email',''), d.get('career_summary',''),
         json.dumps(d.get('key_skills') or []), 'worth_opening', 'Manually added', 'Screening', ts(), ts()))
    cid = c.lastrowid
    c.execute('UPDATE candidates SET qualification=?, preferred_location=? WHERE id=?',
              (d.get('qualification',''), d.get('preferred_location',''), cid))
    c.execute('INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) VALUES (?,?,?,?,?)',
              (cid, '', 'Screening', 'Manually added to pipeline', ts()))
    conn.commit(); conn.close()
    return jsonify({'ok': True, 'id': cid})

# CV
@app.route('/api/candidates/<int:cid>/cv', methods=['POST', 'OPTIONS'])
def upload_cv(cid):
    if request.method == 'OPTIONS':
        return ('', 204)
    if not session.get('user_id'):
        return jsonify({'error': 'auth_required', 'message': 'Please log into HireLab first.'}), 401
    if 'cv' not in request.files: return jsonify({'error': 'No file uploaded'}), 400
    f = request.files['cv']
    ext = Path(f.filename).suffix.lower()
    if ext not in ['.pdf', '.doc', '.docx']: return jsonify({'error': 'PDF or Word files only'}), 400
    fname = 'c' + str(cid) + '_' + datetime.datetime.now().strftime('%Y%m%d%H%M%S') + ext
    f.save(os.path.join(CV_DIR, fname))
    conn = get_db()
    old = conn.execute('SELECT cv_path FROM candidates WHERE id=?', (cid,)).fetchone()
    if old and old['cv_path']:
        op = os.path.join(CV_DIR, old['cv_path'])
        if os.path.exists(op): os.remove(op)
    conn.execute('UPDATE candidates SET cv_path=?,cv_original_name=?,updated_at=? WHERE id=?', (fname, f.filename, ts(), cid))
    conn.execute('INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) VALUES (?,?,?,?,?)',
                 (cid, '', '', 'CV uploaded: ' + f.filename, ts()))
    conn.commit(); conn.close()
    return jsonify({'ok': True, 'filename': fname, 'original': f.filename})

def _resolve_tenant_cv(conn, filename):
    """Return a safe absolute path to a CV file IF (a) it resolves INSIDE
    CV_DIR (blocks ../ path-traversal) and (b) it belongs to a candidate
    owned by the current tenant. Otherwise None. owner_id stores the company id."""
    safe = os.path.basename(filename or '')          # strip any directory parts
    if not safe:
        return None
    fp = os.path.abspath(os.path.join(CV_DIR, safe))
    if not fp.startswith(os.path.abspath(CV_DIR) + os.sep):
        return None                                   # escaped CV_DIR
    owns = conn.execute('SELECT 1 FROM candidates WHERE cv_path=? AND owner_id=? LIMIT 1',
                        (safe, effective_company_id())).fetchone()
    return fp if owns else None


@app.route('/api/cv/<path:filename>')
@login_required
def serve_cv(filename):
    conn = get_db()
    fp = _resolve_tenant_cv(conn, filename)
    conn.close()
    if not fp or not os.path.exists(fp):
        return jsonify({'error': 'Not found'}), 404
    return send_file(fp)


@app.route('/api/cv-view/<path:filename>')
@login_required
def view_cv_html(filename):
    """Render a .docx CV as HTML so it can be shown inline in the browser
    (browsers can show PDF in an iframe natively, but not Word files)."""
    conn = get_db()
    fp = _resolve_tenant_cv(conn, filename)
    conn.close()
    if not fp or not os.path.exists(fp):
        return ('<p style="font-family:sans-serif;padding:20px;color:#888">CV file not found.</p>', 404)
    ext = os.path.splitext(filename)[1].lower()
    if ext != '.docx':
        return ('<p style="font-family:sans-serif;padding:20px;color:#888">Preview only supports .docx. Please download to view.</p>', 200)
    try:
        import mammoth
        with open(fp, 'rb') as f:
            result = mammoth.convert_to_html(f)
        body = result.value or '<p style="color:#888">(Empty document)</p>'
        page = (
            '<!doctype html><html><head><meta charset="utf-8">'
            '<style>'
            'body{font-family:-apple-system,Segoe UI,Roboto,sans-serif;font-size:13px;'
            'line-height:1.6;color:#222;max-width:800px;margin:0 auto;padding:28px 32px;background:#fff}'
            'h1,h2,h3{color:#0E2A47;margin:16px 0 8px} p{margin:6px 0} '
            'table{border-collapse:collapse;width:100%;margin:10px 0} '
            'td,th{border:1px solid #ddd;padding:6px 8px;font-size:12px} '
            'ul,ol{margin:6px 0 6px 22px} img{max-width:100%}'
            '</style></head><body>' + body + '</body></html>'
        )
        return (page, 200, {'Content-Type': 'text/html; charset=utf-8'})
    except Exception as e:
        return ('<p style="font-family:sans-serif;padding:20px;color:#C0522B">Could not render this Word file: '
                + str(e) + '. Please download to view.</p>', 200)

@app.route('/api/candidates/<int:cid>/cv', methods=['DELETE'])
@login_required
def delete_cv(cid):
    conn = get_db()
    r = conn.execute('SELECT cv_path FROM candidates WHERE id=? AND owner_id=?',
                     (cid, effective_company_id())).fetchone()
    if not r:
        conn.close(); return jsonify({'error': 'Not found'}), 404
    if r and r['cv_path']:
        fp = os.path.join(CV_DIR, r['cv_path'])
        if os.path.exists(fp): os.remove(fp)
        conn.execute('UPDATE candidates SET cv_path="",cv_original_name="",updated_at=? WHERE id=?', (ts(), cid))
        conn.execute('INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) VALUES (?,?,?,?,?)',
                     (cid, '', '', 'CV removed', ts()))
        conn.commit()
    conn.close()
    return jsonify({'ok': True})


# Delete candidate
@app.route('/api/candidates/<int:cid>/work-history', methods=['POST'])
@login_required
def save_work_history(cid):
    """Replace the full work-history list for a candidate."""
    d = request.json or {}
    items = d.get('items', [])
    conn = get_db()
    # ownership check
    own = conn.execute('SELECT owner_id FROM candidates WHERE id=?', (cid,)).fetchone()
    if not own or own['owner_id'] != effective_user_id():
        conn.close(); return jsonify({'error': 'Not found'}), 404
    conn.execute('DELETE FROM work_history WHERE candidate_id=?', (cid,))
    for i, it in enumerate(items):
        conn.execute(
            'INSERT INTO work_history (candidate_id,company,designation,start_date,end_date,is_current,description,sort_order) '
            'VALUES (?,?,?,?,?,?,?,?)',
            (cid, (it.get('company') or '').strip(), (it.get('designation') or '').strip(),
             (it.get('start_date') or '').strip(), (it.get('end_date') or '').strip(),
             1 if it.get('is_current') else 0, (it.get('description') or '').strip(), i)
        )
    conn.commit(); conn.close()
    _xp_recompute_safe(cid)
    return jsonify({'ok': True, 'count': len(items)})


def _xp_recompute_safe(cid):
    """Best-effort recompute of experience intelligence; never breaks the caller
    if the xp module isn't loaded or the engine errors on one candidate."""
    try:
        from modules.xp_engine import feedback_loop as _fb
        _conn = get_db()
        _fb.recompute_candidate(_conn, cid)
        _conn.close()
    except Exception as _e:
        print(f'[xp] recompute skipped for {cid}: {_e}')


@app.route('/api/candidates/<int:cid>', methods=['DELETE'])
@login_required
def delete_candidate(cid):
    conn = get_db()
    r = conn.execute('SELECT cv_path FROM candidates WHERE id=? AND owner_id=?',
                     (cid, effective_company_id())).fetchone()
    if not r:
        conn.close(); return jsonify({'error': 'Not found'}), 404
    if r:
        if r['cv_path']:
            fp = os.path.join(CV_DIR, r['cv_path'])
            if os.path.exists(fp):
                try: os.remove(fp)
                except: pass
        conn.execute('DELETE FROM stage_history WHERE candidate_id=?', (cid,))
        conn.execute('DELETE FROM candidates WHERE id=?', (cid,))
        conn.commit()
    conn.close()
    return jsonify({'ok': True})

# DeepSeek parse
@app.route('/api/parse-naukri', methods=['POST'])
@login_required
def parse_naukri():
    d = request.json or {}
    key = get_setting('deepseek_api_key') or d.get('deepseek_api_key', '')
    raw = d.get('raw', '').strip()
    if not key: return jsonify({'error': 'DeepSeek API key not set. Go to Settings.'}), 400
    if not raw: return jsonify({'error': 'No text provided'}), 400
    system_msg = ('Extract candidate details from recruiter text. Return ONLY valid JSON with these fields: '
                  'name, phone, email, company, designation, experience (float years), '
                  'ctc_current (float LPA), ctc_expected (float LPA), notice_period (int days), '
                  'location (current city), preferred_location (preferred/desired job location, if mentioned), '
                  'qualification (highest education degree e.g. B.Tech, MBA), key_skills (array max 6), secondary_skills (array), '
                  'career_summary (2 sentences), is_mnc (bool). '
                  'Use null for missing strings, 0 for missing numbers.')
    try:
        resp = call_deepseek(key,
            {'model': 'deepseek-chat', 'temperature': 0, 'max_tokens': 800,
                  'messages': [{'role': 'system', 'content': system_msg}, {'role': 'user', 'content': raw}],
                  'response_format': {'type': 'json_object'}},
            timeout=30, endpoint='parse')
    except requests.Timeout: return jsonify({'error': 'DeepSeek timeout. Try again.'}), 504
    except Exception as e: return jsonify({'error': str(e)}), 500
    if resp.status_code == 401: return jsonify({'error': 'Invalid DeepSeek API key'}), 401
    if resp.status_code != 200: return jsonify({'error': 'DeepSeek error: ' + resp.text[:150]}), 500
    text = resp.json()['choices'][0]['message']['content']
    parsed = parse_json(text)
    return jsonify({'ok': True, 'data': parsed}) if parsed else (jsonify({'error': 'Parse failed'}), 500)


# ── Resume Text Extraction ───────────────────────────────────────────────────
def extract_text_from_file(file_bytes, filename):
    """Extract plain text from PDF or Word file."""
    ext = Path(filename).suffix.lower()
    text = ''
    try:
        if ext == '.pdf':
            if HAS_PDF:
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
            else:
                return None, 'pdfplumber not installed. Run: pip install pdfplumber'
        elif ext in ['.docx']:
            if HAS_DOCX:
                doc = DocxDocument(io.BytesIO(file_bytes))
                text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
            else:
                return None, 'python-docx not installed. Run: pip install python-docx'
        elif ext == '.doc':
            return None, '.doc format not supported. Please convert to .docx or .pdf'
        else:
            return None, 'Unsupported file type'
        return text.strip() if text.strip() else None, None
    except Exception as e:
        return None, str(e)

# ── Parse Resume (PDF/Word → DeepSeek → candidate fields) ────────────────────
@app.route('/api/parse-resume', methods=['POST'])
@login_required
def parse_resume():
    if 'resume' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    f = request.files['resume']
    ds_key = get_setting('deepseek_api_key') or request.form.get('deepseek_api_key', '')
    if not ds_key:
        return jsonify({'error': 'DeepSeek API key required. Add in Settings.'}), 400

    file_bytes = f.read()
    text, err = extract_text_from_file(file_bytes, f.filename)
    if err:
        return jsonify({'error': err}), 400
    if not text or len(text) < 50:
        return jsonify({'error': 'Could not extract text from file. Try PDF or DOCX format.'}), 400

    system_msg = ('Extract candidate details from this resume/CV text. Return ONLY valid JSON with: '
                  'name, phone, email, company (current), designation (current title), '
                  'experience (float years total), ctc_current (float LPA, 0 if not found), '
                  'ctc_expected (float LPA, 0 if not found), notice_period (int days, 0 if not found), '
                  'location (current city), preferred_location (preferred/desired job location if mentioned), qualification (highest degree), '
                  'key_skills (array of top 8 technical/domain skills), '
                  'secondary_skills (array of other skills), '
                  'career_summary (2-3 sentences about background and strengths), '
                  'industry_background (e.g. FMCG, Manufacturing, IT), is_mnc (bool). '
                  'Use null for missing strings, 0 for missing numbers.')
    try:
        resp = call_deepseek(ds_key,
            {'model': 'deepseek-chat', 'temperature': 0, 'max_tokens': 1000,
                  'messages': [{'role': 'system', 'content': system_msg},
                                {'role': 'user', 'content': 'Extract from this resume:\n\n' + text[:8000]}],
                  'response_format': {'type': 'json_object'}},
            timeout=45, endpoint='resume-parse')
    except requests.Timeout:
        return jsonify({'error': 'DeepSeek timeout — try again'}), 504
    except Exception as e:
        return jsonify({'error': str(e)}), 500

    if resp.status_code == 401: return jsonify({'error': 'Invalid DeepSeek API key'}), 401
    if resp.status_code != 200: return jsonify({'error': 'DeepSeek error: ' + resp.text[:150]}), 500

    raw = resp.json()['choices'][0]['message']['content']
    parsed = parse_json(raw)
    return jsonify({'ok': True, 'data': parsed, 'text_length': len(text)}) if parsed else (jsonify({'error': 'Parse failed', 'raw': raw[:300]}), 500)

# ── Bulk Parse Multiple Naukri Snippets ───────────────────────────────────────
@app.route('/api/parse-naukri-bulk', methods=['POST'])
@login_required
def parse_naukri_bulk():
    d = request.json or {}
    ds_key = get_setting('deepseek_api_key') or d.get('deepseek_api_key', '')
    raw = d.get('raw', '').strip()
    if not ds_key: return jsonify({'error': 'DeepSeek API key required'}), 400
    if not raw:    return jsonify({'error': 'No content provided'}), 400

    system_msg = (
        'You are parsing multiple candidate profiles from Naukri or recruiter notes. '
        'Extract each candidate and return a JSON ARRAY (not object). '
        'Each element must have: name, phone, email, company, designation, '
        'experience (float years), ctc_current (float LPA), ctc_expected (float LPA), '
        'notice_period (int days), location, qualification, '
        'key_skills (array max 6), career_summary (1-2 sentences), is_mnc (bool). '
        'Use null for missing strings, 0 for missing numbers. '
        'IMPORTANT: Return an ARRAY even if there is only one candidate. '
        'Separate candidates by looking for new profile headers, numbers, or clear breaks.'
    )
    try:
        resp = call_deepseek(ds_key,
            {'model': 'deepseek-chat', 'temperature': 0, 'max_tokens': 3000,
                  'messages': [{'role': 'system', 'content': system_msg},
                                {'role': 'user', 'content': 'Extract all candidates from:\n\n' + raw}]},
            timeout=60, endpoint='bulk-parse')
    except requests.Timeout:
        return jsonify({'error': 'DeepSeek timeout — try again'}), 504
    except Exception as e:
        return jsonify({'error': str(e)}), 500

    if resp.status_code == 401: return jsonify({'error': 'Invalid DeepSeek API key'}), 401
    if resp.status_code != 200: return jsonify({'error': 'DeepSeek error: ' + resp.text[:150]}), 500

    raw_resp = resp.json()['choices'][0]['message']['content']
    parsed = parse_json(raw_resp)
    if isinstance(parsed, dict): parsed = [parsed]   # single candidate returned as object
    if not isinstance(parsed, list): return jsonify({'error': 'Could not parse response', 'raw': raw_resp[:300]}), 500
    return jsonify({'ok': True, 'candidates': parsed, 'count': len(parsed)})

# ── Bulk Add Candidates ────────────────────────────────────────────────────────
@app.route('/api/mandates/<int:mid>/candidates/bulk', methods=['POST'])
@login_required
def bulk_add_candidates(mid):
    d = request.json or {}
    candidates = d.get('candidates', [])
    if not candidates: return jsonify({'error': 'No candidates provided'}), 400

    conn = get_db(); c = conn.cursor()
    m = conn.execute('SELECT * FROM mandates WHERE id=? AND owner_id=?',
                     (mid, effective_company_id())).fetchone()
    if not m: conn.close(); return jsonify({'error': 'Mandate not found'}), 404

    added = 0
    ids = []
    for cand in candidates:
        name    = str(cand.get('name') or '').strip()
        company = str(cand.get('company') or '').strip()
        if not name: continue
        c.execute(
            'INSERT INTO candidates (mandate_id,name,company,designation,experience,ctc_current,'
            'ctc_expected,notice_period,location,phone,email,career_summary,key_skills,'
            'screening_decision,ai_reasoning,stage,created_at,updated_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
            (mid, name, company, cand.get('designation',''), float(cand.get('experience') or 0),
             float(cand.get('ctc_current') or 0), float(cand.get('ctc_expected') or 0),
             int(cand.get('notice_period') or 0), cand.get('location',''),
             cand.get('phone',''), cand.get('email',''), cand.get('career_summary',''),
             json.dumps(cand.get('key_skills') or []),
             'worth_opening', 'Manually added (bulk)', 'Screening', ts(), ts()))
        cid = c.lastrowid
        c.execute('INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) VALUES (?,?,?,?,?)',
                  (cid, '', 'Screening', 'Bulk added to pipeline', ts()))
        ids.append(cid)
        added += 1

    conn.commit(); conn.close()
    return jsonify({'ok': True, 'added': added, 'ids': ids})


# ── CALL RECORDING ANALYSIS ──────────────────────────────────────────────────

CALL_DIR = os.path.join(DATA_DIR, 'calls')
os.makedirs(CALL_DIR, exist_ok=True)

@app.route('/api/candidates/<int:cid>/analyse-call', methods=['POST'])
@login_required
def analyse_call(cid):
    # TENANT GUARD: the candidate must belong to the caller's company (tenant).
    _g = get_db()
    _own = _g.execute('SELECT owner_id FROM candidates WHERE id=?', (cid,)).fetchone()
    _g.close()
    if not _own or _own['owner_id'] != effective_company_id():
        return jsonify({'error': 'Candidate not in your workspace'}), 403

    language    = request.form.get('language', 'hi')   # hi = Hindi, en = English
    # Server keys (env var first, then DB) take priority over anything from frontend
    groq_key    = get_setting('groq_api_key') or request.form.get('groq_api_key', '').strip()
    claude_key  = get_setting('claude_api_key') or request.form.get('claude_api_key', '').strip()

    if not groq_key: return jsonify({'error': 'Groq API key required (for transcription). Add in Settings.'}), 400
    if not claude_key: return jsonify({'error': 'Claude API key required (for analysis). Add in Settings.'}), 400
    if 'recording' not in request.files: return jsonify({'error': 'No recording file uploaded'}), 400

    f = request.files['recording']
    ext = Path(f.filename).suffix.lower()
    allowed = ['.mp3', '.m4a', '.mp4', '.wav', '.ogg', '.webm', '.flac']
    if ext not in allowed:
        return jsonify({'error': f'Unsupported format. Use: {", ".join(allowed)}'}), 400

    # Save recording
    fname = f'call_{cid}_{datetime.datetime.now().strftime("%Y%m%d%H%M%S")}{ext}'
    fpath = os.path.join(CALL_DIR, fname)
    file_bytes = f.read()
    with open(fpath, 'wb') as out:
        out.write(file_bytes)

    # ── Step 1: Transcribe with Whisper ──────────────────────────────────────
    try:
        mime = {
            '.mp3': 'audio/mpeg', '.m4a': 'audio/mp4', '.mp4': 'audio/mp4',
            '.wav': 'audio/wav',  '.ogg': 'audio/ogg', '.webm': 'audio/webm',
            '.flac': 'audio/flac'
        }.get(ext, 'audio/mpeg')

        whisper_resp = requests.post(
            'https://api.groq.com/openai/v1/audio/transcriptions',
            headers={'Authorization': 'Bearer ' + groq_key},
            files={'file': (f.filename, file_bytes, mime)},
            data={'model': 'whisper-large-v3', 'language': language,
                  'response_format': 'verbose_json',
                  'prompt': 'This is a recruiter call with a candidate discussing a job opportunity. '
                            'The conversation may be in Hindi, English, or Hinglish.'},
            timeout=120
        )
    except requests.Timeout:
        return jsonify({'error': 'Whisper transcription timed out. Try a shorter recording.'}), 504
    except Exception as e:
        return jsonify({'error': 'Transcription error: ' + str(e)}), 500

    if whisper_resp.status_code == 401:
        return jsonify({'error': 'Invalid Groq API key'}), 401
    if whisper_resp.status_code != 200:
        try:
            err = whisper_resp.json().get('error', {}).get('message', whisper_resp.text[:200])
        except Exception:
            err = whisper_resp.text[:200]
        return jsonify({'error': 'Groq transcription error: ' + err}), 500

    _wjson = whisper_resp.json()
    transcript = _wjson.get('text', '').strip()
    # Log transcription cost per tenant (Whisper bills by audio duration).
    try:
        log_api_usage('groq', 'whisper-large-v3',
                      audio_seconds=float(_wjson.get('duration', 0) or 0),
                      endpoint='transcription')
    except Exception:
        pass
    if not transcript:
        return jsonify({'error': 'Whisper returned empty transcript. Check recording quality.'}), 400

    # ── Step 2: Get candidate + mandate context ───────────────────────────────
    conn = get_db()
    cand = conn.execute('SELECT * FROM candidates WHERE id=?', (cid,)).fetchone()
    if not cand: conn.close(); return jsonify({'error': 'Candidate not found'}), 404
    mandate = conn.execute('SELECT * FROM mandates WHERE id=?', (cand['mandate_id'],)).fetchone()
    conn.close()

    # Get CV text if available
    cv_text = ''
    if cand['cv_path']:
        cv_path = os.path.join(CV_DIR, cand['cv_path'])
        if os.path.exists(cv_path):
            cv_ext = Path(cv_path).suffix.lower()
            try:
                if cv_ext == '.pdf' and HAS_PDF:
                    import pdfplumber
                    with pdfplumber.open(cv_path) as pdf:
                        cv_text = '\n'.join(p.extract_text() or '' for p in pdf.pages)[:4000]
                elif cv_ext in ['.docx'] and HAS_DOCX:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(cv_path)
                    cv_text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())[:4000]
            except Exception:
                pass

    jd_or_sop = (mandate['sop_text'] or html_to_text(mandate['jd']) or '') if mandate else ''
    cand_name  = cand['name'] or 'Candidate'
    role       = mandate['role'] if mandate else 'the position'
    client     = mandate['client'] if mandate else ''

    # ── Step 3: Claude Analysis ───────────────────────────────────────────────
    system_msg = (
        'You are an expert recruitment analyst. Analyse a recruiter-candidate call.\n'
        'Return ONLY valid JSON — no markdown, no explanation.\n\n'
        'JSON structure:\n'
        '{\n'
        '  "interest_level": "HIGH" | "MEDIUM" | "LOW",\n'
        '  "interest_reason": "one sentence why",\n'
        '  "ctc_discussed": null or float (current CTC in LPA the candidate states they earn now),\n'
        '  "ctc_expected_discussed": null or float (expected/asking CTC in LPA, if mentioned),\n'
        '  "current_company_discussed": null or "the company the candidate currently works at, if stated",\n'
        '  "notice_negotiable": true | false | null,\n'
        '  "notice_discussed_days": null or int,\n'
        '  "key_concerns": ["concern1", "concern2"],\n'
        '  "candidate_strengths": ["strength1", "strength2"],\n'
        '  "red_flags": ["flag1"] or [],\n'
        '  "fit_vs_jd": "STRONG" | "MODERATE" | "WEAK",\n'
        '  "fit_reason": "one sentence",\n'
        '  "next_step": "specific action recruiter should take",\n'
        '  "next_step_deadline": "e.g. by Wednesday" or null,\n'
        '  "recommendation": "PROCEED" | "HOLD" | "REJECT",\n'
        '  "recommendation_reason": "one sentence",\n'
        '  "call_summary": "3-4 sentences covering the full conversation",\n'
        '  "key_quotes": ["notable quote 1", "notable quote 2"],\n'
        '  "languages_detected": "Hindi / English / Hinglish"\n'
        '}'
    )

    user_msg = (
        'CANDIDATE: ' + cand_name + '\n'
        'ROLE: ' + role + ((' at ' + client) if client else '') + '\n\n'
        + ('JD / SOP:\n' + jd_or_sop[:2000] + '\n\n' if jd_or_sop else '')
        + ('CV / RESUME (extracted text):\n' + cv_text[:2000] + '\n\n' if cv_text else '')
        + 'CALL TRANSCRIPT:\n' + transcript[:6000]
    )

    claude_resp = call_claude(claude_key, system_msg, [{'role': 'user', 'content': user_msg}], max_tokens=1500)
    if claude_resp.status_code != 200:
        try: err = claude_resp.json().get('error', {}).get('message', 'Claude error')
        except Exception: err = claude_resp.text[:200]
        return jsonify({'error': 'Analysis failed: ' + err, 'transcript': transcript}), 500

    analysis_text = claude_resp.json()['content'][0]['text']
    analysis = parse_json(analysis_text)
    if not analysis:
        return jsonify({'error': 'Could not parse analysis', 'transcript': transcript, 'raw': analysis_text[:500]}), 500

    # ── Step 4: Save to DB ────────────────────────────────────────────────────
    analysis_str = json.dumps(analysis, ensure_ascii=False)
    conn = get_db()
    # Auto-update candidate fields from the call (roadmap: update CTC & company
    # from the call). Only overwrite when the call actually surfaced a value.
    updates = {}
    try:
        ctc = analysis.get('ctc_discussed')
        if ctc is not None and float(ctc) > 0:
            updates['ctc_current'] = float(ctc)
    except (TypeError, ValueError):
        pass
    try:
        ctc_e = analysis.get('ctc_expected_discussed')
        if ctc_e is not None and float(ctc_e) > 0:
            updates['ctc_expected'] = float(ctc_e)
    except (TypeError, ValueError):
        pass
    comp = (analysis.get('current_company_discussed') or '').strip()
    if comp:
        updates['company'] = comp
    try:
        nd = analysis.get('notice_discussed_days')
        if nd is not None and int(nd) >= 0:
            updates['notice_period'] = int(nd)
    except (TypeError, ValueError):
        pass

    note = '[CALL ANALYSIS ' + datetime.datetime.now().strftime('%d %b %Y %H:%M') + '] Recorded. Interest: ' + analysis.get('interest_level', '') + '. ' + analysis.get('call_summary', '')[:200]
    if updates:
        set_clause = ', '.join(f'{k}=?' for k in updates) + ', general_comments=?, updated_at=? WHERE id=?'
        conn.execute('UPDATE candidates SET ' + set_clause,
                     tuple(updates.values()) + (note, ts(), cid))
    else:
        conn.execute('UPDATE candidates SET general_comments=?,updated_at=? WHERE id=?',
                     (note, ts(), cid))
    upd_summary = ', '.join(f'{k}→{v}' for k, v in updates.items()) if updates else ''
    conn.execute('INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) VALUES (?,?,?,?,?)',
                 (cid, cand['stage'], cand['stage'],
                  'Call analysed. Interest: ' + analysis.get('interest_level','') + '. Rec: ' + analysis.get('recommendation','') + '. ' + analysis.get('next_step','') + (' | Updated: ' + upd_summary if upd_summary else ''),
                  ts()))
    conn.commit(); conn.close()
    _interest = analysis.get('interest_level', '')
    log_candidate_event(cid, 'call', 'Call analysed' + (' — interest: ' + _interest if _interest else '') + (' · updated ' + upd_summary if upd_summary else ''))

    return jsonify({
        'ok': True,
        'transcript': transcript,
        'analysis': analysis,
        'recording_file': fname,
        'updated_fields': updates,
        'cv_used': bool(cv_text),
        'jd_used': bool(jd_or_sop)
    })

@app.route('/api/calls/<path:filename>')
@login_required
def serve_call(filename):
    # Recordings are named call_<candidateId>_<timestamp>.<ext>. Verify the
    # candidate belongs to the caller's tenant before serving the audio.
    m = re.match(r'call_(\d+)_', os.path.basename(filename))
    if not m:
        return jsonify({'error': 'Not found'}), 404
    cid = int(m.group(1))
    conn = get_db()
    own = conn.execute('SELECT owner_id FROM candidates WHERE id=?', (cid,)).fetchone()
    conn.close()
    if not own or own['owner_id'] != effective_company_id():
        return jsonify({'error': 'Not found'}), 404
    fp = os.path.join(CALL_DIR, os.path.basename(filename))
    return send_file(fp) if os.path.exists(fp) else (jsonify({'error': 'Not found'}), 404)



# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  CENTRAL DATABASE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def get_or_create_central_mandate():
    """Returns the ID of the Central Database mandate."""
    conn = get_db()
    r = conn.execute("SELECT value FROM settings WHERE key='central_mandate_id'").fetchone()
    if r and r['value']:
        conn.close()
        return int(r['value'])
    # Create central mandate
    c = conn.cursor()
    c.execute("INSERT INTO mandates (client,role,location,ctc_min,ctc_max,status,created_at) VALUES (?,?,?,?,?,?,?)",
              ('HireLab', 'Central Database', 'All', 0, 99, 'active', ts()))
    mid = c.lastrowid
    c.execute("INSERT OR REPLACE INTO settings (key,value) VALUES ('central_mandate_id',?)", (str(mid),))
    conn.commit(); conn.close()
    return mid

@app.route('/api/central-db/search')
@login_required
def central_search():
    q        = request.args.get('q', '').strip().lower()
    company  = request.args.get('company', '').strip().lower()
    location = request.args.get('location', '').strip().lower()
    phone    = request.args.get('phone', '').strip()
    ctc_min  = request.args.get('ctc_min', '')
    ctc_max  = request.args.get('ctc_max', '')
    exp_min  = request.args.get('exp_min', '')
    exp_max  = request.args.get('exp_max', '')
    notice   = request.args.get('notice', '')
    page     = int(request.args.get('page', 1))
    per_page = 30

    conn = get_db()
    # TENANT ISOLATION: only this company's candidates. owner_id stores the
    # tenant (company) id, so this scopes the Central Database to the current
    # agency. Without this filter, agencies would see each other's candidates.
    rows = conn.execute(
        'SELECT c.*, m.role as mandate_role, m.client as mandate_client '
        'FROM candidates c LEFT JOIN mandates m ON c.mandate_id = m.id '
        'WHERE c.owner_id = ? '
        'ORDER BY c.created_at DESC',
        (effective_company_id(),)
    ).fetchall()
    conn.close()

    results = []
    for r in rows:
        d = _cand_public(r)   # drop embedding / embedding_text / embedding_vec (not needed by UI)
        # Apply filters
        if q:
            searchable = ' '.join([
                str(d.get('name') or ''),
                str(d.get('company') or ''),
                str(d.get('designation') or ''),
                str(d.get('key_skills') or ''),
                str(d.get('career_summary') or ''),
                str(d.get('location') or ''),
                str(d.get('industry_background') or ''),
                str(d.get('email') or ''),
                str(d.get('phone') or ''),
            ]).lower()
            if q not in searchable: continue
        if company and company not in (d.get('company') or '').lower(): continue
        if location and location not in (d.get('location') or '').lower(): continue
        if phone and phone not in (d.get('phone') or ''): continue
        if ctc_min:
            try:
                if (d.get('ctc_current') or 0) < float(ctc_min): continue
            except: pass
        if ctc_max:
            try:
                if (d.get('ctc_current') or 0) > float(ctc_max): continue
            except: pass
        if exp_min:
            try:
                if (d.get('experience') or 0) < float(exp_min): continue
            except: pass
        if exp_max:
            try:
                if (d.get('experience') or 0) > float(exp_max): continue
            except: pass
        if notice:
            try:
                if (d.get('notice_period') or 0) > int(notice): continue
            except: pass
        try: d['key_skills'] = json.loads(d['key_skills'] or '[]')
        except: d['key_skills'] = []
        results.append(d)

    total = len(results)
    start = (page - 1) * per_page
    paginated = results[start:start + per_page]
    return jsonify({'ok': True, 'total': total, 'page': page, 'candidates': paginated})

@app.route('/api/central-db/add', methods=['POST'])
@login_required
def central_db_add():
    d   = request.json or {}
    mid = get_or_create_central_mandate()
    if not d.get('name') or not d.get('company'):
        return jsonify({'error': 'Name and Company required'}), 400
    tenant = effective_company_id()
    conn = get_db(); c = conn.cursor()
    c.execute(
        'INSERT INTO candidates (mandate_id,name,company,designation,experience,ctc_current,'
        'ctc_expected,notice_period,location,phone,email,career_summary,key_skills,'
        'screening_decision,ai_reasoning,stage,created_at,updated_at,owner_id) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
        (mid, d['name'], d['company'], d.get('designation',''), float(d.get('experience') or 0),
         float(d.get('ctc_current') or 0), float(d.get('ctc_expected') or 0),
         int(d.get('notice_period') or 0), d.get('location',''), d.get('phone',''), d.get('email',''),
         d.get('career_summary',''), json.dumps(d.get('key_skills') or []),
         'worth_opening', 'Added to Central Database', 'Central DB', ts(), ts(), tenant))
    cid = c.lastrowid
    c.execute('INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) VALUES (?,?,?,?,?)',
              (cid, '', 'Central DB', 'Added to Central Database', ts()))
    conn.commit(); conn.close()
    return jsonify({'ok': True, 'id': cid})

@app.route('/api/central-db/bulk', methods=['POST'])
@login_required
def central_db_bulk():
    d   = request.json or {}
    mid = get_or_create_central_mandate()
    tenant = effective_company_id()
    candidates = d.get('candidates', [])
    conn = get_db(); c = conn.cursor()
    added = 0
    for cand in candidates:
        name = str(cand.get('name') or '').strip()
        company = str(cand.get('company') or '').strip()
        if not name: continue
        c.execute(
            'INSERT INTO candidates (mandate_id,name,company,designation,experience,ctc_current,'
            'ctc_expected,notice_period,location,phone,email,career_summary,key_skills,'
            'screening_decision,ai_reasoning,stage,created_at,updated_at,owner_id) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
            (mid, name, company, cand.get('designation',''), float(cand.get('experience') or 0),
             float(cand.get('ctc_current') or 0), float(cand.get('ctc_expected') or 0),
             int(cand.get('notice_period') or 0), cand.get('location',''),
             cand.get('phone',''), cand.get('email',''), cand.get('career_summary',''),
             json.dumps(cand.get('key_skills') or []),
             'worth_opening', 'Bulk added to Central Database', 'Central DB', ts(), ts(), tenant))
        cid = c.lastrowid
        c.execute('INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) VALUES (?,?,?,?,?)',
                  (cid, '', 'Central DB', 'Bulk added', ts()))
        added += 1
    conn.commit(); conn.close()
    return jsonify({'ok': True, 'added': added})



# ── WhatsApp Response Tracking ────────────────────────────────────────────────
@app.route('/api/candidates/<int:cid>/wa-response', methods=['POST'])
@login_required
def mark_wa_response(cid):
    d = request.json or {}
    response   = d.get('response', '')   # interested / callback / not_interested / no_reply
    note       = d.get('note', '')
    conn = get_db()
    cand = conn.execute('SELECT * FROM candidates WHERE id=? AND owner_id=?',
                        (cid, effective_company_id())).fetchone()
    if not cand: conn.close(); return jsonify({'error': 'Not found'}), 404

    conn.execute('UPDATE candidates SET wa_response=?, wa_response_note=?, wa_response_at=?, updated_at=? WHERE id=?',
                 (response, note, ts(), ts(), cid))

    # Also update stage based on response
    stage_map = {
        'interested':     'Interested',
        'callback':       'Follow Up 1',
        'not_interested': 'Not Interested',
        'no_reply':       cand['stage'],  # keep current stage
    }
    new_stage = stage_map.get(response, cand['stage'])
    if new_stage != cand['stage']:
        conn.execute('UPDATE candidates SET stage=? WHERE id=?', (new_stage, cid))
        conn.execute('INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) VALUES (?,?,?,?,?)',
                     (cid, cand['stage'], new_stage,
                      'WhatsApp response: ' + response + ((' — ' + note) if note else ''), ts()))
    else:
        conn.execute('INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) VALUES (?,?,?,?,?)',
                     (cid, cand['stage'], cand['stage'],
                      'WhatsApp response logged: ' + response + ((' — ' + note) if note else ''), ts()))

    conn.commit(); conn.close()
    return jsonify({'ok': True, 'stage': new_stage})

# ── WhatsApp Follow-up Queue ───────────────────────────────────────────────────
@app.route('/api/mandates/<int:mid>/wa-queue')
@login_required
def wa_queue(mid):
    """Returns candidates needing WA action — sent but no response logged."""
    import datetime as dt
    now = dt.datetime.now()
    conn = get_db()
    if not _tenant_owns_mandate(conn, mid):
        conn.close(); return jsonify({'error': 'Not found'}), 404
    cands = conn.execute(
        'SELECT * FROM candidates WHERE mandate_id=? AND stage NOT IN (?,?,?,?)',
        (mid, 'Screened-Out', 'Not Interested', 'Placed', 'Central DB')
    ).fetchall()
    conn.close()

    fu_due = []
    for c in cands:
        d = dict(c)
        try: d['key_skills'] = json.loads(d['key_skills'] or '[]')
        except: d['key_skills'] = []

        # Intro sent but no response
        if d.get('msg1_sent_at') and not d.get('wa_response'):
            sent = dt.datetime.fromisoformat(d['msg1_sent_at'])
            days_since = (now - sent).days
            msg_type = 'msg1'
            if d.get('fu2_sent_at'):
                sent = dt.datetime.fromisoformat(d['fu2_sent_at'])
                days_since = (now - sent).days
                msg_type = 'fu2'
            elif d.get('fu1_sent_at'):
                sent = dt.datetime.fromisoformat(d['fu1_sent_at'])
                days_since = (now - sent).days
                msg_type = 'fu1'
            d['days_since_last_msg'] = days_since
            d['last_msg_type'] = msg_type
            d['last_msg_sent_at'] = sent.strftime('%d %b')
            fu_due.append(d)

    # Sort by days_since (longest first — most overdue)
    fu_due.sort(key=lambda x: x['days_since_last_msg'], reverse=True)
    return jsonify({'ok': True, 'queue': fu_due, 'count': len(fu_due)})



# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  CANDIDATE SUBMISSION FORM
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.route('/apply')
def apply_form():
    return send_file('apply.html')

@app.route('/api/public/parse-resume', methods=['POST'])
def public_parse_resume():
    if not _rate_ok('pub_parse', 30, 3600):
        return jsonify({'error': 'Too many requests. Please try again later.'}), 429
    if 'resume' not in request.files:
        return jsonify({'error': 'No file'}), 400
    f = request.files['resume']
    text, err = extract_text_from_file(f.read(), f.filename)
    if err or not text:
        return jsonify({'error': err or 'Cannot extract text'}), 400
    ds_key = get_setting('deepseek_api_key')
    if not ds_key:
        return jsonify({'error': 'Resume parsing not configured. Please fill manually.'}), 400
    sys_msg = ('Extract candidate details. Return ONLY JSON: name, phone, email, company, designation, '
               'experience (float), ctc_current (float LPA), ctc_expected (float LPA), '
               'notice_period (int days), location, key_skills (array max 8). '
               'null for missing strings, 0 for missing numbers.')
    try:
        resp = call_deepseek(ds_key,
            {'model': 'deepseek-chat', 'temperature': 0, 'max_tokens': 800,
                  'messages': [{'role': 'system', 'content': sys_msg},
                                {'role': 'user', 'content': 'Extract:\n\n' + text[:8000]}],
                  'response_format': {'type': 'json_object'}},
            timeout=45, endpoint='extract')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    if resp.status_code != 200:
        return jsonify({'error': 'Parse unavailable. Fill manually.'}), 500
    parsed = parse_json(resp.json()['choices'][0]['message']['content'])
    return jsonify({'ok': True, 'data': parsed}) if parsed else (jsonify({'error': 'Parse failed'}), 500)

@app.route('/api/submit', methods=['POST'])
def submit_form():
    # Spam protection for this PUBLIC endpoint:
    # (a) per-IP rate limit, (b) honeypot field that only bots fill in.
    if not _rate_ok('submit', 20, 3600):
        return jsonify({'error': 'Too many submissions. Please try again later.'}), 429
    if (request.form.get('website') or '').strip():
        # Honeypot tripped — pretend success, save nothing.
        return jsonify({'ok': True})
    name    = request.form.get('name', '').strip()
    phone   = request.form.get('phone', '').strip()
    email   = request.form.get('email', '').strip()
    company = request.form.get('company', '').strip()
    if not name or not phone or not email or not company:
        return jsonify({'error': 'Required fields missing'}), 400
    cv_path = ''; cv_name = ''
    if 'resume' in request.files:
        f = request.files['resume']
        if f.filename:
            ext  = Path(f.filename).suffix.lower()
            safe = str(int(datetime.datetime.now().timestamp())) + '_sub' + ext
            f.save(os.path.join(CV_DIR, safe))
            cv_path = safe; cv_name = f.filename
    conn = get_db(); c = conn.cursor()
    c.execute('INSERT INTO submissions (name,phone,email,company,designation,experience,'
              'ctc_current,ctc_expected,notice_period,location,key_skills,custom_fields,'
              'cv_path,cv_original_name,resume_parsed,status,created_at) '
              'VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
        (name, phone, email, company,
         request.form.get('designation', ''),
         float(request.form.get('experience') or 0),
         float(request.form.get('ctc_current') or 0),
         float(request.form.get('ctc_expected') or 0),
         int(request.form.get('notice_period') or 0),
         request.form.get('location', ''),
         request.form.get('key_skills', '[]'),
         request.form.get('custom_fields', '{}'),
         cv_path, cv_name, 1 if cv_path else 0, 'new', ts()))
    conn.commit(); conn.close()
    return jsonify({'ok': True})

@app.route('/api/submissions')
def get_submissions():
    q        = request.args.get('q', '').strip().lower()
    sf       = request.args.get('status', '')
    exp_r    = request.args.get('exp', '')      # e.g. "3-6"
    ctc_r    = request.args.get('ctc', '')      # e.g. "10-15"
    notice_r = request.args.get('notice', '')   # e.g. "1-30"
    loc_f    = request.args.get('loc', '').strip().lower()
    page     = int(request.args.get('page', 1)); per = 30

    conn  = get_db()
    rows  = conn.execute('SELECT * FROM submissions ORDER BY created_at DESC').fetchall()
    conn.close()

    def parse_range(s):
        try: lo, hi = s.split('-'); return float(lo), float(hi)
        except: return None, None

    results = []
    for r in rows:
        d = dict(r)
        try: d['key_skills'] = json.loads(d.get('key_skills') or '[]')
        except: d['key_skills'] = []
        try: d['custom_fields'] = json.loads(d.get('custom_fields') or '{}')
        except: d['custom_fields'] = {}

        # Status filter
        if sf and d.get('status') != sf: continue

        # Experience filter
        if exp_r:
            lo, hi = parse_range(exp_r)
            if lo is not None:
                exp = float(d.get('experience') or 0)
                if not (lo <= exp <= hi): continue

        # CTC filter
        if ctc_r:
            lo, hi = parse_range(ctc_r)
            if lo is not None:
                ctc = float(d.get('ctc_current') or 0)
                if not (lo <= ctc <= hi): continue

        # Notice period filter
        if notice_r:
            lo, hi = parse_range(notice_r)
            if lo is not None:
                notice = float(d.get('notice_period') or 0)
                if not (lo <= notice <= hi): continue

        # Location filter
        if loc_f and loc_f not in (d.get('location') or '').lower(): continue

        # Boolean text search
        if q:
            blob = ' '.join([d.get('name',''), d.get('company',''), d.get('designation',''),
                             d.get('location',''), d.get('email',''),
                             ' '.join(d.get('key_skills',[]))]).lower()
            if ' or ' in q:
                if not any(t.strip() in blob for t in q.split(' or ')): continue
            elif ' and ' in q:
                if not all(t.strip() in blob for t in q.split(' and ')): continue
            else:
                if q not in blob: continue

        results.append(d)

    total = len(results)
    return jsonify({'ok': True, 'total': total, 'page': page,
                    'submissions': results[(page-1)*per : page*per]})

@app.route('/api/submissions/<int:sid>', methods=['PUT'])
def update_submission(sid):
    d = request.json or {}
    conn = get_db()
    if 'status' in d:     conn.execute('UPDATE submissions SET status=? WHERE id=?',     (d['status'], sid))
    if 'notes' in d:      conn.execute('UPDATE submissions SET notes=? WHERE id=?',      (d['notes'], sid))
    if 'domain_tags' in d: conn.execute('UPDATE submissions SET domain_tags=? WHERE id=?', (json.dumps(d['domain_tags']), sid))
    conn.commit(); conn.close()
    return jsonify({'ok': True})

@app.route('/api/submissions/<int:sid>/add-to-pipeline', methods=['POST'])
def add_submission_to_pipeline(sid):
    d   = request.json or {}
    mid = d.get('mandate_id')
    if not mid: return jsonify({'error': 'mandate_id required'}), 400
    conn = get_db()
    sub  = conn.execute('SELECT * FROM submissions WHERE id=?', (sid,)).fetchone()
    if not sub: conn.close(); return jsonify({'error': 'Not found'}), 404
    c = conn.cursor()
    c.execute('INSERT INTO candidates (mandate_id,name,company,designation,experience,ctc_current,'
              'ctc_expected,notice_period,location,phone,email,key_skills,screening_decision,'
              'ai_reasoning,stage,cv_path,cv_original_name,created_at,updated_at) '
              'VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
        (mid, sub['name'], sub['company'], sub['designation'], sub['experience'],
         sub['ctc_current'], sub['ctc_expected'], sub['notice_period'], sub['location'],
         sub['phone'], sub['email'], sub['key_skills'],
         'worth_opening', 'Added from submission form', 'Screening',
         sub['cv_path'], sub['cv_original_name'], ts(), ts()))
    cid = c.lastrowid
    c.execute('INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) '
              'VALUES (?,?,?,?,?)', (cid, '', 'Screening', 'Added from submission form', ts()))
    conn.execute('UPDATE submissions SET status=? WHERE id=?', ('added_to_pipeline', sid))
    conn.commit(); conn.close()
    return jsonify({'ok': True, 'candidate_id': cid})

@app.route('/api/form-config', methods=['GET', 'POST'])
def form_config():
    conn = get_db()
    if request.method == 'POST':
        if not current_user():          # only logged-in staff may change the form
            conn.close(); return jsonify({'error': 'Unauthorized'}), 401
        conn.execute('INSERT OR REPLACE INTO settings (key,value) VALUES (?,?)',
                     ('form_config', json.dumps(request.json or {})))
        conn.commit(); conn.close()
        return jsonify({'ok': True})
    r = conn.execute("SELECT value FROM settings WHERE key='form_config'").fetchone()
    conn.close()
    return jsonify({'ok': True, 'config': json.loads(r['value'] if r else '{}')})


# Intelligence

# Client Submission Sheet
# ══════════════════════════════════════════════════════════════════════════
#  WhatsApp — outbound templates + inbound reply capture (via personal-number
#  listener). Sending is done by the recruiter (wa.me / their phone); a small
#  linked-device listener POSTs every message (sent + received) to
#  /api/wa-inbound so the full thread appears per candidate. Receive-only:
#  the ATS never auto-sends here.
# ══════════════════════════════════════════════════════════════════════════

WA_DEFAULT_TEMPLATES = [
    {"cat": "📥 Sourcing / Early stage", "items": [
        {"id": "intro",       "title": "Intro / First outreach",
         "body": "Hi {{Name}}, {{You}} here from HireLab. We do senior hiring in Solar, Electrical & Automation.\nThere's a {{Role}} role{{ in {{Location}}}} that looks like a good fit for your profile. Interested in knowing more? 🙂"},
        {"id": "fu1",         "title": "Follow-up 1 (2-3 days)",
         "body": "Hi {{Name}}, just following up on my earlier message about the {{Role}} role 🙂 No rush — even if the timing isn't right, that's totally fine. Want me to share the details?"},
        {"id": "fu2",         "title": "Follow-up 2 (final nudge)",
         "body": "Hi {{Name}}, last message from my side, promise 🙏 If this or any future role interests you, just ping me anytime. Take care!"},
        {"id": "share_jd",    "title": "Interested → share JD",
         "body": "Awesome, {{Name}}! Here's the JD for the {{Role}} role 👇\n\n{{JD}}\n\nHave a look — ask me anything if something's unclear, and we can cover the rest on a quick call."},
        {"id": "cv_request",  "title": "CV / resume request",
         "body": "Hi {{Name}}, your profile looks interesting 🙂 Could you send me an updated resume? That way I can match you with the right roles."},
        {"id": "availability","title": "Quick availability check",
         "body": "Hi {{Name}}, wanted to have a quick 5-10 min chat about a role. What time works for you today? I'll give you a call. 🙂"},
        {"id": "reactivate",  "title": "Reactivate old candidate",
         "body": "Hi {{Name}}, it's been a while! How are things going — happy where you are, or open to exploring something new? Just touching base. 🙂"},
    ]},
    {"cat": "🗣️ Objection handling", "items": [
        {"id": "ctc_ask",     "title": "Candidate asks salary/CTC",
         "body": "Fair question 🙂 The package is competitive, but it's best to discuss once I understand your current fitment. Could you share your current & expected CTC? Then I can position you properly."},
        {"id": "which_co",    "title": "Which company?",
         "body": "Totally understand you'd want to know 🙂 The client is a well-known name in this space — I'll share it on a call once there's a mutual fit. Shall we connect for 10 mins?"},
        {"id": "notice",      "title": "Notice period too long",
         "body": "Got it, {{Name}}. For the {{Role}} role, notice can be a bit flexible — is a buyout or early release possible on your end? We can work it out with the client."},
        {"id": "ctc_high",    "title": "Expected CTC mismatch",
         "body": "Thanks for being upfront, {{Name}} 🙂 The number's slightly above the band for this role, but for a strong profile the gap can often be bridged. Is there some flexibility, or is this firm?"},
        {"id": "remote",      "title": "Remote / WFH / hybrid",
         "body": "Good question 🙂 This role is based in {{Location}}. If relocation or commute is a concern, let me know and we'll figure it out."},
        {"id": "relocate",    "title": "Location / relocation",
         "body": "Hi {{Name}}, the role is in {{Location}}. Are you open to relocating, or would you prefer to stay in your current city? Either way, I'll guide you."},
    ]},
    {"cat": "🎯 Interview stage", "items": [
        {"id": "sched_call",  "title": "Interested → schedule call",
         "body": "Perfect, {{Name}}! When works for a 10-min call — today or tomorrow? I'll call you, just share a slot. 📞"},
        {"id": "qual",        "title": "Qualification ask",
         "body": "Thanks, {{Name}}! To position your profile well, could you quickly share:\n• Current company\n• Current & expected CTC (LPA)\n• Notice period\n• Location\nTakes 2 mins 🙏"},
        {"id": "iv_invite",   "title": "Interview invite",
         "body": "Good news, {{Name}}! 🎉 The client wants to take this forward. Your {{Role}} interview is proposed for {{Date}}, {{Time}} ({{Mode}}). Does this time work? Confirm and I'll send the full details."},
        {"id": "iv_reminder", "title": "Interview reminder",
         "body": "Hi {{Name}}, quick reminder — your {{Role}} interview is on {{Date}} at {{Time}}. All set? Let me know if you need anything. All the best! 💪"},
        {"id": "iv_reschedule","title": "Reschedule interview",
         "body": "No problem at all, {{Name}} 🙂 Let's reschedule — which day/time would work better for you? I'll coordinate with the client."},
        {"id": "iv_noshow",   "title": "No-show follow-up",
         "body": "Hi {{Name}}, we couldn't connect for today's interview — everything okay? 🙂 If something came up, no worries, we can set it up again. Just tell me when suits you."},
        {"id": "next_round",  "title": "Positive feedback → next round",
         "body": "Great news, {{Name}}! 🎉 The first round feedback was good and the client is keen on a next round. {{Date}}, {{Time}} is proposed — does that work?"},
    ]},
    {"cat": "🤝 Offer & closing", "items": [
        {"id": "next_steps",  "title": "Post-call → next steps",
         "body": "Great talking, {{Name}}! As discussed, I'm putting your profile forward for {{Role}}. Send your updated resume here so I can move fast. I'll keep you posted at every step. 🙏"},
        {"id": "reject",      "title": "Rejection (graceful)",
         "body": "Hi {{Name}}, the client decided to go in a different direction this time 🙏 The feedback wasn't about your profile — it was purely a role-fit call. I'll keep reaching out for strongly relevant roles — you're on my radar. 🙂"},
        {"id": "offer",       "title": "Offer rolled out 🎉",
         "body": "{{Name}}, congratulations! 🎉 The client has rolled out the {{Role}} offer. Sending the detailed offer across — have a look and let's talk through any questions. Really happy for you! 🙌"},
        {"id": "counter",     "title": "Candidate got a counteroffer",
         "body": "I understand, {{Name}}, this is an important decision 🙂 Counteroffers are often short-term — weigh the long-term growth, role and company, not just the number. Shall we do a quick call? No pressure, just to help you think it through."},
        {"id": "prejoin",     "title": "Pre-joining keep-warm",
         "body": "Hi {{Name}}, hope the prep is going smoothly 🙂 With {{DOJ}} approaching, if you need anything around paperwork or joining, I'm right here. Excited for you! 💪"},
        {"id": "not_interested","title": "Not interested → graceful close",
         "body": "No worries at all, {{Name}} 🙌 Thanks for your time! I'll keep your profile on record and only reach out if something strongly relevant comes up. Take care!"},
    ]},
]


def _wa_norm_phone10(phone):
    """Return the last-10-digit form of a phone number for matching."""
    pc = re.sub(r'[^0-9]', '', phone or '')
    if pc.startswith('91') and len(pc) == 12:
        pc = pc[2:]
    return pc[-10:] if len(pc) >= 10 else pc


def _wa_get_inbound_token(company_id):
    """Get (or lazily create) this company's WhatsApp listener token."""
    conn = get_db()
    r = conn.execute('SELECT value FROM tenant_settings WHERE company_id=? AND key=?',
                     (company_id, 'wa_inbound_token')).fetchone()
    if r and r['value']:
        conn.close(); return r['value']
    tok = secrets.token_urlsafe(24)
    conn.execute('INSERT OR REPLACE INTO tenant_settings (company_id,key,value) VALUES (?,?,?)',
                 (company_id, 'wa_inbound_token', tok))
    conn.commit(); conn.close()
    return tok


def _wa_company_for_token(tok):
    if not tok:
        return None
    conn = get_db()
    r = conn.execute("SELECT company_id FROM tenant_settings WHERE key='wa_inbound_token' AND value=?",
                     (tok,)).fetchone()
    conn.close()
    return r['company_id'] if r else None


@app.route('/api/wa-inbound', methods=['POST'])
def wa_inbound():
    """The personal-number listener POSTs every message here (sent + received).
    Auth is by the per-company listener token (NOT a user session). We match the
    sender to a candidate in that company and store the message; unmatched
    numbers are ignored (by design)."""
    d = request.json or {}
    tok = d.get('token') or request.headers.get('X-WA-Token', '')
    company_id = _wa_company_for_token(tok)
    if not company_id:
        return jsonify({'error': 'invalid listener token'}), 401
    phone = d.get('phone') or ''
    text = (d.get('text') or '').strip()
    direction = 'inbound' if (d.get('direction') or 'inbound') == 'inbound' else 'outbound'
    wamid = (d.get('wa_message_id') or '').strip()
    when = d.get('ts') or ts()
    if not phone or not text:
        return jsonify({'error': 'phone and text required'}), 400

    p10 = _wa_norm_phone10(phone)
    if not p10:
        return jsonify({'ok': True, 'ignored': 'bad_phone'})

    conn = get_db()
    cand = conn.execute(
        "SELECT * FROM candidates WHERE owner_id=? AND "
        "replace(replace(replace(replace(phone,'+',''),' ',''),'-',''),'(','') LIKE ? LIMIT 1",
        (company_id, '%' + p10)).fetchone()
    if not cand:
        conn.close()
        return jsonify({'ok': True, 'ignored': 'no_candidate_match'})

    # Dedupe re-syncs by WhatsApp's own message id
    if wamid:
        dup = conn.execute("SELECT 1 FROM wa_messages WHERE wa_message_id=? LIMIT 1", (wamid,)).fetchone()
        if dup:
            conn.close()
            return jsonify({'ok': True, 'duplicate': True})

    conv = conn.execute(
        "SELECT * FROM wa_conversations WHERE company_id=? AND candidate_id=? ORDER BY id DESC LIMIT 1",
        (company_id, cand['id'])).fetchone()
    if not conv:
        conv_id = conn.execute(
            'INSERT INTO wa_conversations (company_id,candidate_id,mandate_id,candidate_phone,'
            'candidate_name,status,auto_mode,last_message_at,created_at,updated_at) '
            'VALUES (?,?,?,?,?,?,?,?,?,?)',
            (company_id, cand['id'], cand['mandate_id'], cand['phone'] or phone,
             cand['name'] or '', 'active', 0, when, ts(), ts())).lastrowid
    else:
        conv_id = conv['id']

    # If this is an OUTBOUND message the ATS already logged instantly (when the
    # recruiter clicked Send), don't insert a duplicate — just backfill the real
    # WhatsApp message id onto that row so future syncs dedupe by id.
    if direction == 'outbound':
        prior = conn.execute(
            "SELECT id FROM wa_messages WHERE conversation_id=? AND direction='outbound' "
            "AND content=? AND (wa_message_id IS NULL OR wa_message_id='') "
            "ORDER BY id DESC LIMIT 1", (conv_id, text)).fetchone()
        if prior:
            if wamid:
                conn.execute("UPDATE wa_messages SET wa_message_id=? WHERE id=?", (wamid, prior['id']))
            conn.commit(); conn.close()
            return jsonify({'ok': True, 'merged': True})

    conn.execute(
        'INSERT INTO wa_messages (conversation_id,direction,sender,content,message_type,'
        'wa_message_id,created_at) VALUES (?,?,?,?,?,?,?)',
        (conv_id, direction, (cand['name'] if direction == 'inbound' else 'You'),
         text, 'text', wamid, when))
    conn.execute('UPDATE wa_conversations SET last_message_at=?, updated_at=? WHERE id=?',
                 (when, ts(), conv_id))
    conn.commit(); conn.close()
    return jsonify({'ok': True, 'candidate_id': cand['id']})


@app.route('/api/candidates/<int:cid>/wa-thread')
@login_required
def candidate_wa_thread(cid):
    """This candidate's WhatsApp thread only (tenant-scoped)."""
    conn = get_db()
    if not _tenant_owns_candidate(conn, cid):
        conn.close(); return jsonify({'error': 'Not found'}), 404
    conv = conn.execute(
        "SELECT id FROM wa_conversations WHERE company_id=? AND candidate_id=? ORDER BY id DESC LIMIT 1",
        (effective_company_id(), cid)).fetchone()
    if not conv:
        conn.close(); return jsonify({'ok': True, 'messages': []})
    msgs = conn.execute(
        'SELECT direction,sender,content,message_type,created_at FROM wa_messages '
        'WHERE conversation_id=? ORDER BY id ASC', (conv['id'],)).fetchall()
    conn.close()
    return jsonify({'ok': True, 'messages': [dict(m) for m in msgs]})


@app.route('/api/candidates/<int:cid>/wa-send-log', methods=['POST'])
@login_required
def wa_send_log(cid):
    """Log an outbound WhatsApp message instantly (when the recruiter clicks
    Send in the ATS) so the thread updates immediately. The listener later sees
    the same message and dedupes it against this row (see wa_inbound)."""
    d = request.json or {}
    text = (d.get('text') or '').strip()
    if not text:
        return jsonify({'error': 'empty'}), 400
    company_id = effective_company_id()
    conn = get_db()
    cand = conn.execute('SELECT * FROM candidates WHERE id=? AND owner_id=?', (cid, company_id)).fetchone()
    if not cand:
        conn.close(); return jsonify({'error': 'Not found'}), 404
    conv = conn.execute(
        "SELECT id FROM wa_conversations WHERE company_id=? AND candidate_id=? ORDER BY id DESC LIMIT 1",
        (company_id, cid)).fetchone()
    if conv:
        conv_id = conv['id']
    else:
        conv_id = conn.execute(
            'INSERT INTO wa_conversations (company_id,candidate_id,mandate_id,candidate_phone,'
            'candidate_name,status,auto_mode,last_message_at,created_at,updated_at) '
            'VALUES (?,?,?,?,?,?,?,?,?,?)',
            (company_id, cid, cand['mandate_id'], cand['phone'] or '', cand['name'] or '',
             'active', 0, ts(), ts(), ts())).lastrowid
    conn.execute(
        'INSERT INTO wa_messages (conversation_id,direction,sender,content,message_type,'
        'wa_message_id,created_at) VALUES (?,?,?,?,?,?,?)',
        (conv_id, 'outbound', 'You', text, 'text', '', ts()))
    conn.execute('UPDATE wa_conversations SET last_message_at=?, updated_at=? WHERE id=?', (ts(), ts(), conv_id))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/wa-learn-style', methods=['POST'])
@login_required
def wa_learn_style():
    """Read this company's OUTBOUND WhatsApp messages and have DeepSeek write a
    concise 'how this recruiter communicates' style guide, saved for AI drafts."""
    ds_key = get_setting('deepseek_api_key')
    if not ds_key:
        return jsonify({'error': 'DeepSeek API key not set in Settings.'}), 400
    company_id = effective_company_id()
    conn = get_db()
    rows = conn.execute(
        "SELECT m.content FROM wa_messages m JOIN wa_conversations c ON c.id=m.conversation_id "
        "WHERE c.company_id=? AND m.direction='outbound' AND length(trim(m.content))>0 "
        "ORDER BY m.id DESC LIMIT 120", (company_id,)).fetchall()
    conn.close()
    msgs = [r['content'] for r in rows if r['content']]
    if len(msgs) < 5:
        return jsonify({'error': 'Not enough messages yet. Kuch WhatsApp messages bhejo, phir try karo.', 'count': len(msgs)}), 400
    sample = "\n---\n".join(msgs[:120])
    system_msg = (
        "You analyse how a recruiter writes WhatsApp messages, so another AI can "
        "later write messages that sound EXACTLY like them. Read the recruiter's "
        "real sent messages and produce a concise style guide (6-9 short bullet "
        "points) covering: language (Hindi/English/Hinglish mix), tone & warmth, "
        "typical greetings and sign-offs, emoji usage, message length, formality, "
        "and any recurring phrases or habits. Be specific and practical. Output "
        "ONLY the bullet-point style guide, nothing else.")
    try:
        resp = call_deepseek(ds_key, {
            'model': 'deepseek-chat', 'temperature': 0.3, 'max_tokens': 500,
            'messages': [{'role': 'system', 'content': system_msg},
                         {'role': 'user', 'content': "Recruiter's sent messages:\n\n" + sample}]})
        profile = resp.json()['choices'][0]['message']['content'].strip()
    except TokenCapError:
        return jsonify({'error': 'Token limit reached.'}), 429
    except Exception as e:
        return jsonify({'error': 'AI call failed: ' + str(e)[:100]}), 500
    if not profile:
        return jsonify({'error': 'Could not generate a style profile.'}), 500
    set_setting('wa_style_profile', profile)
    return jsonify({'ok': True, 'profile': profile, 'learned_from': len(msgs)})


@app.route('/api/candidates/<int:cid>/wa-draft', methods=['POST'])
@login_required
def wa_draft_reply(cid):
    """Draft the next WhatsApp reply for a candidate, in the recruiter's own
    learned style + recent conversation context."""
    ds_key = get_setting('deepseek_api_key')
    if not ds_key:
        return jsonify({'error': 'DeepSeek API key not set in Settings.'}), 400
    company_id = effective_company_id()
    conn = get_db()
    cand = conn.execute('SELECT * FROM candidates WHERE id=? AND owner_id=?', (cid, company_id)).fetchone()
    if not cand:
        conn.close(); return jsonify({'error': 'Not found'}), 404
    conv = conn.execute(
        "SELECT id FROM wa_conversations WHERE company_id=? AND candidate_id=? ORDER BY id DESC LIMIT 1",
        (company_id, cid)).fetchone()
    thread = []
    if conv:
        rows = conn.execute(
            "SELECT direction, content FROM wa_messages WHERE conversation_id=? ORDER BY id DESC LIMIT 12",
            (conv['id'],)).fetchall()
        thread = list(reversed([dict(r) for r in rows]))
    conn.close()
    if not thread:
        return jsonify({'error': 'No conversation yet to reply to.'}), 400
    convo_txt = "\n".join([('Candidate: ' if m['direction'] == 'inbound' else 'Me: ') + (m['content'] or '') for m in thread])
    style = get_setting('wa_style_profile', '')
    style_block = ("Write EXACTLY in my personal style:\n" + style + "\n\n") if style else ""
    # classify the situation so we can pull the right learning + track confidence
    conn2 = get_db()
    category = 'general'
    try:
        cat_ids = ", ".join(WA_CATEGORY_IDS)
        cresp = call_deepseek(ds_key, {'model': 'deepseek-chat', 'temperature': 0, 'max_tokens': 12,
            'messages': [{'role': 'system', 'content': 'Classify the recruitment WhatsApp conversation into ONE category id from: ' + cat_ids + '. Reply with ONLY the id.'},
                         {'role': 'user', 'content': convo_txt[-1500:]}]})
        c = cresp.json()['choices'][0]['message']['content'].strip().lower()
        category = c if c in WA_CATEGORY_IDS else 'general'
    except Exception:
        category = 'general'
    learn = _wa_learning_block(conn2, company_id, category)
    system_msg = (
        "You are helping a recruiter draft the next WhatsApp reply to a candidate. "
        + style_block + learn +
        "Rules: natural and human, match my style above, never invent salary/CTC "
        "numbers or a client name, keep it concise. Output ONLY the reply text.")
    try:
        resp = call_deepseek(ds_key, {
            'model': 'deepseek-chat', 'temperature': 0.5, 'max_tokens': 220,
            'messages': [{'role': 'system', 'content': system_msg},
                         {'role': 'user', 'content': "Conversation so far:\n" + convo_txt + "\n\nDraft my next reply:"}]})
        draft = resp.json()['choices'][0]['message']['content'].strip()
    except TokenCapError:
        conn2.close(); return jsonify({'error': 'Token limit reached.'}), 429
    except Exception as e:
        conn2.close(); return jsonify({'error': 'AI call failed: ' + str(e)[:100]}), 500
    # record a 'drafted' row; queue-send later resolves it (approved/edited)
    _ensure_wa_feedback(conn2)
    draft_id = conn2.execute(
        "INSERT INTO wa_agent_feedback (company_id,candidate_id,conversation_id,category,"
        "agent_draft,sent_text,action,comment,created_at) VALUES (?,?,?,?,?,?,?,?,?)",
        (company_id, cid, (conv['id'] if conv else None), category, draft, '', 'drafted', '', ts())).lastrowid
    conn2.commit(); conn2.close()
    return jsonify({'ok': True, 'draft': draft, 'styled': bool(style), 'draft_id': draft_id, 'category': category})


def _ensure_wa_outbox(conn):
    conn.execute('''CREATE TABLE IF NOT EXISTS wa_outbox (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        company_id INTEGER, candidate_id INTEGER, phone TEXT, text TEXT,
        status TEXT DEFAULT 'pending', wa_message_id TEXT DEFAULT '',
        created_at TEXT, sent_at TEXT)''')
    conn.execute('CREATE INDEX IF NOT EXISTS idx_wa_outbox_status ON wa_outbox(company_id,status)')


@app.route('/api/candidates/<int:cid>/wa-queue-send', methods=['POST'])
@login_required
def wa_queue_send(cid):
    """Queue an outbound message for the listener to send from the recruiter's
    WhatsApp in the background (no wa.me tab). Also logs it to the thread now so
    it appears instantly."""
    d = request.json or {}
    text = (d.get('text') or '').strip()
    if not text:
        return jsonify({'error': 'empty'}), 400
    company_id = effective_company_id()
    conn = get_db()
    cand = conn.execute('SELECT * FROM candidates WHERE id=? AND owner_id=?', (cid, company_id)).fetchone()
    if not cand:
        conn.close(); return jsonify({'error': 'Not found'}), 404
    digits = re.sub(r'[^0-9]', '', cand['phone'] or '')
    if not digits:
        conn.close(); return jsonify({'error': 'No phone number'}), 400
    if len(digits) == 10:
        digits = '91' + digits
    _ensure_wa_outbox(conn)
    conv = conn.execute(
        "SELECT id FROM wa_conversations WHERE company_id=? AND candidate_id=? ORDER BY id DESC LIMIT 1",
        (company_id, cid)).fetchone()
    if conv:
        conv_id = conv['id']
    else:
        conv_id = conn.execute(
            'INSERT INTO wa_conversations (company_id,candidate_id,mandate_id,candidate_phone,'
            'candidate_name,status,auto_mode,last_message_at,created_at,updated_at) '
            'VALUES (?,?,?,?,?,?,?,?,?,?)',
            (company_id, cid, cand['mandate_id'], cand['phone'] or '', cand['name'] or '',
             'active', 0, ts(), ts(), ts())).lastrowid
    conn.execute(
        'INSERT INTO wa_messages (conversation_id,direction,sender,content,message_type,'
        'wa_message_id,created_at) VALUES (?,?,?,?,?,?,?)',
        (conv_id, 'outbound', 'You', text, 'text', '', ts()))
    conn.execute('UPDATE wa_conversations SET last_message_at=?, updated_at=? WHERE id=?', (ts(), ts(), conv_id))
    conn.execute('INSERT INTO wa_outbox (company_id,candidate_id,phone,text,status,created_at) '
                 'VALUES (?,?,?,?,?,?)', (company_id, cid, digits, text, 'pending', ts()))
    # LEARNING: if this send came from an AI draft, record whether I kept or edited it
    draft_id = d.get('draft_id')
    if draft_id:
        _ensure_wa_feedback(conn)
        fb = conn.execute("SELECT agent_draft FROM wa_agent_feedback WHERE id=? AND company_id=?",
                          (draft_id, company_id)).fetchone()
        if fb:
            act = 'approved' if text == (fb['agent_draft'] or '').strip() else 'edited'
            conn.execute("UPDATE wa_agent_feedback SET sent_text=?, action=?, comment=?, resolved_at=? WHERE id=?",
                         (text, act, (d.get('comment') or '').strip(), ts(), draft_id))
    conn.commit(); conn.close()
    return jsonify({'ok': True, 'queued': True})


@app.route('/api/wa-outbox')
def wa_outbox_pending():
    """Listener polls this (with its token) for messages to send."""
    tok = request.args.get('token') or request.headers.get('X-WA-Token', '')
    company_id = _wa_company_for_token(tok)
    if not company_id:
        return jsonify({'error': 'invalid token'}), 401
    conn = get_db(); _ensure_wa_outbox(conn)
    rows = conn.execute(
        "SELECT id, phone, text FROM wa_outbox WHERE company_id=? AND status='pending' "
        "ORDER BY id ASC LIMIT 20", (company_id,)).fetchall()
    conn.close()
    return jsonify({'ok': True, 'messages': [dict(r) for r in rows]})


@app.route('/api/wa-outbox/ack', methods=['POST'])
def wa_outbox_ack():
    """Listener reports a queued message as sent (or failed)."""
    d = request.json or {}
    tok = d.get('token') or request.headers.get('X-WA-Token', '')
    company_id = _wa_company_for_token(tok)
    if not company_id:
        return jsonify({'error': 'invalid token'}), 401
    oid = d.get('id')
    if not oid:
        return jsonify({'error': 'id required'}), 400
    status = d.get('status', 'sent')
    if status not in ('sent', 'failed'):
        status = 'sent'
    conn = get_db(); _ensure_wa_outbox(conn)
    conn.execute("UPDATE wa_outbox SET status=?, wa_message_id=?, sent_at=? WHERE id=? AND company_id=?",
                 (status, d.get('wa_message_id', ''), ts(), oid, company_id))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


def _ensure_wa_suggestions(conn):
    conn.execute('''CREATE TABLE IF NOT EXISTS wa_suggestions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        company_id INTEGER, candidate_id INTEGER, conversation_id INTEGER,
        kind TEXT, message TEXT, reason TEXT, based_on TEXT,
        status TEXT DEFAULT 'pending', created_at TEXT)''')
    conn.execute('CREATE INDEX IF NOT EXISTS idx_wa_sugg ON wa_suggestions(company_id,status)')
    # additive column for the learning agent (safe if it already exists)
    try:
        conn.execute("ALTER TABLE wa_suggestions ADD COLUMN category TEXT DEFAULT 'followup'")
    except Exception:
        pass


def _tenant_setting_raw(conn, company_id, key, default=''):
    """Read a tenant setting by explicit company_id (no session needed — used by
    the background scheduler)."""
    r = conn.execute('SELECT value FROM tenant_settings WHERE company_id=? AND key=?',
                     (company_id, key)).fetchone()
    if r and (r['value'] not in (None, '')):
        return r['value']
    g = conn.execute('SELECT value FROM settings WHERE key=?', (key,)).fetchone()
    return (g['value'] if (g and g['value']) else default)


def _hours_since(ts_str):
    if not ts_str:
        return None
    try:
        import datetime as _dt
        t = _dt.datetime.fromisoformat(ts_str)
        return (_ist_now() - t).total_seconds() / 3600.0
    except Exception:
        return None


def _wa_analyze_conversation(thread_txt, hours_since, style, stage, wait_hours, learning=''):
    """Ask DeepSeek for the single best next action + situation category."""
    ds_key = get_setting('deepseek_api_key')
    if not ds_key:
        return None
    style_block = ("The recruiter writes like this — match it in any message:\n" + style + "\n\n") if style else ""
    cat_ids = ", ".join(WA_CATEGORY_IDS)
    system_msg = (
        "You assist a recruiter running WhatsApp chats with candidates. Read the FULL "
        "conversation (Me = recruiter, Candidate = candidate) and how long since the last "
        "message, then choose the single best next action.\n\n" + style_block + learning +
        "Respond with ONLY a JSON object:\n"
        '{"action":"followup"|"not_interested"|"none","category":"<one of: ' + cat_ids + '>",'
        '"message":"<the WhatsApp follow-up to send, ONLY when action=followup, in the recruiter\'s style>",'
        '"reason":"<one short line>"}\n\n'
        "Rules:\n"
        f"- 'followup' ONLY if the candidate showed interest but a needed action is still pending "
        f"(e.g. hasn't sent an updated resume/profile) AND at least {wait_hours} hours passed since the "
        "last message. Write a gentle, non-pushy nudge.\n"
        "- 'not_interested' ONLY if the candidate clearly declined or said not interested / not looking.\n"
        "- 'none' if it's too soon, the candidate already provided what was needed, the recruiter already "
        "acknowledged it (e.g. said thanks for sharing), or nothing is needed.\n"
        "- 'category' = the situation this conversation is currently in.\n"
        "- NEVER invent salary/CTC numbers or a client name.")
    user_msg = ("Hours since last message: " + str(hours_since) + "\nCandidate stage: " + (stage or '') +
                "\n\nConversation:\n" + thread_txt + "\n\nReturn the decision as JSON.")
    try:
        resp = call_deepseek(ds_key, {'model': 'deepseek-chat', 'temperature': 0.3, 'max_tokens': 340,
            'messages': [{'role': 'system', 'content': system_msg}, {'role': 'user', 'content': user_msg}]})
        data = parse_json(resp.json()['choices'][0]['message']['content'])
        return data if isinstance(data, dict) else None
    except Exception:
        return None


_last_wa_scan = 0

def _wa_followup_scan(force=False):
    """Background pass: for each active WhatsApp conversation that's gone quiet,
    ask DeepSeek whether to suggest a follow-up or a 'move to Not Interested'.
    Self-gated to run at most every ~3 hours (unless force)."""
    global _last_wa_scan
    now = _time.time()
    if not force and (now - _last_wa_scan) < 3 * 3600:
        return
    _last_wa_scan = now
    if not get_setting('deepseek_api_key'):
        return
    TERMINAL = ('Not Interested', 'Placed', 'Not Suitable',
                'Client Rejected on Paper', 'Client Rejected After Interview')
    conn = get_db()
    _ensure_wa_suggestions(conn)
    try:
        rows = conn.execute(
            "SELECT cv.id conv_id, cv.company_id, cv.candidate_id, cv.last_message_at, ca.stage "
            "FROM wa_conversations cv JOIN candidates ca ON ca.id=cv.candidate_id "
            "WHERE cv.status='active' ORDER BY cv.last_message_at DESC LIMIT 200").fetchall()
    except Exception:
        conn.close(); return
    analyzed = 0
    for r in rows:
        if analyzed >= 12:
            break
        stage = r['stage'] or ''
        if stage in TERMINAL:
            continue
        last_at = r['last_message_at'] or ''
        try:
            wait_hours = int(_tenant_setting_raw(conn, r['company_id'], 'wa_followup_hours', '24'))
        except Exception:
            wait_hours = 24
        hs = _hours_since(last_at)
        if hs is None or hs < wait_hours:
            continue
        # already looked at this exact conversation state? skip.
        if conn.execute("SELECT 1 FROM wa_suggestions WHERE conversation_id=? AND based_on=? LIMIT 1",
                        (r['conv_id'], last_at)).fetchone():
            continue
        # a pending suggestion already waiting? don't pile up.
        if conn.execute("SELECT 1 FROM wa_suggestions WHERE conversation_id=? AND status='pending' LIMIT 1",
                        (r['conv_id'],)).fetchone():
            continue
        msgs = conn.execute("SELECT direction, content FROM wa_messages WHERE conversation_id=? "
                            "ORDER BY id ASC LIMIT 30", (r['conv_id'],)).fetchall()
        if not msgs:
            continue
        thread_txt = "\n".join([('Candidate: ' if m['direction'] == 'inbound' else 'Me: ') + (m['content'] or '') for m in msgs])
        style = _tenant_setting_raw(conn, r['company_id'], 'wa_style_profile', '')
        analyzed += 1
        # first pass to know the category, then a learning-aware draft
        data = _wa_analyze_conversation(thread_txt, int(hs), style, stage, wait_hours)
        if not data:
            continue
        category = (data.get('category') or 'followup').strip()
        if category not in WA_CATEGORY_IDS:
            category = 'followup'
        action = (data.get('action') or 'none').strip()
        reason = (data.get('reason') or '')[:200]
        # re-draft the follow-up using what the recruiter taught us in this category
        if action == 'followup':
            learn = _wa_learning_block(conn, r['company_id'], category)
            if learn:
                d2 = _wa_analyze_conversation(thread_txt, int(hs), style, stage, wait_hours, learning=learn)
                if d2 and (d2.get('action') == 'followup') and (d2.get('message') or '').strip():
                    data = d2
        if action == 'followup' and (data.get('message') or '').strip():
            kind, msg, status = 'followup', data['message'].strip(), 'pending'
        elif action == 'not_interested':
            kind, msg, status = 'not_interested', '', 'pending'
        else:
            kind, msg, status = 'none', '', 'skipped'   # record so we don't re-analyze same state
        conn.execute("INSERT INTO wa_suggestions (company_id,candidate_id,conversation_id,kind,message,"
                     "reason,based_on,status,created_at,category) VALUES (?,?,?,?,?,?,?,?,?,?)",
                     (r['company_id'], r['candidate_id'], r['conv_id'], kind, msg, reason, last_at, status, ts(), category))
        conn.commit()
    conn.close()


@app.route('/api/wa-suggestions')
@login_required
def wa_suggestions_list():
    company_id = effective_company_id()
    conn = get_db(); _ensure_wa_suggestions(conn)
    rows = conn.execute(
        "SELECT s.id, s.candidate_id, s.kind, s.message, s.reason, s.created_at, "
        "COALESCE(s.category,'followup') category, c.name cand_name, c.phone cand_phone "
        "FROM wa_suggestions s JOIN candidates c ON c.id=s.candidate_id "
        "WHERE s.company_id=? AND s.status='pending' ORDER BY s.id DESC LIMIT 100",
        (company_id,)).fetchall()
    conn.close()
    return jsonify({'ok': True, 'suggestions': [dict(r) for r in rows]})


@app.route('/api/wa-suggestions/<int:sid>/approve', methods=['POST'])
@login_required
def wa_suggestion_approve(sid):
    company_id = effective_company_id()
    conn = get_db(); _ensure_wa_suggestions(conn)
    s = conn.execute("SELECT * FROM wa_suggestions WHERE id=? AND company_id=? AND status='pending'",
                     (sid, company_id)).fetchone()
    if not s:
        conn.close(); return jsonify({'error': 'Not found'}), 404
    cand = conn.execute('SELECT * FROM candidates WHERE id=? AND owner_id=?',
                        (s['candidate_id'], company_id)).fetchone()
    if not cand:
        conn.close(); return jsonify({'error': 'Candidate not found'}), 404

    category = None
    try:
        category = s['category']
    except Exception:
        category = 'followup'
    comment = ((request.json or {}).get('comment') or '').strip()

    if s['kind'] == 'followup':
        # allow the user to tweak the message before approving
        text = ((request.json or {}).get('message') or s['message'] or '').strip()
        digits = re.sub(r'[^0-9]', '', cand['phone'] or '')
        if not digits:
            conn.close(); return jsonify({'error': 'No phone number'}), 400
        if len(digits) == 10:
            digits = '91' + digits
        _ensure_wa_outbox(conn)
        conn.execute(
            'INSERT INTO wa_messages (conversation_id,direction,sender,content,message_type,'
            'wa_message_id,created_at) VALUES (?,?,?,?,?,?,?)',
            (s['conversation_id'], 'outbound', 'You', text, 'text', '', ts()))
        conn.execute('UPDATE wa_conversations SET last_message_at=?, updated_at=? WHERE id=?',
                     (ts(), ts(), s['conversation_id']))
        conn.execute('INSERT INTO wa_outbox (company_id,candidate_id,phone,text,status,created_at) '
                     'VALUES (?,?,?,?,?,?)', (company_id, s['candidate_id'], digits, text, 'pending', ts()))
        # LEARNING: did the recruiter send as-is or edit it?
        act = 'approved' if text == (s['message'] or '').strip() else 'edited'
        _wa_record_feedback(conn, company_id, s['candidate_id'], s['conversation_id'],
                            category or 'followup', s['message'] or '', text, act, comment)
    elif s['kind'] == 'not_interested':
        old = cand['stage']
        conn.execute('UPDATE candidates SET stage=?, updated_at=? WHERE id=? AND owner_id=?',
                     ('Not Interested', ts(), s['candidate_id'], company_id))
        conn.execute('INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) '
                     'VALUES (?,?,?,?,?)', (s['candidate_id'], old, 'Not Interested',
                                            'Auto-moved (AI: ' + (s['reason'] or 'not interested') + ')', ts()))
        _wa_record_feedback(conn, company_id, s['candidate_id'], s['conversation_id'],
                            'not_interested', 'move to Not Interested', 'moved', 'approved', comment)
    conn.execute("UPDATE wa_suggestions SET status='sent' WHERE id=?", (sid,))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/wa-suggestions/<int:sid>/dismiss', methods=['POST'])
@login_required
def wa_suggestion_dismiss(sid):
    company_id = effective_company_id()
    conn = get_db(); _ensure_wa_suggestions(conn)
    s = conn.execute("SELECT * FROM wa_suggestions WHERE id=? AND company_id=?", (sid, company_id)).fetchone()
    if s:
        try: cat = s['category'] or 'followup'
        except Exception: cat = 'followup'
        _wa_record_feedback(conn, company_id, s['candidate_id'], s['conversation_id'],
                            cat, s['message'] or '', '', 'rejected',
                            ((request.json or {}).get('comment') or '').strip())
    conn.execute("UPDATE wa_suggestions SET status='dismissed' WHERE id=? AND company_id=?",
                 (sid, company_id))
    conn.commit(); conn.close()
    return jsonify({'ok': True})


@app.route('/api/wa-suggestions/scan-now', methods=['POST'])
@login_required
def wa_suggestions_scan_now():
    """Manually trigger a scan (for testing / on-demand)."""
    try:
        _wa_followup_scan(force=True)
    except Exception as e:
        return jsonify({'error': str(e)[:120]}), 500
    return wa_suggestions_list()


@app.route('/api/candidates/<int:cid>/jd-text')
@login_required
def candidate_jd_text(cid):
    """The candidate's mandate JD as plain text, with CTC/salary lines stripped
    (for the 'Share JD' WhatsApp template)."""
    conn = get_db()
    cand = conn.execute('SELECT mandate_id FROM candidates WHERE id=? AND owner_id=?',
                        (cid, effective_company_id())).fetchone()
    if not cand or not cand['mandate_id']:
        conn.close(); return jsonify({'ok': True, 'jd': ''})
    m = conn.execute('SELECT * FROM mandates WHERE id=? AND owner_id=?',
                     (cand['mandate_id'], effective_company_id())).fetchone()
    conn.close()
    if not m:
        return jsonify({'ok': True, 'jd': ''})
    jd = html_to_text(m['jd']) if m['jd'] else ''
    if jd.strip():
        # Sentence/clause-level strip: drop only the bits mentioning
        # compensation, keep the rest of the JD intact.
        _kws = ['ctc', 'salary', 'compensation', 'package', 'lpa', 'lakh', 'stipend',
                'remuneration', 'pay range', 'budget', 'per annum', 'in-hand', 'take home']
        parts = re.split(r'(?<=[.\n;])\s+', jd)
        jd = ' '.join([p for p in parts if not any(k in p.lower() for k in _kws)]).strip()
    return jsonify({'ok': True, 'jd': jd.strip()})


# ── Agent learning: capture how the recruiter approves/edits AI drafts, so the
#    agent improves over time and we can track per-situation confidence. ──────
WA_CATEGORIES = [
    ('intro',                'Intro / first outreach'),
    ('followup',             'Follow-up / nudge'),
    ('share_jd',             'Sharing the JD'),
    ('cv_request',           'CV / profile request'),
    ('availability',         'Availability / schedule a call'),
    ('ctc_salary',           'CTC / salary questions'),
    ('company_query',        'Which company / client'),
    ('notice_period',        'Notice period'),
    ('location_remote',      'Location / relocation / remote'),
    ('interview_invite',     'Interview invite / scheduling'),
    ('interview_reminder',   'Interview reminder / reschedule / no-show'),
    ('offer',                'Offer / closing'),
    ('counteroffer',         'Counteroffer handling'),
    ('rejection',            'Rejection (graceful)'),
    ('not_interested',       'Not interested / decline'),
    ('pre_joining',          'Pre-joining / keep-warm'),
    ('reactivation',         'Re-engage old candidate'),
    ('general',              'General / other'),
]
WA_CATEGORY_IDS = [c[0] for c in WA_CATEGORIES]


def _ensure_wa_feedback(conn):
    conn.execute('''CREATE TABLE IF NOT EXISTS wa_agent_feedback (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        company_id INTEGER, candidate_id INTEGER, conversation_id INTEGER,
        category TEXT, agent_draft TEXT, sent_text TEXT,
        action TEXT, comment TEXT, created_at TEXT, resolved_at TEXT)''')
    conn.execute('CREATE INDEX IF NOT EXISTS idx_wa_fb ON wa_agent_feedback(company_id,category,action)')


def _wa_learning_block(conn, company_id, category):
    """Recent corrections the recruiter made in this category — fed back into the
    prompt so the agent 'learns' from every edit immediately."""
    try:
        rows = conn.execute(
            "SELECT agent_draft, sent_text, comment FROM wa_agent_feedback "
            "WHERE company_id=? AND category=? AND action IN ('edited','approved') "
            "AND sent_text IS NOT NULL AND sent_text!='' ORDER BY id DESC LIMIT 6",
            (company_id, category or 'general')).fetchall()
    except Exception:
        return ''
    ex = []
    for r in rows:
        drafted = (r['agent_draft'] or '').strip()
        sent = (r['sent_text'] or '').strip()
        if not sent:
            continue
        if drafted and drafted != sent:
            block = "You drafted: " + drafted + "\nI actually sent: " + sent
        else:
            block = "Good draft (I sent as-is): " + sent
        if r['comment']:
            block += "\n(Why: " + r['comment'] + ")"
        ex.append(block)
    if not ex:
        return ''
    return ("\n\nLearn from how I corrected earlier drafts in this situation — match "
            "these patterns:\n" + "\n---\n".join(ex) + "\n")


def _wa_record_feedback(conn, company_id, candidate_id, conv_id, category,
                        agent_draft, sent_text, action, comment=''):
    _ensure_wa_feedback(conn)
    conn.execute(
        "INSERT INTO wa_agent_feedback (company_id,candidate_id,conversation_id,category,"
        "agent_draft,sent_text,action,comment,created_at,resolved_at) VALUES (?,?,?,?,?,?,?,?,?,?)",
        (company_id, candidate_id, conv_id, category or 'general', agent_draft or '',
         sent_text or '', action, comment or '', ts(), ts()))


@app.route('/api/wa-skillboard')
@login_required
def wa_skillboard():
    """Per-situation confidence: how often the recruiter approved the agent's
    draft as-is vs edited vs rejected. Drives the 'auto-ready' indicator."""
    company_id = effective_company_id()
    conn = get_db(); _ensure_wa_feedback(conn)
    rows = conn.execute(
        "SELECT category, action, COUNT(*) n FROM wa_agent_feedback WHERE company_id=? "
        "AND action IN ('approved','edited','rejected') GROUP BY category, action",
        (company_id,)).fetchall()
    agg = {}
    for r in rows:
        c = r['category'] or 'general'
        agg.setdefault(c, {'approved': 0, 'edited': 0, 'rejected': 0})
        agg[c][r['action']] = r['n']
    auto = []
    raw = _tenant_setting_raw(conn, company_id, 'wa_auto_categories', '')
    if raw:
        try: auto = json.loads(raw)
        except Exception: auto = []
    conn.close()
    board = []
    for cid_, label in WA_CATEGORIES:
        a = agg.get(cid_, {'approved': 0, 'edited': 0, 'rejected': 0})
        total = a['approved'] + a['edited'] + a['rejected']
        conf = round(100 * a['approved'] / total) if total else 0
        board.append({'id': cid_, 'label': label, 'approved': a['approved'], 'edited': a['edited'],
                      'rejected': a['rejected'], 'total': total, 'confidence': conf,
                      'auto_ready': (total >= 5 and conf >= 90), 'auto_on': (cid_ in auto)})
    return jsonify({'ok': True, 'board': board})


@app.route('/api/wa-auto-categories', methods=['POST'])
@login_required
def wa_auto_categories():
    cats = (request.json or {}).get('categories')
    if not isinstance(cats, list):
        return jsonify({'error': 'categories must be a list'}), 400
    cats = [c for c in cats if c in WA_CATEGORY_IDS]
    set_setting('wa_auto_categories', json.dumps(cats))
    return jsonify({'ok': True})


@app.route('/api/wa-inbound-config')
@login_required
def wa_inbound_config():
    """Config the recruiter pastes into their listener (URL + token)."""
    tok = _wa_get_inbound_token(effective_company_id())
    base = request.host_url.rstrip('/')
    return jsonify({'ok': True, 'url': base + '/api/wa-inbound', 'token': tok})


@app.route('/api/wa-templates', methods=['GET', 'POST'])
@login_required
def wa_templates():
    if request.method == 'POST':
        data = request.json or {}
        tpls = data.get('templates')
        if not isinstance(tpls, list):
            return jsonify({'error': 'templates must be a list'}), 400
        set_setting('wa_templates', json.dumps(tpls))
        return jsonify({'ok': True})
    raw = get_setting('wa_templates', '')
    if raw:
        try:
            return jsonify({'ok': True, 'templates': json.loads(raw)})
        except Exception:
            pass
    return jsonify({'ok': True, 'templates': WA_DEFAULT_TEMPLATES})


def _submission_token(mid):
    """Unguessable per-mandate token for the public client-submission link.
    Derived via HMAC(secret, mid) — no storage needed, and an attacker cannot
    forge it without the server secret."""
    key = app.secret_key or 'fallback'
    if isinstance(key, str):
        key = key.encode()
    return hmac.new(key, ('submission:%d' % mid).encode(), hashlib.sha256).hexdigest()[:24]


@app.route('/api/mandates/<int:mid>/share-link')
@login_required
def mandate_share_link(mid):
    """Return the public, tokenised client-submission URL for a mandate the
    current tenant owns. Recruiters share THIS link with clients."""
    conn = get_db()
    if not _tenant_owns_mandate(conn, mid):
        conn.close(); return jsonify({'error': 'Not found'}), 404
    conn.close()
    base = request.host_url.rstrip('/')
    url = base + '/api/mandates/' + str(mid) + '/submission?stage=Shared+with+Client&t=' + _submission_token(mid)
    return jsonify({'ok': True, 'url': url})


@app.route('/api/mandates/<int:mid>/submission')
def client_submission(mid):
    conn = get_db()
    m = conn.execute('SELECT * FROM mandates WHERE id=?', (mid,)).fetchone()
    if not m:
        conn.close(); return ('Not found', 404)
    # Access = logged-in owner (recruiter preview) OR a valid share token
    # (client link). Anything else 404s, so mandate ids can't be enumerated to
    # read another agency's shared candidates.
    tok = request.args.get('t', '')
    is_owner = bool(current_user()) and m['owner_id'] == effective_company_id()
    if not is_owner and not (tok and secrets.compare_digest(tok, _submission_token(mid))):
        conn.close(); return ('Not found', 404)
    stage_filter = request.args.get('stage', 'Shared with Client')
    cands = conn.execute(
        'SELECT * FROM candidates WHERE mandate_id=? AND stage=? ORDER BY ai_score DESC',
        (mid, stage_filter)).fetchall()
    conn.close()

    date_str = datetime.date.today().strftime('%d %b %Y')
    rows_html = ''
    for i, c in enumerate(cands):
        skills_list = json.loads(c['key_skills'] or '[]')
        skills_str = ', '.join(skills_list[:5]) if skills_list else '--'
        bg = '#fafafa' if i % 2 else '#ffffff'
        ctc_curr = ('Rs ' + str(int(c['ctc_current'])) + 'L') if c['ctc_current'] else '--'
        ctc_exp  = ('Rs ' + str(int(c['ctc_expected'])) + 'L') if c['ctc_expected'] else '--'
        notice   = (str(c['notice_period']) + 'd') if c['notice_period'] else '--'
        rows_html += (
            '<tr style="background:' + bg + ';border-bottom:0.5px solid #f0f0f0">'
            '<td style="text-align:center;padding:9px 8px;font-weight:500;color:#888">' + str(i+1) + '</td>'
            '<td style="padding:9px 8px"><div style="font-weight:600;font-size:12px">' + (c['name'] or '--') + '</div>'
            '<div style="font-size:10px;color:#666;margin-top:2px">' + (c['designation'] or '--') + '</div></td>'
            '<td style="padding:9px 8px">' + (c['company'] or '--') + '</td>'
            '<td style="padding:9px 8px;text-align:center;font-weight:500;white-space:nowrap">' + ctc_curr + '</td>'
            '<td style="padding:9px 8px;text-align:center;font-weight:500;white-space:nowrap;color:#1D9E75">' + ctc_exp + '</td>'
            '<td style="padding:9px 8px;text-align:center;white-space:nowrap">' + notice + '</td>'
            '<td style="padding:9px 8px">' + (c['location'] or '--') + '</td>'
            '<td style="padding:9px 8px;font-size:10px;color:#555">' + skills_str + '</td>'
            '<td style="padding:9px 8px;font-size:10px;color:#444;max-width:180px">' + (c['career_summary'] or c['ai_reasoning'] or '--') + '</td>'
            '</tr>'
        )

    no_cands = '<div style="padding:20px;text-align:center;color:#888;font-size:12px;border:1px dashed #ddd;border-radius:6px">No candidates in <strong>' + stage_filter + '</strong> stage yet.</div>' if not cands else ''

    html = (
        '<!DOCTYPE html><html><head><meta charset="UTF-8">'
        '<title>HireLab Client Submission</title>'
        '<style>'
        '*{box-sizing:border-box;margin:0;padding:0}'
        'body{font-family:-apple-system,"Segoe UI",sans-serif;font-size:11px;color:#1a1a1a;background:#fff;padding:32px}'
        'table{width:100%;border-collapse:collapse;margin-bottom:20px}'
        'thead{background:#0a2540;color:#fff;-webkit-print-color-adjust:exact;print-color-adjust:exact}'
        'th{padding:9px 8px;text-align:left;font-size:10px;font-weight:500;letter-spacing:.3px}'
        '@media print{.no-print{display:none}body{padding:12px}}'
        '</style></head><body>'

        # Print bar
        '<div class="no-print" style="background:#0a2540;color:#fff;padding:10px 32px;margin:-32px -32px 24px;display:flex;align-items:center;gap:12px">'
        '<span style="font-size:13px;font-weight:500">Client Submission Sheet</span>'
        '<button onclick="window.print()" style="margin-left:auto;background:#1D9E75;color:#fff;border:none;border-radius:5px;padding:7px 16px;font-size:12px;cursor:pointer">Print / Save PDF</button>'
        '<button onclick="window.close()" style="background:rgba(255,255,255,.1);color:#fff;border:0.5px solid rgba(255,255,255,.3);border-radius:5px;padding:7px 14px;font-size:12px;cursor:pointer;margin-left:6px">Close</button>'
        '</div>'

        # Header
        '<div style="display:flex;align-items:flex-start;justify-content:space-between;margin-bottom:20px;padding-bottom:14px;border-bottom:2px solid #1D9E75">'
        '<div>'
        '<div style="font-size:18px;font-weight:700;color:#0a2540">HireLab <span style="color:#1D9E75">Talent Resource</span></div>'
        '<div style="font-size:10px;color:#888;margin-top:2px">Intelligence-Led Recruitment  |  Ghaziabad, NCR</div>'
        '</div>'
        '<div style="text-align:right">'
        '<div style="font-size:20px;font-weight:700;color:#0a2540">Candidate Submission</div>'
        '<div style="font-size:12px;color:#666;margin-top:3px">' + (m['role'] or '') + '  |  ' + (m['client'] or '') + '</div>'
        '<div style="font-size:10px;color:#aaa;margin-top:2px">Date: ' + date_str + '  |  CONFIDENTIAL</div>'
        '</div></div>'

        # Mandate info
        '<div style="background:#f9f9f9;border:0.5px solid #e8e8e8;border-left:3px solid #1D9E75;border-radius:6px;padding:12px 16px;margin-bottom:18px;display:grid;grid-template-columns:repeat(5,1fr);gap:12px">'
        '<div><div style="font-size:9px;color:#888;text-transform:uppercase;letter-spacing:.5px;margin-bottom:3px">Position</div><div style="font-size:12px;font-weight:500">' + (m['role'] or '') + '</div></div>'
        '<div><div style="font-size:9px;color:#888;text-transform:uppercase;letter-spacing:.5px;margin-bottom:3px">Client</div><div style="font-size:12px;font-weight:500">' + (m['client'] or '') + '</div></div>'
        '<div><div style="font-size:9px;color:#888;text-transform:uppercase;letter-spacing:.5px;margin-bottom:3px">Location</div><div style="font-size:12px;font-weight:500">' + (m['location'] or '--') + '</div></div>'
        '<div><div style="font-size:9px;color:#888;text-transform:uppercase;letter-spacing:.5px;margin-bottom:3px">CTC Range</div><div style="font-size:12px;font-weight:500">Rs ' + str(int(m['ctc_min'] or 0)) + '--' + str(int(m['ctc_max'] or 0)) + ' LPA</div></div>'
        '<div><div style="font-size:9px;color:#888;text-transform:uppercase;letter-spacing:.5px;margin-bottom:3px">Profiles Shared</div><div style="font-size:22px;font-weight:700;color:#1D9E75">' + str(len(cands)) + '</div></div>'
        '</div>'

        # Table
        '<table>'
        '<thead><tr>'
        '<th style="width:30px;text-align:center">#</th>'
        '<th style="min-width:120px">Candidate</th>'
        '<th>Current Company</th>'
        '<th style="text-align:center">Curr CTC</th>'
        '<th style="text-align:center">Exp CTC</th>'
        '<th style="text-align:center">Notice</th>'
        '<th>Location</th>'
        '<th>Key Skills</th>'
        '<th style="min-width:150px">Summary</th>'
        '</tr></thead>'
        '<tbody>' + rows_html + '</tbody>'
        '</table>'
        + no_cands +

        # Footer
        '<div style="display:flex;align-items:center;justify-content:space-between;padding-top:14px;border-top:0.5px solid #e8e8e8;font-size:10px;color:#aaa">'
        '<span style="background:#FAEEDA;color:#854F0B;padding:3px 10px;border-radius:4px;font-weight:500">CONFIDENTIAL | For ' + (m['client'] or '') + ' use only</span>'
        '<span>HireLab Talent Resource | GSTIN: 09ECWPP1647A1Z9 | UDYAM: UP-29-0178859</span>'
        '<span>' + date_str + '</span>'
        '</div>'
        '</body></html>'
    )
    return Response(html, mimetype='text/html')


@app.route('/api/export')
@admin_required
def export_data():
    conn = get_db()
    candidates = [dict(r) for r in conn.execute('SELECT * FROM candidates').fetchall()]
    # Strip platform secrets (API keys / passwords / tokens) from the backup so a
    # shared/leaked export file never exposes credentials.
    def _is_secret(k):
        kl = (k or '').lower()
        if kl.endswith('api_key') or kl.endswith('_apikey'): return True
        if 'secret' in kl or 'password' in kl or kl.endswith('_pass') or kl.endswith('_token'): return True
        return kl in ('claude_api_key','deepseek_api_key','groq_api_key','openai_api_key',
                      'gemini_api_key','anthropic_api_key','smtp_pass','smtp_password',
                      'wa_token','wa_access_token','verify_token','flask_secret_key','secret_key')
    settings_out = {r['key']: r['value'] for r in conn.execute('SELECT * FROM settings').fetchall()
                    if not _is_secret(r['key'])}
    data = {
        'exported_at': ts(), 'app': 'HireLab Screener', 'version': '2.1',
        'mandates':   [dict(r) for r in conn.execute('SELECT * FROM mandates').fetchall()],
        'candidates': candidates,
        'history':    [dict(r) for r in conn.execute('SELECT * FROM stage_history').fetchall()],
        'settings':   settings_out,
    }
    conn.close()
    # Include actual CV files (PDF/Word) as base64 so they transfer with the backup.
    import base64 as _b64
    cv_files = {}
    for cand in candidates:
        cvp = cand.get('cv_path')
        if cvp:
            fpath = os.path.join(CV_DIR, cvp)
            if os.path.exists(fpath):
                try:
                    with open(fpath, 'rb') as _cf:
                        cv_files[cvp] = _b64.b64encode(_cf.read()).decode('ascii')
                except Exception:
                    pass
    data['cv_files'] = cv_files
    fname = 'hirelab_' + str(datetime.date.today()) + '.json'
    return Response(json.dumps(data, indent=2, ensure_ascii=False), mimetype='application/json',
                    headers={'Content-Disposition': 'attachment; filename=' + fname})

@app.route('/api/import', methods=['POST'])
@login_required
def import_data():
    import time
    # Ensure DB is initialized before import
    try:
        init_db()
    except Exception:
        pass

    for _attempt in range(5):
        try:
            data = request.json or {}
            if not data.get('mandates') and not data.get('candidates'):
                return jsonify({'error': 'Invalid backup file. Must be a HireLab JSON export.'}), 400
            conn = get_db(); c = conn.cursor()
            n = ts(); mid_map = {}; cid_map = {}; m_done = cand_done = hist_done = 0
            for k, v in (data.get('settings') or {}).items():
                c.execute('INSERT OR REPLACE INTO settings (key,value) VALUES (?,?)', (k, str(v)))
            for m in (data.get('mandates') or []):
                old_id = m.get('id')
                c.execute('INSERT INTO mandates (client,role,location,division,ctc_min,ctc_max,jd,sop_text,sop_version,sop_changelog,status,created_at,owner_id) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)',
                          (m.get('client',''), m.get('role',''), m.get('location',''), m.get('division',''),
                           float(m.get('ctc_min') or 0), float(m.get('ctc_max') or 0), m.get('jd',''),
                           m.get('sop_text',''), m.get('sop_version', 1), m.get('sop_changelog', '[]'),
                           m.get('status', 'active'), m.get('created_at') or n, effective_user_id()))
                mid_map[old_id] = c.lastrowid; m_done += 1
            for cand in (data.get('candidates') or []):
                old_id = cand.get('id')
                new_mid = mid_map.get(cand.get('mandate_id'), cand.get('mandate_id'))
                c.execute(
                    'INSERT INTO candidates (mandate_id,name,company,designation,experience,ctc_current,'
                    'ctc_expected,notice_period,location,phone,email,qualification,key_skills,secondary_skills,'
                    'career_summary,industry_background,is_mnc,screening_decision,ai_score,ai_reasoning,'
                    'stage,recruiter_feedback,client_feedback,general_comments,cv_path,cv_original_name,'
                    'msg1_sent_at,fu1_sent_at,fu2_sent_at,created_at,updated_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
                    (new_mid, cand.get('name',''), cand.get('company',''), cand.get('designation',''),
                     float(cand.get('experience') or 0), float(cand.get('ctc_current') or 0),
                     float(cand.get('ctc_expected') or 0), int(cand.get('notice_period') or 0),
                     cand.get('location',''), cand.get('phone',''), cand.get('email',''), cand.get('qualification',''),
                     json.dumps(cand.get('key_skills') if isinstance(cand.get('key_skills'), list) else json.loads(cand.get('key_skills') or '[]')),
                     json.dumps(cand.get('secondary_skills') if isinstance(cand.get('secondary_skills'), list) else json.loads(cand.get('secondary_skills') or '[]')),
                     cand.get('career_summary',''),
                     cand.get('industry_background',''), cand.get('is_mnc', 0), cand.get('screening_decision',''),
                     float(cand.get('ai_score') or 0), cand.get('ai_reasoning',''), cand.get('stage','Screening'),
                     cand.get('recruiter_feedback',''), cand.get('client_feedback',''), cand.get('general_comments',''),
                     cand.get('cv_path',''), cand.get('cv_original_name',''),
                     cand.get('msg1_sent_at',''), cand.get('fu1_sent_at',''), cand.get('fu2_sent_at',''),
                     cand.get('created_at') or n, cand.get('updated_at') or n))
                cid_map[old_id] = c.lastrowid; cand_done += 1
            for h in (data.get('history') or []):
                new_cid = cid_map.get(h.get('candidate_id'))
                if new_cid:
                    c.execute('INSERT INTO stage_history (candidate_id,from_stage,to_stage,note,created_at) VALUES (?,?,?,?,?)',
                              (new_cid, h.get('from_stage',''), h.get('to_stage',''), h.get('note',''), h.get('created_at') or n))
                    hist_done += 1
            conn.commit(); conn.close()
            # Restore CV files (PDF/Word) that were embedded in the backup as base64.
            cv_restored = 0
            cv_files = data.get('cv_files') or {}
            if cv_files:
                import base64 as _b64
                os.makedirs(CV_DIR, exist_ok=True)
                for _fname, _b64data in cv_files.items():
                    try:
                        _dest = os.path.join(CV_DIR, _fname)
                        if not os.path.exists(_dest):   # don't overwrite existing
                            with open(_dest, 'wb') as _wf:
                                _wf.write(_b64.b64decode(_b64data))
                        cv_restored += 1
                    except Exception:
                        pass
            return jsonify({'ok': True, 'mandates': m_done, 'candidates': cand_done,
                            'history': hist_done, 'cvs': cv_restored})
        except sqlite3.OperationalError as e:
            if 'locked' in str(e).lower() and _attempt < 4:
                time.sleep(2)
                continue
            return jsonify({'error': 'DB locked: ' + str(e)}), 503
        except Exception as e:
            import traceback
            return jsonify({'error': str(e), 'detail': traceback.format_exc()[-500:]}), 500
    return jsonify({'error': 'Import failed after 5 retries'}), 500




# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  CENTRAL DATABASE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


# ── Startup: runs both with gunicorn AND python server.py ──────────────────────
# This ensures DB tables exist regardless of how the app is started
try:
    migrate_old()
    init_db()
    # ── Mount RecruitOS platform modules (CRM, etc.) as Flask blueprints ──
    try:
        import modules
        modules.register_all(app)
    except Exception as _mod_err:
        print(f'[modules] registration skipped: {_mod_err}')
    # Safety net order matters:
    # 1) check persistence (writes/reads a marker that proves the disk survives restarts)
    # 2) auto-restore if the live DB came up empty but a backup has users
    # 3) take a fresh backup of the (now-healthy) DB
    _PERSISTENCE = check_storage_persistence()
    auto_restore_if_empty()
    daily_backup()
    # Start the reminder push-notification scheduler (background thread)
    try:
        _start_reminder_scheduler()
    except Exception as _sched_err:
        print(f'[reminder-scheduler] failed to start: {_sched_err}')
    # Start the async embedding queue worker (background thread)
    try:
        _start_embedding_worker()
    except Exception as _emb_err:
        print(f'[embed-worker] failed to start: {_emb_err}')
    _ucount = _db_user_count(DB_PATH)
    print('\n' + '=' * 56)
    print('  HireLab Screener — startup')
    print('  DATA_DIR : ' + DATA_DIR)
    print('  DB_PATH  : ' + DB_PATH)
    print('  Users in DB: ' + str(_ucount))
    if _PERSISTENCE.get('persistent') is True:
        print(f'  Storage  : PERSISTENT ✓ (survived {_PERSISTENCE.get("boots_seen")} restarts)')
    else:
        print('  Storage  : NOT YET CONFIRMED persistent (first boot, or marker was wiped)')
        if DATA_DIR.rstrip('/').endswith('HireLab') or 'expanduser' in DATA_DIR:
            print('  *** WARNING: DATA_DIR is NOT the mounted disk. On Render you must set')
            print('  *** DATA_DIR=/data AND attach a persistent disk mounted at /data,')
            print('  *** otherwise ALL data is lost on every restart/spin-down. ***')
    print('=' * 56 + '\n')
except Exception as _startup_err:
    print(f'Startup init warning: {_startup_err}')

# ── SAFE DATA RESET (controlled by env var) ────────────────────────────────────
# Set RESET_DATA=yes in the host environment (e.g. Render) to wipe all mandates,
# candidates, reminders, work history and stage history on the next start.
# IMPORTANT: remove the variable again right after, so it does not wipe on every
# restart. User accounts are preserved unless RESET_DATA=all is used.
try:
    _reset = (os.environ.get('RESET_DATA') or '').strip().lower()
    if _reset in ('yes', 'all', '1', 'true'):
        # SAFETY: only run a given reset value ONCE, ever. We record which reset
        # token was last executed in a marker file on disk. If the env var still
        # holds the same value on the next restart, we SKIP it — so forgetting to
        # remove the variable can never wipe data again.
        _marker = os.path.join(DATA_DIR, '.last_reset')
        _already = ''
        try:
            if os.path.exists(_marker):
                with open(_marker) as _f: _already = _f.read().strip()
        except Exception: pass
        # Build a unique token: value + a user-supplied tag so the same 'yes' won't
        # re-run unless the user changes RESET_TAG too.
        _tag = (os.environ.get('RESET_TAG') or '').strip()
        _token = _reset + '|' + _tag
        if _token == _already:
            print(f'*** RESET_DATA={_reset} SKIPPED — already executed (token unchanged). Safe. ***')
        else:
            _conn = get_db(); _c = _conn.cursor()
            for _tbl in ['candidates', 'mandates', 'reminders', 'work_history',
                         'stage_history', 'submissions', 'activity_log']:
                try: _c.execute(f'DELETE FROM {_tbl}')
                except Exception: pass
            if _reset == 'all':
                try: _c.execute('DELETE FROM users')
                except Exception: pass
            _conn.commit(); _conn.close()
            try:
                with open(_marker, 'w') as _f: _f.write(_token)
            except Exception: pass
            print(f'*** RESET_DATA={_reset} executed ONCE — data cleared. Will NOT repeat. ***')
except Exception as _reset_err:
    print(f'Reset warning: {_reset_err}')

if __name__ == '__main__':
    check_timers()
    print('Local server: http://localhost:' + str(os.environ.get('PORT', 5000)))
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)), debug=False)
